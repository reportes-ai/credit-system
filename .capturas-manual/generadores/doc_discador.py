# -*- coding: utf-8 -*-
"""Documento técnico-funcional — Discador AF Dialer (predictivo/asistido/potencia)."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'AF DIALER', bold=True, color=AZUL_OSCURO, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'Discador predictivo, asistido y de potencia — servicio independiente, plug-in de Business Suite', color=AZUL, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Documento técnico y funcional', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Septiembre 2026 — uso interno', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Decisión de diseño ──
    h1(doc, '1. Decisión de diseño: servicio aparte, no un módulo')
    p(doc, 'El discador se construye como un SERVICIO INDEPENDIENTE (AF Dialer) que conversa con Business '
           'Suite por API y webhooks — no como código dentro del monolito. Razones:')
    vineta(doc, 'la telefonía (SIP, WebRTC, detección de contestador, pacing) es un dominio propio con su propio ciclo de vida y sus propios proveedores; mezclarla con el core de crédito acopla dos mundos que evolucionan distinto.', bold_hasta='Dominios separados: ')
    vineta(doc, 'el mismo discador sirve para los nodos de Ecuador y Bolivia, para otras apps del grupo, e incluso como producto comercializable — solo cambia la app que le manda listas y recibe resultados.', bold_hasta='Reutilizable: ')
    vineta(doc, 'si el proveedor de voz cambia (precio, calidad), se cambia dentro del Dialer sin tocar ninguna app cliente.', bold_hasta='Proveedor intercambiable: ')
    vineta(doc, 'una caída del discador no toca la operación de crédito, y viceversa.', bold_hasta='Aislamiento de fallas: ')
    regla(doc, 'El contrato de integración es el mismo estándar que ya usa la Suite: API REST con X-API-Key '
               '(mantenedor de APIs) hacia el Dialer, y webhooks firmados del Dialer hacia la app cliente. '
               'Business Suite es el PRIMER cliente, no el único.')

    # ── 2. Los tres modos ──
    h1(doc, '2. Funcional: los tres modos de discado')
    tabla(doc, ('Modo', 'Cómo funciona', 'Cuándo usarlo'),
          (('PREVIEW (asistido)', 'El agente ve la ficha del contacto ANTES de llamar y dispara la llamada con un clic. Ritmo lo pone el agente.', 'Cobranza dura, casos delicados, clientes de alto valor'),
           ('POWER (potencia)', 'El sistema marca automáticamente el siguiente de la lista apenas el agente queda libre (1 llamada por agente, ratio 1:1). Sin tiempo muerto entre llamadas.', 'Campañas comerciales y cobranza temprana — el mejor punto de partida'),
           ('PREDICTIVO', 'El sistema marca MÁS llamadas que agentes libres (ratio dinámico 1.2–2.5:1) prediciendo cuántas contestarán; conecta solo las contestadas. Máxima productividad, riesgo de llamadas abandonadas.', 'Listas grandes de baja contactabilidad, con equipo de 5+ agentes')),
          (3.2, 8.3, 5.0))
    advertencia(doc, 'El modo predictivo genera LLAMADAS ABANDONADAS (contestó una persona y no había agente '
                     'libre). La tasa de abandono se limita por parámetro (estándar internacional: ≤3%) y '
                     'en cobranza cada intento debe respetar los topes legales de contacto. Por eso el '
                     'roadmap parte con Preview + Power y deja Predictivo para la fase 3, con métricas reales.')
    h2(doc, '2.1 Funcionalidades de la consola del agente')
    vineta(doc, 'softphone WebRTC EN EL NAVEGADOR: sin teléfonos físicos ni anexos — audífono y listo.')
    vineta(doc, 'ficha del contacto en pantalla al conectar (nombre, deuda/campaña, historial de gestiones) con tipificación obligatoria al cortar (contactado, compromiso de pago, no contesta, número malo, volver a llamar…).')
    vineta(doc, 'agenda de rellamadas personales, transferencia a supervisor, y modo escucha/susurro para supervisión.')
    vineta(doc, 'grabación de llamadas con retención configurable y acceso auditado.')
    h2(doc, '2.2 Funcionalidades de campaña (supervisor)')
    vineta(doc, 'campañas con lista, horario permitido, intentos máximos por contacto, orden de discado (prioridad, deciles) y reciclaje de no contestados.')
    vineta(doc, 'tablero en vivo: agentes conectados, llamadas en curso, contactabilidad, abandono, tiempos medios; y reporte por campaña/agente/día.')
    vineta(doc, 'detección de contestador (AMD) con política configurable: cortar, dejar mensaje grabado o conectar igual.')
    vineta(doc, 'lista de exclusión (no llamar) global y por campaña.')

    # ── 3. Integración con Business Suite ──
    h1(doc, '3. Integración con Business Suite (el plug-in)')
    flujo(doc, 'BS Campañas/Cobranza → API Dialer (lista + reglas) → discado → webhook resultado → gestión registrada en BS')
    vineta(doc, 'las Campañas Masivas y la Segmentación por Tramos de Cobranza EXPORTAN la lista al Dialer por API (contacto, teléfono, prioridad, datos de ficha). La Suite sigue siendo la fuente única de a quién llamar.', bold_hasta='Origen de listas: ')
    vineta(doc, 'cada llamada terminada vuelve por webhook: resultado, tipificación, duración, grabación. La Suite la registra como GESTIÓN en la bitácora del crédito/campaña — cuenta contra los topes legales de contacto igual que una gestión manual.', bold_hasta='Resultados: ')
    vineta(doc, 'la Suite calcula la disponibilidad legal (gestiones de la semana) ANTES de exportar la lista: al Dialer solo llegan contactos que se pueden llamar.', bold_hasta='Cumplimiento: ')
    vineta(doc, 'el agente puede abrir la ficha completa en la Suite con un clic desde la consola (deep-link con el token de su sesión).', bold_hasta='Ficha: ')

    # ── 4. Arquitectura técnica ──
    h1(doc, '4. Arquitectura técnica')
    h2(doc, '4.1 Componentes')
    tabla(doc, ('Componente', 'Tecnología propuesta', 'Rol'),
          (('Dialer Core (API)', 'Node.js + base propia (TiDB/Postgres)', 'Campañas, colas, pacing, tipificaciones, reportes, webhooks'),
           ('Capa de voz', 'Proveedor CPaaS (Twilio Voice o similar) — fase 1', 'Originación de llamadas, AMD, grabación, WebRTC del agente'),
           ('Alternativa de voz', 'Asterisk/FreeSWITCH + troncal SIP local — fase 3 (optimización de costo)', 'Mismo rol con minutos hasta 5–10 veces más baratos, a cambio de administrar el conmutador'),
           ('Consola del agente', 'Web (mismo stack visual de la Suite) + SDK WebRTC del proveedor', 'Softphone, ficha, tipificación, agenda'),
           ('Tablero supervisor', 'Web + WebSocket', 'Tiempo real de la operación'),
           ('Integración', 'REST con X-API-Key + webhooks firmados (HMAC)', 'Contrato estándar para cualquier app cliente')),
          (3.6, 6.2, 6.5))
    p(doc, 'Con CPaaS (fase 1) NO se administra ninguna central telefónica: el proveedor entrega llamadas, '
           'WebRTC, AMD y grabación por API. Es la misma filosofía sin-fierros de la Suite: pagar por uso, '
           'partir en días y cambiar de proveedor si conviene.')
    h2(doc, '4.2 El motor de pacing (corazón del predictivo)')
    vineta(doc, 'en Power: 1 marcación por agente libre. Punto.', bold_hasta='Regla base: ')
    vineta(doc, 'en Predictivo: ratio = f(contactabilidad última hora, duración media de llamada, agentes conectados), recalculado cada 30–60 segundos, con techo por la tasa de abandono objetivo (≤3%). Si el abandono sube, el ratio baja solo.', bold_hasta='Ratio dinámico: ')
    vineta(doc, 'toda llamada, intento y abandono queda en el log del Dialer — el reporte de cumplimiento sale de ahí.', bold_hasta='Trazabilidad: ')
    h2(doc, '4.4 Detección de tonos y voz (AMD) — el problema más complejo, resuelto por capas')
    p(doc, 'Distinguir en segundos si contestó una persona, un contestador, un fax o un tono de red es lo '
           'más difícil de un discador — y donde mueren los desarrollos caseros. AF Dialer NO construye '
           'procesamiento de señal propio: lo resuelve en tres capas, igual que las plataformas enterprise '
           'modernas.')
    tabla(doc, ('Capa', 'Qué detecta', 'Cómo se resuelve'),
          (('1. Tonos de red', 'Ocupado, número fuera de servicio (tonos SIT), fax, congestión, no contesta; DTMF (teclas)', 'Viene RESUELTO en la señalización del proveedor: llega como estado de la llamada. Cero desarrollo propio.'),
           ('2. AMD clásico (humano vs contestador)', 'Silencio inicial, duración del saludo ("¿aló?" corto y silencio = humano; parlamento largo y corrido = máquina), detección del beep', 'Integrado en el CPaaS (~US$ 0,0075 por llamada, modo asíncrono) con ~85–90% de acierto; en Asterisk, la aplicación AMD() con los mismos parámetros. Fase 1: comprado, no construido.'),
           ('3. AMD con IA (fase 2–3)', 'Clasificación por CONTENIDO: streaming de los primeros segundos a un modelo de voz — "deje su mensaje después del tono" es máquina; "¿aló, quién habla?" es humano. Detecta además el fin del beep para dejar mensajes grabados en el momento exacto', 'Modelos de audio en tiempo real (VAD + transcripción parcial): 95%+ de acierto. Es donde un discador propio de hoy supera a las licencias enterprise antiguas.')),
          (3.4, 5.9, 7.0))
    regla(doc, 'Principio de diseño: el pacing NO espera certeza. Conecta al agente con la clasificación '
               'probabilística y el agente confirma en un segundo — el AMD tiene que ser bueno, no perfecto. '
               'Perseguir el 100% de acierto es el error clásico de los discadores caseros.')
    h2(doc, '4.5 Asterisk: cuándo y por qué (la decisión de la fase 3)')
    p(doc, 'Asterisk (o FreeSWITCH) es la central open source con troncal SIP local — la alternativa al '
           'CPaaS. No es un "sí o no", es un "cuándo", y el número que gatilla la decisión sale del piloto:')
    vineta(doc, 'minuto a móvil ~50–80% más barato que el CPaaS, cero licencias, números de la telefónica local (mejor contactabilidad), AMD y grabación propios.', bold_hasta='A favor: ')
    vineta(doc, 'es infraestructura que ADMINISTRAR (va contra la filosofía sin-fierros): seguridad SIP crítica (el fraude telefónico por centrales mal aseguradas es el ataque N°1 del rubro), WebRTC a configurar a mano, contrato y enrolamiento con la telefónica local, y dependencia de alguien que sepa operarla.', bold_hasta='En contra: ')
    vineta(doc, 'si el piloto en CPaaS factura sobre ~US$ 600–800/mes en minutos, la troncal SIP se paga sola y se monta en fase 3 SIN tocar el Dialer Core — la capa de voz es intercambiable por diseño. Si el consumo es bajo, el CPaaS se queda y nunca se administró una central.', bold_hasta='Regla de decisión: ')
    h2(doc, '4.6 Seguridad y datos')
    vineta(doc, 'el Dialer guarda SOLO lo mínimo del contacto (nombre, teléfono, referencia externa); la ficha completa vive en la app cliente. Grabaciones en bucket con retención y acceso auditado.')
    vineta(doc, 'API keys por app cliente, webhooks firmados, usuarios de agente propios del Dialer o federados desde la Suite.')

    # ── 5. Roadmap y tiempos ──
    h1(doc, '5. Roadmap y tiempos (con holgura)')
    tabla(doc, ('Fase', 'Alcance', 'Duración'),
          (('1. MVP Preview + Power', 'Servicio core + CPaaS + consola de agente + campañas + webhook a BS + tablero básico. Piloto con cobranza temprana (2–3 agentes).', '6–8 semanas'),
           ('2. Supervisión y reportería', 'Escucha/susurro, grabaciones con auditoría, reportes por campaña/agente, AMD con mensaje de voz, integración completa con Campañas Masivas.', '3–4 semanas'),
           ('3. Predictivo + costo', 'Motor de pacing dinámico con techo de abandono; evaluación de troncal SIP local para bajar el costo por minuto.', '4–6 semanas'),
           ('TOTAL a discador completo', '', '≈ 3,5 a 4,5 meses')),
          (4.4, 8.6, 3.1))

    # ── 6. Costos ──
    h1(doc, '6. Costos estimados')
    tabla(doc, ('Concepto', 'Estimación', 'Nota'),
          (('Infraestructura del servicio', 'US$ 15–40 /mes', 'Servidor + BD + bucket de grabaciones (misma dieta de la Suite)'),
           ('Minutos de voz (CPaaS)', '~US$ 0,05–0,12 por minuto saliente a móvil en Chile', 'El costo dominante: 5 agentes × 3 h efectivas/día ≈ US$ 450–1.100 /mes'),
           ('Números telefónicos', 'US$ 1–5 /mes por número', 'Con identificador local (recomendado varios números rotativos)'),
           ('Optimización fase 3 (SIP local)', 'baja el minuto en 50–80%', 'A cambio de administrar Asterisk; se evalúa con el volumen real del piloto'),
           ('Licencias de software de discador', 'US$ 0', 'Se construye propio — un discador comercial cuesta US$ 25–60 por agente/mes')),
          (5.0, 4.8, 6.5))
    p(doc, 'Referencia de mercado: un discador comercial para 10 agentes cuesta US$ 250–600 mensuales SOLO '
           'de licencia, más los minutos. AF Dialer elimina la licencia, deja el costo casi puro en minutos '
           'de voz y queda como activo del grupo, reutilizable en cada país.')

    # ── 7. Riesgos ──
    h1(doc, '7. Riesgos y mitigaciones')
    tabla(doc, ('Riesgo', 'Mitigación'),
          (('Marcado de números como SPAM por las telefónicas', 'Varios números de origen rotativos, identificador local, volumen gradual y monitoreo de contactabilidad por número'),
           ('Abandono del predictivo fuera de norma', 'Techo paramétrico de abandono + partir en Power; Predictivo recién con métricas reales'),
           ('Cumplimiento en cobranza (topes de contacto)', 'La Suite filtra la lista ANTES de exportar; cada llamada vuelve como gestión y cuenta contra el tope'),
           ('Calidad de audio del WebRTC en la oficina', 'Prueba de ancho de banda en el piloto; audífonos USB decentes; el CPaaS trae diagnóstico de calidad'),
           ('Subestimar la telefonía local (fase SIP)', 'La fase 3 es opcional y se decide con datos del piloto — el CPaaS funciona desde el día uno')),
          (7.0, 9.5))

    # ── 8. Recomendación ──
    h1(doc, '8. Recomendación')
    p(doc, 'Construir AF Dialer como servicio independiente con CPaaS, partir con Preview + Power en un '
           'piloto de cobranza temprana (fase 1, 6–8 semanas), y decidir Predictivo y troncal SIP con los '
           'números reales del piloto. Business Suite se integra por el estándar de APIs que ya existe — y '
           'el discador queda listo para enchufarse a los nodos de Ecuador y Bolivia o a cualquier otra '
           'aplicación del grupo.', bold=True)

    return doc
