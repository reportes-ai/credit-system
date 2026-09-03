# -*- coding: utf-8 -*-
"""Documento Técnico de Seguridad de la Información, Arquitectura y Trazabilidad — Business Suite.
Versión extensa para auditoría externa. Describe ARQUITECTURA y CONTROLES; jamás secretos,
credenciales, cuentas especiales ni su resguardo."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'SEGURIDAD DE LA INFORMACIÓN, ARQUITECTURA Y TRAZABILIDAD', bold=True, color=AZUL_OSCURO, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'Documento Técnico para Auditoría', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Marco de control alineado a ISO/IEC 27001:2022, NIST CSF 2.0, OWASP ASVS,', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'CIS Controls v8, ISO 22301 y la legislación chilena aplicable.', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(5): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Septiembre 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz. Distribución restringida al equipo auditor.', size=10,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones y alcance')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Septiembre 2026', 'Business Suite', 'Emisión inicial para auditoría externa'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Objetivo')
    p(doc, 'Describir, con nivel técnico verificable, cómo AutoFácil Business Suite protege la '
           'confidencialidad, integridad y disponibilidad de la información que procesa; cómo está '
           'construida la plataforma; qué modelos y normas de referencia gobiernan su diseño; y cómo '
           'se garantiza la trazabilidad de cada hecho de negocio desde su origen hasta la '
           'contabilidad. Cada afirmación de este documento apunta a un artefacto concreto '
           '(módulo, tabla, middleware, motor o runbook) que el auditor puede solicitar como evidencia.')
    h2(doc, 'Alcance')
    vineta(doc, 'la aplicación web AutoFácil Business Suite (afbs.autofacilchile.cl), su API, sus 28 motores automáticos, sus integraciones externas y su infraestructura en nube.', bold_hasta='Dentro del alcance: ')
    vineta(doc, 'los sistemas de las financieras (AutoFin, Unidad de Crédito), los ERP contables heredados (AVSOFT) y los dispositivos de usuario final, que se tratan como terceros o entorno.', bold_hasta='Fuera del alcance: ')
    regla(doc, 'Regla de redacción: este documento describe ARQUITECTURA y CONTROLES. No contiene '
               'credenciales, claves de API, nombres de cuentas especiales, rutas de resguardo de '
               'códigos ni ningún dato que por sí mismo habilite un acceso. Esa información se '
               'gestiona por canales separados y no se documenta por escrito.')
    h2(doc, 'Cómo leer este documento')
    p(doc, 'La Parte I presenta los modelos de referencia y cómo se mapean al sistema. La Parte II '
           'describe la arquitectura técnica. La Parte III detalla los controles de seguridad por '
           'dominio. La Parte IV explica la trazabilidad y los registros incorruptibles. La Parte V '
           'cubre continuidad, monitoreo y gestión del cambio. La Parte VI presenta la matriz de '
           'riesgos, las brechas conocidas y el plan de tratamiento. Los anexos contienen las '
           'matrices de mapeo normativo y el índice de evidencias.')
    h1(doc, 'Índice')
    toc(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE I — MODELOS DE REFERENCIA
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE I — Modelos y marcos de referencia')
    p(doc, 'El sistema no fue diseñado "contra" una norma en particular, sino sobre un conjunto de '
           'principios internos (las máximas de diseño, sección 1.2) que luego se contrastaron con los '
           'marcos de referencia de la industria. El resultado es un mapeo natural: la mayoría de los '
           'controles existen porque el negocio los necesitaba, y los marcos sirven para nombrarlos, '
           'ordenarlos y detectar lo que falta.')

    h2(doc, '1.1 Marcos adoptados y su rol')
    tabla(doc, ('Marco', 'Versión', 'Rol en el sistema', 'Dónde se mapea'),
          (('ISO/IEC 27001', '2022', 'Sistema de gestión: los 93 controles del Anexo A como lista de verificación de cobertura', 'Anexo A de este documento'),
           ('ISO/IEC 27002', '2022', 'Guía de implementación de cada control', 'Parte III'),
           ('NIST Cybersecurity Framework', '2.0 (2024)', 'Estructura funcional: Gobernar, Identificar, Proteger, Detectar, Responder, Recuperar', 'Anexo B'),
           ('NIST SP 800-53', 'Rev. 5', 'Catálogo de controles por familia (AC, AU, IA, SC, SI, CP)', 'Referencia cruzada en Parte III'),
           ('OWASP ASVS', '4.0', 'Verificación de seguridad de la aplicación web (nivel 2 como meta)', 'Sección 9'),
           ('OWASP Top 10', '2021', 'Riesgos de aplicación y su mitigación concreta', 'Anexo C'),
           ('CIS Controls', 'v8', 'Salvaguardas priorizadas (IG1/IG2)', 'Anexo D'),
           ('ISO 22301', '2019', 'Continuidad del negocio: RTO, RPO, ensayos', 'Sección 13'),
           ('COBIT', '2019', 'Gobierno de TI: separación de funciones, gestión del cambio', 'Secciones 11 y 15'),
           ('Basilea (riesgo operacional)', 'II/III', 'Registro de eventos, controles de proceso, líneas de defensa', 'Sección 11'),
           ('IFRS 9 / NIIF', 'vigente', 'Provisiones por riesgo de crédito, devengo, reconocimiento', 'Sección 12'),
           ('Ley 19.628 / Ley 21.719', 'Chile', 'Protección de datos personales (régimen actual y nuevo)', 'Sección 16'),
           ('Ley 19.799', 'Chile', 'Firma electrónica: validez de la Firma Electrónica Simple', 'Sección 10'),
           ('Ley 21.459', 'Chile', 'Delitos informáticos: acceso ilícito, interceptación, daño', 'Sección 16'),
           ('Ley 21.320 / Ley 21.643 (Karin)', 'Chile', 'Cobranza extrajudicial · canal de denuncias', 'Sección 16')),
          (3.6, 2.0, 6.0, 4.9))

    h2(doc, '1.2 Las máximas de diseño (el marco interno)')
    p(doc, 'Cuatro reglas no negociables gobiernan cada línea de código. Son el equivalente interno '
           'de una política de seguridad de la información, y se aplican en la revisión de cada cambio:')
    tabla(doc, ('Máxima', 'Enunciado', 'Control que produce', 'Norma equivalente'),
          (('1. Un solo motor por cálculo', 'Cada magnitud de negocio se calcula en UN solo lugar que usan todos los consumidores', 'Integridad: no existen dos versiones del mismo número', 'ISO 27001 A.8.9 (gestión de configuración) · NIST PR.DS'),
           ('2. Una sola fuente de datos', 'Cada dato vive en un lugar autoritativo; el resto lo referencia, nunca lo copia', 'Integridad y calidad del dato; snapshots deliberados y explícitos', 'ISO 27001 A.5.12 (clasificación) · A.8.10'),
           ('3. Todo proceso nace documentado', 'La documentación es parte del mismo cambio, no una tarea posterior', 'Trazabilidad documental; el auditor encuentra el proceso descrito', 'ISO 27001 A.5.37 (procedimientos operativos)'),
           ('4. Todo movimiento de dinero se contabiliza', 'Cada ingreso y egreso genera asiento automático por el motor único', 'Integridad financiera; conciliación operación↔contabilidad', 'Basilea (riesgo operacional) · IFRS reconocimiento')),
          (3.4, 5.0, 4.4, 3.7))
    p(doc, 'A estas se suma el Principio Rector Paramétrico: el contenido del negocio (montos, tramos, '
           'textos, umbrales) se administra desde mantenedores sin tocar código, pero la ESTRUCTURA '
           '(orden de etapas, validaciones, atribuciones) está protegida en el servidor. Es la '
           'traducción práctica del principio de separación entre configuración y control.')

    h2(doc, '1.3 Modelo de tres líneas de defensa')
    p(doc, 'Siguiendo el modelo del IIA (Instituto de Auditores Internos) y la práctica de Basilea, '
           'el sistema materializa las tres líneas dentro de la plataforma:')
    vineta(doc, 'las validaciones en el servidor, la segregación de funciones, la doble firma y los candados de mes cerrado. Operan en el momento del hecho y no dependen de que alguien revise después.', bold_hasta='Primera línea (controles en la operación): ')
    vineta(doc, 'los motores de revisión automática (Revisor de Cartas, Certificación de Operaciones, checklist de Cierre de Mes), el log de eventos contables sin regla y los escalamientos por plazo. Vigilan lo que la primera línea dejó pasar.', bold_hasta='Segunda línea (supervisión y cumplimiento): ')
    vineta(doc, 'el módulo de Auditoría, las bitácoras inmutables, la consola SQL de solo lectura auditada y los documentos verificables con QR. Permiten a un tercero reconstruir cualquier hecho sin confiar en el operador.', bold_hasta='Tercera línea (auditoría independiente): ')

    # ══════════════════════════════════════════════════════════════════════
    # PARTE II — ARQUITECTURA
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE II — Arquitectura de la plataforma')

    h2(doc, '2. Visión general')
    p(doc, 'Business Suite es una aplicación web de tres capas desplegada íntegramente en nube, sin '
           'infraestructura propia. La decisión es deliberada: elimina la clase entera de riesgos '
           'físicos (acceso a sala de servidores, robo de discos, corte eléctrico) y traslada al '
           'proveedor, bajo contrato, la seguridad del centro de datos, el cifrado en reposo y la '
           'redundancia de hardware.')
    flujo(doc, 'NAVEGADOR (HTTPS) → API GATEWAY (Node/Express) → SERVICIOS DE NEGOCIO → BASE DE DATOS (TiDB Cloud) + ALMACÉN DE DOCUMENTOS (GCS)')
    tabla(doc, ('Capa', 'Tecnología', 'Proveedor / ubicación', 'Función'),
          (('Presentación', 'HTML5 / CSS / JavaScript vanilla, PWA instalable', 'Servida por el gateway', 'Interfaz por módulo; sin framework pesado, sin dependencias de terceros en el cliente'),
           ('Gateway y API', 'Node.js 24 · Express.js', 'Render (región Virginia, EE.UU.)', 'Enrutamiento, autenticación, autorización, rate limiting, cabeceras de seguridad, proxy a servicios'),
           ('Servicios de negocio', 'Módulos Express independientes bajo /services/', 'Mismo proceso Node (monolito modular)', 'Créditos, clientes, comisiones, cartas, post venta, tesorería, contabilidad, RRHH, cobranza, CRM, etc.'),
           ('Motores automáticos', 'Programador central (shared/scheduler.js)', 'Mismo proceso', '28 tareas de fondo gobernadas por un interruptor único'),
           ('Base de datos', 'TiDB Cloud (compatible MySQL), pool compartido', 'TiDB Cloud (nube gestionada)', 'Toda la operación, configuración y auditoría; ~390 tablas'),
           ('Documentos', 'Google Cloud Storage, bucket versionado', 'Google Cloud', 'Todo archivo subido por usuarios; la base guarda solo la referencia'),
           ('Contingencia de servicio', 'Google Cloud Run (servicio dormido)', 'Google Cloud, us-east4', 'Host alternativo apuntando a la misma base, motores apagados'),
           ('Contingencia de datos', 'Cloud SQL (instancia detenida)', 'Google Cloud', 'Réplica de respaldo de la base para escenario de pérdida del proveedor principal'),
           ('Código y CI', 'Git · GitHub · GitHub Actions', 'GitHub', 'Repositorio único, despliegue automático a producción, respaldo nocturno programado'),
           ('Correo transaccional', 'SMTP (Brevo)', 'Brevo', 'Todo correo del sistema, con remitentes paramétricos por plantilla')),
          (2.8, 4.2, 3.8, 5.7))

    h2(doc, '3. Estructura del código')
    p(doc, 'El repositorio sigue una estructura por dominio de negocio. Cada servicio expone sus rutas '
           'HTTP y sus controladores; el gateway los monta bajo /api/<servicio>. La regla de '
           'convenciones (tablas snake_case, rutas kebab-case, funciones camelCase, archivos '
           'kebab-case.js) se aplica en revisión de código.')
    tabla(doc, ('Directorio', 'Contenido', 'Relevancia para seguridad'),
          (('api-gateway/src/index.js', 'Arranque, cabeceras, montaje de rutas, proxy estático', 'Punto único de entrada: HSTS, trust proxy, rate limiting global'),
           ('api-gateway/public/', 'Frontend por módulo', 'Sin lógica de autorización confiable: toda decisión se repite en el servidor'),
           ('services/<dominio>/src/controllers/', 'Lógica de negocio y migraciones idempotentes', 'Validación de entrada, requireFunc, auditar(), affectedRows'),
           ('services/<dominio>/src/routes/', 'Definición de endpoints', 'Cada ruta declara verifyToken + requireFunc'),
           ('shared/config/database.js', 'Pool único de conexiones', 'Credenciales desde entorno; keep-alive; timeouts'),
           ('shared/middleware/auth.js', 'Emisión y verificación de JWT', 'Tipos de token, expiración, modo solo lectura, modo "ver como"'),
           ('shared/middleware/permisos.js', 'requireFunc / tieneFunc', 'Autorización paramétrica contra la matriz en BD'),
           ('shared/audit.js', 'Bitácora transversal', 'auditoria_movimientos: quién, qué, cuándo, IP'),
           ('shared/verificacion.js', 'Documentos verificables y FES', 'QR público + hash SHA-256 del contenido firmado'),
           ('shared/scheduler.js', 'Programador de motores', 'Interruptor MOTORES / ENTORNO; visible en /api/health'),
           ('shared/almacen-docs.js', 'Motor único de archivos', 'Bucket GCS; nunca LONGBLOB nuevo'),
           ('shared/migrate.js', '"Capataz" de migraciones', 'Cambios de esquema idempotentes, en fila, al arrancar'),
           ('shared/rate-limit.js', 'Limitador de peticiones', 'Ventana deslizante por usuario o por IP'),
           ('shared/mailer.js', 'Correo del sistema', 'Remitentes paramétricos; modo desarrollo redirige todo a prueba'),
           ('docs/', 'Runbooks, auditorías, certificaciones, inventarios', 'Evidencia documental para el auditor'),
           ('tests/', 'Pruebas automatizadas (node --test)', '233 pruebas sin base de datos; guardas contra patrones peligrosos')),
          (5.0, 5.5, 6.0))

    h2(doc, '4. Modelo de identidades y fronteras')
    p(doc, 'Coexisten tres poblaciones de usuarios con superficies de ataque distintas. La frontera '
           'entre ellas está en el middleware de autenticación, no en la interfaz:')
    tabla(doc, ('Población', 'Cómo entra', 'Tipo de credencial', 'Qué ve', 'Frontera'),
          (('Staff interno (~40 personas)', 'Login con correo corporativo y clave', 'JWT firmado, 8 horas, en sessionStorage de la pestaña', 'Lo que su perfil y sus overrides permiten; visibilidad de ejecutivos acotada por asignación', 'Endpoints /api/* internos'),
           ('Dealers (concesionarios)', 'Portal del Dealer con enrolamiento por RUT', 'JWT de tipo dealer; llaves portadoras hasheadas para acceso directo', 'SOLO su cartera: operaciones, cartolas, saldos, pre-aprobaciones', 'Un token dealer es RECHAZADO por los endpoints del staff'),
           ('Clientes finales', 'Portal /mis-creditos con OTP', 'Código de un solo uso, expirable', 'SOLO su RUT: créditos, cuotas, certificados', 'Endpoints del portal cliente exclusivamente'),
           ('Integraciones externas', 'API Pública', 'X-API-Key emitida desde el mantenedor de APIs', 'Solo los endpoints publicados para esa llave', 'Llaves revocables, con documentación por llave')),
          (3.2, 3.4, 3.6, 3.6, 2.7))
    regla(doc, 'La identidad de quien pregunta sale SIEMPRE del token, jamás del cuerpo de la petición. '
               'Ningún endpoint del portal acepta un RUT o un id de dealer "declarado" por el cliente: '
               'lo toma del token verificado. La auditoría de esta frontera (julio 2026) concluyó con '
               'cero hallazgos de fuga de datos entre dealers.')

    h2(doc, '5. Modelo de datos y clasificación de la información')
    p(doc, 'La base contiene cerca de 390 tablas. Para efectos de seguridad se clasifican por '
           'sensibilidad, lo que determina qué controles aplican a cada una:')
    tabla(doc, ('Clase', 'Ejemplos', 'Sensibilidad', 'Controles específicos'),
          (('Datos personales de clientes', 'clientes, antecedentes_laborales, informacion_comercial, documentos de evaluación', 'ALTA — datos personales y financieros (Ley 19.628 / 21.719)', 'Acceso por perfil; visibilidad por asignación; informes comerciales con bloqueo de repetición; documentos en bucket con referencia'),
           ('Operación crediticia', 'creditos, cartas_aprobacion, postventa_*, cobranza_*', 'ALTA — núcleo del negocio', 'Inmutabilidad al otorgar; mes cerrado; motor único de etapa; auditoría de cambios'),
           ('Financiera y contable', 'ctb_comprobantes, ctb_movimientos, ordenes_pago, op_correlativos, cajas', 'ALTA — integridad financiera', 'Doble firma; segregación emite/paga; correlativo central; cierre contable con acta'),
           ('Recursos humanos', 'rh_fichas, rh_liquidaciones, rh_documentos, rh_firmas', 'ALTA — datos laborales y de salud', 'Universo canónico; firmas FES; carpeta digital por colaborador; solo RRHH y jefaturas'),
           ('Configuración (mantenedores)', 'tasas, uf, parametros_credito, comisiones_variables, estados_credito', 'MEDIA — afecta cálculos', 'Permiso por mantenedor; bitácoras de versiones de variables; vigencia desde/hasta'),
           ('Identidad y permisos', 'usuarios, perfiles, permisos_perfil, permisos_usuario, funcionalidades', 'ALTA — control de acceso', 'Claves con hash; política de claves en servidor; cambios auditados; caché de 60 s'),
           ('Auditoría y bitácoras', 'auditoria_movimientos, ctb_eventos_log, wf_log, documentos_verificables, revisor bitácora', 'CRÍTICA — evidencia', 'Solo inserción desde el sistema; sin endpoints de edición o borrado'),
           ('Documentos', 'cartas_documentos, rh_documentos, fichas dealer, fundantes', 'ALTA', 'Bucket versionado con espejo; SHA-256 verificado en migración; referencia en BD')),
          (3.4, 4.6, 3.6, 4.9))
    caso(doc, 'Los documentos NO viven en la base. Medido el 04-08-2026, los archivos eran el 95% del '
              'tamaño de la base (112,9 MB en 293 archivos contra ~6 MB de TODA la operación). Se '
              'migraron al bucket con verificación de tamaño y SHA-256 archivo por archivo, y la base '
              'pasó de 118 MB a 44,4 MB. Además de costo, esto reduce la superficie: el respaldo de '
              'la base ya no arrastra documentos, y el bucket tiene su propio versionado y espejo.')

    h2(doc, '6. Integraciones externas y su gobierno')
    p(doc, 'Cada integración tiene un dueño de credencial, un mantenedor donde se gobierna y un log. '
           'Las credenciales viven exclusivamente como variables de entorno del servidor de '
           'producción; el ambiente local y el host de contingencia NO las tienen, por diseño '
           '(degradación deliberada, sección 13).')
    tabla(doc, ('Integración', 'Propósito', 'Dirección', 'Gobierno'),
          (('DealerNet (SOAP)', 'Informes comerciales por RUT', 'Saliente', 'Mantenedor DealerNet: productos, costos, bloqueo de repetición; sello "SIN INFORMACIÓN" cuando la respuesta llega vacía'),
           ('CMF (API)', 'Indicadores UF/UTM/dólar/TMC', 'Saliente', 'Sincronización programada; mantenedor de indicadores'),
           ('SII / SimpleAPI', 'RCV, libro de compras, DTE', 'Saliente', 'Certificado gestionado fuera del repositorio; documentado en runbook de rescate'),
           ('Meta WhatsApp Cloud', 'Bot Facilito, cobranza, campañas', 'Bidireccional (webhook)', 'Plantillas HSM aprobadas; bandeja con historial; grupo de control en campañas'),
           ('Brevo SMTP', 'Correo transaccional', 'Saliente', 'Remitente paramétrico por plantilla; todo correo queda en Auditoría → Correos Enviados'),
           ('Fintoc', 'Conexión bancaria y conciliación', 'Saliente', 'Sandbox validado; matching automático con revisión'),
           ('Workera', 'Reloj control, asistencia, vacaciones', 'Bidireccional', 'Espejo de vacaciones aprobadas; atrasos en fase 2'),
           ('Anthropic (IA)', 'Análisis de perfiles, lectura de documentos, reportes', 'Saliente', 'Subsistema paramétrico: interruptor maestro, funcionalidades habilitadas, modelo, costo por análisis trazado'),
           ('Google Cloud (GCS, Cloud Run, Cloud SQL)', 'Documentos y contingencia', 'Saliente', 'Cuenta con verificación en dos pasos; bucket versionado; espejo en otra región'),
           ('GitHub', 'Código y respaldos', 'Saliente', 'Repositorio privado; artefactos de respaldo con retención')),
          (3.6, 4.0, 2.6, 6.3))

    # ══════════════════════════════════════════════════════════════════════
    # PARTE III — CONTROLES POR DOMINIO
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE III — Controles de seguridad por dominio')

    h2(doc, '7. Control de acceso (ISO 27001 A.5.15–A.5.18, A.8.2–A.8.5 · NIST PR.AA)')
    h3(doc, '7.1 Autenticación')
    vineta(doc, 'el servidor emite un JSON Web Token firmado con expiración de 8 horas. El secreto de firma es una variable de entorno obligatoria: el proceso NO arranca sin ella. El secreto de producción es distinto al de cualquier otro ambiente, de modo que un token de desarrollo jamás es válido en producción.', bold_hasta='Token de sesión: ')
    vineta(doc, 'el token vive en sessionStorage, no en localStorage: cerrar la pestaña cierra la sesión, y no persiste entre reinicios del navegador. La única excepción documentada es la PWA de terreno, que necesita continuidad para su uso en ruta.', bold_hasta='Almacenamiento en cliente: ')
    vineta(doc, 'longitud mínima, mayúsculas, números, caracteres especiales, historial de claves, vencimiento con aviso previo por correo: todo paramétrico desde Usuarios → Seguridad y aplicado EN EL SERVIDOR al cambiar la clave. La política no es una sugerencia de pantalla.', bold_hasta='Política de contraseñas: ')
    vineta(doc, 'tras N intentos fallidos (paramétrico) la cuenta se bloquea y solo el Administrador la desbloquea. Independientemente, el endpoint de login tiene un tope de 10 intentos por minuto por dirección IP.', bold_hasta='Bloqueo por intentos: ')
    vineta(doc, 'la inactividad cierra la sesión (minutos paramétricos). Un 401 en cualquier llamada limpia el token del cliente y redirige al login con aviso, para no dejar sesiones "zombies".', bold_hasta='Cierre de sesión: ')
    vineta(doc, 'al cambiar la clave el servidor reemite el token, invalidando el estado "clave por cambiar".', bold_hasta='Rotación al cambiar clave: ')
    vineta(doc, 'activa en la cuenta que administra la infraestructura en nube (Google). La extensión a todo el equipo está planificada con aviso, plazo y acompañamiento (sección 17).', bold_hasta='Verificación en dos pasos: ')
    h3(doc, '7.2 Autorización: la matriz manda')
    p(doc, 'El modelo es RBAC (control de acceso basado en roles) con overrides individuales, '
           'implementado en cuatro tablas: perfiles, funcionalidades, permisos_perfil y '
           'permisos_usuario. Cada endpoint sensible declara el código de funcionalidad que exige '
           'mediante el middleware requireFunc(código). La resolución sigue este orden:')
    paso(doc, 1, 'Perfil Administrador', 'Pasa siempre (bypass explícito y único).')
    paso(doc, 2, 'Permisos del perfil', 'Se cargan los códigos habilitados en permisos_perfil para el perfil del usuario, leyendo el perfil DESDE LA BASE (no del token), para que un cambio de perfil aplique sin re-login.')
    paso(doc, 3, 'Override individual', 'permisos_usuario puede AGREGAR o QUITAR códigos a una persona concreta sin tocar su perfil. Es lo que permite, por ejemplo, habilitar la anulación de órdenes de pago a un analista sin abrirla a todo su perfil.')
    paso(doc, 4, 'Suplencias', 'Si el usuario es suplente activo de un titular, hereda las funciones del titular de forma aditiva. Toda acción hecha en suplencia queda marcada en auditoría con el nombre del titular.')
    paso(doc, 5, 'Caché', 'El resultado se cachea 60 segundos por usuario para no castigar la base; guardar permisos limpia el caché para efecto inmediato.')
    regla(doc, 'Está PROHIBIDO en código nuevo comparar contra nombres de perfil ("si es Administrador"). '
               'Todo control de acceso pasa por requireFunc contra la matriz. Esta regla se aplica en '
               'revisión de código y se corrige retroactivamente cuando se detecta (ejemplo: la anulación '
               'de órdenes de pago pasó de un chequeo por nombre de perfil al permiso paramétrico '
               'ordenes_pago_anular en septiembre 2026).')
    h3(doc, '7.3 Visibilidad de datos (acotamiento horizontal)')
    p(doc, 'Además del "qué puede hacer", el sistema controla "qué puede ver": un ejecutivo ve solo '
           'los ejecutivos que tiene asignados (tabla usuario_ejecutivos, motor '
           'shared/visibilidad-ejecutivos.js); un jefe comercial ve a su gente vigente según la '
           'ficha de Usuarios; un dealer ve solo su RUT; un cliente solo el suyo. Los módulos de '
           'seguimiento (fundantes, historial, bitácoras) aplican el mismo filtro en el servidor.')
    h3(doc, '7.4 Modos especiales de sesión')
    tabla(doc, ('Modo', 'Quién', 'Qué permite', 'Control'),
          (('Solo lectura', 'Perfiles de dirección/auditoría', 'Navegar todo, no operar', 'El middleware bloquea toda escritura; única excepción: cambiar la propia clave'),
           ('Ver como', 'Administrador', 'Abrir una pestaña como otro usuario para diagnosticar', 'Token marcado; escritura bloqueada por completo; banner permanente; expira a los 30 minutos; se audita'),
           ('Ver como dealer', 'Administrador', 'Ver el portal exactamente como lo ve un dealer', 'Solo lectura; token en hash de URL, 30 minutos'),
           ('Cuenta TV', 'Pantalla del Cuadro de Mando', 'Mostrar indicadores', 'Cuenta dedicada sin permisos operativos')),
          (3.0, 3.4, 4.6, 5.5))

    h2(doc, '8. Seguridad perimetral y de red (A.8.20–A.8.23 · NIST PR.PS)')
    vineta(doc, 'todo el tráfico va cifrado en tránsito. El proveedor de hosting redirige HTTP→HTTPS y la aplicación emite Strict-Transport-Security con max-age de un año e includeSubDomains, de modo que el navegador no vuelve a intentar HTTP.', bold_hasta='HTTPS forzado y HSTS: ')
    vineta(doc, 'TODA la API tiene techo de peticiones. Límite general de 600/min por usuario autenticado; 120/min en reportería, tablas dinámicas, diseño de consulta y contabilidad; login 10/min por IP. Se limita por usuario y no por IP en lo general porque la oficina sale por una sola IP pública. La respuesta es 429 con Retry-After.', bold_hasta='Rate limiting: ')
    vineta(doc, 'la aplicación declara trust proxy para leer la IP real detrás del balanceador del proveedor (X-Forwarded-For), de modo que el rate limiting por IP y la auditoría registran la IP del cliente y no la del proxy.', bold_hasta='IP real: ')
    vineta(doc, 'no existe acceso SSH ni RDP a servidores propios: no hay servidores propios. El acceso administrativo es a consolas de proveedores con verificación en dos pasos.', bold_hasta='Sin puertos de administración: ')
    vineta(doc, 'la base de datos solo acepta conexiones cifradas desde el pool de la aplicación con credenciales de entorno; no está expuesta a Internet para consulta directa.', bold_hasta='Base no expuesta: ')

    h2(doc, '9. Seguridad de la aplicación (OWASP ASVS · A.8.25–A.8.29)')
    h3(doc, '9.1 Validación de entrada y consultas parametrizadas')
    p(doc, 'Regla de proyecto: validar entrada en TODAS las rutas. Las consultas a la base usan '
           'siempre parámetros posicionales del driver (mysql2), nunca concatenación de texto. Los '
           'identificadores de negocio (RUT, patente, montos) pasan por normalizadores canónicos '
           'compartidos (rut-core, formateadores de moneda) antes de tocar la base.')
    h3(doc, '9.2 Respuesta uniforme y manejo de errores')
    p(doc, 'Toda respuesta de API sigue el contrato {success, data, error}. Los errores internos '
           'devuelven un mensaje genérico al cliente ("Error interno del servidor") y el detalle va '
           'al log y a la alerta por correo, evitando la fuga de rutas internas o estructura de base '
           'en mensajes de error.')
    h3(doc, '9.3 Escritura verificada')
    regla(doc, 'Un UPDATE que no afecta filas NO es un error para el driver. En operaciones críticas '
               '(otorgar, pagar, cerrar, anular) el código revisa affectedRows y aborta con aviso si es '
               'cero. Sin esto, un cambio puede fallar en silencio y descubrirse semanas después. La '
               'auditoría de este patrón está en docs/AUDITORIA-2026-08-05.md.')
    h3(doc, '9.4 Cabeceras y contenido')
    vineta(doc, 'HSTS (descrito arriba); tipo de contenido declarado; sin inline eval en el cliente para las páginas críticas.', bold_hasta='Cabeceras: ')
    vineta(doc, 'todo texto de usuario que vuelve a la pantalla pasa por escapado HTML (escH/esc) antes de insertarse; el motor único del documento de carta escapa cada campo.', bold_hasta='Escapado de salida (XSS): ')
    vineta(doc, 'los archivos subidos tienen tope de tamaño (por ejemplo 7 MB en órdenes de pago), se renombran de forma determinista y se almacenan fuera del árbol servible, en el bucket, accesibles solo vía endpoint autenticado.', bold_hasta='Carga de archivos: ')
    h3(doc, '9.5 Secretos')
    vineta(doc, 'el repositorio contiene únicamente un .env.example con marcadores; la auditoría de julio 2026 confirmó cero secretos en el historial de Git. El inventario de qué secretos existen (sin sus valores) está en docs/SECRETOS-inventario.md.', bold_hasta='Fuera del código: ')
    vineta(doc, 'las claves críticas (firma JWT, base de datos) son obligatorias al arrancar: sin ellas el proceso termina con error explícito en vez de operar degradado.', bold_hasta='Exigidos al boot: ')
    vineta(doc, 'ningún documento escrito del sistema contiene credenciales; los canales de intercambio de secretos son verbales o de gestor de claves, nunca chat ni capturas de pantalla.', bold_hasta='Nunca en documentos: ')
    h3(doc, '9.6 Pruebas automatizadas como control')
    p(doc, 'La suite de pruebas (node --test, 233 casos, sin base de datos) corre antes de cada '
           'despliegue e incluye pruebas de GUARDA: patrones prohibidos en el código (por ejemplo el '
           'uso de toISOString().slice(0,10) para fechas, que produce desfases de zona horaria) '
           'hacen fallar la suite. Así una regla de seguridad o integridad queda codificada como '
           'prueba y no como recomendación.')

    h2(doc, '10. Criptografía y firma electrónica (A.8.24 · Ley 19.799)')
    tabla(doc, ('Uso', 'Mecanismo', 'Detalle'),
          (('Contraseñas de usuario', 'Hash unidireccional con sal', 'Nunca se almacena ni se transmite la clave en claro; el servidor no puede recuperarla'),
           ('Sesiones', 'JWT firmado (HMAC)', 'Firma verificada en cada petición; expiración 8 h; tipos de token por población'),
           ('Tránsito', 'TLS 1.2+', 'Certificados gestionados por el proveedor; HSTS de un año'),
           ('Reposo', 'Cifrado del proveedor', 'TiDB Cloud y Google Cloud Storage cifran en reposo por defecto; delegado bajo contrato'),
           ('Documentos verificables', 'Código único + SHA-256 del contenido', 'Cartas, certificados, comprobantes y checklists llevan QR a /verificar/<código>'),
           ('Firma Electrónica Simple', 'Registro de firmante + hash', 'Identidad de sesión, nombre, cargo, fecha/hora, IP y huella SHA-256 del contenido firmado (tabla documentos_verificables, columnas de firmante; rh_firmas para RRHH)'),
           ('Llaves portadoras', 'Hash en base', 'Los links de acceso directo del portal se guardan hasheados y son regenerables'),
           ('Integridad de migración', 'SHA-256 por archivo', 'La migración de documentos al bucket verificó tamaño y hash uno a uno antes de liberar el blob')),
          (3.6, 4.4, 8.5))
    p(doc, 'La Firma Electrónica Simple cumple con la Ley 19.799: identifica al firmante, es '
           'verificable y permite detectar cualquier alteración posterior del documento (si el '
           'contenido cambia, el hash deja de calzar). Se usa en cartas de aprobación, checklists '
           'del revisor automático, comprobantes de vacaciones (cadena trabajador → empleador → '
           'RRHH) y actas de cierre. La verificación pública en /verificar/<código> muestra '
           'vigencia, y un documento reemplazado o anulado aparece como NO VIGENTE con su motivo.')

    h2(doc, '11. Integridad de procesos y segregación de funciones (COBIT · Basilea)')
    p(doc, 'La parametrización abre el CONTENIDO, nunca la ESTRUCTURA. Los siguientes controles '
           'existen en el servidor y no pueden desactivarse desde la interfaz:')
    tabla(doc, ('Control', 'Dónde aplica', 'Qué impide', 'Evidencia'),
          (('Doble firma', 'Anulación de operaciones, castigo de saldo, aplicación de fondos, incorporaciones con excepción', 'Que una persona sola deje sin efecto plata comprometida', 'Tablas de firmas con dos usuarios distintos; auditoría'),
           ('Segregación emite/paga', 'Todas las órdenes de pago', 'Que quien emite una ODP también la pague; paramétrico en Cajas', 'segregacion-pagos.js; bloqueo por id de usuario, no por nombre'),
           ('Segregación solicita/aprueba', 'Workflows de compras, RRHH, liquidez', 'Que alguien se apruebe a sí mismo', 'Bloqueo en servidor por usuario'),
           ('Mes cerrado', 'Toda operación retroactiva', 'Que un mes ya informado y pagado cambie', 'meses_cerrados; candado con acta; editores por nivel'),
           ('Inmutabilidad del otorgado', 'Créditos otorgados', 'Que calendario, tasa o comisión se reescriban', 'Todo FIJO al otorgar; TMC de mora fija al otorgamiento'),
           ('Snapshot deliberado', 'Cartola aprobada, finiquito, código de excepción, comisión pactada', 'Que el documento pagado cambie porque el origen avanzó', 'Columnas de snapshot con un solo hogar'),
           ('Correlativo central', 'Órdenes de pago de cualquier origen', 'Números duplicados o huecos silenciosos', 'op_correlativos: emitir/anular con quién y cuándo; el número anulado no se libera'),
           ('Una carta vigente por operación', 'Cartas de aprobación', 'Comisiones duplicadas por cartas paralelas', 'anularCartasPrevias al nacer una nueva'),
           ('Corrección por reemplazo', 'Cartas emitidas', 'Editar un documento ya firmado y entregado', 'Sufijo -C1/-C2; la anterior REEMPLAZADA; QR y FES no vigentes'),
           ('Máquina de estados', 'Estados del crédito', 'Transiciones fuera del mapa', 'estados_credito / estados_transicion (fase actual: configura y dibuja; enforcement al validar el mapa)'),
           ('Motor único', 'Todos los cálculos', 'Dos pantallas con números distintos', 'Módulos isomorfos compartidos (rentabilidad-core, uf, cotizador, comisión dealer)'),
           ('Escritura verificada', 'Otorgar, pagar, cerrar', 'Fallas silenciosas', 'affectedRows revisado')),
          (3.2, 4.2, 4.4, 4.7))

    h2(doc, '12. Integridad financiera y contable (IFRS · Basilea riesgo operacional)')
    p(doc, 'La cuarta máxima —todo movimiento de dinero se contabiliza— se implementa con un motor '
           'único de centralización (services/contabilidad/src/motor-asientos.js → contabilizar()) '
           'alimentado por reglas paramétricas del mantenedor Reglas de Centralización. Cada hecho '
           'de negocio que mueve plata (pago de saldo precio, comisión devengada, orden de pago, '
           'liquidación de sueldo, provisión) dispara el motor, que genera el asiento con sus '
           'cuentas, impuestos (IVA débito/crédito, retenciones) y período.')
    vineta(doc, 'el gasto o ingreso se reconoce cuando nace la obligación (documento recibido, liquidación emitida); el pago solo rebaja el pasivo contra banco. Consistente con el principio de devengo de las NIIF.', bold_hasta='Devengo y pago separados: ')
    vineta(doc, 'el motor nunca lanza excepción hacia la operación. Registra el resultado en ctb_eventos_log con estado CONTABILIZADO, SIN_REGLA, DESCUADRE o ERROR. Un evento SIN_REGLA es un pendiente visible para el contador, no un hecho perdido.', bold_hasta='La contabilidad nunca bloquea la operación: ')
    vineta(doc, 'las cuentas de las reglas nunca llevan año; la separación por ejercicio la da el período contable.', bold_hasta='Sin cuentas por ejercicio: ')
    vineta(doc, 'provisiones automáticas por riesgo de crédito con cierre mensual; castigo de saldo con doble firma gerencial. Alineado con el enfoque de pérdida esperada de IFRS 9 en su versión simplificada para cartera de consumo.', bold_hasta='Provisiones: ')
    vineta(doc, 'el Cierre Contable mensual valida saldos pendientes contra el seguimiento de post venta, genera el acta con firma y candado, y produce el informe a casa matriz.', bold_hasta='Cierre con acta: ')
    vineta(doc, 'el libro diario histórico de AVSOFT (2020 en adelante) se importa con un script idempotente que valida debe=haber por comprobante y no duplica; los balances de 8 columnas oficiales por año sirven para cuadrar la importación cuenta por cuenta.', bold_hasta='Migración contable trazable: ')
    p(doc, 'Trazabilidad extremo a extremo: la Solicitud de Pago de una comisión imprime al pie la '
           'cadena carta → aprobación → otorgamiento → fundantes → factura → orden → pago, y la '
           'glosa del asiento lleva documento, beneficiario y número de ODP. Un auditor puede partir '
           'de un asiento y llegar al crédito, o al revés, sin salir del sistema.')

    # ══════════════════════════════════════════════════════════════════════
    # PARTE IV — TRAZABILIDAD
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE IV — Trazabilidad y registros de auditoría')

    h2(doc, '13. Modelo de trazabilidad (A.8.15 registro · A.8.16 monitoreo · NIST DE.CM)')
    p(doc, 'El principio: la pregunta "¿quién hizo esto, cuándo, desde dónde y qué cambió?" siempre '
           'tiene respuesta, y esa respuesta no depende de la buena fe del operador. Se logra con '
           'registros que SOLO AGREGAN FILAS: el sistema no expone ninguna operación de edición o '
           'borrado sobre ellos.')
    h3(doc, '13.1 Registro transversal de acciones')
    p(doc, 'Tabla auditoria_movimientos, alimentada por shared/audit.js con el patrón fire-and-forget '
           '(nunca frena ni hace fallar la operación principal). Estructura:')
    tabla(doc, ('Columna', 'Contenido', 'Propósito'),
          (('fecha', 'Marca de tiempo del servidor', 'Nunca del navegador: las marcas las pone el servidor'),
           ('id_usuario / usuario / perfil', 'Quién, resuelto desde el token', 'Identidad verificada, no declarada'),
           ('id_titular / titular_nombre', 'Si actuó como suplente, a nombre de quién', 'Distingue acción propia de acción en suplencia'),
           ('modulo / accion / entidad / entidad_id', 'Qué se hizo y sobre qué', 'Vocabulario controlado: CREAR, EDITAR, APROBAR, RECHAZAR, ANULAR, ELIMINAR, LOGIN…'),
           ('detalle', 'Texto legible', 'Lo que un humano necesita leer'),
           ('rut', 'RUT del cliente o dealer afectado', 'Permite reconstruir la historia de un tercero'),
           ('meta', 'JSON con contexto', 'Valores antes/después cuando aplica'),
           ('ip', 'IP real del cliente', 'Vía X-Forwarded-For con trust proxy')),
          (4.2, 5.6, 6.7))
    h3(doc, '13.2 Bitácoras especializadas')
    tabla(doc, ('Bitácora', 'Qué registra', 'Inmutabilidad'),
          (('ctb_eventos_log', 'Cada hecho contabilizable con su resultado (CONTABILIZADO, SIN_REGLA, DESCUADRE, ERROR)', 'Solo inserción; es la lista de pendientes del contador'),
           ('documentos_verificables', 'Cada documento emitido con QR: tipo, referencia, emisor, firmante, hash, vigencia, motivo de anulación', 'Solo inserción y marca de anulación; nunca borrado'),
           ('rh_firmas', 'Cadena de firmas FES de RRHH (trabajador → empleador → RRHH)', 'Solo inserción'),
           ('wf_log', 'Escalamientos de los 15 flujos de workflow con horas hábiles', 'Solo inserción'),
           ('Bitácora Revisor Automático', 'Cada carta revisada por el motor: qué comparó, qué falló, resultado', 'Solo inserción; filtros por mes y resultado'),
           ('Versiones de variables', 'Cambios en variables de comisiones y Bono Jefe Comercial, con vigencia desde/hasta', 'Bitácora inmutable; el valor histórico nunca se pisa'),
           ('op_correlativos', 'Emisión y anulación de cada número de orden de pago', 'El número anulado queda reservado y marcado, no se libera'),
           ('Correos Enviados', 'Todo correo del sistema: destinatario, remitente, plantilla, resultado', 'Auditoría → Correos Enviados'),
           ('Consultas SQL', 'Cada consulta ejecutada en la Consola SQL (solo lectura)', 'Auditada con usuario y texto de la consulta'),
           ('Bitácora de dealers', 'Accesos al portal, qué miraron, qué preguntaron', 'Separada de la del staff'),
           ('Bitácora del crédito', 'Línea de tiempo por RUT/OP: ingreso, otorgamiento, documentos, pagos, anulación', 'Reconstrucción de solo lectura desde las fuentes')),
          (4.0, 7.3, 5.2))
    h3(doc, '13.3 Documentos verificables y verificación pública')
    p(doc, 'Cada documento que sale del sistema hacia un tercero (carta de aprobación, certificado '
           'de deuda o prepago, comprobante de vacaciones, checklist del revisor, acta) recibe un '
           'código único y un QR que apunta a /verificar/<código>. La página pública muestra si el '
           'documento es auténtico, quién lo emitió, cuándo y si sigue vigente — sin necesidad de '
           'entrar al sistema. Un documento reemplazado (por ejemplo una carta corregida) aparece '
           'como NO VIGENTE con el texto "Reemplazada por la carta N° X".')
    h3(doc, '13.4 Trazabilidad de la configuración')
    p(doc, 'Los cambios de esquema de la base pasan por un "capataz" de migraciones '
           '(shared/migrate.js) que las ejecuta en fila, de forma idempotente, al arrancar cada '
           'servicio. Cada migración vive en el código versionado en Git, junto al cambio funcional '
           'que la motivó, con su comentario explicativo. No existen cambios de esquema manuales '
           'fuera del repositorio en el flujo normal.')
    h3(doc, '13.5 Trazabilidad del despliegue')
    p(doc, 'Cada commit lleva un mensaje descriptivo y una versión global (APP_VERSION) que aparece '
           'en el badge de todas las páginas. El auditor puede tomar el badge de producción, '
           'buscarlo en el historial de Git y leer exactamente qué cambió en esa versión y por qué. '
           'El hash del commit desplegado aparece en los logs del proveedor de hosting.')

    # ══════════════════════════════════════════════════════════════════════
    # PARTE V — CONTINUIDAD, MONITOREO Y CAMBIO
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE V — Continuidad, monitoreo y gestión del cambio')

    h2(doc, '14. Continuidad del negocio y recuperación (ISO 22301 · A.5.29–A.5.30 · A.8.13–A.8.14 · NIST RC)')
    h3(doc, '14.1 Respaldos: la regla de las tres copias')
    tabla(doc, ('Activo', 'Copia 1', 'Copia 2', 'Copia 3', 'Retención'),
          (('Base de datos', 'Respaldo automático del proveedor (TiDB)', 'Volcado nocturno propio a las 02:17 (hora Chile) vía GitHub Actions, artefacto privado', 'Copia nocturna adicional en Google', '1 día · 30 días · 90 días'),
           ('Documentos', 'Bucket GCS con versionado de objetos', 'Bucket espejo en otra región', 'Paquete semanal a GitHub', 'Permanente'),
           ('Código', 'Repositorio GitHub (historial completo)', 'Clones locales del equipo', '—', 'Permanente'),
           ('Configuración de negocio', 'Vive en la base (mantenedores)', 'Viaja dentro del respaldo de la base', '—', 'La de la base')),
          (2.8, 3.8, 4.0, 3.2, 2.7))
    p(doc, 'El respaldo nocturno propio nació porque el plan del proveedor retiene solo 1 día. Se '
           'probó restaurando en un ambiente aparte (22 MB comprimidos en su momento). La '
           'restauración está documentada en el propio workflow. Pendiente menor: prueba de '
           'restauración real periódica en un branch de la base (sección 17).')
    h3(doc, '14.2 Contingencia del servicio')
    p(doc, 'Existe un host alternativo DORMIDO en Google Cloud Run (servicio afbs-standby, región '
           'us-east4, accesible en afbs2.autofacilchile.cl) apuntando a la MISMA base de producción. '
           'Reemplazar al principal no exige migrar datos. Métricas medidas: arranque en frío de '
           '5 segundos; costo en reposo de aproximadamente US$1/mes; se reconstruye solo cada día a '
           'las 05:00 para tener siempre la última versión del código.')
    regla(doc, 'El standby duerme con MOTORES=off: atiende peticiones pero NO ejecuta ninguno de los '
               '28 motores automáticos. Sin ese interruptor no se podría tener un standby desplegado: '
               'dos procesos contra la misma base disparan cada reloj dos veces, y el daño es '
               'silencioso (una comisión aprobada dos veces no se queja; un devengo duplicado tampoco). '
               'El orden de vuelta atrás —primero apagar los motores del standby— está en el runbook.')
    tabla(doc, ('Parámetro', 'Valor', 'Comentario'),
          (('RTO (objetivo de tiempo de recuperación)', 'Minutos', 'Promoción con un comando siguiendo docs/CONTINGENCIA-cloud-run.md; arranque en frío 5 s'),
           ('RPO (objetivo de punto de recuperación) — caída del servicio', '0', 'El standby usa la misma base: no hay pérdida de datos'),
           ('RPO — pérdida del proveedor de base', 'Hasta 24 h', 'Última copia nocturna; runbook docs/RUNBOOK-contingencia-bd.md §11-bis con Cloud SQL como destino'),
           ('Último ensayo', '04-08-2026', 'Promoción completa ensayada contra la base de staging'),
           ('Degradación aceptada', 'Integraciones apagadas', 'Promovido funciona toda la operación del negocio; IA, WhatsApp, DealerNet y SII quedan apagados hasta cargar sus credenciales (pendiente, sección 17)')),
          (5.4, 2.6, 8.5))
    h3(doc, '14.3 Contingencia de la base')
    p(doc, 'Además del respaldo, existe una instancia de Cloud SQL DETENIDA en Google Cloud '
           '(costo medido: US$2,4/mes detenida, US$4,8/día encendida) como destino de restauración '
           'si el proveedor principal de base no estuviera disponible por un período prolongado.')
    h3(doc, '14.4 El acceso a la consola es parte de la contingencia')
    p(doc, 'Sin acceso a la consola del proveedor no se puede promover el standby. Por eso la '
           'cuenta que administra la infraestructura tiene verificación en dos pasos con múltiples '
           'factores y códigos de respaldo resguardados físicamente en más de un lugar. Ni la '
           'ubicación ni los códigos se documentan por escrito.')

    h2(doc, '15. Monitoreo y detección (A.8.15–A.8.16 · NIST DE)')
    tabla(doc, ('Mecanismo', 'Frecuencia', 'Qué detecta', 'A quién avisa'),
          (('Uptime por servicio (shared/uptime.js)', 'Cada 5 minutos', 'Caída o lentitud de cada servicio, con historial', 'Mantenedor Salud y Uptime'),
           ('Alerta de errores (shared/alerta-errores.js)', 'En cada error 500', 'Fallas de servidor en producción, con freno de 10 minutos por ruta', 'Correo al administrador — el sistema avisa antes que el usuario'),
           ('Endpoint /api/health', 'Bajo demanda y por el proveedor', 'Base viva, uptime, estado de documentos, motores apagados', 'Proveedor de hosting (reinicio automático) y operador'),
           ('Badge de versión', 'En cada página', 'Qué build corre en producción', 'Todo usuario; confirma despliegues'),
           ('Vigía de relojes', 'Diario', 'Incoherencias de zona horaria y fechas futuras', 'Correo'),
           ('Log de eventos contables', 'En cada hecho', 'Hechos sin regla, descuadres', 'Dashboard contable'),
           ('Escalamientos (wf_log)', 'Horas hábiles', 'Pendientes vencidos sin decisión', 'Jefatura del responsable'),
           ('Correos programados', 'Diario/semanal', 'Checklist de cierre vencido, denuncias vencidas, informe diario', 'Responsables paramétricos'),
           ('Alerta del proveedor de correo', 'Automática', 'Errores de entrega', 'Administrador')),
          (4.2, 2.6, 5.2, 4.5))
    p(doc, 'El programador central expone en /api/health la lista de motores apagados. Un reloj '
           'suelto (setInterval fuera del programador) está prohibido por regla de proyecto '
           'precisamente porque sería invisible para el monitoreo y no podría apagarse en el standby.')

    h2(doc, '16. Gestión del cambio y ciclo de vida seguro (A.8.25–A.8.32 · COBIT BAI06)')
    paso(doc, 1, 'Cambio en rama principal', 'El flujo es Git → push a main → despliegue automático a producción por el proveedor (2 a 3 minutos). Cada commit describe qué cambió y por qué, en lenguaje de negocio.')
    paso(doc, 2, 'Pruebas antes del push', 'La suite de 233 pruebas corre localmente; un fallo bloquea el push por disciplina de equipo y por las pruebas de guarda.')
    paso(doc, 3, 'Ambiente de staging', 'Existe un ambiente de pruebas (credit-system-staging) con base propia, correos redirigidos a un dominio inválido y motores que salen al mundo apagados (ENTORNO=staging).')
    paso(doc, 4, 'Modo desarrollo', 'Un interruptor global redirige TODO el correo a una casilla de prueba, para operar contra datos reales sin molestar a clientes.')
    paso(doc, 5, 'Versión visible', 'APP_VERSION sube en cada cambio de frontend; el badge en producción confirma que el despliegue se aplicó.')
    paso(doc, 6, 'Cambios quirúrgicos', 'Regla de proyecto: editar solo las líneas necesarias, leer antes de tocar, un cambio a la vez, nunca cambiar lo que no se pidió. El caso que originó la regla (un cambio no solicitado que rompió un módulo crítico) está documentado.')
    paso(doc, 7, 'Documentación en el mismo cambio', 'Máxima 3: el manual de procesos y el de usuario se actualizan en el mismo commit que la funcionalidad.')
    paso(doc, 8, 'Auditorías de código', 'Se realizan auditorías adversariales por dimensión (permisos, inyección SQL, affectedRows, validación de entrada) con hallazgos confirmados documentados en docs/AUDITORIA-*.md, y una certificación pre-producción (docs/CERTIFICACION-PRE-PRODUCCION-2026-08-03.md).')
    h3(doc, '16.1 Dependencias')
    p(doc, 'El backend usa un conjunto acotado de dependencias (Express, mysql2, jsonwebtoken, '
           'pdfkit, nodemailer y utilitarios). El frontend no carga librerías de terceros desde '
           'CDN externos para su operación crítica: Bootstrap y los íconos se sirven desde el propio '
           'gateway. Las dependencias se actualizan con el cambio funcional que las requiere y '
           'pasan por la suite de pruebas.')

    h2(doc, '17. Cumplimiento legal y regulatorio (Chile)')
    tabla(doc, ('Norma', 'Obligación', 'Cómo se cumple en el sistema'),
          (('Ley 19.628 (vigente) y Ley 21.719 (protección de datos personales, plena vigencia dic-2026)', 'Tratamiento lícito, finalidad, seguridad, derechos del titular', 'Acceso por perfil y visibilidad acotada; documentos fuera de la base con acceso autenticado; informes comerciales con bloqueo de repetición (15 días) y sello cuando llegan vacíos; portal del cliente con OTP para que el titular vea solo su información; auditoría de accesos'),
           ('Ley 19.799 (firma electrónica)', 'Validez de la firma electrónica simple con identificación del firmante y detección de alteraciones', 'Firma FES con identidad de sesión, fecha, IP y hash SHA-256; verificación pública por QR'),
           ('Ley 21.459 (delitos informáticos)', 'Prevenir acceso ilícito, interceptación, daño', 'Cifrado en tránsito; autenticación y autorización en servidor; rate limiting; bloqueo por intentos; auditoría con IP'),
           ('Ley 21.320 (cobranza extrajudicial)', 'Horarios, frecuencia y forma de las gestiones', 'Automatizaciones de cobranza paramétricas que respetan tramos y horarios; toda gestión queda en el CRM'),
           ('Ley 21.643 (Ley Karin)', 'Canal de denuncias confidencial', 'Canal de Compliance con denuncias anónimas identificadas por código; alegato semanal de vencidas'),
           ('Código del Trabajo (art. 32, 45, 71)', 'Horas extras, semana corrida, feriado con renta variable', 'Motores únicos de horas extras, semana corrida por mes y feriado variable (promedio de 3 liquidaciones)'),
           ('SII (DTE, RCV, F29)', 'Libros y declaraciones', 'RCV importado al auxiliar; IVA y retenciones como parte del asiento; cuentas sin año'),
           ('Normativa CMF (referencia)', 'Gestión de seguridad de la información y ciberseguridad (RAN 20-10 para entidades supervisadas)', 'Adoptada como referencia de buenas prácticas aunque la empresa no sea entidad supervisada: gobierno, inventario de activos, gestión de incidentes, continuidad')),
          (4.6, 4.6, 7.3))

    # ══════════════════════════════════════════════════════════════════════
    # PARTE VI — RIESGOS, BRECHAS Y PLAN
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'PARTE VI — Riesgos, brechas conocidas y plan de tratamiento')

    h2(doc, '18. Matriz de riesgos')
    p(doc, 'Probabilidad e impacto en escala 1 (bajo) a 5 (alto). El riesgo residual considera los '
           'controles vigentes descritos en este documento.')
    tabla(doc, ('#', 'Riesgo', 'P', 'I', 'Controles vigentes', 'Residual'),
          (('R1', 'Acceso no autorizado por credencial robada', '2', '5', 'Sesión 8 h en pestaña; bloqueo por intentos; rate limit login; auditoría con IP; MFA en infraestructura', 'Medio (MFA a todo el equipo pendiente)'),
           ('R2', 'Escalada de privilegios', '1', '5', 'RBAC en servidor; requireFunc; prohibición de chequeo por nombre de perfil; auditoría de cambios de permisos', 'Bajo'),
           ('R3', 'Fuga de datos entre dealers/clientes', '1', '5', 'Identidad desde el token; frontera en middleware; auditoría julio 2026 sin hallazgos', 'Bajo'),
           ('R4', 'Inyección SQL / XSS', '1', '4', 'Consultas parametrizadas; escapado de salida; validación de entrada; auditorías adversariales', 'Bajo'),
           ('R5', 'Fraude interno en pagos', '2', '5', 'Segregación emite/paga; doble firma; correlativo central; mes cerrado; trazabilidad extremo a extremo', 'Bajo'),
           ('R6', 'Alteración de documentos emitidos', '1', '4', 'Hash SHA-256; QR público; corrección solo por reemplazo', 'Bajo'),
           ('R7', 'Caída del proveedor de hosting', '2', '3', 'Standby ensayado; RTO minutos; RPO 0', 'Bajo'),
           ('R8', 'Pérdida de la base de datos', '1', '5', 'Tres copias; retención 30/90 días; Cloud SQL como destino', 'Medio (prueba de restauración periódica pendiente)'),
           ('R9', 'Doble ejecución de motores', '1', '4', 'Programador central con interruptor; standby con MOTORES=off; runbook de vuelta atrás', 'Bajo'),
           ('R10', 'Secreto expuesto en código o documento', '1', '5', 'Cero secretos en Git (auditado); regla de no documentar; inventario sin valores', 'Bajo'),
           ('R11', 'Error silencioso en escritura crítica', '2', '4', 'affectedRows verificado; alerta por 500; log de eventos contables', 'Bajo'),
           ('R12', 'Integración externa caída (DealerNet, CMF, WhatsApp)', '3', '2', 'Sello de respuesta vacía; reintentos; uptime por servicio; operación no depende de ellas', 'Bajo'),
           ('R13', 'Dependencia de una sola persona técnica', '3', '4', 'Documentación viva (5 documentos + runbooks); código comentado en lenguaje de negocio; manuales por rol', 'Medio'),
           ('R14', 'Ingeniería social sobre usuarios', '3', '3', 'Anti-enumeración; correo trazado; política de claves; capacitación (Academia)', 'Medio')),
          (0.9, 4.2, 0.7, 0.7, 6.9, 3.1))

    h2(doc, '19. Brechas conocidas y plan de tratamiento')
    p(doc, 'Este documento no presenta un sistema perfecto. Las siguientes brechas están '
           'identificadas, priorizadas y registradas en docs/PENDIENTES.md (fuente única de lo '
           'abierto), que se actualiza en el momento en que se detecta cada pendiente:')
    tabla(doc, ('Brecha', 'Riesgo asociado', 'Tratamiento', 'Estado'),
          (('MFA para todo el equipo (~40 personas)', 'R1', 'Exigir verificación en dos pasos desde la consola de Workspace con aviso, plazo y acompañamiento', 'Planificado'),
           ('2FA aplicativo (TOTP) para perfiles administradores', 'R1, R2', 'Segundo factor dentro de Business Suite para perfiles críticos', 'Backlog'),
           ('Prueba periódica de restauración de la base', 'R8', 'Restaurar el volcado nocturno en un branch de la base cada trimestre', 'Pendiente menor'),
           ('Claves de integración en el host de contingencia', 'R7 (degradación)', 'Cargar los secretos de integraciones en el standby según §6 del manual de contingencia', 'Pendiente'),
           ('Tabla audit_log unificada para acciones administrativas', 'Trazabilidad', 'auditoria_movimientos ya cubre lo crítico; falta unificar cierre de mes y carga masiva bajo el mismo vocabulario', 'Parcial'),
           ('Enforcement de la máquina de estados', 'R11', 'Activar el bloqueo de transiciones cuando el mapa esté validado', 'Fase siguiente'),
           ('Pruebas de penetración externas', 'R1–R4', 'Contratar un pentest independiente anual sobre la superficie pública (login, portales)', 'No realizado'),
           ('Listado consolidado de servicios pagados y costos', 'Gobierno', 'Vista única de proveedores con costo por escenario', 'Pendiente'),
           ('WAF / protección DDoS dedicada', 'Disponibilidad', 'Hoy delegado al proveedor de hosting; evaluar capa adicional si el volumen lo justifica', 'Evaluación'),
           ('Documentar reglas de negocio en el código', 'R13', 'Comentario "Regla negocio" en cada cálculo no obvio', 'En curso')),
          (4.8, 2.4, 6.5, 2.8))

    h2(doc, '20. Gestión de incidentes (A.5.24–A.5.28 · NIST RS)')
    paso(doc, 1, 'Detección', 'Alerta automática por correo en cada error 500; uptime cada 5 minutos; reclamo de usuario; alerta del proveedor.')
    paso(doc, 2, 'Clasificación', 'Disponibilidad (servicio caído), integridad (dato incorrecto), confidencialidad (acceso indebido). La integridad de datos financieros tiene prioridad máxima por su efecto silencioso.')
    paso(doc, 3, 'Contención', 'Disponibilidad: promover el standby por runbook. Confidencialidad: suspender la cuenta (Usuarios → Suspender) y revocar llaves; el token expira en máximo 8 horas. Integridad: candado del mes afectado y punto de restauración UAT si aplica.')
    paso(doc, 4, 'Investigación', 'Reconstrucción desde auditoria_movimientos, bitácora del crédito, log de eventos contables y correos enviados. La consola SQL de solo lectura permite consultas ad hoc auditadas.')
    paso(doc, 5, 'Corrección y lección', 'La corrección va al código con su prueba de guarda cuando corresponde, y el caso queda documentado en el proceso ("Caso real") para que la regla tenga memoria.')
    caso(doc, 'Ejemplo de ciclo completo: una operación quedó con la carta OTORGADA y el crédito en '
              'estado "Digitado" (minúsculas) por la naturaleza case-sensitive de la base, invisible '
              'en dashboard y comisiones. Se corrigió el dato, se estableció la regla de comparar '
              'siempre con UPPER() y escribir estados en mayúsculas, y la regla quedó en la guía del '
              'proyecto como error frecuente aprendido en producción.')

    # ══════════════════════════════════════════════════════════════════════
    # ANEXOS
    # ══════════════════════════════════════════════════════════════════════
    h1(doc, 'Anexo A — Mapeo ISO/IEC 27001:2022 (Anexo A) → controles del sistema')
    p(doc, 'Cobertura de los controles del Anexo A agrupados por tema. "Implementado" significa que '
           'existe un artefacto verificable; "Delegado" que lo cubre el proveedor bajo contrato; '
           '"Parcial" que existe con brechas registradas; "N/A" que no aplica a una operación sin '
           'infraestructura propia.')
    tabla(doc, ('Control', 'Nombre', 'Estado', 'Evidencia'),
          (('A.5.1', 'Políticas de seguridad', 'Implementado', 'Máximas de diseño; CLAUDE.md como guía de proyecto; este documento'),
           ('A.5.9', 'Inventario de activos', 'Implementado', 'Diccionario de datos (389 tablas); inventario de secretos sin valores; catálogo de endpoints'),
           ('A.5.10', 'Uso aceptable de activos', 'Parcial', 'Perfiles y permisos; falta política escrita de uso aceptable para usuarios finales'),
           ('A.5.12', 'Clasificación de la información', 'Implementado', 'Sección 5 de este documento'),
           ('A.5.15', 'Control de acceso', 'Implementado', 'RBAC con overrides; requireFunc'),
           ('A.5.16', 'Gestión de identidades', 'Implementado', 'Módulo Usuarios; suspensión; suplencias con vigencia'),
           ('A.5.17', 'Información de autenticación', 'Implementado', 'Política de claves en servidor; hash; reemisión de token'),
           ('A.5.18', 'Derechos de acceso', 'Implementado', 'Matriz auditable; script audit-permisos.js'),
           ('A.5.19–5.22', 'Seguridad en relaciones con proveedores', 'Parcial', 'Inventario de integraciones; falta evaluación formal de cada proveedor'),
           ('A.5.23', 'Seguridad en servicios en nube', 'Implementado', 'Toda la plataforma en nube gestionada; contingencia multi-nube'),
           ('A.5.24–5.28', 'Gestión de incidentes', 'Implementado', 'Sección 20; alertas automáticas; runbooks'),
           ('A.5.29–5.30', 'Continuidad', 'Implementado', 'Sección 14; ensayo 04-08-2026'),
           ('A.5.31–5.36', 'Cumplimiento legal y revisión', 'Implementado', 'Sección 17'),
           ('A.5.37', 'Procedimientos operativos documentados', 'Implementado', '5 documentos vivos + runbooks + manuales por rol'),
           ('A.6.1–6.8', 'Controles de personas', 'Parcial', 'Canal de Compliance; Academia; falta MFA universal y política formal de uso'),
           ('A.7.x', 'Controles físicos', 'Delegado / N/A', 'Sin infraestructura propia; centros de datos del proveedor'),
           ('A.8.1', 'Dispositivos de usuario', 'Parcial', 'Sesión en pestaña; sin MDM'),
           ('A.8.2–8.5', 'Acceso privilegiado y autenticación segura', 'Implementado', 'Bypass único del Administrador; modos solo lectura y ver-como; MFA en infraestructura'),
           ('A.8.6', 'Gestión de capacidad', 'Implementado', 'Rate limiting; paginación en todos los endpoints; límites defensivos'),
           ('A.8.7', 'Protección contra malware', 'Delegado', 'Proveedores de nube y estaciones con Workspace'),
           ('A.8.8', 'Gestión de vulnerabilidades técnicas', 'Parcial', 'Auditorías adversariales internas; pentest externo pendiente'),
           ('A.8.9', 'Gestión de configuración', 'Implementado', 'Migraciones versionadas; mantenedores; APP_VERSION'),
           ('A.8.10', 'Borrado de información', 'Parcial', 'Borrado por lista explícita verificada; punto de restauración UAT; falta política de retención por clase'),
           ('A.8.12', 'Prevención de fuga de datos', 'Parcial', 'Fronteras por token; sin DLP'),
           ('A.8.13', 'Respaldo', 'Implementado', 'Tres copias; sección 14.1'),
           ('A.8.14', 'Redundancia', 'Implementado', 'Standby Cloud Run; bucket espejo; Cloud SQL'),
           ('A.8.15', 'Registro (logging)', 'Implementado', 'auditoria_movimientos y bitácoras especializadas'),
           ('A.8.16', 'Monitoreo', 'Implementado', 'Sección 15'),
           ('A.8.17', 'Sincronización de relojes', 'Implementado', 'Zona horaria del pool; vigía diario; marcas de tiempo del servidor'),
           ('A.8.18', 'Programas utilitarios privilegiados', 'Implementado', 'Consola SQL de solo lectura, perfil máximo, auditada'),
           ('A.8.19', 'Instalación de software', 'Delegado', 'Despliegue automatizado por el proveedor desde Git'),
           ('A.8.20–8.23', 'Seguridad de redes', 'Implementado / Delegado', 'HSTS; base no expuesta; red del proveedor'),
           ('A.8.24', 'Criptografía', 'Implementado', 'Sección 10'),
           ('A.8.25–8.29', 'Desarrollo seguro y pruebas', 'Implementado', 'Sección 9 y 16; 233 pruebas; guardas'),
           ('A.8.30', 'Desarrollo externalizado', 'N/A', 'Desarrollo interno'),
           ('A.8.31', 'Separación de ambientes', 'Implementado', 'Producción, staging, local; secretos distintos por ambiente'),
           ('A.8.32', 'Gestión del cambio', 'Implementado', 'Sección 16'),
           ('A.8.33', 'Información de prueba', 'Implementado', 'Staging con correos a dominio inválido; modo desarrollo'),
           ('A.8.34', 'Protección de sistemas durante auditoría', 'Implementado', 'Consola de solo lectura; modo ver-como sin escritura')),
          (2.2, 4.6, 2.8, 6.9))

    h1(doc, 'Anexo B — Mapeo NIST CSF 2.0')
    tabla(doc, ('Función', 'Categoría', 'Implementación en Business Suite'),
          (('GOVERN (GV)', 'Contexto, política, roles, riesgo', 'Máximas de diseño; principio paramétrico; tres líneas de defensa; matriz de riesgos; PENDIENTES.md como registro vivo'),
           ('IDENTIFY (ID)', 'Activos, riesgos, mejora', 'Diccionario de datos; inventario de integraciones y secretos (sin valores); clasificación de información; auditorías internas'),
           ('PROTECT (PR)', 'Identidad, concienciación, datos, plataforma', 'JWT 8 h; RBAC + overrides; política de claves; HSTS; rate limiting; consultas parametrizadas; documentos fuera de la base; Academia'),
           ('DETECT (DE)', 'Monitoreo continuo, análisis', 'Uptime 5 min; alerta por 500; log de eventos contables; vigía de relojes; bitácoras inmutables'),
           ('RESPOND (RS)', 'Gestión de incidentes, análisis, mitigación', 'Sección 20; suspensión de cuentas; candado de mes; punto de restauración UAT'),
           ('RECOVER (RC)', 'Recuperación, comunicación', 'Standby ensayado; tres copias; runbooks; aviso de mantención global')),
          (2.8, 4.4, 9.3))

    h1(doc, 'Anexo C — OWASP Top 10 (2021) y su mitigación')
    tabla(doc, ('Riesgo', 'Mitigación en el sistema'),
          (('A01 Control de acceso roto', 'Autorización en servidor por requireFunc; identidad desde el token; visibilidad horizontal por asignación; auditoría de fronteras sin hallazgos'),
           ('A02 Fallas criptográficas', 'TLS + HSTS; hash de claves; JWT firmado; cifrado en reposo del proveedor; SHA-256 en documentos'),
           ('A03 Inyección', 'Consultas parametrizadas; validación de entrada; normalizadores canónicos; auditoría adversarial de SQL'),
           ('A04 Diseño inseguro', 'Máximas de diseño; segregación de funciones; doble firma; motor único; máquina de estados'),
           ('A05 Configuración insegura', 'Secretos por entorno obligatorios; errores genéricos al cliente; ambientes separados; sin puertos de administración'),
           ('A06 Componentes vulnerables', 'Dependencias acotadas; sin CDN externos para operación crítica; actualización con pruebas'),
           ('A07 Fallas de autenticación', 'Política de claves en servidor; bloqueo por intentos; rate limit de login; sesión corta en pestaña; anti-enumeración'),
           ('A08 Fallas de integridad', 'Despliegue desde Git firmado por historial; migraciones versionadas; hash en documentos; correlativo central'),
           ('A09 Fallas de registro y monitoreo', 'Bitácoras solo-inserción; alerta por 500; uptime; correos trazados; consola SQL auditada'),
           ('A10 SSRF', 'Las integraciones salientes apuntan a hosts fijos de configuración; ningún endpoint toma una URL del usuario para consultarla')),
          (4.6, 11.9))

    h1(doc, 'Anexo D — CIS Controls v8 (IG1/IG2) — cobertura resumida')
    tabla(doc, ('Control CIS', 'Cobertura', 'Comentario'),
          (('1–2 Inventario de activos y software', 'Alta', 'Todo el software es propio o de proveedor gestionado; diccionario de datos'),
           ('3 Protección de datos', 'Alta', 'Clasificación; documentos fuera de la base; fronteras'),
           ('4 Configuración segura', 'Alta', 'Cabeceras; secretos por entorno; ambientes'),
           ('5–6 Gestión de cuentas y accesos', 'Alta', 'RBAC; overrides; suspensión; suplencias; MFA parcial'),
           ('7 Gestión de vulnerabilidades', 'Media', 'Auditorías internas; pentest externo pendiente'),
           ('8 Gestión de registros', 'Alta', 'Bitácoras inmutables con IP y marca del servidor'),
           ('9 Protección de correo y navegador', 'Media', 'Delegado a Workspace; correo saliente trazado'),
           ('10 Defensa contra malware', 'Delegado', 'Proveedores'),
           ('11 Recuperación de datos', 'Alta', 'Tres copias; ensayo; runbooks'),
           ('12 Infraestructura de red', 'Delegado', 'Sin red propia'),
           ('13 Monitoreo de red', 'Media', 'Uptime y alertas de aplicación; sin IDS propio'),
           ('14 Concienciación', 'Media', 'Academia; Canal de Compliance; falta programa formal de seguridad'),
           ('15 Proveedores', 'Media', 'Inventario; falta evaluación formal'),
           ('16 Seguridad de aplicaciones', 'Alta', 'Sección 9; pruebas de guarda'),
           ('17 Respuesta a incidentes', 'Alta', 'Sección 20'),
           ('18 Pruebas de penetración', 'Baja', 'Pendiente contratar')),
          (5.0, 2.6, 8.9))

    h1(doc, 'Anexo E — Índice de evidencias para el auditor')
    p(doc, 'Cada ítem puede solicitarse y se entrega en la forma indicada. Ninguna evidencia '
           'contiene secretos.')
    tabla(doc, ('Evidencia', 'Forma', 'Ubicación'),
          (('Matriz de perfiles y funcionalidades', 'Exportación desde Usuarios → Perfiles y Permisos; salida de scripts/audit-permisos.js', 'Sistema / repositorio'),
           ('Bitácora de auditoría de un período', 'Filtro por fechas en el módulo Auditoría; exportación', 'Sistema'),
           ('Log de eventos contables', 'Dashboard Contable; consulta de ctb_eventos_log', 'Sistema'),
           ('Verificación de un documento', 'URL pública /verificar/<código> del documento', 'Público'),
           ('Historial de cambios de código', 'git log con mensajes descriptivos; badge de versión', 'Repositorio'),
           ('Resultado de pruebas', 'Salida de node --test (233 casos)', 'Repositorio'),
           ('Auditorías internas de seguridad', 'docs/AUDITORIA-2026-08-03.md, docs/AUDITORIA-2026-08-05.md', 'Repositorio'),
           ('Certificación pre-producción', 'docs/CERTIFICACION-PRE-PRODUCCION-2026-08-03.md', 'Repositorio'),
           ('Plan y ensayo de contingencia', 'docs/CONTINGENCIA-cloud-run.md, docs/RUNBOOK-contingencia-bd.md', 'Repositorio'),
           ('Inventario de secretos (sin valores)', 'docs/SECRETOS-inventario.md', 'Repositorio'),
           ('Almacén de documentos y migración verificada', 'docs/ALMACEN-documentos.md', 'Repositorio'),
           ('Pendientes y brechas', 'docs/PENDIENTES.md y card Pendientes Abiertos', 'Repositorio / sistema'),
           ('Estado de salud y motores', 'GET /api/health', 'Sistema'),
           ('Manuales de proceso y usuario', 'Suite de Documentación (5 documentos vivos + Biblioteca Word)', 'Sistema'),
           ('Diccionario de datos', 'AutoFacil_Diccionario_Datos.docx (389 tablas, 4.284 columnas)', 'Biblioteca Word')),
          (5.2, 6.8, 4.5))

    h1(doc, 'Anexo F — Glosario')
    tabla(doc, ('Término', 'Definición'),
          (('requireFunc', 'Middleware de autorización que exige un código de funcionalidad de la matriz de permisos'),
           ('Motor único', 'Función pura compartida donde se calcula una magnitud de negocio para todos los consumidores'),
           ('Snapshot deliberado', 'Copia intencional e inmutable de un valor en un momento (ej. comisión pactada al emitir la carta)'),
           ('Mes cerrado', 'Período contable con candado: la operación retroactiva queda bloqueada'),
           ('FES', 'Firma Electrónica Simple según Ley 19.799, con identidad del firmante y hash del contenido'),
           ('Documento verificable', 'Documento con código único y QR que se valida públicamente en /verificar'),
           ('Standby', 'Host de contingencia dormido apuntando a la misma base, con motores apagados'),
           ('MOTORES=off', 'Interruptor de entorno que apaga todos los motores de negocio en un proceso'),
           ('Capataz de migraciones', 'Componente que ejecuta cambios de esquema idempotentes en fila al arrancar'),
           ('Override individual', 'Permiso agregado o quitado a una persona concreta sin cambiar su perfil'),
           ('Suplencia', 'Herencia temporal y aditiva de funciones de un titular a un suplente, auditada'),
           ('Ver como', 'Sesión de solo lectura del Administrador como otro usuario, con banner y expiración'),
           ('RTO / RPO', 'Tiempo objetivo de recuperación / punto objetivo de recuperación'),
           ('Tres líneas de defensa', 'Modelo de control: operación, supervisión, auditoría independiente')),
          (4.0, 12.5))

    return doc


if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Seguridad-Informacion-Arquitectura-Trazabilidad-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
