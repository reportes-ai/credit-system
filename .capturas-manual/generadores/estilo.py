# -*- coding: utf-8 -*-
"""Estilos y helpers del Manual de Operaciones — Business Suite."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x01, 0x41, 0xA2)
AZUL_OSCURO = RGBColor(0x01, 0x2D, 0x70)
GRIS = RGBColor(0x47, 0x55, 0x69)
GRIS_SUAVE = RGBColor(0x9A, 0xA0, 0xA6)
AMBAR = RGBColor(0xB4, 0x5F, 0x06)
ROJO = RGBColor(0xB9, 0x1C, 0x1C)
VERDE = RGBColor(0x15, 0x80, 0x3D)

LOGO = r'C:\Users\patri\Documents\credit-system\api-gateway\public\img\logo-bs.png'

CAPTURAS = []          # (numero, titulo, pantalla) — para el anexo

def nuevo_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)   # carta
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.2)
    n = doc.styles['Normal']
    n.font.name = 'Calibri'; n.font.size = Pt(11); n.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15
    for nombre, tam, color, antes in (('Heading 1', 20, AZUL_OSCURO, 18),
                                      ('Heading 2', 15, AZUL, 14),
                                      ('Heading 3', 12.5, AZUL_OSCURO, 10)):
        s = doc.styles[nombre]
        s.font.name = 'Calibri'; s.font.size = Pt(tam); s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(antes); s.paragraph_format.space_after = Pt(6)
        # que Word no los pinte con el azul de tema
        s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return doc

def sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def p(doc, texto='', bold=False, italic=False, color=None, size=None, align=None, antes=None, despues=None):
    par = doc.add_paragraph()
    if align is not None: par.alignment = align
    if antes is not None: par.paragraph_format.space_before = Pt(antes)
    if despues is not None: par.paragraph_format.space_after = Pt(despues)
    if texto:
        r = par.add_run(texto)
        r.bold, r.italic = bold, italic
        if color: r.font.color.rgb = color
        if size: r.font.size = Pt(size)
    return par

def runs(doc, partes, align=None):
    """partes = [(texto, {'bold':True,...}), ...]"""
    par = doc.add_paragraph()
    if align is not None: par.alignment = align
    for texto, fmt in partes:
        r = par.add_run(texto)
        r.bold = fmt.get('bold', False); r.italic = fmt.get('italic', False)
        if 'color' in fmt: r.font.color.rgb = fmt['color']
        if 'size' in fmt: r.font.size = Pt(fmt['size'])
    return par

def h1(doc, texto, salto=True):
    if salto:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(texto, level=1)

def h2(doc, texto): doc.add_heading(texto, level=2)
def h3(doc, texto): doc.add_heading(texto, level=3)

def paso(doc, num, titulo, cuerpo):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8); par.paragraph_format.space_after = Pt(2)
    r = par.add_run(f'Paso {num} — {titulo}')
    r.bold = True; r.font.color.rgb = AZUL_OSCURO
    if cuerpo:
        q = doc.add_paragraph(cuerpo)
        q.paragraph_format.left_indent = Cm(0.6)

def vineta(doc, texto, nivel=0, bold_hasta=None):
    par = doc.add_paragraph(style='List Bullet' if nivel == 0 else 'List Bullet 2')
    if bold_hasta:
        r = par.add_run(bold_hasta); r.bold = True
        par.add_run(texto)
    else:
        par.add_run(texto)
    return par

def caja(doc, titulo, cuerpo, fill='EFF6FF', color_titulo=AZUL_OSCURO):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(16.5); t.rows[0].cells[0].width = Cm(16.5)
    c = t.rows[0].cells[0]; sombrear(c, fill)
    par = c.paragraphs[0]
    r = par.add_run(titulo); r.bold = True; r.font.color.rgb = color_titulo; r.font.size = Pt(10.5)
    if cuerpo:
        q = c.add_paragraph(cuerpo); q.paragraph_format.space_after = Pt(2)
        for rr in q.runs: rr.font.size = Pt(10.5)
    p(doc, '', despues=2)

def advertencia(doc, cuerpo, titulo='⚠ OJO'):
    caja(doc, titulo, cuerpo, fill='FEF3C7', color_titulo=AMBAR)

def regla(doc, cuerpo, titulo='🔒 Regla del sistema'):
    caja(doc, titulo, cuerpo, fill='FEE2E2', color_titulo=ROJO)

def caso(doc, cuerpo, titulo='🧾 Caso real'):
    caja(doc, titulo, cuerpo, fill='ECFDF5', color_titulo=VERDE)

def captura(doc, pantalla, detalle=''):
    n = len(CAPTURAS) + 1
    CAPTURAS.append((n, pantalla, detalle))
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(16.5); t.rows[0].cells[0].width = Cm(16.5)
    c = t.rows[0].cells[0]; sombrear(c, 'F1F5F9')
    par = c.paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(f'📸 CAPTURA {n} · {pantalla}')
    r.bold = True; r.font.color.rgb = GRIS; r.font.size = Pt(10)
    if detalle:
        q = c.add_paragraph(detalle); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in q.runs: rr.font.size = Pt(9); rr.font.color.rgb = GRIS_SUAVE; rr.italic = True
    p(doc, '', despues=2)

def ficha(doc, quien, permisos, prereq, pantalla):
    t = doc.add_table(rows=4, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchos = (Cm(3.4), Cm(13.1))
    filas = (('Pantalla', pantalla), ('Quién', quien), ('Permisos', permisos), ('Prerequisitos', prereq))
    for i, (k, v) in enumerate(filas):
        t.rows[i].cells[0].width, t.rows[i].cells[1].width = anchos
        c0, c1 = t.rows[i].cells
        sombrear(c0, 'EFF6FF')
        r = c0.paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUL_OSCURO
        r = c1.paragraphs[0].add_run(v); r.font.size = Pt(10)
    p(doc, '', despues=2)

def flujo(doc, texto):
    par = p(doc, texto, bold=True, color=AZUL, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, antes=6, despues=6)
    return par

def tabla(doc, encabezados, filas, anchos_cm):
    t = doc.add_table(rows=1, cols=len(encabezados)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (enc, w) in enumerate(zip(encabezados, anchos_cm)):
        c = t.rows[0].cells[i]; c.width = Cm(w); sombrear(c, '012D70')
        r = c.paragraphs[0].add_run(enc); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for fila in filas:
        cells = t.add_row().cells
        for i, (val, w) in enumerate(zip(fila, anchos_cm)):
            cells[i].width = Cm(w)
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
    p(doc, '', despues=2)

def toc(doc):
    par = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
    inner = OxmlElement('w:r'); t = OxmlElement('w:t')
    t.text = 'Índice — clic derecho → Actualizar campos'
    inner.append(t); fld.append(inner)
    par._p.append(fld)
