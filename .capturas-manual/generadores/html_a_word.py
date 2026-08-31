# -*- coding: utf-8 -*-
"""Conversor de los documentos vivos (HTML de /mantenedores/documentacion/) a Word.

Regenera las copias Word de la Suite desde su fuente HTML, con el estilo de la
biblioteca (estilo.py). Uso:

    python html_a_word.py procesos "Manual de Procesos" salida.docx
    python html_a_word.py manual-usuario "Manual de Usuario" salida.docx
    python html_a_word.py config-maestro "Configuración Maestro" salida.docx

Mapeo: h1=portada · h2=capítulo (salto de página) · h3=sección · p/li=texto con
negritas · .meta=fichas Quién/Prerequisitos · .flujo=cinta de estados ·
.afecta/.tip/.nota=cajas · .max=regla · .paso=paso · .shot=captura pendiente ·
tablas HTML=tablas Word · .toc → campo TOC de Word.
"""
import io, os, re, sys
from html.parser import HTMLParser
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

DOCS_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__),
    '..', '..', 'api-gateway', 'public', 'mantenedores', 'documentacion'))

BLOQUES = {'meta', 'flujo', 'afecta', 'max', 'paso', 'tip', 'nota', 'shot', 'toc', 'rev', 'tag'}

class Conv(HTMLParser):
    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.pila = []            # tags abiertos
        self.bloque = None        # clase del bloque especial activo
        self.bloque_prof = 0
        self.partes = []          # [(texto, bold)] del párrafo en curso
        self.modo = None          # 'p' | 'li' | 'h1' | 'h2' | 'h3'
        self.bold = 0
        self.saltar = 0           # dentro de style/script/topnav
        self.primer_h2 = True
        self.li_nivel = 0
        # tablas
        self.tabla = None; self.fila = None; self.celda = None

    # ── helpers ──────────────────────────────────────────────────────────
    def flush(self, tipo=None):
        partes = [(t, b) for t, b in self.partes if t]
        self.partes = []
        texto = ''.join(t for t, _ in partes)
        if not texto.strip(): return
        tipo = tipo or self.modo
        if self.bloque == 'shot':
            captura(self.doc, texto.strip(), ''); return
        if tipo == 'h2':
            h1(self.doc, texto.strip(), salto=not self.primer_h2); self.primer_h2 = False; return
        if tipo == 'h3': h2(self.doc, texto.strip()); return
        if tipo == 'li':
            # vineta con negritas: aprovechar bold_hasta si el primer run es bold
            if partes and partes[0][1] and len(partes) > 1:
                vineta(self.doc, partes[0][0] + ''.join(t for t, _ in partes[1:]),
                       nivel=self.li_nivel, bold_hasta=partes[0][0])
            else:
                vineta(self.doc, texto, nivel=self.li_nivel)
            return
        if self.bloque == 'flujo':
            flujo(self.doc, re.sub(r'\s+', ' ', texto).strip()); return
        if self.bloque == 'afecta':
            caja(self.doc, '', texto.strip(), fill='FFFBEB', color_titulo=AZUL_OSCURO); return
        if self.bloque == 'max':
            regla(self.doc, texto.strip(), titulo='🔒 Regla del sistema'); return
        if self.bloque in ('tip', 'nota'):
            caja(self.doc, '💡 ' + ('Tip' if self.bloque == 'tip' else 'Nota'), texto.strip()); return
        if self.bloque == 'paso':
            if partes and partes[0][1]:
                runs(self.doc, [(partes[0][0], {'bold': True, 'color': AZUL_OSCURO})] +
                               [(t, {'bold': b}) for t, b in partes[1:]])
            else:
                p(self.doc, texto)
            return
        if self.bloque == 'meta':
            if partes and partes[0][1] and len(partes) > 1:
                vineta(self.doc, partes[0][0] + ': ' + ''.join(t for t, _ in partes[1:]).strip(),
                       bold_hasta=partes[0][0] + ': ')
            else:
                vineta(self.doc, texto)
            return
        runs(self.doc, [(t, {'bold': b}) for t, b in partes])

    # ── parser ───────────────────────────────────────────────────────────
    def handle_starttag(self, tag, attrs):
        a = dict(attrs); cls = (a.get('class') or '').split()
        if tag in ('style', 'script'): self.saltar += 1; return
        if self.saltar: return
        if tag == 'i' and any(c.startswith('bi') for c in cls): return   # íconos
        if tag == 'div':
            c = next((c for c in cls if c in BLOQUES), None)
            if c and not self.bloque:
                self.flush()
                self.bloque = c; self.bloque_prof = 1
                if c == 'toc':
                    h1(self.doc, 'Índice', salto=False); toc(self.doc)
                return
            if self.bloque: self.bloque_prof += 1
            return
        if tag == 'table':
            self.flush(); self.tabla = []; return
        if self.tabla is not None:
            if tag == 'tr': self.fila = []
            elif tag in ('td', 'th'): self.celda = ''
            return
        if tag in ('h1', 'h2', 'h3'): self.flush(); self.modo = tag; return
        if tag == 'p': self.flush(); self.modo = 'p'; return
        if tag == 'li': self.flush(); self.modo = 'li'; return
        if tag in ('ul', 'ol'):
            if self.modo == 'li' or self.li_nivel: self.li_nivel += 1
            return
        if tag in ('b', 'strong', 'code'): self.bold += 1
        if tag == 'br': self.flush('p')

    def handle_endtag(self, tag):
        if tag in ('style', 'script'): self.saltar = max(0, self.saltar - 1); return
        if self.saltar: return
        if self.tabla is not None:
            if tag in ('td', 'th'):
                self.fila.append(re.sub(r'\s+', ' ', self.celda or '').strip()); self.celda = None
            elif tag == 'tr':
                if self.fila: self.tabla.append(tuple(self.fila))
                self.fila = None
            elif tag == 'table':
                if len(self.tabla) > 1:
                    n = len(self.tabla[0]); ancho = 16.5 / max(1, n)
                    filas = [f + ('',) * (n - len(f)) for f in self.tabla[1:]]
                    tabla(self.doc, self.tabla[0], tuple(filas), tuple([ancho] * n))
                self.tabla = None
            return
        if tag == 'div':
            if self.bloque:
                self.bloque_prof -= 1
                if self.bloque_prof <= 0:
                    if self.bloque != 'toc': self.flush()
                    self.bloque = None
            return
        if tag in ('h1', 'h2', 'h3', 'p', 'li'):
            self.flush(); self.modo = None; return
        if tag in ('ul', 'ol') and self.li_nivel: self.li_nivel -= 1
        if tag in ('b', 'strong', 'code'): self.bold = max(0, self.bold - 1)

    def handle_data(self, data):
        if self.saltar: return
        if self.celda is not None: self.celda += data; return
        if self.bloque == 'toc': return
        t = re.sub(r'\s+', ' ', data)
        if not t: return
        if self.partes and self.partes[-1][1] == bool(self.bold):
            self.partes[-1] = (self.partes[-1][0] + t, bool(self.bold))
        else:
            self.partes.append((t, bool(self.bold)))


def convertir(archivo_html, titulo, subtitulo=''):
    html = io.open(os.path.join(DOCS_DIR, archivo_html + '.html'), encoding='utf-8').read()
    # solo el cuerpo del documento
    m = re.search(r'<div class="doc">(.*)</div>\s*(?:<script|</body)', html, re.S)
    cuerpo = m.group(1) if m else html
    ver = re.search(r'<!--\s*(v[\d.]+)', html)

    doc = nuevo_doc()
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, titulo.upper(), bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    if subtitulo:
        p(doc, '', despues=20)
        p(doc, subtitulo, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    import datetime
    hoy = datetime.date.today()
    MESES = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
    p(doc, f'Copia Word del documento vivo ({ver.group(1) if ver else "s/v"}) · {hoy.day} de {MESES[hoy.month-1]} de {hoy.year}',
      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'La versión de referencia vive en el sistema: Mantenedores → Documentación',
      size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    c = Conv(doc)
    c.feed(cuerpo)
    c.flush()
    return doc


DOCS = {
    'procesos': ('Manual de Procesos',
                 'Los procesos del negocio de punta a punta — cruzando módulos.',
                 'Manual-de-Procesos-Business-Suite.docx'),
    'manual-usuario': ('Manual de Usuario',
                       'Pantalla por pantalla: qué hace cada módulo y cómo se usa.',
                       'Manual-de-Usuario-Business-Suite.docx'),
    'config-maestro': ('Configuración Maestro',
                       'Todos los mantenedores y sus variables: qué gobierna cada parámetro.',
                       'Configuracion-Maestro-Business-Suite.docx'),
}

if __name__ == '__main__':
    destino = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\patri\Documents'
    for clave, (titulo, sub, fname) in DOCS.items():
        d = convertir(clave, titulo, sub)
        out = os.path.join(destino, fname)
        d.save(out)
        print('OK ->', out)
