# -*- coding: utf-8 -*-
"""Kit de Levantamiento — Internacionalización Business Suite (compartible)."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def preguntas(doc, items):
    for i, q in enumerate(items, 1):
        p(doc, f'{i}. {q}', size=10.5)

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'KIT DE LEVANTAMIENTO', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'Internacionalización Business Suite — Ecuador y Bolivia', color=AZUL, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Cuestionarios por área para el levantamiento en terreno,', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'y el caso de negocio de operar en la nube sin infraestructura propia.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Septiembre 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'AutoFácil Crédito Automotriz — documento de trabajo para el Grupo', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Índice')
    toc(doc)

    h1(doc, '1. Cómo usar este kit')
    p(doc, 'Cada cuestionario se trabaja en una o dos sesiones con el responsable del área en el país. Las '
           'respuestas se documentan y se firman al cierre del levantamiento en el Documento de Alcance del '
           'MVP. Regla práctica: si una respuesta es "depende", se anota el caso y de quién depende — los '
           '"depende" sin dueño son los que atrasan los proyectos.')
    vineta(doc, 'quién responde cada área, con nombre y cargo, acordado antes del viaje.', bold_hasta='Contraparte: ')
    vineta(doc, 'pedir SIEMPRE ejemplos reales: un crédito completo, una factura, un cierre de mes impreso.', bold_hasta='Evidencia: ')
    vineta(doc, 'lo que no exista hoy (ej.: no hay tabla de amortización formal) también es una respuesta — define lo que la Suite va a aportar.', bold_hasta='Vacíos: ')

    # ── 2. Producto crediticio ──
    h1(doc, '2. Cuestionario A — El producto crediticio')
    ficha(doc, 'Gerencia Comercial / de Crédito', '—', 'Fichas o contratos de cada producto vigente', 'Semana 1 del levantamiento')
    preguntas(doc, [
        '¿Qué tipos de crédito ofrecen hoy (consumo automotriz, comercial, con garantía, leasing)? Ficha de cada uno.',
        '¿Moneda de los créditos? ¿Existe alguna unidad reajustable o son montos nominales?',
        'Tasas: ¿fijas o variables? ¿Cómo se determinan? ¿Existe tasa máxima legal y quién la fija?',
        'Plazos típicos y montos mínimo/máximo por producto. ¿Pie mínimo? ¿Se financian accesorios (GPS, gastos)?',
        'Cuota: ¿sistema francés? ¿Hay períodos de gracia, cuotas especiales o balloon?',
        'Seguros asociados: ¿cuáles, con qué compañías, obligatorios u opcionales, cómo se cobran (contado/financiado)?',
        'Garantías: ¿prenda del vehículo? ¿Cómo se constituye y ante quién? ¿Costo y plazo del trámite?',
        '¿Quién origina? (fuerza propia, concesionarios, brokers) ¿Cómo se les paga la comisión y contra qué documento?',
        'Política de crédito: ¿score, buró (cuál), reglas de aprobación, niveles de atribución, excepciones?',
        '¿Existe pre-aprobación o cotización en el punto de venta? ¿Con qué herramienta hoy?',
        'Prepago: ¿está permitido? ¿Con qué comisión o castigo? ¿Cómo se calcula el saldo insoluto?',
        'Volumen: operaciones por mes, cartera vigente (número y monto), ticket promedio, mora por tramos.',
    ])

    # ── 3. Operación ──
    h1(doc, '3. Cuestionario B — El flujo operativo')
    ficha(doc, 'Gerencia de Operaciones', '—', 'Organigrama del área y un expediente de crédito completo', 'Semana 2')
    preguntas(doc, [
        'Organigrama real: quién hace qué desde la solicitud hasta el pago al concesionario. Nombres y cargos.',
        'Paso a paso de una operación: solicitud → evaluación → aprobación → documentación → desembolso. Tiempos reales de cada paso.',
        '¿Qué documentos componen el expediente (fundantes)? ¿Cuáles son obligatorios y quién los valida?',
        '¿Cómo se paga al concesionario? (saldo de precio, comisión) ¿Contra qué documento y con qué plazos?',
        '¿Cómo se registran hoy los pagos de los clientes? (caja propia, bancos, corresponsales, débito automático)',
        'Cobranza: ¿quién la hace, con qué herramientas, hay normativa local de contacto (horarios, frecuencia)?',
        '¿Hay castigo de cartera y venta de cartera? ¿Con qué políticas?',
        '¿Qué sistemas usan hoy (core, planillas)? ¿Qué datos históricos hay que migrar y en qué formato están?',
        '¿Cuántos usuarios operarían la Suite y con qué roles?',
        '¿Qué reportes recibe hoy la gerencia y con qué frecuencia? Ejemplos impresos.',
    ])

    # ── 4. Normativa e impuestos ──
    h1(doc, '4. Cuestionario C — Normativa e impuestos')
    ficha(doc, 'Gerencia de Finanzas + asesor legal/tributario local', '—', 'Última declaración de impuestos y un talonario/factura', 'Semana 3')
    preguntas(doc, [
        '¿La empresa está regulada por el supervisor financiero (ASFI / Superintendencia de Bancos)? Si sí: ¿qué reportes, provisiones normadas y exigencias de capital le aplican?',
        'IVA: tasa vigente, ¿el interés del crédito está gravado o exento? ¿Y las comisiones y seguros?',
        'Retenciones: ¿a proveedores, a profesionales, a intereses? Tasas y formularios.',
        'Facturación electrónica: ¿obligatoria? ¿Proveedor autorizado o conexión directa (SRI / SIN-SIAT)? ¿Qué documentos (factura, nota de crédito, retención)?',
        'Tasa máxima de interés: ¿quién la fija, con qué periodicidad y dónde se publica?',
        'Normativa de cobranza y protección al consumidor: ¿límites de contacto, intereses moratorios máximos, condonaciones?',
        'Identificación tributaria y de personas: formato de cédula/RUC/NIT y su dígito verificador.',
        'Feriados nacionales y bancarios del año.',
        '¿Qué libros y declaraciones exige el fisco (equivalente al F29 chileno)? ¿Con qué periodicidad?',
        'Protección de datos personales: ¿ley vigente y requisitos para datos de clientes en la nube?',
    ])

    # ── 5. Contabilidad ──
    h1(doc, '5. Cuestionario D — Contabilidad')
    ficha(doc, 'Contador General local', '—', 'Plan de cuentas vigente y el último balance', 'Semanas 3–4')
    preguntas(doc, [
        'Plan de cuentas vigente (archivo). ¿Bajo qué marco: IFRS, norma local, manual del regulador?',
        '¿Cómo se contabiliza hoy la colocación de un crédito, el devengo de intereses y el pago de cuotas?',
        'Provisiones de incobrables: ¿política propia o tramos normados por el regulador? Tabla vigente.',
        '¿Cómo se contabilizan los pagos a concesionarios (saldo de precio y comisión) y sus impuestos?',
        'Cierre de mes: checklist actual, quién lo hace, cuántos días tarda y qué reportes salen a casa matriz.',
        'Moneda funcional y de presentación. En Bolivia: ¿mantienen ajuste por inflación / UFV en alguna cuenta?',
        '¿Con qué sistema contable trabajan hoy y qué historia habría que migrar (saldos de apertura o libros completos)?',
        '¿Auditoría externa? ¿Quién y con qué exigencias de respaldo documental?',
    ])

    # ── 6. TI y datos ──
    h1(doc, '6. Cuestionario E — TI, datos e infraestructura')
    ficha(doc, 'Responsable de TI (o quien haga sus veces)', '—', 'Inventario de sistemas y accesos', 'Transversal')
    preguntas(doc, [
        '¿Qué infraestructura tienen hoy (servidores propios, hosting, nube)? ¿Quién la administra?',
        'Calidad de la conectividad a internet en las oficinas y puntos de venta.',
        'Correo corporativo: ¿dominio propio, con qué proveedor?',
        '¿Existen respaldos hoy? ¿Con qué frecuencia y dónde? ¿Se ha probado una restauración?',
        'Datos a migrar: sistemas de origen, formatos exportables, volumen y calidad (duplicados, campos vacíos).',
        '¿Restricciones legales o de política de grupo para datos en la nube (país de residencia de los datos)?',
        '¿Quién sería el administrador local de la Suite (usuarios, permisos, mantenedores)?',
    ])

    # ── 7. La nube: el caso de negocio ──
    h1(doc, '7. Sin infraestructura propia: el caso de negocio')
    p(doc, 'La Suite opera 100% en la nube: servidor administrado (Render), base de datos administrada '
           '(TiDB Cloud), documentos en bucket (Google Cloud) y respaldos automáticos. La empresa NO compra '
           'ni administra ningún equipo. Esto no es un detalle técnico: es una decisión económica y de riesgo.')
    h2(doc, '7.1 Ventajas de operar sin fierros propios')
    vineta(doc, 'el nodo completo de un país se levanta en horas, no en semanas de compras e instalación.', bold_hasta='Velocidad: ')
    vineta(doc, 'no hay inversión inicial (CAPEX cero): todo es gasto mensual pequeño y cancelable.', bold_hasta='Capital: ')
    vineta(doc, 'actualizaciones, parches de seguridad, monitoreo y escalamiento los hace el proveedor; no se necesita personal de sistemas dedicado.', bold_hasta='Administración: ')
    vineta(doc, 'si el negocio crece, se sube de plan con un clic; si un país no prospera, se apaga y el costo desaparece.', bold_hasta='Elasticidad: ')
    vineta(doc, 'respaldo diario del proveedor + respaldo propio nocturno con 30 días de retención + host de contingencia que cuesta ~US$1/mes dormido y se promueve en minutos. Esta redundancia ya está diseñada, probada y documentada en Chile.', bold_hasta='Continuidad: ')
    h2(doc, '7.2 Comparación de costos: nube vs. infraestructura local equivalente')
    tabla(doc, ('Concepto', 'Nube (lo nuestro)', 'Infraestructura local equivalente'),
          (('Inversión inicial', 'US$ 0', 'US$ 8.000 – 20.000 (2 servidores, UPS, red, licencias de SO y base de datos)'),
           ('Servidor de aplicación', 'US$ 7–25 /mes', 'Depreciación + energía + sala: US$ 150–300 /mes'),
           ('Base de datos', 'US$ 5–25 /mes (administrada)', 'Licencia + DBA a demanda: US$ 200–500 /mes'),
           ('Respaldos', 'US$ 0 (incluido + GitHub)', 'Equipo/servicio de backup + cintas o segundo sitio: US$ 50–150 /mes'),
           ('Personal de sistemas', 'No requerido', '½ a 1 jornada de TI: US$ 400 – 1.200 /mes'),
           ('Contingencia (redundancia)', '~US$ 1 /mes (standby dormido)', 'Segundo servidor en otro sitio + enlace + pruebas: US$ 200–500 /mes'),
           ('TOTAL mensual', 'US$ 40 – 90', 'US$ 1.000 – 2.650 (más la inversión inicial)')),
          (4.2, 5.0, 7.2))
    p(doc, 'Orden de magnitud: la nube cuesta entre un 5% y un 10% de lo que costaría montar y mantener una '
           'infraestructura local con un nivel de servicio comparable. En dos países, el ahorro anual '
           'estimado supera los US$ 25.000 – 60.000, sin contar la inversión inicial evitada.')
    h2(doc, '7.3 Riesgos de intentar la redundancia en forma local')
    tabla(doc, ('Riesgo local', 'Detalle'),
          (('Redundancia de papel', 'El segundo servidor local rara vez se prueba: cuando se necesita, no enciende o está desactualizado. Nuestro standby en la nube se reconstruye SOLO todos los días y su promoción está ensayada.'),
           ('Dependencia de una persona', 'La infraestructura local depende del técnico que la armó; su salida es un riesgo operacional completo.'),
           ('Seguridad física y eléctrica', 'Cortes de energía, robo, incendio, humedad: el servidor vive en la oficina, con la operación.'),
           ('Obsolescencia', 'Los equipos se deprecian y exigen recambio cada 4–5 años (nueva inversión); la nube se renueva sola.'),
           ('Respaldo no verificado', 'Un backup local sin prueba de restauración es una ilusión de seguridad. El esquema de la Suite incluye restauración documentada y ensayada.'),
           ('Escalar tarde o de más', 'Con fierros se compra para el peak (capital ocioso) o se queda corto en el peak (operación caída).')),
          (4.6, 11.9))
    regla(doc, 'Conclusión: replicar localmente la disponibilidad y redundancia que la nube entrega por '
               'US$ 40–90 mensuales costaría 10 a 20 veces más, con MENOR confiabilidad real. La '
               'recomendación del plan es no comprar ningún equipo en ningún país.')

    # ── 8. Cierre ──
    h1(doc, '8. Entregable del levantamiento')
    p(doc, 'Al terminar las 3–4 semanas, este kit respondido se consolida en el DOCUMENTO DE ALCANCE DEL MVP: '
           'qué productos y procesos entran, el plan de cuentas local, los parámetros normativos, los datos a '
           'migrar y la contraparte comprometida — firmado por la gerencia local. Con eso se crea el nodo y '
           'parte la localización.')

    return doc
