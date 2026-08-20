# -*- coding: utf-8 -*-
"""Manual del Administrador — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DEL ADMINISTRADOR', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Usuarios, permisos, mantenedores y salud del sistema:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'gobernar la plataforma sin tocar código.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Administrador'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo: para qué existe, quién lo hace, prerequisitos, paso a paso y los recuadros '
           '⚠ OJO / 🔒 Regla del sistema / 🧾 Caso real / 📸 CAPTURA (pendientes en esta versión).')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Filosofía ────────────────────────────────────────────────────────
    h1(doc, '1. La filosofía: aplicación paramétrica')
    p(doc, 'El principio rector del sistema es que el Administrador pueda hacer la mayor cantidad de '
           'modificaciones posible desde los mantenedores, sin tocar código y sin alterar el espíritu '
           'ni el flujo de los procesos. Así está construida la mayoría de la plataforma: módulos, '
           'permisos, etapas, tasas, plantillas, avisos, umbrales.')
    vineta(doc, 'montos, tramos, nombres de estados, plazos, textos de plantillas, listas de opciones y mapeos viven en mantenedores. El código solo tiene lógica.', bold_hasta='Datos de negocio = mantenedor. ')
    vineta(doc, 'parametrizar no significa poder romper el proceso: el orden de las etapas, las validaciones y las atribuciones siguen protegidas. Lo que se abre es el contenido (valores, textos, umbrales), no la estructura.', bold_hasta='El flujo se respeta, los parámetros se ajustan. ')
    vineta(doc, 'cada página de mantenedor termina con el recuadro "Qué afecta este mantenedor", que lista qué proceso impacta cada variable, con link. Antes de cambiar un valor, leerlo.', bold_hasta='El footer de trazabilidad. ')
    regla(doc, 'Si un cambio de negocio parece requerir programador, primero preguntarse: ¿existe el '
               'mantenedor y no lo encontré? La respuesta suele estar en el buscador de Mantenedores o '
               'en el footer "Qué afecta" de un mantenedor vecino.')
    captura(doc, 'Mantenedores → home', 'La grilla de cards de mantenedores.')

    # ── 2. Usuarios y permisos ──────────────────────────────────────────────
    h1(doc, '2. Usuarios, perfiles y permisos')
    p(doc, 'El acceso se gobierna por la matriz de Perfiles y Permisos: cada funcionalidad tiene su '
           'código, cada perfil marca qué funcionalidades ve, y cada usuario tiene un perfil (más '
           'overrides individuales cuando se justifican).')
    ficha(doc,
          'Administrador',
          'Administrador (bypass total) — usar con juicio',
          'Perfil definido para el rol antes de crear el usuario',
          'Usuarios → Gestión de usuarios y permisos')
    h2(doc, '2.1 Conceptos')
    vineta(doc, 'con href definido generan sub-items en los menús de su sección; con href NULL son permisos de acción (crear, editar, eliminar, aprobar).', bold_hasta='Funcionalidades: ')
    vineta(doc, 'el Administrador ve todo sin restricciones. Los demás perfiles, solo lo marcado en la matriz.', bold_hasta='Perfiles: ')
    vineta(doc, 'controla qué ejecutivos ve cada usuario en comisiones y otras vistas con visibilidad acotada.', bold_hasta='Usuario-ejecutivos: ')
    vineta(doc, 'las cards de las landing se ocultan según los permisos del usuario: si alguien no ve una card, el problema es de permisos, no de la página.', bold_hasta='Cards por permiso: ')
    h2(doc, '2.2 Alta de un usuario')
    paso(doc, 1, 'Crear el usuario', 'Con su perfil. La ficha alimenta también RRHH (fecha de ingreso '
         'es dato de cálculo para bonos y equipo).')
    paso(doc, 2, 'Revisar visibilidad', 'Si el rol lo requiere, asignar qué ejecutivos ve.')
    paso(doc, 3, 'Auditar', 'Tras cambios masivos de permisos o perfiles nuevos, correr la auditoría '
         'de permisos (script de integridad: duplicados, huérfanos, matriz por perfil).')
    captura(doc, 'Usuarios → matriz de permisos de un perfil', 'La matriz con funcionalidades marcadas.')
    h2(doc, '2.3 Seguridad de acceso')
    vineta(doc, 'tras N intentos fallidos la cuenta se bloquea (paramétrico).', bold_hasta='Bloqueo de login: ')
    vineta(doc, 'existe una cuenta protegida de emergencia (el Administrador sabe cuál); no se documenta más que esto.', bold_hasta='Emergencia: ')
    vineta(doc, 'el token de sesión vive en la pestaña (cerrar la pestaña = cerrar sesión); solo la PWA de Terreno persiste.', bold_hasta='Sesiones: ')
    vineta(doc, 'el login tiene tope de intentos por minuto y TODA el API tiene techo de peticiones por usuario.', bold_hasta='Rate limiting: ')

    # ── 3. Módulos desde BD ─────────────────────────────────────────────────
    h1(doc, '3. Módulos y menús: todo desde la base')
    p(doc, 'Los módulos del home y los sub-items de cada sección NO están escritos en el código: salen '
           'de la base (tablas de módulos y funcionalidades) según los permisos de quien mira.')
    paso(doc, 1, 'Agregar un módulo nuevo', 'Insertar el módulo (nombre, ícono, ruta, orden) y sus '
         'funcionalidades (nombre, código, href, ícono) — desde la Consola SQL o pidiendo el insert a '
         'TI.')
    paso(doc, 2, 'Asignar permisos', 'Marcar la funcionalidad en los perfiles que corresponda. Sin '
         'este paso, la card NUEVA NO APARECE para nadie (error clásico).')
    paso(doc, 3, 'Ordenar', 'El botón Reordenar del home permite arreglar la disposición de cards por '
         'perfil (placement).')
    advertencia(doc, 'Si una card nueva no aparece, el problema está en la base o en los permisos del '
                     'perfil — no en el código. Revisar: ¿existe la funcionalidad? ¿tiene href? ¿está '
                     'marcada en el perfil del usuario?')

    # ── 4. Mantenedores clave ───────────────────────────────────────────────
    h1(doc, '4. Los mantenedores que mueven el negocio')
    p(doc, 'El catálogo completo vive en el módulo Mantenedores; estos son los que definen plata y '
           'flujo, con lo que toca cada uno:')
    tabla(doc, ('Mantenedor', 'Qué gobierna', 'Ojo con'),
          (('Tasas', 'Pizarra de tasas por tramo', 'Rige cotizador, cartas y revisor automático'),
           ('Parámetros del Crédito', 'Umbrales de preferencia, rentabilidad mínima, tolerancias', 'Cambia veredictos de ¿Dónde Curso?'),
           ('Preferencia Financiera', 'Elegibilidad por plazo/saldo de cada financiera', 'Bloquea combinaciones en la carta'),
           ('Dealers', 'Ficha, banco, tabla pactada, parque/calle', 'La tabla pactada manda sobre la pizarra'),
           ('Parques (Arriendos y Comisiones)', 'Arriendo mensual y comisión por parque', 'Alimenta la cartola del parque'),
           ('Excepciones Comerciales', 'Piso de rentabilidad, escalera, revisor automático', 'Switches del revisor parten apagados'),
           ('Estado Créditos', 'Máquina de estados y transiciones', 'Hoy dibuja el flujo; aún no bloquea'),
           ('Estado Cartera', 'Estados de pago de cartera propia', 'Segunda dimensión, no confundir con etapa'),
           ('Parámetros de Cobranza', 'Tramos, gastos, tasa de mora, provisiones', 'La TMC queda fija al otorgar'),
           ('Comisiones Variables', 'Modelo de comisión de ejecutivos', 'Versionado con vigencia y bitácora'),
           ('Reglas de Centralización', 'Qué asiento genera cada hecho económico', 'Un hecho sin regla queda SIN_REGLA en el log'),
           ('Tipos de Documento', 'Matriz de fundantes por financiera', 'Define qué exige el curse'),
           ('Avisos', 'Campanitas: eventos y destinatarios', 'Ver capítulo 5'),
           ('Correos Programados', 'Correos automáticos y pop-ups', 'Incluye la rendición semanal de fundantes'),
           ('UF / Indicadores', 'UF, UTM, dólar, TMC desde la API oficial', 'Se sincronizan solos; revisar si fallan')),
          (4.2, 6.6, 5.7))
    captura(doc, 'Un mantenedor con su footer "Qué afecta"', 'El recuadro de trazabilidad al pie, con sus links.')
    regla(doc, 'Regla de oro antes de cambiar un parámetro: leer el footer "Qué afecta este '
               'mantenedor". Si la variable que vas a tocar no está listada ahí, el cambio está '
               'incompleto — avisar a TI para que lo agregue.')

    # ── 5. Avisos y correos ─────────────────────────────────────────────────
    h1(doc, '5. Avisos, campanitas y correos')
    p(doc, 'Todos los avisos del sistema — campanitas, correos automáticos, pop-ups — son paramétricos: '
           'qué evento avisa, a quién, con qué texto y si está encendido.')
    vineta(doc, 'cada evento (carta pendiente, anulación solicitada, incorporación por aprobar…) tiene sus destinatarios configurables por perfil o usuario.', bold_hasta='Mantenedor de Avisos: ')
    vineta(doc, 'los envíos programados (informes, recordatorios, pop-up de rendición) con su frecuencia, texto e interruptor.', bold_hasta='Correos Programados: ')
    vineta(doc, 'una campanita retirada obedece reglas de retiro — el aviso muere cuando el hecho se resuelve, no cuando alguien lo lee.', bold_hasta='Retiro: ')
    vineta(doc, 'todo correo enviado queda en Auditoría → Correos Enviados.', bold_hasta='Trazabilidad: ')
    advertencia(doc, 'El Modo Desarrollo redirige TODO el correo del ambiente a cuentas de prueba. Si '
                     'los correos "no llegan" en staging, es eso — no un bug.')
    captura(doc, 'Mantenedores → Avisos', 'La grilla de eventos con destinatarios e interruptores.')

    # ── 6. Salud del sistema ────────────────────────────────────────────────
    h1(doc, '6. Salud del sistema')
    p(doc, 'La plataforma corre en la nube con auto-deploy: un cambio publicado por TI queda en '
           'producción en 2-3 minutos. El Administrador vigila la salud con estas herramientas:')
    vineta(doc, 'el badge de versión (esquina de todas las páginas) confirma qué versión está desplegada — si no coincide con la anunciada, el deploy no se aplicó.', bold_hasta='Versión: ')
    vineta(doc, 'el endpoint de salud reporta base de datos, uptime, documentos y motores apagados.', bold_hasta='/api/health: ')
    vineta(doc, '28 motores automáticos actúan solos (aprobar comisiones, desistir vencidos, devengos, correos). Todos pasan por un interruptor central: en el host de contingencia corren apagados a propósito.', bold_hasta='Motores: ')
    vineta(doc, 'cada error 500 en producción dispara correo al administrador (con freno de 10 minutos por ruta).', bold_hasta='Alertas de error: ')
    vineta(doc, 'sondeo de servicios cada 5 minutos con su historial.', bold_hasta='Uptime: ')
    captura(doc, 'Monitoreo / Salud del sistema', 'El panel con los servicios y su estado.')
    h2(doc, '6.1 Respaldos y contingencia')
    vineta(doc, 'la nube guarda 1 día; el respaldo propio nocturno guarda 30 días.', bold_hasta='Base de datos: ')
    vineta(doc, 'los archivos viven en el bucket de documentos, no en la base (eran el 95% del tamaño).', bold_hasta='Documentos: ')
    vineta(doc, 'existe un host de contingencia dormido con la misma base, que se promueve con un comando siguiendo el runbook. Sus motores están apagados: NUNCA encenderlos con el principal vivo — cada tarea correría dos veces y el daño es silencioso.', bold_hasta='Contingencia: ')
    regla(doc, 'Dos procesos con los motores encendidos contra la misma base duplican comisiones y '
               'devengos sin que nada reclame. El orden de la vuelta atrás está en el runbook: primero '
               'apagar los motores del standby.')

    # ── 7. Auditoría ────────────────────────────────────────────────────────
    h1(doc, '7. Auditoría y Consola SQL')
    vineta(doc, 'las acciones críticas quedan registradas: quién, qué, cuándo. Módulo Auditoría con sus vistas (acciones, correos enviados, bitácoras de revisores).', bold_hasta='Auditoría: ')
    vineta(doc, 'consulta directa de solo lectura a la base, restringida al perfil más alto y auditada. Para responder preguntas puntuales sin pedir informes a TI.', bold_hasta='Consola SQL: ')
    vineta(doc, 'los registros inmutables (bitácora del revisor, bitácoras de variables) solo agregan filas: no existe editar ni borrar desde el sistema.', bold_hasta='Inmutables: ')
    captura(doc, 'Auditoría', 'El log de acciones con sus filtros.')

    # ── 8. Documentación ────────────────────────────────────────────────────
    h1(doc, '8. La Suite de Documentación')
    p(doc, 'Mantenedores → Documentación reúne los documentos vivos del sistema: Manual de Usuario, '
           'Procesos (un capítulo por proceso end-to-end), Configuración Maestra, Documentación '
           'Técnica, Brochure y Pendientes Abiertos.')
    regla(doc, 'Máxima del proyecto: todo proceso nace documentado. Cada funcionalidad nueva debe '
               'actualizar el documento que corresponda EN el mismo trabajo que la crea — un proceso '
               'sin capítulo está incompleto. Como Administrador, exigirlo en cada entrega.')
    vineta(doc, 'la lista completa y consolidada de todo lo abierto (técnico, contable, backlog) vive en Pendientes Abiertos. Un pendiente nuevo se anota en el momento en que se detecta.', bold_hasta='Pendientes: ')
    captura(doc, 'Mantenedores → Documentación', 'La suite con sus cards.')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('Una card nueva no aparece', 'Falta el permiso en el perfil (o el href)', 'Revisar funcionalidad y matriz (cap. 3)'),
           ('Un usuario no ve un botón que otro sí ve', 'Permiso de acción no marcado', 'Matriz del perfil, funcionalidades sin href (cap. 2)'),
           ('Los correos no llegan en staging', 'Modo Desarrollo los redirige', 'Es intencional (cap. 5)'),
           ('El badge no muestra la versión anunciada', 'El deploy no se aplicó', 'Deploy manual desde el panel del hosting (cap. 6)'),
           ('Un cambio de parámetro no surtió efecto', 'Caché de permisos (60 s) o parámetro equivocado', 'Esperar 1 minuto; leer el footer "Qué afecta" (cap. 4)'),
           ('Aparecen movimientos sin asiento', 'Hecho económico sin regla de centralización', 'Crear la regla; el log es la lista de pendientes (cap. 4)'),
           ('Dos valores distintos del mismo dato en dos pantallas', 'Violación de motor único / fuente única', 'Reportar a TI: es deuda técnica, no configuración')),
          (5.4, 5.8, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('Mantenedor', 'Pantalla de configuración de un dato de negocio'),
           ('Funcionalidad', 'Unidad de permiso; con href genera menú, sin href es permiso de acción'),
           ('Matriz de permisos', 'Cruce perfil × funcionalidad que gobierna el acceso'),
           ('Placement', 'Disposición de las cards del home, ordenable por perfil'),
           ('Motor automático', 'Tarea de fondo que actúa sola (28 en total), con interruptor central'),
           ('Host de contingencia', 'Réplica dormida del sistema, con los motores apagados'),
           ('Runbook', 'Procedimiento paso a paso para la emergencia'),
           ('Footer "Qué afecta"', 'Recuadro de trazabilidad al pie de cada mantenedor'),
           ('Modo Desarrollo', 'Redirección de todo el correo a cuentas de prueba'),
           ('Staging', 'Ambiente de pruebas separado de producción'),
           ('Consola SQL', 'Consulta directa de solo lectura, restringida y auditada'),
           ('Suite de Documentación', 'Los documentos vivos del sistema, en Mantenedores → Documentación')),
          (4.3, 12.2))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-Administrador-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
