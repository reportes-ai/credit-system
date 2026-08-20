# -*- coding: utf-8 -*-
"""Manual de Canales Digitales y Atención al Cliente — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE CANALES DIGITALES', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'y Atención al Cliente — AutoFácil Business Suite', color=AZUL, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'WhatsApp, Atención Remota, portales y campañas:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'atender a clientes y dealers sin que nadie quede esperando.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Canales Digitales'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo: para qué existe, quién lo hace, prerequisitos, paso a paso y los recuadros '
           '⚠ OJO / 🔒 Regla del sistema / 🧾 Caso real / 📸 CAPTURA (pendientes en esta versión). '
           'Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Mapa de canales ──────────────────────────────────────────────────
    h1(doc, '1. El mapa de canales')
    p(doc, 'Cinco puertas digitales, cada una con su público y su propósito. Todas leen de los mismos '
           'motores del sistema, así que el valor que ve un cliente por WhatsApp es el mismo que ve el '
           'cajero en caja.')
    tabla(doc, ('Canal', 'Para quién', 'Qué resuelve'),
          (('WhatsApp "Facilito"', 'Clientes y prospectos', 'Cotiza, pre-evalúa, informa dónde pagar y deriva a ejecutivos (cap. 4-6)'),
           ('Atención Remota', 'Dealers', 'Chat en vivo con la mesa, documentos y videollamada (cap. 2)'),
           ('Portal del Dealer', 'Dealers', 'Autoconsulta de operaciones, pagos, cartolas y pre-aprobación (cap. 3)'),
           ('Portal del Cliente', 'Clientes finales', 'Autoconsulta de créditos, cuotas y comprobantes (cap. 7)'),
           ('Campañas Masivas', 'Marketing / Cobranza', 'Contacto masivo por mail o WhatsApp con grupo de control (cap. 8)')),
          (3.8, 3.6, 9.1))
    regla(doc, 'Principio transversal: los canales NO calculan nada propio. Cuota, mora, saldo y '
               'estado salen de los motores únicos del sistema. Si un canal muestra un valor distinto '
               'al de otra pantalla, es un bug — reportarlo, no explicarlo al cliente.')

    # ── 2. Atención Remota ──────────────────────────────────────────────────
    h1(doc, '2. Atención Remota: la mesa digital de dealers')
    p(doc, 'Chat en tiempo real entre los dealers y la mesa de atención, con envío de documentos y '
           'videollamada. El dealer entra por su portal con cuenta propia; el ejecutivo atiende desde '
           'la consola, hasta 3 conversaciones en paralelo (configurable).')
    ficha(doc,
          'Ejecutivos de la mesa (consola) · el dealer (portal externo)',
          'atencion_remota · configuración con atencion_remota_config',
          'Cuenta del dealer creada (o su autoregistro aprobado)',
          '/atencion-remota/ (consola) · /portal-dealer/ (dealer)')
    h2(doc, '2.1 Las cuentas del dealer')
    vineta(doc, 'el dealer pide cuenta desde el portal (RUT, razón social, contacto); la solicitud llega a la pestaña Solicitudes y el ejecutivo la aprueba (crea la cuenta y la liga al dealer por RUT) o rechaza. Hay link público de registro con QR para difundir.', bold_hasta='Autoregistro: ')
    vineta(doc, 'cada cuenta tiene un link con llave única: el dealer entra directo al chat sin clave, y la sesión queda recordada. El link se copia o regenera desde la pestaña Cuentas. El login con correo y clave queda de respaldo.', bold_hasta='Acceso por link directo: ')
    advertencia(doc, 'El link de acceso es una llave portadora: quien lo tenga entra como ese dealer. '
                     'Si se filtró, regenerarlo desde la consola — el viejo muere al instante.')
    h2(doc, '2.2 Atender')
    paso(doc, 1, 'Tomar la conversación', 'La consola muestra las conversaciones entrantes; cada '
         'ejecutivo puede llevar hasta 3 en paralelo.')
    paso(doc, 2, 'Identificar al interlocutor', 'Se anota a mano la persona del dealer con quien se '
         'habla; al reabrir ese dealer aparecen sus interlocutores históricos para seleccionar.')
    paso(doc, 3, 'Responder con plantillas', 'El panel de respuestas rápidas (mantenedor propio) carga '
         'el texto en el input sin enviarlo — se ajusta y se manda.')
    paso(doc, 4, 'Recibir documentos', 'El dealer envía documentos escaneados por el mismo chat; '
         'quedan adjuntos a la conversación.')
    paso(doc, 5, 'Videollamada', 'Se inicia desde el chat: prueba de cámara y micrófono, el dealer '
         'acepta, y la ventana flotante muestra ambos videos con controles de micrófono, cámara y '
         'compartir pantalla. Soporta Picture-in-Picture para seguir viendo al interlocutor mientras '
         'se trabaja en otra pantalla.')
    captura(doc, 'Atención Remota → consola', 'La consola con una conversación activa y el panel de respuestas rápidas.')

    # ── 3. Portal del Dealer ────────────────────────────────────────────────
    h1(doc, '3. Portal del Dealer')
    p(doc, 'Cada dealer entra a su portal y ve SOLO sus datos: operaciones, el estado de sus pagos, '
           'sus cartolas y una pre-aprobación en línea. Es autoconsulta de solo lectura montada sobre '
           'los mismos datos del sistema — menos llamadas preguntando "¿cuándo me pagan?".')
    ficha(doc,
          'El dealer (autoconsulta) · el staff gestiona cuentas desde Atención Remota',
          'Sin perfil interno: el acceso se acota por pertenencia (su cuenta define qué dealer es)',
          'Cuenta de dealer creada y ligada al dealer (RUT / id)',
          '/portal-dealer/')
    h2(doc, '3.1 Qué ve el dealer')
    vineta(doc, 'las 4 tarjetas del inicio son su pipeline de pago: Fundantes pendientes → Fundantes recibidos → Liberado a pago → Saldo precio pagado.', bold_hasta='El resumen: ')
    vineta(doc, 'listado con filtro por estado, buscador y detalle por operación: línea de tiempo, qué documento falta (en rojo) y el estado del pago con sus montos ("Saldo Precio: $X · estado" y "Comisión: $Y · estado"). De los reparos solo ve el plazo.', bold_hasta='Sus operaciones: ')
    vineta(doc, 'las cartolas que se le han enviado, con sus plazos.', bold_hasta='Sus cartolas: ')
    vineta(doc, 'el chat con la mesa es el mismo de Atención Remota (cap. 2).', bold_hasta='Chat: ')
    captura(doc, 'Portal del Dealer → inicio', 'Las 4 tarjetas del pipeline y el listado de operaciones de un dealer.')
    h2(doc, '3.2 Pre-aprobación en línea')
    p(doc, 'El dealer ingresa RUT del cliente + precio + pie + año del vehículo, y el sistema evalúa '
           'al instante: antigüedad del vehículo contra la política, renta contra el 30% (manda la '
           'declarada por el dealer; la interna es respaldo), informes limpios y elegibilidad de la '
           'financiera. El resultado es PREAPROBADO (con las cuotas por plazo, del motor único) o '
           '"necesito más información".')
    regla(doc, 'Al dealer NUNCA se le muestran datos del cliente — solo el veredicto. El detalle '
               '(motivos, renta, deudas) va por correo al Jefe Comercial cuando el dealer pide '
               'contacto, y ahí se abre el chat con el asunto prellenado.')
    captura(doc, 'Portal del Dealer → Pre-Aprobación', 'Una pre-aprobación con su veredicto y las cuotas por plazo.')
    h2(doc, '3.3 Cuentas, onboarding y soporte')
    vineta(doc, 'el dealer entra con su EMAIL; si el correo coincide con el registrado en su ficha y el RUT calza, el alta es automática (clave al correo). Si no coincide, queda como solicitud para el staff. El primer ingreso muestra un tour de bienvenida con datos de demostración.', bold_hasta='Onboarding self-service: ')
    vineta(doc, 'en Atención Remota → Cuentas el staff crea, desactiva, resetea clave y regenera el link de acceso. Las cuentas con 6 meses de inactividad se desactivan solas; al volver, el dealer re-onboardea. "¿Olvidaste tu clave?" funciona self-service con link al correo.', bold_hasta='Gestión de cuentas: ')
    vineta(doc, 'para revisar qué ve un dealer antes de invitarlo — o mirar su misma pantalla cuando llama con dudas — el botón morado en Cuentas abre el portal exactamente como él lo ve, en solo lectura y con cinta morada. El chat queda bloqueado: chatear sería escribir a nombre del dealer. Todo auditado, expira a los 30 minutos.', bold_hasta='"Ver como dealer" (staff): ')
    vineta(doc, 'la pestaña Dealers del módulo Auditoría registra accesos, qué miró cada dealer y sus preguntas al asistente.', bold_hasta='Bitácora: ')
    advertencia(doc, 'Una cuenta sin dealer ligado (sin RUT ni id) entra pero no ve NADA — le pasó a '
                     'una cuenta de prueba. Si un dealer reclama portal vacío, revisar el vínculo de '
                     'su cuenta en Atención Remota → Cuentas.')
    h2(doc, '3.4 Asistente IA del portal')
    p(doc, 'Si está activado en el mantenedor de IA, el dealer puede preguntarle al asistente por sus '
           'operaciones. El asistente solo recibe los datos DE ESE dealer ya filtrados — no genera '
           'consultas propias, así que no puede ver a otro. Tiene cuota diaria paramétrica por dealer.')

    # ── 4. WhatsApp: la bandeja ─────────────────────────────────────────────
    h1(doc, '4. WhatsApp: la bandeja del equipo')
    p(doc, 'El número oficial del negocio (+56 9 3246 9071) vive en el sistema: el bot Facilito '
           'atiende primero y las personas toman las conversaciones que lo requieren, todas desde la '
           'misma bandeja — nadie atiende desde su teléfono personal.')
    ficha(doc,
          'Ejecutivos con permiso de atender · administradores del canal',
          'wsp_panel (ver) · wsp_atender (tomar) · wsp_campanas / wsp_config (administrar)',
          'PWA "AF WhatsApp" instalada y push activado (para recibir ofertas de conversación)',
          '/whatsapp/ — app instalable propia, separada del Suite')
    h2(doc, '3.1 Cómo llega una conversación a una persona')
    paso(doc, 1, 'El bot deriva', 'Cuando el cliente pide hablar con alguien o el bot no puede '
         'resolver, se dispara la oferta: push "¿quién lo toma?" a todos los ejecutivos con permiso.')
    paso(doc, 2, 'TOMAR es atómico', 'El primero que toca Tomar se la queda; al resto le avisa que ya '
         'fue tomada. El que la toma responde por el MISMO número del negocio desde la bandeja.')
    paso(doc, 3, 'Fuera de horario', 'Si la oportunidad llega fuera de jornada (o el cliente prefiere '
         'que lo llamen), va por correo a un ejecutivo elegido por turno equitativo, con copia al '
         'Jefe Comercial y el detalle de la cotización. La conversación queda asignada a él.')
    captura(doc, 'AF WhatsApp → bandeja', 'La bandeja con conversaciones, una ficha de cliente abierta al costado.')
    h2(doc, '3.2 Trabajar en la bandeja')
    vineta(doc, 'contacto y créditos del cliente con su estado de cartera, al costado del chat.', bold_hasta='Ficha lateral: ')
    vineta(doc, 'escribir "/" ofrece las respuestas configuradas.', bold_hasta='Respuestas rápidas: ')
    vineta(doc, 'quienes solo tienen permiso de atender ven SOLO sus conversaciones y las derivadas sin tomar — el resto está bloqueado por el servidor.', bold_hasta='Visibilidad: ')
    regla(doc, 'La ventana de 24 horas es de Meta, no nuestra: pasado ese plazo desde el último '
               'mensaje DEL CLIENTE, solo se puede escribir con plantilla HSM aprobada. El sistema '
               'bloquea el envío libre y avisa cuando quedan menos de 3 horas.')

    # ── 5. Facilito ─────────────────────────────────────────────────────────
    h1(doc, '5. Facilito: el bot que cotiza y pre-evalúa')
    p(doc, 'Facilito conversa con IA usando las Respuestas configuradas como conocimiento oficial. '
           'Pero las cifras NUNCA las inventa la IA: cuota, tasa y evaluación las calcula el código '
           'con los motores del sistema, y la IA solo redacta.')
    h2(doc, '4.1 El guion de cotización')
    p(doc, 'Cinco preguntas en orden: valor aproximado del auto → pie → cuotas o pago mensual '
           'deseado → nombre y RUT → confirmar teléfono. Plazos solo 12/24/36/48 (máximo 48; siempre '
           'se redondea al tramo superior). La cuota sale del cotizador del sistema con gastos, '
           'seguros y la tasa vigente del tramo — con su disclaimer de valor aproximado. Las '
           'cotizaciones del bot quedan en la tabla de cotizaciones, visibles en Evaluación '
           'Crediticia.')
    h2(doc, '4.2 La preevaluación')
    paso(doc, 1, 'Pide el RUT tras la cuota', 'El código valida el dígito verificador y trae los '
         'informes (con caché de 15 días — si hay informe vigente, no gasta consulta).')
    paso(doc, 2, 'El código decide, la IA redacta', 'Antecedentes buenos + pie ≥ 40% → vía exprés '
         '(cédula, acreditar domicilio, 3 referencias). Buenos con menos pie → sugiere subir el pie o '
         'hablar con ejecutivo. Malos → ofrece llamada de un ejecutivo, sin revelar detalles del '
         'informe.')
    paso(doc, 3, 'Reporte automático', 'Cada preevaluación genera en segundo plano el análisis de IA '
         'del RUT, disponible para el ejecutivo que tome el caso.')
    regla(doc, 'Los informes tienen límite anti-abuso paramétrico: por conversación y global diario. '
               'Superado, el bot ofrece ejecutivo en vez de consultar de nuevo.')
    h2(doc, '4.3 Otros servicios del bot')
    vineta(doc, 'pide el RUT, busca los créditos otorgados y arma la respuesta según la financiera de cada uno (cuentas y canales de pago de AutoFácil, Unidad y AutoFin — datos paramétricos).', bold_hasta='"¿Dónde pago?": ')
    vineta(doc, 'el switch 24/7 define si la IA atiende siempre (default) o solo en horario; fuera de horario los humanos no reciben ofertas, van por correo.', bold_hasta='Horario: ')
    vineta(doc, 'juegos, insolencia o pesca de información terminan en despedida amable y conversación cerrada.', bold_hasta='Mal uso: ')
    captura(doc, 'AF WhatsApp → conversación del bot', 'Una cotización completa de Facilito con la cuota y el disclaimer.')

    # ── 6. Configurar el canal ──────────────────────────────────────────────
    h1(doc, '6. Configurar el canal WhatsApp')
    p(doc, 'Todo lo que el bot dice y hace es paramétrico, en las pestañas del módulo:')
    vineta(doc, 'el conocimiento oficial del bot (keywords → respuesta). La IA las lee como material de estudio: mantenerlas es educar al bot.', bold_hasta='Respuestas: ')
    vineta(doc, 'reglas duras que corren ANTES de la IA — PROBLEMA/RIESGO/OPORTUNIDAD disparan alerta y derivación por área.', bold_hasta='Triggers: ')
    vineta(doc, 'horario, mensajes de bienvenida/fuera de horario/derivación, personalidad del bot (prompt editable) y el simulador integrado para probar sin enviar.', bold_hasta='Configuración: ')
    vineta(doc, 'se crean y envían a aprobación de Meta desde la misma app; la lista muestra en vivo APROBADA/PENDIENTE/RECHAZADA con el motivo. Con plantilla aprobada se puede escribir fuera de la ventana de 24 h (campañas, cobranza).', bold_hasta='Plantillas HSM: ')
    captura(doc, 'AF WhatsApp → Plantillas', 'El gestor con plantillas y sus estados de aprobación de Meta.')
    advertencia(doc, 'El Modo Desarrollo deja los envíos como SIMULADOS (registrados pero no salen). '
                     'Si una campaña "no llegó" en pruebas, es eso.')
    h2(doc, '6.1 Seguimiento de cartas por WhatsApp')
    p(doc, 'Motor diario (11:00): busca cartas aprobadas no otorgadas que VENCEN MAÑANA y le escribe '
           'al móvil del dealer con la plantilla aprobada ("aprobamos un crédito a nombre de X… vence '
           'mañana, ¿está vigente o ya vendiste el auto?"). La conversación la sigue Facilito en modo '
           'seguimiento: si está vigente, entrega el contacto del ejecutivo de la carta; si se vendió '
           'con financiamiento de otro, pregunta quién financió y por qué, y tabula el motivo '
           '(calidad, tiempo de respuesta, comisión, tasa u otro). El panel muestra candidatos, '
           'resultados y el toggle de activación.')
    captura(doc, 'WhatsApp → Seguimiento de Cartas', 'El panel con candidatos del día y resultados tabulados.')

    # ── 7. Portal del cliente ───────────────────────────────────────────────
    h1(doc, '7. Portal del Cliente')
    p(doc, 'Autoconsulta de créditos para el cliente final, estilo banca en línea, de solo lectura. '
           'Menos llamadas preguntando el saldo.')
    ficha(doc,
          'El cliente final · Operaciones/Cobranza responden dudas de enrolamiento',
          'Público (con enrolamiento OTP)',
          'Cliente con correo registrado en la base',
          '/mis-creditos (alias /portal-cliente)')
    paso(doc, 1, 'Enrolarse', 'RUT → código OTP de 6 dígitos al correo REGISTRADO en el sistema '
         '(10 minutos de vigencia, máximo 3 por hora) → el cliente define su clave.')
    paso(doc, 2, 'Consultar', 'Sus créditos con estado, la tabla de desarrollo del calendario '
         'congelado, los valores AL DÍA de cada cuota vencida (mora y gastos del motor único), sus '
         'comprobantes de pago y dónde/cómo pagar (datos paramétricos).')
    captura(doc, 'Portal del Cliente', 'El dashboard de un cliente con sus créditos y el banner de deuda.')
    advertencia(doc, 'Si el cliente no recibe el OTP, el correo registrado está malo o vacío: se '
                     'corrige en la ficha del cliente, no hay vía alternativa — es una decisión de '
                     'seguridad (la respuesta del portal es genérica a propósito, para no revelar '
                     'qué RUT tiene cuenta).')

    # ── 8. Campañas ─────────────────────────────────────────────────────────
    h1(doc, '8. Campañas masivas')
    p(doc, 'Contacto masivo por mail o WhatsApp con mensajes personalizados, deciles de control y '
           'medición de conversión: saber si la campaña movió la aguja o habría pasado igual.')
    ficha(doc,
          'Marketing/Comercial (venta) · Cobranza (mora)',
          'Acceso a Campañas Masivas',
          'Datos de contacto (BD o CSV) · para WhatsApp frío: plantilla HSM aprobada por Meta',
          'Campañas Masivas')
    paso(doc, 1, 'Crear', 'Objetivo VENTA o COBRANZA → variables → data (CSV o desde la base por '
         'parámetros) → deciles de control → contenido → vista previa uno a uno → enviar por lotes.')
    paso(doc, 2, 'Analizar antes de enviar (opcional)', 'Análisis de política de crédito y de '
         'informes con exclusión automática de quienes no cumplen.')
    paso(doc, 3, 'Medir', 'Recalcular conversión cada 2-3 días: tasa de la campaña vs el grupo de '
         'control = el uplift real. Cada mensaje queda como gestión CRM.')
    captura(doc, 'Campañas Masivas', 'Una campaña con su audiencia, deciles de control y conversión medida.')
    flujo(doc, 'BORRADOR → ENVIADA → CONVERSIÓN MEDIDA')
    regla(doc, 'El grupo de control no es opcional por deporte: sin control no hay forma de saber si '
               'la campaña funcionó. Los deciles se definen al crear y no se contactan.')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('No puedo escribirle a un cliente por WhatsApp', 'Ventana de 24 h vencida', 'Usar plantilla HSM aprobada (cap. 4)'),
           ('No me llegan ofertas de conversación', 'PWA sin instalar o push sin activar', 'Instalar AF WhatsApp y activar la campana (cap. 4)'),
           ('El bot dio una cuota distinta al simulador', 'No debería (mismo motor)', 'Reportar a TI con el pantallazo (cap. 5)'),
           ('Los envíos de una campaña salen SIMULADOS', 'Modo Desarrollo activo', 'Es intencional en pruebas (cap. 6)'),
           ('El cliente no recibe el OTP del portal', 'Correo registrado malo o vacío', 'Corregir la ficha del cliente (cap. 7)'),
           ('El dealer no puede entrar a Atención Remota', 'Link de acceso regenerado o cuenta pendiente', 'Copiar el link vigente desde Cuentas (cap. 2)'),
           ('Un dealer ve el portal vacío', 'Cuenta sin dealer ligado (RUT / id)', 'Revisar el vínculo en Atención Remota → Cuentas (cap. 3)'),
           ('Una plantilla no se puede usar', 'Meta aún no la aprueba (o la rechazó)', 'Ver el estado y el motivo en Plantillas (cap. 6)')),
          (5.6, 5.6, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('Facilito', 'El bot de WhatsApp del negocio; conversa con IA, calcula con los motores'),
           ('Ventana 24 h', 'Regla de Meta: mensaje libre solo dentro de 24 h del último mensaje del cliente'),
           ('Plantilla HSM', 'Mensaje pre-aprobado por Meta; único permitido fuera de la ventana'),
           ('Derivar', 'Pasar la conversación del bot a una persona'),
           ('TOMAR', 'Asignarse una conversación derivada; el primero gana'),
           ('OTP', 'Código de un solo uso para enrolarse en el portal del cliente'),
           ('Uplift', 'Conversión de la campaña menos la del grupo de control'),
           ('Decil de control', 'Porción de la audiencia que NO se contacta, para medir'),
           ('PWA', 'App instalable desde el navegador; AF WhatsApp es una, separada del Suite'),
           ('Picture-in-Picture', 'Video flotante que sigue visible al cambiar de aplicación'),
           ('Llave portadora', 'Link que da acceso por sí solo; se regenera si se filtra')),
          (4.3, 12.2))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-Canales-Digitales-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
