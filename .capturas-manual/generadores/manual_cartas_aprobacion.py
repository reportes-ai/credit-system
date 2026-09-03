# -*- coding: utf-8 -*-
"""Manual de Usuario — Cartas de Aprobación (Business Suite), para ejecutivos comerciales."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE CARTAS DE APROBACIÓN', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'El documento que compromete las condiciones al dealer:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'cómo emitirla, corregirla, otorgarla o desistirla.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Septiembre 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Septiembre 2026', 'Business Suite',
            'Emisión inicial: emitir, revisión automática, rentabilidad, "Compra para" un tercero, '
            'Cartas Vigentes (otorgar/desistir), corrección por reemplazo, pantallazo AutoFin, '
            'vendedores por dealer y códigos de excepción'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Está pensado para leerse de corrido la primera vez y usarse como referencia después. Los '
           'recuadros de color: ⚠ OJO (lo que confunde), 🔒 Regla del sistema (lo que no se puede '
           'saltar), 🧾 Caso real (el episodio que originó la regla) y 📸 CAPTURA (dónde va el '
           'pantallazo, pendiente en esta versión). Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Qué es la carta ──────────────────────────────────────────────────
    h1(doc, '1. Qué es la Carta de Aprobación')
    p(doc, 'La carta es el documento que compromete las condiciones de un crédito al dealer: monto, '
           'tasa, cuotas, vehículo y comisión. Sale con correlativo propio, código de verificación QR '
           'y Firma Electrónica Simple (Ley 19.799) — por eso, una vez emitida, no se edita: si algo '
           'hay que cambiar, se corrige por reemplazo (cap. 6).')
    flujo(doc, 'EVALUACIÓN → EMITIR CARTA → REVISIÓN → APROBADA → OTORGAR / DESISTIR')
    vineta(doc, 'ejecutivo comercial. Emite la carta y adjunta los documentos de respaldo.', bold_hasta='Quién la crea: ')
    vineta(doc, 'el pool de analistas de crédito, o Business Suite en el momento si el Revisor Automático está encendido para esa financiera.', bold_hasta='Quién la aprueba: ')
    vineta(doc, 'el ejecutivo o Operaciones, dentro del plazo de vigencia (5 días corridos, configurable).', bold_hasta='Quién la otorga: ')

    # ── 2. Emitir una carta ─────────────────────────────────────────────────
    h1(doc, '2. Emitir una carta')
    ficha(doc,
          'Ejecutivo comercial',
          'Emisión de cartas (ordenes_pago_emitir no aplica acá — es el permiso propio del módulo Cartas)',
          'Evaluación crediticia hecha · dealer con ficha vigente · vendedor en la base del dealer',
          'Cartas de Aprobación → Generador de Carta de Aprobación')
    h2(doc, '2.1 Autocompletar desde documentos')
    p(doc, 'Antes de llenar nada a mano, sube los PDF de la financiera: el sistema lee los datos y '
           'llena el formulario solo. Quedan adjuntos a la carta para la revisión del analista.')
    tabla(doc, ('Financiera', 'Documentos a subir'),
          (('UNIDAD', 'Carta Compromiso + Cotización (dos PDF, dos botones separados)'),
           ('AUTOFIN', 'Carta de Aprobación (PDF)')),
          (4.0, 12.5))
    advertencia(doc, 'El autofill SUGIERE, tú respondes por lo emitido. Revisa siempre contra el '
                     'documento original antes de enviar a revisión — sobre todo monto, tasa y plazo.')
    h2(doc, '2.2 Datos de la operación')
    paso(doc, 1, 'Tipo de carta', 'DEALER PARQUE o DEALER CALLE — define qué tabla de comisión aplica.')
    paso(doc, 2, 'ID Financiera y fecha', 'El "Operación carta" se arma solo a partir del ID Financiera.')
    paso(doc, 3, 'Ejecutivo', 'Se elige de la lista; teléfono y correo se completan solos.')
    h2(doc, '2.3 Datos del cliente')
    paso(doc, 1, 'RUT del cliente', 'Escribe el RUT y presiona Buscar: si el cliente existe en la base, '
         'nombres y apellidos se llenan solos. Si no existe, el sistema ofrece llevarte a crearlo sin '
         'perder lo ya digitado (tu carta queda guardada como borrador).')
    paso(doc, 2, 'Compra para (opcional)', 'Ver capítulo 3 — botón junto al título "Datos del cliente".')
    h2(doc, '2.4 Vehículo, montos y plazo')
    vineta(doc, 'tipo, marca, modelo, año, patente (vacía si es 0 km) y si lleva prenda.', bold_hasta='Vehículo: ')
    vineta(doc, 'precio de venta, pie y saldo (se calcula solo), plazo en cuotas, monto y tasa del crédito.', bold_hasta='Montos: ')
    vineta(doc, 'la Preferencia Financiera bloquea combinaciones fuera de política y muestra abajo del '
                'selector el veredicto de ¿Dónde Curso? — cursar en la financiera conveniente no es '
                'excepción; con veredicto DECIDES TÚ tampoco lo es cursar en cualquiera de las dos.', bold_hasta='Financiera / acreedor: ')
    h2(doc, '2.5 Dealer, vendedor y participación')
    p(doc, 'El vendedor se elige de la base propia del dealer (cap. 7); si no está, se agrega en el '
           'momento con Nombre, RUT y Mail. Elegir bien el vendedor evita que la venta quede sin dueño '
           'para el reporte de comisiones del dealer.')
    regla(doc, 'La comisión del dealer que rige es la de la carta (negociación especial, si la trae). '
               'Si la carta no la trae, se calcula por la tabla de parámetros según el plazo.')
    h2(doc, '2.6 Enviar a revisión')
    p(doc, 'Al guardar, la carta queda PENDIENTE en el pool de analistas — o Business Suite la revisa '
           'en el mismo instante si el Revisor Automático está encendido para esa financiera (cap. 4).')
    captura(doc, 'Cartas de Aprobación → generador', 'El formulario completo con los datos autocompletados, el "Compra para" cargado y los documentos adjuntos.')

    # ── 3. Compra para un tercero ───────────────────────────────────────────
    h1(doc, '3. Compra para un tercero')
    p(doc, 'Para los casos en que se autoriza comprar el vehículo a nombre de una persona natural o '
           'jurídica DISTINTA del cliente que solicita el crédito (ej. una empresa que compra para uno '
           'de sus trabajadores). El concesionario deberá inscribir la Limitación al Dominio a nombre '
           'del tercero, a favor de la entidad financiera que corresponda.')
    ficha(doc, 'Ejecutivo comercial', 'El mismo de emisión de cartas', 'Ninguno adicional',
          'Cartas de Aprobación → Generador, botón "Compra para" junto al título "Datos del cliente"')
    paso(doc, 1, 'Abrir el popup', 'Botón "Compra para" al lado del título "Datos del cliente".')
    paso(doc, 2, 'RUT del tercero', 'Se valida el dígito verificador al tipear. Botón Buscar consulta '
         'la base de Clientes — si existe, llena Nombres, Apellido paterno y materno solo; si no está '
         '(puede ser una empresa), se digita a mano.')
    paso(doc, 3, 'Nombres o razón social', 'Obligatorio junto con el RUT. Apellidos son opcionales '
         '(una empresa no los necesita).')
    paso(doc, 4, 'Grabar', 'Queda un párrafo verde no editable bajo los campos del cliente con el '
         'resumen del tercero. Reabrir el botón permite corregir los datos antes de grabar la carta.')
    regla(doc, 'Dejar todos los campos vacíos y presionar Grabar QUITA el "Compra para" de la carta — '
               'no hace falta un botón aparte para eso.')
    p(doc, 'Con el dato cargado, la carta impresa agrega la fila COMPRA PARA / RUT justo bajo el '
           'nombre y RUT del cliente titular — mismo formato que la carta antigua en papel. El '
           'revisor también lo ve en la ficha de revisión antes de aprobar. Es opcional: sin '
           '"Compra para" la carta sale exactamente igual que siempre.')
    captura(doc, 'Popup Compra para', 'RUT validado, nombre encontrado en Clientes y el resumen listo para grabar.')

    # ── 4. La revisión ──────────────────────────────────────────────────────
    h1(doc, '4. La revisión de la carta')
    h2(doc, '4.1 El pool de analistas')
    p(doc, 'Sin Revisor Automático, la carta cae al pool de analistas de crédito: cualquiera de ellos '
           'puede tomarla, aprobarla o rechazarla con motivo.')
    h2(doc, '4.2 Revisor Automático')
    p(doc, 'Si el motor está encendido para tu financiera (Mantenedores → Excepciones Comerciales), '
           'Business Suite compara el documento contra lo digitado (ID, RUT, saldo, plazo, monto, '
           'tasa), valida la comisión del dealer contra el motor y el código de excepción si lo hay. '
           'Todo cuadra → carta APROBADA al instante, con checklist firmado (timbre + QR verificable) '
           'y campanita. Algo no cuadra → sigue con el analista, con el detalle de qué falló en rojo.')
    regla(doc, 'Para AUTOFIN, el pantallazo debe mostrar la solicitud en estado Revisión Firma o '
               'Cursado — ningún otro estado emite carta, y su ID debe calzar con la carta (cap. 5).')
    p(doc, 'Todo queda en la card Bitácora Revisor Automático (registro inmutable, filtros por mes y '
           'resultado) — sirve para entender por qué una carta pasó o no pasó sola.')

    # ── 5. Pantallazo AutoFin ───────────────────────────────────────────────
    h1(doc, '5. Pantallazo AutoFin (obligatorio)')
    p(doc, 'Al digitar una carta AutoFin, el generador incluye el recuadro "Pantallazo AutoFin" con '
           'dos vías: el botón 📸 Capturar pantalla (el navegador pregunta qué ventana compartir — '
           'elige la de AutoFin y el sistema toma el cuadro completo) o pegar con Ctrl+V (Alt+ImprPant '
           'copia la ventana entera).')
    advertencia(doc, 'Captura la SOLICITUD PUNTUAL — búscala por su ID en el sistema AutoFin para que '
                     'salga sola con su estado. No sirve un listado ni una pantalla donde no se lea '
                     'el estado: el sistema te avisa al tiro si no puede leerlos.')
    p(doc, 'Al guardar, la IA lee el ID y el Estado del pantallazo; si algo no calza con la carta, '
           'avisa de inmediato en vez de dejarlo pasar para que lo descubra el analista.')

    # ── 6. Cartas Vigentes ──────────────────────────────────────────────────
    h1(doc, '6. Cartas Vigentes: otorgar o desistir')
    p(doc, 'La carta aprobada vive 5 días corridos (configurable). En ese plazo el negocio se concreta '
           '(OTORGAR) o se cae (DESISTIR). Si nadie hace nada, el desistimiento automático la vence.')
    ficha(doc, 'Ejecutivo comercial / Operaciones', 'Otorgamiento de cartas',
          'Carta APROBADA vigente', 'Cartas de Aprobación → Cartas Vigentes')
    paso(doc, 1, 'Otorgar', 'Crea el crédito en etapa OTORGADO con todo congelado: calendario, tasa y '
         'la comisión de la carta — la carta manda. Nace el movimiento de cartola del dealer.')
    paso(doc, 2, 'Desistir', 'Deja la carta DESISTIDA con motivo. Queda en la estadística de conversión.')
    captura(doc, 'Cartas Vigentes', 'La bandeja con una carta por vencer, mostrando los días restantes.')
    advertencia(doc, 'Vencida la carta, no revive: hay que emitir una nueva. Por eso conviene otorgar '
                     'apenas el curse se confirma — una carta otorgada tarde deja la comisión para la '
                     'cartola del mes siguiente.')

    # ── 7. Corregir una carta emitida ───────────────────────────────────────
    h1(doc, '7. Corregir una carta ya emitida')
    p(doc, 'Una carta emitida ya salió al dealer con su QR y su firma electrónica, así que NO se '
           'edita. Corregirla emite una carta NUEVA con sufijo -C1, -C2… y deja la anterior '
           'REEMPLAZADA; en la verificación pública su QR y firma pasan a no vigentes con el mensaje '
           '"Reemplazada por la carta N° X".')
    regla(doc, 'Monto del crédito, saldo precio, tasa y cuotas NO se pueden corregir por esta vía: '
               'definen la operación, y son lo que permite que la carta nueva siga colgando del mismo '
               'crédito. Si alguno cambió, ya no es la misma operación — se anula por el flujo de '
               'Operaciones y se emite una carta nueva.')
    p(doc, 'La comisión que aún no salió en una cartola sigue a la carta nueva; una ya enviada es '
           'historia y no se toca. Todo queda auditado (quién, cuándo, qué campos y el motivo).')
    vineta(doc, 'Administración / Operaciones (módulo Corrección de Cartas de Aprobación, hoy solo Administrador).', bold_hasta='Quién la hace: ')
    vineta(doc, 'la carta debe estar APROBADA u OTORGADA — una PENDIENTE o RECHAZADA se corrige con el lápiz normal, que no necesita reemplazo.', bold_hasta='Prerequisito: ')

    # ── 8. Vendedores por dealer ─────────────────────────────────────────────
    h1(doc, '8. Vendedores por Dealer')
    p(doc, 'El campo Vendedor de la carta es un selector que muestra los vendedores registrados del '
           'dealer elegido. Si el vendedor no está en la lista, elige "➕ Agregar vendedor" y completa '
           'Nombre, RUT y Mail — queda agregado a la base al guardar la carta.')
    p(doc, 'Dónde se mide el aporte de cada vendedor: Post Venta → Vendedores con Ventas (abre en el '
           'último mes cerrado) y Cartas → Detalle Mensual, que lista cada carta con su vendedor, RUT '
           'y correo.')

    # ── 9. Códigos de excepción ──────────────────────────────────────────────
    h1(doc, '9. Códigos de excepción (referencia)')
    p(doc, 'Cuando el negocio necesita salirse de la pizarra —bajar la tasa o la comisión del dealer— '
           'el Simulador de Excepciones (¿Dónde Curso? 2.0) entrega un código que aprueba la excepción '
           'sola si respeta el piso de rentabilidad. El detalle completo del simulador vive en su '
           'propio manual; acá solo el uso dentro de la carta.')
    regla(doc, 'Un código = un solo uso, del ejecutivo dueño (no sirve el de un colega; los de '
               'Gerencia sí son transferibles). El sistema valida saldo precio, primeros dígitos del '
               'RUT y vigencia (24 horas) antes de aceptarlo. La carta queda "aprobada por código del '
               'sistema" y lo estampa impreso.')
    flujo(doc, 'SIMULAR → CUMPLE PISO → CÓDIGO → PEGAR CÓDIGO EN LA CARTA → USADO')

    # ── 10. Preguntas frecuentes ─────────────────────────────────────────────
    h1(doc, '10. Preguntas frecuentes')
    vineta(doc, 'no — se corrige por reemplazo (cap. 7) o, si cambió monto/tasa/plazo, se anula y se emite una nueva.', bold_hasta='¿Puedo editar una carta ya APROBADA? ')
    vineta(doc, 'no revive: hay que emitir una carta nueva desde cero.', bold_hasta='¿Qué pasa si la carta vence sin otorgar ni desistir? ')
    vineta(doc, 'sí, es opcional y se puede dejar vacío sin afectar el resto de la carta.', bold_hasta='¿"Compra para" es obligatorio? ')
    vineta(doc, 'no — el "compra para" puede ser persona natural o jurídica (empresa); si no está en Clientes, se digita a mano.', bold_hasta='¿El "compra para" tiene que estar en la base de Clientes? ')
    vineta(doc, 'sí, siempre — es la única forma en que la financiera valida la solicitud puntual antes de liberar la carta.', bold_hasta='¿El pantallazo AutoFin es obligatorio en toda carta AutoFin? ')

    return doc
