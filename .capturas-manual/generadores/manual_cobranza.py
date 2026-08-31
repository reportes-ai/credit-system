# -*- coding: utf-8 -*-
"""Manual de Cobranza — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE COBRANZA', bold=True, color=AZUL_OSCURO, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'De la mora temprana al castigo:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'recuperar la cartera cumpliendo la ley, con cada gestión trazada.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 2.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Cobranza'),
           ('2.0', 'Agosto 2026', 'Business Suite',
            'Gestión manual con visita en terreno / visita del cliente; bitácora exportable (Excel e '
            'informe PDF); capítulos nuevos de gastos e interés de mora, provisiones en detalle, '
            'contabilización e informes; judicial ampliado con catálogos de abogados')),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo: para qué existe el proceso, quién lo hace, prerequisitos, paso a paso y los '
           'recuadros ⚠ OJO / 🔒 Regla del sistema / 🧾 Caso real / 📸 CAPTURA (pendientes en esta '
           'versión). Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. El marco ─────────────────────────────────────────────────────────
    h1(doc, '1. El marco: estados, motores y ley')
    p(doc, 'La cobranza trabaja sobre la cartera propia, con dos coordenadas por crédito: la ETAPA '
           '(OTORGADO — no cambia por la mora) y el ESTADO DE CARTERA, que describe el pago:')
    flujo(doc, 'VIGENTE → EN MORA → VENCIDO → JUDICIAL → CASTIGADO · AL DÍA')
    p(doc, 'Los estados de cartera y sus transiciones son paramétricos (mantenedor Estado Cartera). '
           'Confundir etapa con estado de cartera es el error conceptual más común: un crédito EN MORA '
           'sigue siendo, en etapa, OTORGADO.')
    h2(doc, '1.1 Los motores únicos')
    vineta(doc, 'la tasa de mora queda FIJA al otorgamiento (TMC de ese momento, según el modo configurado). El interés de mora que muestra la caja, el portal del cliente, el certificado y esta área es EL MISMO: sale de un solo motor.', bold_hasta='Interés por mora: ')
    vineta(doc, 'corren recién desde el día 21 de atraso, según los tramos de Parámetros de Cobranza (ver cap. 5).', bold_hasta='Gastos de cobranza: ')
    vineta(doc, 'los valores del crédito quedaron congelados al otorgar; el calendario de cuotas es inmutable.', bold_hasta='El calendario: ')
    vineta(doc, 'el universo "en mora" que ven la Segmentación por Tramos, Pre-judicial, las automatizaciones y el Dashboard de Cobranza sale de la MISMA consulta (motor MORA_SQL): imposible que dos pantallas muestren carteras distintas.', bold_hasta='El universo: ')
    h2(doc, '1.2 La ley')
    regla(doc, 'Ley 21.320 y Ley del Consumidor: por semana calendario, 1 gestión telefónica o '
               'presencial y 2 gestiones remotas (WhatsApp/SMS/Email) separadas por al menos 2 días. '
               'El tope está BLOQUEADO por el sistema — el panel de cada crédito muestra la '
               'disponibilidad de la semana antes de gestionar. Además rigen horarios permitidos y '
               'el registro de cada contacto: lo que no está en la bitácora no existió.')

    # ── 2. Mora temprana automática ─────────────────────────────────────────
    h1(doc, '2. Mora temprana: los motores automáticos')
    p(doc, 'En los primeros tramos de atraso el sistema gestiona solo: correo por tramo y secuencia de '
           'WhatsApp, cada envío registrado en la bitácora CRM del crédito. El analista entra cuando '
           'la automatización no alcanza.')
    ficha(doc,
          'Los motores automáticos (correo y WhatsApp) · el Analista de Cobranza supervisa',
          'Configuración en Parámetros de Cobranza y Mantenedores → Correos Programados',
          'Automatizaciones activadas · tramos y plantillas configurados',
          'Cobranza → Automatizaciones · Mantenedores → Parámetros de Cobranza')
    vineta(doc, 'cada tramo de atraso tiene su plantilla y su momento. Se apagan y prenden por parámetro, sin tocar código.', bold_hasta='Correo por tramo: ')
    vineta(doc, 'secuencia de plantillas aprobadas por Meta, en orden (1°, 2°, 3° aviso…). A quiénes: filtro configurable por tramo de días de mora y monto mínimo. Respeta el cupo semanal legal y el horario permitido.', bold_hasta='WhatsApp: ')
    vineta(doc, 'cada {{n}} de la plantilla se mapea a un campo del crédito: nombre, RUT, días de mora, cuotas en mora, monto en mora, saldo insoluto, N° de operación, y los datos de la CUOTA impaga más antigua (N° de cuota, día de vencimiento y monto de la cuota).', bold_hasta='Variables: ')
    vineta(doc, 'todo envío queda en la bitácora CRM del crédito, visible para cualquier analista que tome el caso.', bold_hasta='Trazabilidad: ')
    captura(doc, 'Cobranza → bitácora CRM de un crédito', 'La línea de tiempo con gestiones automáticas y humanas mezcladas.')
    regla(doc, 'El tope semanal de gestiones está bloqueado por sistema: si un contacto más viola el '
               'tope, el sistema no lo deja salir. No es optativo. Los motores automáticos también lo '
               'respetan: nunca compiten con la gestión humana por el cupo.')

    # ── 3. Pre-judicial ─────────────────────────────────────────────────────
    h1(doc, '3. Cobranza pre-judicial: la gestión humana')
    p(doc, 'Cuando la mora avanza, la gestión pasa a ser humana: llamadas, visitas y compromisos de '
           'pago, todos registrados.')
    ficha(doc,
          'Analista de Cobranza',
          'Acceso a Cobranza → Pre-judicial',
          'Crédito en mora fuera del tramo temprano',
          'Cobranza → Pre-judicial')
    paso(doc, 1, 'Tomar el caso', 'La bandeja segmenta la cartera por tramos (1–15, 16–30, 31–60, '
         '61–90 días) con atraso, saldo y última gestión. Al abrir un crédito, el panel muestra la '
         'mora al día, el contacto del cliente y la DISPONIBILIDAD legal de la semana.')
    paso(doc, 2, 'Registrar la gestión manual', 'Botón "Registrar Llamada/Visita": elegir el tipo — '
         'Llamada telefónica, Visita en terreno o Visita del cliente (en oficina) — y el resultado '
         '(contactado, no contesta, promesa de pago, rechaza pago, número errado). Las notas y la '
         'promesa (fecha + monto) quedan en la misma gestión.')
    paso(doc, 3, 'Gestiones remotas', 'WhatsApp, SMS o Email desde el mismo panel, con el mensaje de '
         'la plantilla de Parámetros de Cobranza. Al confirmar el envío, la gestión se registra y '
         'consume cupo remoto.')
    paso(doc, 4, 'Seguir el compromiso', 'Un compromiso cumplido normaliza; uno incumplido escala la '
         'gestión.')
    captura(doc, 'Cobranza → Pre-judicial → panel de un crédito', 'El formulario de gestión manual con los tres tipos y la disponibilidad semanal.')
    p(doc, 'Herramientas de apoyo: la PWA de Terreno (ruta del día, check-in GPS y fotos) para las '
           'visitas, y el Score de Mora (mora histórica por segmento) para priorizar la cartera.')
    h2(doc, '3.1 La bitácora y sus exportaciones')
    p(doc, 'La pestaña Bitácora del crédito muestra TODAS las gestiones — automáticas y humanas — en '
           'orden, con tipo, canal, resultado, gestor y notas. Desde ahí:')
    vineta(doc, 'descarga las gestiones del crédito en planilla (fecha, tipo, canal, resultado, gestor, notas y promesas).', bold_hasta='Exportar Excel: ')
    vineta(doc, 'genera el informe formal con la ficha del deudor (cliente, RUT, mora, saldo) y la tabla completa de gestiones — el documento que se adjunta a una carpeta judicial o se envía a un estudio de abogados.', bold_hasta='Informe PDF de gestiones: ')
    captura(doc, 'Pre-judicial → Bitácora', 'Los botones Excel e Informe PDF sobre la línea de tiempo.')

    # ── 4. Judicial ─────────────────────────────────────────────────────────
    h1(doc, '4. Cobranza judicial y estudios de abogados')
    p(doc, 'La demanda con abogado, tribunal y etapa procesal. La ficha judicial congela un SNAPSHOT '
           'financiero del crédito al entrar (deuda, cuotas, saldo), y encima registra el avance '
           'procesal: abogado, status legal, status del crédito, juzgado, rol y comentarios.')
    ficha(doc,
          'Analista de Cobranza · abogados y estudios externos',
          'Acceso a Cobranza → Judicial',
          'Decisión de demandar tomada · antecedentes del crédito completos',
          'Cobranza → Judicial')
    paso(doc, 1, 'Abrir la causa', 'Crédito, abogado (del catálogo), juzgado, rol y status inicial.')
    paso(doc, 2, 'Actualizar el avance', 'Cada cambio de status legal o del crédito se registra; la '
         'historia completa queda en la causa. La cartera judicial se filtra por abogado y por status.')
    paso(doc, 3, 'Adjuntar el respaldo', 'El pagaré y los antecedentes quedan referenciados en la '
         'ficha; el Informe PDF de gestiones (cap. 3.1) acredita la gestión pre-judicial realizada.')
    h2(doc, '4.1 Los catálogos son paramétricos')
    p(doc, 'En Mantenedores → Cobranza Judicial se administran los catálogos que alimentan la ficha: '
           'ABOGADOS (los estudios y profesionales con que trabaja la empresa), STATUS LEGAL (etapas '
           'procesales) y STATUS DEL CRÉDITO. Agregar un estudio de abogados nuevo es agregar una fila '
           'ahí — sin programador.')
    captura(doc, 'Mantenedores → Cobranza Judicial', 'Los tres catálogos y la cartera judicial filtrable por abogado.')
    regla(doc, 'Un crédito en cobranza judicial sigue devengando su interés de mora por el motor '
               'único. El snapshot de la ficha es la foto de ENTRADA a judicial; la deuda viva se '
               'consulta siempre en el sistema.')

    # ── 5. Gastos de cobranza e interés de mora ─────────────────────────────
    h1(doc, '5. Gastos de cobranza e interés de mora')
    p(doc, 'Los dos recargos del atraso, ambos de motor único y ambos paramétricos:')
    h2(doc, '5.1 Interés por mora')
    vineta(doc, 'la tasa aplicable (TMC) queda fija al OTORGAMIENTO del crédito, según el modo configurado — no cambia aunque la TMC vigente suba o baje después.', bold_hasta='Tasa fija: ')
    vineta(doc, 'se calcula por cuota vencida, por los días efectivos de atraso.', bold_hasta='Devengo: ')
    h2(doc, '5.2 Gastos de cobranza (Ley 19.496)')
    vineta(doc, 'solo después de 20 días corridos desde el vencimiento — corren desde el día 21. Antes de eso, cero gasto.', bold_hasta='Cuándo: ')
    vineta(doc, 'tramos MARGINALES sobre la deuda en UF: hasta 10 UF → 9% · entre 10 y 50 UF → 6% · sobre 50 UF → 3%. Los tramos y porcentajes viven en Mantenedores → Parámetros de Cobranza.', bold_hasta='Cuánto: ')
    vineta(doc, 'caja, portal del cliente, certificados y prepago muestran el mismo gasto, porque lo calcula el mismo motor por cuota al día D.', bold_hasta='Dónde se ve: ')
    h2(doc, '5.3 Condonación')
    regla(doc, 'El orden de condonación es fijo y está protegido en el servidor: primero GASTOS, '
               'luego INTERESES; el CAPITAL nunca se condona. Los descuentos negociados van por '
               'Aplicación de Fondos, con su cadena de firmas.')

    # ── 6. Provisiones ──────────────────────────────────────────────────────
    h1(doc, '6. Provisiones: el riesgo reconocido mes a mes')
    p(doc, 'Cada tramo de días de mora provisiona un porcentaje del saldo insoluto (paramétrico en '
           'Parámetros de Cobranza). La provisión se calcula y contabiliza sola al cierre de cada mes.')
    vineta(doc, 'al día configurado del mes (paramétrico), el motor recorre la cartera, clasifica cada crédito en su tramo y calcula su provisión. No requiere que nadie lo ejecute.', bold_hasta='Cierre automático: ')
    vineta(doc, 'queda el detalle auditable crédito a crédito: días de mora, tramo, porcentaje y provisión — el respaldo que pide un auditor.', bold_hasta='Detalle: ')
    vineta(doc, 'el asiento va por la VARIACIÓN mensual: si la cartera se deterioró se constituye provisión (PROVISION_CIERRE); si tu gestión normalizó créditos, se libera (PROVISION_LIBERACION). Idempotente por mes.', bold_hasta='Contabilización: ')
    caso(doc, 'Un analista normaliza en agosto un crédito que venía con 45 días de mora. En el cierre '
              'de agosto ese crédito baja de tramo: la diferencia de provisión se LIBERA y mejora el '
              'resultado del mes. La gestión de cobranza se ve directamente en el estado de resultados.')

    # ── 7. Castigo y recupero ───────────────────────────────────────────────
    h1(doc, '7. Castigo y recupero')
    p(doc, 'La baja contable de la operación incobrable. Es la última estación y es deliberadamente '
           'manual.')
    ficha(doc,
          'Gerencia — doble firma gerencial',
          'Atribución de castigo',
          'Gestión agotada · antecedentes de la causa',
          'Tesorería → Provisiones + Castigos')
    regla(doc, 'El castigo NUNCA es automático. Exige doble firma gerencial y genera su asiento '
               'contable por regla de centralización. El cierre automático mensual de provisiones no '
               'castiga: solo provisiona.')
    p(doc, 'Tras el castigo, el crédito queda CASTIGADO en estado de cartera. La recuperación '
           'posterior, si la hay, entra por caja como cualquier pago y genera su asiento de RECUPERO '
           '— la historia del crédito nunca se pierde.')
    captura(doc, 'Castigos', 'Una solicitud de castigo con sus dos firmas.')

    # ── 8. Contabilización ──────────────────────────────────────────────────
    h1(doc, '8. Cómo se contabiliza la cobranza')
    p(doc, 'Máxima del sistema: todo movimiento de dinero genera su asiento, automáticamente, por el '
           'motor de centralización y sus reglas paramétricas (Mantenedores → Reglas de '
           'Centralización). En cobranza:')
    tabla(doc, ('Hecho', 'Qué se contabiliza'),
          (('Pago de cuota en caja', 'Ingreso por caja/banco contra la cartera; el interés de mora y los gastos de cobranza van a sus cuentas de ingreso propias'),
           ('Pago no identificado', 'Cuentas Transitorias hasta aclarar el pagador; al aplicarse, se reclasifica'),
           ('Provisión mensual', 'PROVISION_CIERRE (constitución) o PROVISION_LIBERACION (reverso) por la variación del mes'),
           ('Castigo', 'Asiento de baja del incobrable, gatillado por la doble firma'),
           ('Recupero de castigado', 'El abono entra por caja y genera su asiento de recupero'),
           ('Prepago / Aplicación de Fondos', 'El pago rebaja cartera; las condonaciones aprobadas se reconocen según su regla')),
          (5.5, 11.0))
    regla(doc, 'El motor contable nunca bloquea la operación: si un evento no tiene regla, el pago '
               'sale igual y el evento queda SIN_REGLA en el log contable — esa lista es la deuda '
               'pendiente de cablear, visible para Contabilidad.')

    # ── 9. Informes ─────────────────────────────────────────────────────────
    h1(doc, '9. Los informes del área')
    vineta(doc, 'la vista por tramos de la cartera en mora (cantidad y monto por tramo), la misma segmentación de Pre-judicial.', bold_hasta='Segmentación por tramos: ')
    vineta(doc, 'mora de la cartera en vivo dentro del Dashboard general y el Cuadro de Mando.', bold_hasta='Dashboard: ')
    vineta(doc, 'por crédito: Excel de gestiones e Informe PDF formal desde la bitácora (cap. 3.1).', bold_hasta='Informe de gestiones: ')
    vineta(doc, 'la mora histórica por segmento que alimenta la política de crédito — la cobranza de hoy es la política de admisión de mañana.', bold_hasta='Score de Mora: ')
    vineta(doc, 'para cruces a medida (mora por dealer, por ejecutivo, por cosecha), el constructor visual de Reportería → Diseño de Consulta.', bold_hasta='Diseño de Consulta: ')
    vineta(doc, 'el Resumen Ejecutivo Diario (IA) comenta la mora de la cartera cuando amerita alerta.', bold_hasta='Correo diario: ')

    # ── 10. Pagos y salidas ─────────────────────────────────────────────────
    h1(doc, '10. Cómo paga el cliente: caja, portal y prepago')
    p(doc, 'Las vías de normalización que la cobranza debe conocer (el detalle operativo vive en el '
           'Manual de Contabilidad y Tesorería):')
    vineta(doc, 'cuotas individuales o en lote, con la mora y gastos del motor único. Comprobante con timbre PAGADO.', bold_hasta='Caja (Tesorería → Caja): ')
    vineta(doc, 'el cliente entra con RUT + código OTP a /mis-creditos, ve sus cuotas y su mora — los MISMOS valores que la caja, porque es el mismo motor.', bold_hasta='Portal del cliente: ')
    vineta(doc, 'saldar el crédito completo: al valor del motor (caja) o con descuento negociado (Aplicación de Fondos, con firmas). El orden de condonación es fijo: primero gastos, luego intereses; capital nunca.', bold_hasta='Prepago: ')
    vineta(doc, 'un abono no identificado espera en Cuentas Transitorias y se aplica al aclarar el pagador — si un cliente jura que pagó y no aparece, revisar ahí primero.', bold_hasta='Transitorias: ')
    advertencia(doc, 'Si el cliente reclama que "el sistema le cobra distinto" entre el portal, el '
                     'certificado y la caja: no puede pasar, es el mismo motor. Si de verdad difiere, '
                     'es un bug — reportarlo a TI de inmediato con los pantallazos.')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('No puedo hacer otra gestión esta semana', 'Tope legal de gestiones alcanzado', 'Esperar: el bloqueo es de ley, no del sistema (cap. 1)'),
           ('El cliente dice que pagó y sigue en mora', 'Abono no identificado', 'Buscar en Cuentas Transitorias (cap. 10)'),
           ('La mora del portal difiere de la de caja', 'No puede pasar (mismo motor)', 'Reportar a TI con pantallazos (cap. 10)'),
           ('Los gastos de cobranza no aparecen', 'Menos de 21 días de atraso', 'Corren desde el día 21 (cap. 5)'),
           ('Quiero condonar parte del capital', 'No existe esa vía', 'El capital nunca se condona (cap. 5)'),
           ('El WhatsApp automático no salió', 'Sin cupo semanal, fuera de horario o sin teléfono', 'Revisar disponibilidad del crédito y el registro de envíos (cap. 2)'),
           ('Necesito acreditar la gestión ante el abogado', 'Informe formal', 'Informe PDF de gestiones desde la bitácora (cap. 3.1)'),
           ('El crédito moroso no sale en mi informe de otorgados', 'Confusión etapa / estado de cartera', 'EN MORA sigue siendo OTORGADO en etapa (cap. 1)')),
          (5.6, 5.6, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('Estado de cartera', 'Cómo paga el crédito: VIGENTE, EN MORA, VENCIDO, PREPAGADO, CASTIGADO'),
           ('TMC', 'Tasa máxima convencional; la tasa de mora queda fija al otorgamiento'),
           ('Tramo', 'Rango de días de atraso que define gestiones, gastos y provisión'),
           ('Bitácora CRM', 'Registro único de todas las gestiones de un crédito; exportable a Excel y PDF'),
           ('Gestión remota', 'WhatsApp, SMS o Email — máx. 2 por semana, separadas por 2 días'),
           ('Gestión presencial', 'Llamada, visita en terreno o visita del cliente — máx. 1 por semana'),
           ('Compromiso de pago', 'Promesa registrada con fecha y monto; su cumplimiento se sigue'),
           ('Provisión', 'Reconocimiento contable del riesgo por tramo de mora; cierre automático mensual'),
           ('Castigo', 'Baja contable del incobrable; doble firma gerencial'),
           ('Recupero', 'Pago de un crédito ya castigado; entra por caja con su asiento'),
           ('Aplicación de Fondos', 'Prepago con descuento negociado y cadena de firmas'),
           ('Score de Mora', 'Mora histórica por segmento; alimenta la política de crédito'),
           ('Snapshot judicial', 'Foto financiera del crédito al entrar a cobranza judicial'),
           ('PWA Terreno', 'App móvil de visitas: ruta del día, check-in GPS y fotos')),
          (4.3, 12.2))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-Cobranza-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
