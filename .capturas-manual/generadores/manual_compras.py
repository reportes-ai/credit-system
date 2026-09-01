# -*- coding: utf-8 -*-
"""Manual de Usuario — Compras de Oficina (Business Suite)."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE COMPRAS DE OFICINA', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Del pedido de cada persona a la orden de compra, la orden de pago', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'y el correo de aprobación — todo por el sistema, sin planillas ni correos a mano.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Septiembre 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Septiembre 2026', 'Business Suite',
            'Emisión inicial: catálogo curado con códigos PRISA/Dimeiggs, ciclo mensual con cierre el día 10, '
            'conciliación de cantidades y precios, workflow de firmas, ODP automática y correo al solicitante'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Los capítulos 1 a 3 son para TODOS los usuarios (cómo pedir). Los capítulos 4 a 7 son para '
           'Administración (conciliar, consolidar, firmar) y el 8 para el Administrador (mantenedores). '
           'Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. El ciclo en una mirada ────────────────────────────────────────────
    h1(doc, '1. El ciclo mensual en una mirada')
    p(doc, 'Las compras de materiales de oficina, cafetería y aseo funcionan en un ciclo mensual:')
    flujo(doc, 'PEDIDO (cada persona) → CIERRE día 10 → CONCILIACIÓN → ORDEN DE COMPRA + FIRMAS → ODP + CORREO → COMPRA Y RETIRO')
    vineta(doc, 'cualquier usuario, desde Soporte → Compras de Oficina, durante todo el mes.', bold_hasta='Pedir: ')
    vineta(doc, 'el día 10 de cada mes (o el hábil siguiente si cae fin de semana). La fecha exacta aparece siempre en el banner de la pantalla de Compras. Lo pedido después del cierre entra al ciclo del mes siguiente.', bold_hasta='Cierre: ')
    vineta(doc, 'Administración revisa cada pedido (cantidades y precios), consolida y genera la Orden de Compra.', bold_hasta='Conciliación: ')
    vineta(doc, 'la orden pasa por los niveles de aprobación configurados (Revisión administrativa → Gerencia General).', bold_hasta='Firmas: ')
    vineta(doc, 'con todas las firmas, el sistema genera SOLO la Orden de Pago por proveedor y le avisa por correo a cada persona qué se le aprobó.', bold_hasta='Automático: ')
    regla(doc, 'Todo se despacha a CASA MATRIZ. Los productos se retiran ahí cuando llegan '
               '(el correo de aprobación lo recuerda). Es un parámetro del sistema: si algún día se '
               'abre despacho a sucursales, se activa sin tocar código.')

    # ── 2. Hacer un pedido ───────────────────────────────────────────────────
    h1(doc, '2. Hacer un pedido (todos los usuarios)')
    ficha(doc, 'Cualquier colaborador con acceso a la Suite', 'compras (según perfil)',
          'Tener artículos asignados a tu perfil en el catálogo', 'Soporte → Compras de Oficina')
    paso(doc, 1, 'Entra a Soporte → Compras de Oficina',
         'Arriba verás la FECHA DE CIERRE del pedido de este mes. Si ya se cerró, el banner te avisa '
         'que lo que pidas entra al próximo ciclo.')
    paso(doc, 2, 'Busca tus productos',
         'El catálogo muestra solo lo habilitado para tu perfil, con foto, precio referencial y código. '
         'Usa el buscador o el filtro de categorías (Oficina, PRISA/Cafetería, PRISA/Aseo e Higiene, etc.).')
    paso(doc, 3, 'Agrega al carro con su cantidad',
         'Puedes ajustar cantidades en el carro. El total es referencial: el precio final lo concilia '
         'Administración contra la factura del proveedor.')
    paso(doc, 4, 'Envía el pedido',
         'El despacho va fijo a Casa Matriz. Puedes agregar una observación si necesitas explicar algo '
         '(ej.: "es para el archivo del piso 16").')
    paso(doc, 5, 'Sigue tu pedido en "Mis pedidos"',
         'El botón de la esquina superior muestra tus pedidos y su estado: PENDIENTE (aún no se consolida), '
         'CONSOLIDADO (ya está en una orden de compra) y el resto del avance.')
    advertencia(doc, 'Pide lo del MES. La columna "Mes pasado" que ve Administración compara tu pedido '
                     'actual con lo que pediste el mes anterior: los saltos de cantidad se revisan uno a uno.')
    captura(doc, 'Soporte → Compras de Oficina', 'catálogo con banner de fecha de cierre y carro al costado')

    # ── 3. El correo de aprobación ───────────────────────────────────────────
    h1(doc, '3. El correo "tu pedido fue aprobado"')
    p(doc, 'Cuando la orden de compra del mes completa todas sus firmas, cada persona que pidió recibe '
           'un correo automático con:')
    vineta(doc, 'producto por producto, la cantidad SOLICITADA y la cantidad APROBADA (si Administración la ajustó, dice "(ajustado)").')
    vineta(doc, 'dónde retirar: "tus productos estarán disponibles en los próximos días en Casa Matriz".')
    p(doc, 'El texto del correo es paramétrico (Mantenedores Post Venta → Correos del Sistema, plantilla '
           '"Pedido de materiales aprobado"), con su interruptor, copia y botón de prueba.')

    # ── 4. Administrar los pedidos (Asistente Administrativo) ───────────────
    h1(doc, '4. Administrar los pedidos (Asistente Administrativo)')
    ficha(doc, 'Asistente Administrativo (Cristina Peña)', 'compras_admin',
          'Pedidos pendientes del mes', 'Soporte → Administrar Compras → Pedidos pendientes')
    p(doc, 'La pestaña Pedidos pendientes muestra todos los pedidos del ciclo, con un RESUMEN POR CENTRO '
           'DE COSTO arriba (cuánto está pidiendo cada área). El centro de costo de cada pedido es el del '
           'solicitante (su ficha de Usuarios).')
    h2(doc, '4.1 Revisar y corregir un pedido')
    paso(doc, 1, 'Abre el pedido (clic en el N°)',
         'El detalle muestra cada producto con su precio, cantidad, y la columna "MES PASADO": cuánto de '
         'ese mismo producto pidió esa persona el mes anterior. Si está pidiendo más, el número sale en naranjo.')
    paso(doc, 2, 'Corrige cantidades si corresponde',
         'Cambia el número y listo — queda guardado al tiro. La cantidad SOLICITADA original se conserva: '
         'el correo de aprobación mostrará "solicitado 3 · aprobado 2 (ajustado)".')
    paso(doc, 3, 'Concilia los PRECIOS antes de consolidar',
         'El precio unitario también es editable. Al corregirlo contra la factura o lista real del proveedor: '
         '(a) se recalcula el pedido, (b) se ACTUALIZA EL CATÁLOGO — el precio conciliado queda marcado y '
         'el sync nocturno del proveedor no lo pisa — y (c) se emparejan los demás pedidos pendientes que '
         'llevan el mismo producto, para que toda la orden salga con el precio real.')
    paso(doc, 4, 'Saca productos si es necesario',
         'El botón de basurero elimina un ítem (si era el último, se elimina el pedido completo). '
         'Todo queda en Auditoría.')
    regla(doc, 'Solo se corrigen pedidos PENDIENTES. Una vez consolidados, viven en su orden de compra '
               'y ya no se tocan; si la orden se rechaza o anula, vuelven a pendientes y se pueden corregir de nuevo.')
    captura(doc, 'Administrar Compras → detalle del pedido', 'columnas Precio, Cant. y Mes pasado editables')

    # ── 5. Consolidar y la Orden de Compra ───────────────────────────────────
    h1(doc, '5. Consolidar: la Orden de Compra')
    paso(doc, 1, 'Selecciona los pedidos y pulsa "Consolidar a Casa Matriz"',
         'Sin selección, ofrece consolidar todos los pendientes. Se genera la Orden de Compra (ODC) y los '
         'pedidos pasan a CONSOLIDADO. Desde ese momento el banner de Compras informa el cierre del MES SIGUIENTE.')
    paso(doc, 2, 'Revisa la orden',
         'El detalle trae: el CONSOLIDADO POR CÓDIGO DE PRODUCTO (código, detalle, cantidad total, subtotal '
         '— lo que se le pide al proveedor), las hojas por origen con el N° de pedido de cada línea, y el '
         'texto listo para copiar y pegar al proveedor.')
    paso(doc, 3, 'La orden entra al workflow de firmas',
         'Nace EN APROBACIÓN en el nivel 1. Los firmantes reciben la campanita y firman desde la misma '
         'pantalla (o desde Soporte → Revisión de Compras). Rechazar exige motivo y devuelve los pedidos a pendientes.')

    # ── 6. Lo que pasa solo al aprobarse ─────────────────────────────────────
    h1(doc, '6. Al completarse las firmas: ODP y correos (automático)')
    p(doc, 'Cuando el último nivel firma, el sistema hace tres cosas sin que nadie las pida:')
    vineta(doc, 'la orden queda ABIERTA (lista para comprar) y suena el aviso "Orden de compra aprobada".', bold_hasta='1. ')
    vineta(doc, 'nace UNA ORDEN DE PAGO POR PROVEEDOR (los artículos con código PRISA a PRISA; los del catálogo Dimeiggs a Dimerc), con correlativo ODP oficial, monto referencial con IVA desglosado, y en las observaciones el detalle del pedido por código y LAS FIRMAS de la orden de compra (quién aprobó cada nivel y cuándo). Máxima 4: todo movimiento de dinero se registra.', bold_hasta='2. ')
    vineta(doc, 'cada solicitante recibe su correo de aprobación (capítulo 3).', bold_hasta='3. ')
    advertencia(doc, 'La ODP nace con el monto REFERENCIAL del catálogo. Cuando llegue la factura del '
                     'proveedor, Tesorería la ajusta (o anula y re-emite) antes de pagar — por eso la '
                     'conciliación de precios del capítulo 4 importa: mientras mejor conciliado el catálogo, '
                     'menor el ajuste.')
    h2(doc, '6.1 Comprar y recibir')
    p(doc, 'Con la orden ABIERTA: se envía el pedido al proveedor (botón "Copiar para Dimerc" / correo a PRISA), '
           'se marca COMPRADA al ordenar y RECIBIDA cuando llega. El reporte mensual por sucursal y el gasto '
           'por centro de costo (Órdenes de Pago → Estadísticas) se alimentan solos.')

    # ── 7. El gasto por centro de costo ──────────────────────────────────────
    h1(doc, '7. El gasto por centro de costo')
    p(doc, 'Cada compra queda imputada al centro de costo del solicitante (mantenedor Centros de Costo: '
           'las 6 áreas del organigrama y los parques con oficina). El reporte "Gasto por Centro de Costo — '
           'mes a mes" en Órdenes de Pago → Estadísticas muestra la evolución con el Δ contra el mes anterior.')

    # ── 8. Mantenedores (Administrador) ──────────────────────────────────────
    h1(doc, '8. Mantenedores (Administrador)')
    h2(doc, '8.1 Mantenedor Compras (Mantenedores → Mantenedor Compras)')
    vineta(doc, 'qué productos ve cada perfil. El botón "Solo asignados" muestra únicamente los marcados. El catálogo curado actual: un producto por cada ítem real de las facturas del proveedor, con su código PRISA o Dimeiggs.', bold_hasta='Artículos por perfil: ')
    vineta(doc, 'las oficinas de despacho (hoy solo se usa Casa Matriz).', bold_hasta='Direcciones: ')
    vineta(doc, 'niveles de firma por cargo y/o persona, con recordatorios si no firman.', bold_hasta='Workflow de aprobación: ')
    h2(doc, '8.2 Parámetros del ciclo (compras_config)')
    tabla(doc, ('Parámetro', 'Valor actual', 'Qué controla'),
          (('dia_cierre', '10', 'Día del mes en que cierra el pedido (corre al hábil siguiente si cae fin de semana)'),
           ('despacho_solo_cm', '1', 'Despacho fijo a Casa Matriz; con 0 vuelve la elección de dirección por usuario')),
          (4.5, 3.0, 9.0))
    h2(doc, '8.3 Los precios y el sync')
    regla(doc, 'Los artículos Dimeiggs se sincronizan cada noche desde el proveedor (nombre, foto, stock y '
               'precio). Un precio CONCILIADO a mano queda marcado y el sync NO lo pisa. Los artículos PRISA '
               'no tienen sync: sus precios se mantienen a mano (conciliación o mantenedor).')

    return doc
