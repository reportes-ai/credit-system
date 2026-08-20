# -*- coding: utf-8 -*-
"""Manual de Contabilidad y Tesorería — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE CONTABILIDAD Y TESORERÍA', bold=True, color=AZUL_OSCURO, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'La caja, los pagos, los libros y el cierre:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'una contabilidad que se construye sola desde la operación.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Contabilidad y Tesorería'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo: para qué existe el proceso, quién lo hace, prerequisitos, paso a paso y los '
           'recuadros ⚠ OJO / 🔒 Regla del sistema / 🧾 Caso real / 📸 CAPTURA (pendientes en esta '
           'versión). Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Principios ───────────────────────────────────────────────────────
    h1(doc, '1. Los principios: por qué esta contabilidad es distinta')
    p(doc, 'El objetivo del módulo es reemplazar el sistema contable externo (AVSOFT). Eso solo es '
           'posible si la contabilidad se construye sola desde la operación: un movimiento que no '
           'genera asiento obliga a digitarlo aparte — y ahí volvimos al sistema anterior.')
    regla(doc, 'Cada ingreso y cada egreso que el sistema registra DEBE reflejarse en la contabilidad, '
               'automáticamente y por el motor único de centralización, con su regla paramétrica en el '
               'mantenedor Reglas de Centralización. Sin excepciones.')
    vineta(doc, 'el gasto/ingreso se reconoce cuando nace la obligación (documento recibido, liquidación emitida); el pago solo rebaja el pasivo contra banco. No se confunden.', bold_hasta='Devengo y pago son dos hechos distintos: ')
    vineta(doc, 'IVA crédito/débito fiscal y retenciones van a su cuenta de pasivo o activo, para que el F29 salga del sistema y no a mano.', bold_hasta='Los impuestos son parte del asiento: ')
    vineta(doc, 'la separación por ejercicio la da el período contable, no cuentas distintas ("FACTURAS POR PAGAR 2025" no existe).', bold_hasta='Las cuentas nunca llevan año: ')
    vineta(doc, 'el motor no lanza errores a la operación: registra el resultado en el log de eventos (CONTABILIZADO / SIN_REGLA / DESCUADRE / ERROR). Ese log es la lista de pendientes: lo que aparece SIN_REGLA falta cablearlo.', bold_hasta='La contabilidad nunca bloquea la operación: ')
    captura(doc, 'Contabilidad → log de eventos contables', 'El log con eventos CONTABILIZADO y alguno SIN_REGLA si existe.')

    # ── 2. Caja ─────────────────────────────────────────────────────────────
    h1(doc, '2. Recaudación: la caja')
    p(doc, 'Cobrar la cartera propia — cada peso que entra queda registrado, timbrado y contabilizado '
           'una sola vez.')
    ficha(doc,
          'Cajero (cobra) · Tesorería (aprueba ODP y aplica transitorias)',
          'Caja asignada en Administración de Cajas',
          'Caja creada y asignada · calendario de cuotas del crédito',
          'Tesorería → Caja')
    paso(doc, 1, 'Buscar el crédito y cobrar', 'Cuotas individuales o en lote. La mora y los gastos '
         'salen del motor único — los mismos valores que muestran el portal del cliente y los '
         'certificados. El comprobante sale con timbre PAGADO.')
    paso(doc, 2, 'Pago con aprobación', 'Se emite ODP de cuotas; Tesorería aprueba y el pago queda '
         'registrado con correo de respaldo.')
    paso(doc, 3, 'Abono no identificado', 'Entra a Cuentas Transitorias y se aplica al crédito cuando '
         'se aclara el pagador.')
    paso(doc, 4, 'Cierre de caja', 'Al final del día, cuadratura de lo recaudado.')
    captura(doc, 'Tesorería → Caja', 'Un crédito con sus cuotas, una pagada con timbre y el detalle de mora del motor.')
    flujo(doc, 'CUOTA VIGENTE → PAGADA · TRANSITORIA → APLICADA')
    advertencia(doc, 'El calendario de cuotas se congeló al otorgar el crédito y es inmutable. Si una '
                     'cuota "no calza", el problema está en la operación (mes de mora, gastos), nunca '
                     'se corrige editando el calendario.')

    # ── 3. Prepago ──────────────────────────────────────────────────────────
    h1(doc, '3. Prepago y Aplicación de Fondos')
    p(doc, 'Saldar un crédito completo antes de plazo. Dos vías: al valor del motor (prepago simple, '
           'en caja) o con descuento negociado (Aplicación de Fondos, con cadena de firmas).')
    ficha(doc,
          'Cajero (prepago simple) · Tesorería + cadena de firmas (aplicación con descuento)',
          'Atribuciones de condonación configuradas',
          'Crédito vigente de cartera propia',
          'Caja → botón Prepago · Tesorería → Aplicación de Fondos')
    paso(doc, 1, 'Prepago en caja', 'El mismo motor del certificado calcula capital + accesorios; la '
         'condonación posible depende de las atribuciones del usuario. El crédito queda PREPAGADO.')
    paso(doc, 2, 'Aplicación de Fondos', 'Propuesta ítem por ítem sobre el motor de prepago, que '
         'recorre HECHO → APROBADO → PROCESADO antes de aplicarse. Lo condonado queda registrado.')
    regla(doc, 'El orden de condonación es fijo: primero gastos, luego intereses. El capital NUNCA se '
               'condona.')
    captura(doc, 'Tesorería → Aplicación de Fondos', 'Una propuesta ítem por ítem con sus firmas.')
    flujo(doc, 'HECHO → APROBADO → PROCESADO → PREPAGADO')

    # ── 4. Órdenes de pago ──────────────────────────────────────────────────
    h1(doc, '4. Órdenes de Pago y segregación de funciones')
    p(doc, 'Toda salida de plata pasa por una ODP con correlativo único global: comisiones de dealers '
           'y parques, saldos precio, proveedores, cuotas con aprobación. La ODP es el punto donde se '
           'aplican los controles.')
    regla(doc, 'Segregación de funciones: quien emite o manda a pago NO paga. El bloqueo es por '
               'usuario, está en el servidor y es paramétrico en Administración de Cajas. Vale para '
               'todas las ODP del sistema.')
    vineta(doc, 'con doble firma donde corresponde (anulaciones, castigos, aplicaciones).', bold_hasta='Emitir: ')
    vineta(doc, 'una ODP no pagada se puede anular con motivo: el correlativo queda anulado y nunca se reutiliza.', bold_hasta='Anular: ')
    vineta(doc, 'un pago se puede revertir con motivo auditado: la ODP vuelve a "por pagar".', bold_hasta='Revertir: ')
    captura(doc, 'Órdenes de Pago', 'El ledger de ODP con sus estados y correlativo.')
    p(doc, 'Cada pago genera su asiento automático: devengo al nacer la obligación, pago al '
           'confirmarse (rebaja del pasivo contra banco). Ver capítulo 1.')

    # ── 5. Saldos precio ────────────────────────────────────────────────────
    h1(doc, '5. Saldos precio: plata en tránsito')
    p(doc, 'La financiera transfiere el saldo del precio del vehículo y AutoFácil lo entrega íntegro '
           'al dealer. Tesorería emite y paga la nómina; el compromiso es el SLA con el canal.')
    ficha(doc,
          'Post Venta define fondos y nómina · Tesorería emite y paga',
          'Emisión y pago separados',
          'Fondos recibidos de la financiera · etapas de Saldo Precio al día',
          'Post Venta → Saldos Precio · Tesorería → card Saldo Precio en Proceso de Pago')
    regla(doc, 'El saldo precio es una CUENTA DE PASO: al marcar FONDOS RECIBIDOS entra banco contra '
               'el pasivo transitorio, y al marcar SALDO PRECIO PAGADO se rebaja contra banco. No toca '
               'el resultado — no infla ingresos ni gastos — y el saldo de la cuenta muestra cuánta '
               'plata de dealers hay en tránsito.')
    h2(doc, '5.1 La card de seguimiento en Tesorería')
    p(doc, 'La planilla histórica se reconstruye en vivo desde las etapas de Post Venta: recepción de '
           'fundantes (con AM/PM), vencimiento por SLA — 24/48/72 horas hábiles según la clasificación '
           'del dealer (SUPERPARTNER / PARTNER / SOCIO; recepción PM cuenta desde el día hábil '
           'siguiente) —, N° de ODP, carga al banco y pago.')
    vineta(doc, 'sale de la ODP (saldo + gastos). Sin ODP aún, se muestra el saldo marcado preliminar.', bold_hasta='El monto a pagar: ')
    vineta(doc, 'genera el archivo con el nombre exacto que exige la otra aplicación. Su "CARGADO" es nuestra etapa ENVIADO A PAGO; su "CARGAR" es ORDEN DE PAGO EMITIDA.', bold_hasta='Exportar Excel: ')
    captura(doc, 'Tesorería → Saldo Precio en Proceso de Pago', 'La card con operaciones en distintos estados y sus SLA.')
    flujo(doc, 'SALDO PENDIENTE → EN NÓMINA → ORDEN EMITIDA → PAGADO')

    # ── 6. Comisiones dealer/parque (lado pago) ────────────────────────────
    h1(doc, '6. Pagar comisiones: dealers y parques')
    p(doc, 'El circuito completo de cartolas vive en el Manual de Operaciones (capítulos 12 y 13). '
           'Aquí, lo que le toca a Tesorería y Contabilidad: el documento, el impuesto y el pago.')
    h2(doc, '6.1 Factura o boleta: qué paga la ODP')
    regla(doc, 'La comisión calculada es BRUTA (IVA incluido); se desagrega con el motor único, nunca '
               'se le suma IVA encima. Con FACTURA, la ODP paga el bruto y el IVA es crédito fiscal. '
               'Con BOLETA de honorarios, el honorario bruto es el neto de la factura equivalente y la '
               'ODP paga el líquido (honorario − retención); la retención queda como pasivo con el SII '
               'y va al F29. La boleta entra sola al auxiliar de honorarios.')
    h2(doc, '6.2 Los asientos automáticos')
    tabla(doc, ('Hecho', 'Asiento'),
          (('Se registra el documento del dealer', 'Devengo del gasto contra el pasivo'),
           ('Se marca COMISIÓN PAGADA', 'Rebaja del pasivo contra banco'),
           ('Se aprueba el mes de un parque', 'DEBE Arriendos + DEBE Comisiones por Ventas → HABER pasivo (cuentas separadas)'),
           ('Se paga la ODP del parque', 'DEBE pasivo → HABER banco'),
           ('Se solicita facturar el ingreso AutoFácil', 'Devengo de la comisión por cobrar (neto + IVA débito fiscal)')),
          (7.0, 9.5))
    p(doc, 'Cada asiento es idempotente por referencia: aprobar dos veces no duplica. Las reglas son '
           'paramétricas en Reglas de Centralización.')

    # ── 7. Plan Liquidez ────────────────────────────────────────────────────
    h1(doc, '7. Plan Liquidez: el anticipo de los Super Partner')
    p(doc, 'Renovar cada mes el anticipo de comisiones de los dealers Super Partner, descontando por '
           'ODP sin tocar la cartola.')
    ficha(doc,
          'Jefe/Gerente Comercial (N1 aprueba o modifica) · Gerente General (N2) · Finanzas toma '
          'conocimiento · Tesorería paga',
          'Cadena de niveles del plan',
          'Plan creado (contrato + pagaré) · cartolas del mes enviadas (la comisión sale de ahí)',
          'Dealers → Plan Liquidez → Hojas')
    paso(doc, 1, 'Generar la Hoja de Liquidación', 'Una línea por dealer: comisión C, deuda D, '
         'adelanto A = min(C, tope), descuento D−A y pago neto.')
    paso(doc, 2, 'Cadena N1 → N2 → N3', '48 horas por nivel, auto-avanza al vencer. N1 puede '
         'modificar líneas con motivo.')
    paso(doc, 3, 'Emitir la ODP', 'Con el descuento aplicado (o una segunda ODP si es aumento). Al '
         'pagarse, el abono queda en la cuenta corriente del dealer y la deuda pasa a A.')
    captura(doc, 'Plan Liquidez → Hoja de Liquidación', 'Una hoja con sus líneas por dealer y la cadena de firmas.')
    flujo(doc, 'HOJA GENERADA → N1 → N2 → CONOCIMIENTO N3 → ODP → PAGADA / ABONADA')

    # ── 8. Conciliación ─────────────────────────────────────────────────────
    h1(doc, '8. Conciliación bancaria')
    p(doc, 'La cartola del banco entra por la conexión bancaria (Fintoc) o por archivo, y el matching '
           'automático calza cada abono con su movimiento del sistema: cuotas pagadas, saldos '
           'recibidos, transferencias. Lo que no calza queda en la bandeja para conciliar a mano.')
    vineta(doc, 'un abono de cliente que no calza con ninguna cuota va a Cuentas Transitorias hasta aclarar el pagador (cap. 2).', bold_hasta='Abonos no identificados: ')
    vineta(doc, 'el cierre de mes exige la conciliación al día — es uno de los puntos del checklist (cap. 11).', bold_hasta='Relación con el cierre: ')
    captura(doc, 'Tesorería → Conciliación Bancaria', 'La bandeja con movimientos calzados y alguno pendiente.')

    # ── 9. RCV ──────────────────────────────────────────────────────────────
    h1(doc, '9. Libro de compras: del SII al auxiliar (RCV)')
    p(doc, 'Dejar de digitar facturas de proveedores una por una: el SII ya tiene el detalle (RUT, '
           'folio, fechas, neto, IVA, total); el sistema lo trae y lo convierte en documentos del '
           'auxiliar con su asiento. El F29 sale del sistema porque el libro de compras se construye '
           'solo.')
    ficha(doc,
          'El Contador revisa y confirma · la sincronización la hace el sistema solo (cada 2 días)',
          'ctb_libros_aux',
          'Certificado digital y Clave Tributaria cargados · mes sin candado',
          'Contabilidad → Libros Auxiliares')
    paso(doc, 1, 'Sincronizar', 'Botón "Sincronizar SII (RCV)" o automático cada 2 días (el plan '
         'gratuito da 30 consultas al mes). Deja un espejo de solo lectura del SII; reemplaza el mes '
         'completo, así que re-sincronizar siempre es seguro.')
    paso(doc, 2, 'Ver la brecha', 'El banner muestra documentos del SII vs documentos del auxiliar. '
         'Esa diferencia es exactamente lo que falta por ingresar.')
    paso(doc, 3, 'Traer del SII al auxiliar', 'Lista los documentos del mes que no están en el '
         'auxiliar. Folios y montos vienen del SII y no se editan; lo único que define el contador es '
         'la CUENTA DE GASTO. La cuenta se propone sola: la que ese proveedor usó más veces en el '
         'histórico. Un proveedor nuevo aparece en rojo, sin propuesta.')
    paso(doc, 4, 'Confirmar', 'Cada documento entra con su asiento: DEBE gasto (neto+exento) + DEBE '
         'IVA crédito → HABER proveedores (total). Un documento que falla no bota a los demás: el '
         'resultado dice cuáles entraron y por qué se saltó cada otro.')
    captura(doc, 'Libros Auxiliares → Traer del SII', 'El diálogo con documentos pendientes y la cuenta propuesta por historial.')
    regla(doc, 'No se duplica: se busca por RUT + tipo de DTE + folio en TODO el auxiliar, no solo en '
               'el mes. Las notas de crédito (tipo 61) NO se importan — en el auxiliar van con signo '
               'negativo y se digitan aparte. Ingresar por RCV usa el mismo motor que la digitación '
               'manual, así que las validaciones son idénticas por los dos caminos.')
    flujo(doc, 'SII (RCV) → ESPEJO → REVISIÓN DE CUENTAS → AUXILIAR + ASIENTO → F29')

    # ── 10. Provisiones y castigos ──────────────────────────────────────────
    h1(doc, '10. Provisiones y castigos')
    p(doc, 'La provisión reconoce contablemente el riesgo de la cartera morosa; el castigo da de baja '
           'la operación incobrable. La primera es automática; el segundo, deliberadamente manual.')
    vineta(doc, 'si nadie las guarda a mano, el sistema cierra solo el mes anterior (día configurable), genera el asiento por la variación (constitución o liberación) y deja el detalle por crédito auditable. Los tramos viven en Parámetros de Cobranza.', bold_hasta='Provisiones automáticas: ')
    vineta(doc, 'baja contable de la operación, SIEMPRE manual y con doble firma gerencial. Nunca automático. Genera su asiento por regla de centralización.', bold_hasta='Castigo: ')
    captura(doc, 'Tesorería → Provisiones', 'El mes con su provisión calculada y el detalle por tramo.')

    # ── 11. Cierre de mes ───────────────────────────────────────────────────
    h1(doc, '11. Cierre de Mes')
    p(doc, 'Cerrar el mes con todos los controles dados y las cifras congeladas: el candado que hace '
           'confiable la contabilidad y los informes a la matriz.')
    ficha(doc,
          'Cada responsable de punto (checklist) · quien tiene la atribución CERRAR MES · el sistema '
          '(provisiones y chequeos automáticos)',
          'Atribución CERRAR MES',
          'Checklist configurado (responsables, días hábiles) · tramos de provisión configurados',
          'Tesorería → Cierre de Mes')
    paso(doc, 1, 'Checklist', 'Cada responsable da su OK; varios puntos se chequean solos '
         '(conciliación, transitorias, cartolas, comisiones, castigos). Correo diario con los '
         'vencidos hasta cumplir.')
    paso(doc, 2, 'CERRAR MES', 'Con los obligatorios en verde: candado contable + acta congelada '
         '(quién dio cada OK) + correo del acta a los destinatarios.')
    paso(doc, 3, 'Informe a la matriz', 'El Informe de Cierre Mensual sale de Contabilidad, con la '
         'Bitácora de Cierres y su análisis de IA.')
    captura(doc, 'Tesorería → Cierre de Mes', 'El checklist con puntos en verde y el botón de cierre.')
    regla(doc, 'Cerrado el mes, todo lo retroactivo queda bloqueado: los recálculos respetan meses '
               'cerrados y las operaciones de ese mes no se editan. Lo que haya que corregir de un mes '
               'cerrado va por procesos formales, nunca reabriendo.')
    flujo(doc, 'MES ABIERTO → PUNTOS EN VERDE → CERRADO (CANDADO) → ACTA ENVIADA')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('Un movimiento quedó SIN_REGLA en el log', 'Falta la regla de centralización de ese hecho', 'Crearla en Reglas de Centralización; el log es la lista de pendientes (cap. 1)'),
           ('Un abono del banco no calza con nada', 'Pagador no identificado', 'Cuentas Transitorias; aplicar al aclarar (cap. 2 y 8)'),
           ('La factura del dealer no cuadra con la cartola', 'El dealer facturó otro monto', 'Se rechaza con la diferencia; debe facturar el TOTAL de la cartola'),
           ('No puedo pagar una ODP que emití', 'Segregación de funciones', 'Debe pagarla otro usuario (cap. 4)'),
           ('El RCV no trae documentos nuevos', 'Cuota mensual del SII agotada o certificado vencido', 'Revisar el log de sincronización (cap. 9)'),
           ('Un documento del RCV no entró', 'Duplicado (RUT+tipo+folio ya existe) o folio inválido', 'El resultado de la importación dice el motivo (cap. 9)'),
           ('No se puede cerrar el mes', 'Puntos obligatorios del checklist pendientes', 'El checklist dice cuáles y quién es responsable (cap. 11)'),
           ('Una cuota cobrada muestra otra mora que el certificado', 'No puede pasar: es el mismo motor', 'Si pasa, reportar a TI de inmediato — es un bug (cap. 2)')),
          (5.6, 5.6, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('Devengo', 'Reconocimiento del gasto/ingreso cuando nace la obligación, antes del pago'),
           ('ODP', 'Orden de pago, correlativo único global'),
           ('Cuenta de paso', 'Cuenta que refleja plata en tránsito; no toca resultado'),
           ('RCV', 'Registro de Compras y Ventas del SII'),
           ('Auxiliar', 'Libro operativo de compras/honorarios del sistema'),
           ('F29', 'Declaración mensual de impuestos; se arma con los datos del sistema'),
           ('Regla de centralización', 'Definición paramétrica de qué asiento genera cada hecho económico'),
           ('Idempotente', 'Ejecutar dos veces produce el mismo resultado: no duplica'),
           ('Candado', 'Bloqueo del mes cerrado: nada retroactivo se toca'),
           ('Provisión', 'Reconocimiento contable del riesgo de la cartera morosa'),
           ('Castigo', 'Baja contable de una operación incobrable; doble firma'),
           ('Transitoria', 'Cuenta donde espera un abono cuyo pagador no está claro'),
           ('SLA', 'Plazo comprometido de pago del saldo precio (24/48/72 h hábiles según categoría)')),
          (4.3, 12.2))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-Contabilidad-Tesoreria-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
