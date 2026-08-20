# -*- coding: utf-8 -*-
"""Contabilidad y Cierre de Mes — Business Suite (capítulo conceptual)."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'CONTABILIDAD Y CIERRE DE MES', bold=True, color=AZUL_OSCURO, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'La contabilidad que se construye sola desde la operación,', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'y el candado mensual que la hace confiable.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. La misión ────────────────────────────────────────────────────────
    h1(doc, '1. La misión: reemplazar el sistema contable externo')
    p(doc, 'El módulo de Contabilidad existe para una sola cosa: que AutoFácil deje de pagar y '
           'digitar un sistema contable aparte (AVSOFT). Eso solo es posible si la contabilidad se '
           'construye SOLA desde la operación: cada cartola pagada, cada cuota cobrada, cada '
           'liquidación emitida genera su asiento sin que nadie lo digite. Un movimiento que obliga '
           'a digitar aparte es una regresión al sistema anterior.')
    p(doc, 'El estado actual: plan de cuentas propio, comprobantes, libros diario y mayor, balance, '
           'auxiliares de compras y honorarios, los libros históricos de AVSOFT importados (más de '
           '6.000 comprobantes con su cierre de ejercicio), el Registro de Compras del SII entrando '
           'solo, y el motor de centralización contabilizando los hechos del negocio. Lo que falta '
           'para el reemplazo total es la emisión electrónica de documentos (decisión en Finanzas).')

    # ── 2. Principios ───────────────────────────────────────────────────────
    h1(doc, '2. Los cinco principios contables del sistema')
    tabla(doc, ('Principio', 'Qué significa'),
          (('Un solo motor de asientos', 'Todo asiento automático nace del motor de centralización, según la regla paramétrica de su hecho. Ningún módulo escribe asientos "por su cuenta"'),
           ('Devengo ≠ pago', 'El gasto o ingreso se reconoce cuando nace la obligación (documento recibido, liquidación emitida); el pago solo rebaja el pasivo contra banco. Son dos asientos, dos momentos'),
           ('Los impuestos van en el asiento', 'IVA crédito y débito fiscal, y las retenciones, a su cuenta de activo o pasivo — para que el F29 salga del sistema y no a mano'),
           ('Las cuentas no llevan año', 'La separación por ejercicio la da el período contable, no cuentas como "FACTURAS POR PAGAR 2025"'),
           ('La contabilidad nunca bloquea', 'El motor no detiene la operación: registra el resultado de cada hecho en su log (CONTABILIZADO / SIN_REGLA / DESCUADRE / ERROR). Ese log es la lista de trabajo del contador')),
          (4.4, 12.1))
    regla(doc, 'El corolario para todo el equipo: un feature nuevo que mueve plata está INCOMPLETO '
               'hasta que tiene su regla contable. La pregunta "¿qué asiento genera esto y con qué '
               'cuentas?" es parte de cerrar cualquier funcionalidad.')

    # ── 3. Las piezas ───────────────────────────────────────────────────────
    h1(doc, '3. Las piezas de la contabilidad')
    vineta(doc, 'el catálogo propio, administrable, con cuentas imputables y de agrupación. Toda regla de centralización apunta acá.', bold_hasta='Plan de cuentas: ')
    vineta(doc, 'cada asiento con su correlativo, cuadratura obligatoria (un comprobante descuadrado no entra) y su origen: manual, o automático con la referencia del hecho que lo generó.', bold_hasta='Comprobantes: ')
    vineta(doc, 'diario, mayor por cuenta y balance, con los períodos contables como columna vertebral.', bold_hasta='Libros y balance: ')
    vineta(doc, 'compras y honorarios — alimentados por el RCV del SII y por las boletas que entran solas al pagar comisiones. La comparación contra el SII es el control.', bold_hasta='Auxiliares: ')
    vineta(doc, 'los libros de AVSOFT importados dan la continuidad histórica: el balance no parte de cero.', bold_hasta='La historia: ')
    captura(doc, 'Contabilidad → Comprobantes', 'Un comprobante automático con su referencia al hecho de origen.')

    # ── 4. Centralización ───────────────────────────────────────────────────
    h1(doc, '4. La centralización: del hecho al asiento')
    p(doc, 'El corazón del módulo. Cada hecho económico del sistema dispara un evento; el motor '
           'busca su regla en el mantenedor Reglas de Centralización y arma el asiento. Ejemplos '
           'reales del circuito:')
    tabla(doc, ('Hecho del negocio', 'Asiento (resumen)'),
          (('Se registra la factura del dealer', 'DEBE gasto comisión (neto) + DEBE IVA crédito → HABER proveedores'),
           ('Se paga la comisión (ODP)', 'DEBE proveedores → HABER banco'),
           ('Boleta de honorarios del dealer', 'DEBE gasto → HABER retención SII + HABER por pagar; la retención va al F29'),
           ('Se aprueba el mes de un parque', 'DEBE arriendos + DEBE comisiones por ventas → HABER pasivo del parque'),
           ('Fondos recibidos de la financiera', 'DEBE banco → HABER pasivo transitorio (cuenta de paso)'),
           ('Saldo precio pagado al dealer', 'DEBE pasivo transitorio → HABER banco'),
           ('Se emite el libro de remuneraciones', 'Sueldos, cotizaciones y retenciones, cada una a su cuenta'),
           ('Finiquito emitido / pagado', 'Devengo a finiquitos por pagar; el pago rebaja contra banco'),
           ('Anticipo o préstamo al personal', 'Cuenta por cobrar al personal (activo), NO gasto'),
           ('Documento del RCV confirmado', 'DEBE gasto + DEBE IVA crédito → HABER proveedores'),
           ('Provisión mensual de mora / vacaciones', 'Constitución o liberación por la VARIACIÓN contra el mes anterior')),
          (6.6, 9.9))
    vineta(doc, 'procesar dos veces el mismo hecho NO duplica el asiento: cada uno es idempotente por su referencia.', bold_hasta='Idempotencia: ')
    vineta(doc, 'un hecho sin regla queda SIN_REGLA en el log — visible, no perdido. Cablearlo es crear la regla en el mantenedor, sin tocar código.', bold_hasta='Lo no cableado se ve: ')
    captura(doc, 'Contabilidad → Reglas de Centralización', 'El mantenedor con las reglas por evento y sus cuentas.')

    # ── 5. SII ──────────────────────────────────────────────────────────────
    h1(doc, '5. El SII dentro del sistema')
    vineta(doc, 'las facturas de proveedores entran solas desde el Registro de Compras cada 2 días: folios y montos del SII (no se editan), la cuenta de gasto propuesta por el historial del proveedor, y el asiento en el mismo acto. El banner de brecha (SII vs auxiliar) es el control permanente.', bold_hasta='RCV — compras: ')
    vineta(doc, 'con el IVA crédito y débito en sus cuentas y las retenciones registradas, la declaración mensual se arma con datos del sistema, no a mano.', bold_hasta='F29: ')
    vineta(doc, 'la certificación ante el SII para emitir facturas desde la plataforma es el último tramo del reemplazo — evaluada y con su informe en Finanzas.', bold_hasta='Emisión electrónica (pendiente): ')

    # ── 6. Cierre de mes ────────────────────────────────────────────────────
    h1(doc, '6. El Cierre de Mes: el candado que hace todo confiable')
    p(doc, 'Sin cierre, cualquier informe es provisorio para siempre. El cierre convierte el mes en '
           'un hecho: controles dados, cifras congeladas, acta firmada.')
    paso(doc, 1, 'El checklist', 'Cada punto con su responsable y su día hábil límite: conciliación '
         'bancaria, transitorias aclaradas, cartolas enviadas, comisiones aprobadas, castigos '
         'resueltos. Varios se chequean solos; el correo diario persigue a los vencidos.')
    paso(doc, 2, 'Las provisiones', 'Mora y vacaciones se provisionan automáticamente; si nadie las '
         'guarda a mano, el sistema cierra solo el mes anterior el día configurado.')
    paso(doc, 3, 'CERRAR MES', 'Con los obligatorios en verde: candado contable + acta congelada '
         '(quién dio cada OK) + correo del acta a los destinatarios.')
    paso(doc, 4, 'El informe a la matriz', 'El Informe de Cierre Mensual sale hacia Casa Matriz con '
         'sus secciones validadas contra la base, y la Bitácora de Cierres guarda cada mes con su '
         'análisis.')
    captura(doc, 'Tesorería → Cierre de Mes', 'El checklist con sus responsables y el botón de cierre.')
    regla(doc, 'Cerrado el mes, NADA retroactivo se toca: los recálculos respetan el candado, las '
               'operaciones del mes no se editan y las variables de comisiones no aceptan vigencias '
               'sobre meses cerrados. Lo que haya que corregir va por proceso formal en el mes '
               'abierto — la historia informada no se reescribe.')
    flujo(doc, 'MES ABIERTO → PUNTOS EN VERDE → CERRADO (CANDADO) → ACTA ENVIADA → INFORME A LA MATRIZ')

    # ── 7. Honestidad ───────────────────────────────────────────────────────
    h1(doc, '7. Lo que falta (y está anotado)')
    vineta(doc, 'las reglas contables de las comisiones internas y la colocación están definidas; falta conectar su disparador automático.', bold_hasta='Comisiones internas y colocación: ')
    vineta(doc, 'la convención para operaciones cursadas a caballo entre dos meses (curse el día 1 vs mes contable anterior) está pendiente de decisión — toca meses cerrados, así que se decide con Finanzas, no de facto.', bold_hasta='Mes contable vs fecha de curse: ')
    vineta(doc, 'certificarse ante el SII y emitir desde la plataforma — con costo conocido y decisión en Finanzas.', bold_hasta='Emisión electrónica: ')
    p(doc, 'Todo lo abierto vive en Pendientes Abiertos de la Suite de Documentación — la lista '
           'única, actualizada en el momento en que algo se detecta.')

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Contabilidad-CierreMes-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
