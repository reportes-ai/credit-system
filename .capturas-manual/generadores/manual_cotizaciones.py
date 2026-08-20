# -*- coding: utf-8 -*-
"""Cotizaciones y Evaluación Crediticia — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'COTIZACIONES Y EVALUACIÓN CREDITICIA', bold=True, color=AZUL_OSCURO, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Del precio del auto a la decisión de crédito:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'una cuota que no miente y una evaluación con toda la información.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Cotizar ──────────────────────────────────────────────────────────
    h1(doc, '1. Cotizaciones: la cuota que no miente')
    p(doc, 'La cotización convierte precio + pie + plazo en la cuota real que pagará el cliente — '
           'con gastos operacionales, seguros y CAE, no la cuota "pelada" que después crece. Un solo '
           'motor la calcula para todas las puertas de entrada.')
    ficha(doc,
          'Ejecutivo comercial · el dealer desde su portal · el cliente vía Facilito',
          'Acceso al módulo Cotizaciones',
          'Tasas vigentes · parámetros del crédito · UF del día',
          'Cotizaciones · Simulador Rápido · Portal del Dealer · Facilito (WhatsApp)')
    h2(doc, '1.1 Cómo se arma la cuota')
    flujo(doc, 'PRECIO − PIE = SALDO → + GASTOS OPERACIONALES → + SEGUROS → MONTO FINANCIADO → CUOTA FRANCESA')
    vineta(doc, 'prenda, inscripción, gestión, GPS y demás — paramétricos en Parámetros del Crédito.', bold_hasta='Gastos operacionales: ')
    vineta(doc, 'el paquete completo (desgravamen + RDH + cesantía) con el factor actuarial de la aseguradora por tramo de plazo — mantenedor Factores de Seguros. Desmarcar un seguro recalcula todo.', bold_hasta='Seguros: ')
    vineta(doc, 'la del mantenedor de Tasas para el tramo (mayor/menor 200 UF) a la fecha. La cuota es francesa sobre el monto financiado total.', bold_hasta='Tasa: ')
    vineta(doc, 'el costo total del crédito expresado como exige la ley, calculado por su motor.', bold_hasta='CAE: ')
    regla(doc, 'Las cuatro puertas cotizan con el MISMO motor: la cuota que da Facilito por WhatsApp '
               'es la que da el simulador del ejecutivo y la pre-aprobación del dealer. Una '
               'cotización distinta entre canales es un bug, no una promoción.')
    h2(doc, '1.2 La cotización queda guardada')
    p(doc, 'Toda cotización — incluidas las del bot — queda guardada al RUT en la tabla única de '
           'cotizaciones. Cuando el negocio avanza, la evaluación y la carta parten de ahí; y la '
           'estadística de conversión cotización→carta→curse tiene su materia prima.')
    captura(doc, 'Cotizaciones → simulador', 'Una cotización completa con gastos, seguros, CAE y la cuota final.')
    h2(doc, '1.3 ¿Dónde Curso?: la cotización con veredicto')
    p(doc, 'La PWA del ejecutivo agrega a la cotización el veredicto de financiera: AUTOFIN, UNIDAD '
           'o DECIDES TÚ, con la rentabilidad de cada lado. Primero elegibilidad (el cuadro de '
           'Preferencia Financiera), después rentabilidad, con el umbral de empate. El detalle vive '
           'en el Manual del Ejecutivo (cap. 3) y el motor en el documento de Motores (4.2).')

    # ── 2. El cliente ───────────────────────────────────────────────────────
    h1(doc, '2. La ficha del cliente: la materia prima')
    p(doc, 'La evaluación es tan buena como sus datos. El sistema los organiza en tres piezas con '
           'fuente única:')
    tabla(doc, ('Pieza', 'Qué guarda', 'Regla'),
          (('Datos personales', 'Identidad, contacto, dirección, comuna', 'Un RUT = un cliente; el resto del sistema referencia, nunca copia'),
           ('Antecedentes laborales', 'Situación laboral, renta fija líquida, antigüedad', 'Se renuevan si tienen más de 30 días al evaluar'),
           ('Información comercial', 'Perfil de deudas: morosidades, protestos, castigos', 'Alimentada por los informes; es la foto de riesgo del RUT')),
          (3.6, 6.4, 6.5))
    p(doc, 'A eso se suman los documentos que el cliente aporta (liquidaciones de sueldo, carpeta '
           'tributaria) — leídos por el subsistema de IA cuando está activo, con las reglas '
           'previsionales paramétricas para calcular la renta líquida real.')

    # ── 3. Informes ─────────────────────────────────────────────────────────
    h1(doc, '3. Los informes comerciales (DealerNet)')
    p(doc, 'El informe comercial del RUT se pide en línea y alimenta la información comercial. Sus '
           'reglas de uso están pensadas para no gastar de más y no concluir de menos:')
    vineta(doc, 'un RUT consultado hace menos de 15 días sirve el informe vigente sin pagar otra consulta.', bold_hasta='Caché de 15 días: ')
    vineta(doc, 'la deuda del producto crediticio viene en MILES de pesos; el motor la convierte. Comparar a mano contra el documento original exige recordar el factor.', bold_hasta='La deuda viene en miles: ')
    advertencia(doc, 'Un informe con sello rojo "SIN INFORMACIÓN — Intente más tarde" llegó VACÍO: '
                     'no significa que el cliente no tenga deuda y NO sirve para concluir. No ocupa '
                     'el bloqueo de 15 días — se puede repedir de inmediato. Confundir "sin '
                     'información" con "sin deuda" es el error más caro de esta pantalla.')

    # ── 4. Evaluación ───────────────────────────────────────────────────────
    h1(doc, '4. La Evaluación Crediticia')
    p(doc, 'Donde todo converge: la ficha, los informes, la política y el análisis. El resultado es '
           'la información completa para que el analista decida — o para que el sistema pre-decida '
           'en los canales de autoservicio.')
    ficha(doc,
          'Ejecutivo comercial arma la evaluación · Analista de Crédito decide',
          'Acceso a Evaluación Crediticia',
          'Cliente en la base · saldo de informes disponible',
          'Evaluación Crediticia')
    h2(doc, '4.1 El flujo')
    paso(doc, 1, 'Buscar por RUT', 'Trae la ficha completa; los antecedentes con más de 30 días se '
         'renuevan.')
    paso(doc, 2, 'Pedir informes', 'El informe comercial entra a la ficha (con las reglas del '
         'capítulo 3).')
    paso(doc, 3, 'Leer el análisis', 'La IA cruza renta, deudas y política, y redacta el análisis. '
         'Las cifras son de los motores; la IA ordena y redacta.')
    paso(doc, 4, 'Decidir', 'El analista aprueba, rechaza o pide más antecedentes. Al entrar un caso '
         'a análisis, la campanita avisa al pool de analistas — nadie queda esperando en silencio.')
    captura(doc, 'Evaluación Crediticia', 'Una evaluación completa con ficha, informes y el análisis de IA.')
    h2(doc, '4.2 La política de crédito y el scorecard')
    p(doc, 'La política vigente (V3.0) combina las reglas duras (relación cuota/renta, antigüedad '
           'del vehículo, informes limpios) con un scorecard experto construido sobre la mora '
           'histórica por segmento: los quintiles de riesgo salen de cómo se comportó de verdad la '
           'cartera, no de una intuición. El Score de Mora del sistema es el que alimenta esa '
           'calibración — la cobranza de ayer es la política de admisión de hoy.')
    regla(doc, 'La regla del 30%: la cuota no puede superar el 30% de la renta líquida. En los '
               'canales de autoservicio (pre-aprobación del dealer) manda la renta DECLARADA por '
               'quien pide, con la interna de respaldo — y si difieren mucho, el detalle viaja al '
               'Jefe Comercial, no al solicitante.')

    # ── 5. Los tres niveles de evaluación ───────────────────────────────────
    h1(doc, '5. Un motor, tres niveles de profundidad')
    p(doc, 'La misma lógica de evaluación corre en tres contextos, con profundidad creciente:')
    tabla(doc, ('Nivel', 'Dónde', 'Qué entrega'),
          (('Preevaluación exprés', 'Facilito (WhatsApp)', 'El código valida el RUT, pide informes (con caché y límites) y decide el veredicto: vía exprés con pie ≥ 40% y antecedentes buenos, o derivación a ejecutivo. Sin revelar detalles del informe al cliente'),
           ('Pre-aprobación en línea', 'Portal del Dealer', 'Antigüedad del vehículo, regla del 30%, informes limpios y elegibilidad de financiera → PREAPROBADO con cuotas, o "necesito más información". El dealer nunca ve datos del cliente'),
           ('Evaluación completa', 'Evaluación Crediticia', 'Todo lo anterior más el análisis de IA, el scorecard y el criterio del analista — la única que aprueba de verdad')),
          (3.4, 3.6, 9.5))
    p(doc, 'Los dos primeros niveles filtran y aceleran; el tercero decide. Ninguno contradice a '
           'otro porque comparten los motores y la política.')

    # ── 6. Gobierno ─────────────────────────────────────────────────────────
    h1(doc, '6. Gobierno')
    vineta(doc, 'tasas, gastos, factores de seguros, umbrales de preferencia y reglas de política viven en mantenedores, con su footer de trazabilidad.', bold_hasta='Todo paramétrico: ')
    vineta(doc, 'los informes cuestan plata: caché de 15 días, límites por conversación y por día en el bot, y el mantenedor de costo que recomienda plan según consumo.', bold_hasta='El gasto se cuida: ')
    vineta(doc, 'las evaluaciones, sus informes y sus análisis quedan en el repositorio único de informes — el mismo que consultan la carta y el revisor.', bold_hasta='Fuente única: ')
    vineta(doc, 'quien ve la evaluación completa es porque su perfil lo permite; los canales de autoservicio entregan veredictos, nunca detalles.', bold_hasta='Datos protegidos: ')

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Cotizaciones-Evaluacion-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
