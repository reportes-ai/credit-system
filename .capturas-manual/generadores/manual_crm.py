# -*- coding: utf-8 -*-
"""CRM y Campañas — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'CRM Y CAMPAÑAS', bold=True, color=AZUL_OSCURO, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Cada contacto con un cliente queda registrado, y cada campaña se mide de verdad.',
      size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Concepto ─────────────────────────────────────────────────────────
    h1(doc, '1. El concepto: una sola bitácora de gestiones')
    p(doc, 'El CRM del sistema no es un módulo aparte con su propia base de clientes: es la BITÁCORA '
           'ÚNICA donde cae toda interacción con un cliente, venga de donde venga. Una llamada de '
           'cobranza, un correo automático de mora, un WhatsApp de campaña, una gestión comercial — '
           'todo queda como gestión, con fecha, canal, resultado y autor.')
    caja(doc, 'Por qué importa',
         'Cuando un cliente llama, cualquiera que lo atienda ve la historia completa: qué se le ha '
         'enviado, quién lo llamó, qué se le prometió. Y cuando la ley pregunta cuántas gestiones de '
         'cobranza se hicieron esta semana, la respuesta está en un solo lugar — el tope legal se '
         'controla contra esta bitácora.')
    p(doc, 'Sobre esa bitácora se montan dos artillerías de contacto masivo, cada una con su '
           'módulo: las Campañas Masivas (correo y WhatsApp) y las Campañas de Venta telefónicas '
           '(con discador). Ambas escriben sus contactos como gestiones — nada queda fuera de la '
           'historia.')

    # ── 2. Campañas masivas ─────────────────────────────────────────────────
    h1(doc, '2. Campañas Masivas: correo y WhatsApp')
    p(doc, 'Contacto masivo con mensajes personalizados y — la diferencia con mandar correos a '
           'ciegas — MEDICIÓN contra grupo de control: saber si la campaña movió la aguja o si esas '
           'ventas habrían pasado igual.')
    ficha(doc,
          'Marketing/Comercial (venta) · Cobranza (mora)',
          'campanas_masivas',
          'Datos de contacto (BD o CSV) · para WhatsApp frío, plantilla HSM aprobada por Meta',
          'Campañas Masivas (correlativo CM-####)')
    h2(doc, '2.1 Armar la audiencia')
    vineta(doc, 'con plantilla descargable, hasta 10 campos de datos que después se usan como {{variables}} en el mensaje.', bold_hasta='CSV manual: ')
    vineta(doc, 'cobranza filtra por días de mora y capital; venta, por meses para el término del crédito y renta mínima — los clientes que están terminando de pagar son la mejor audiencia de renovación. Se puede excluir por región.', bold_hasta='Desde la base por parámetros: ')
    vineta(doc, 'dos análisis SEPARADOS y opcionales antes de enviar: el de CRÉDITO (política: relación cuota/renta → cumple o no) y el de INFORMES (riesgo DealerNet → excluir los altos). Se filtra ANTES de gastar el envío.', bold_hasta='Depurar con análisis: ')
    vineta(doc, 'si a un cliente le falta correo o teléfono, el sistema puede buscar sus datos de contacto en el informe comercial y proponerlos para revisión uno a uno — con el teléfono rankeado por frecuencia y los correos con el apellido del cliente primero.', bold_hasta='Completar contactos: ')
    h2(doc, '2.2 El grupo de control (champion–challenger)')
    regla(doc, 'Los deciles de control se definen al crear la campaña y esos clientes NO se '
               'contactan. Sin control no hay forma de saber si la campaña funcionó: la conversión '
               'del grupo contactado se compara contra la del grupo que no recibió nada, y la '
               'diferencia — el uplift — es el efecto real de la campaña.')
    h2(doc, '2.3 Enviar')
    vineta(doc, 'en lotes, con las variables mergeadas y plantilla visual (banner, logo, título). Cada correo lleva un píxel de lectura firmado: al abrirse, el estado pasa de ENVIADO a LEÍDO. La vista previa no lleva píxel, para no inventar lecturas.', bold_hasta='Correo: ')
    vineta(doc, 'dos modos — PLANTILLA APROBADA (llega a cualquier contacto, sin ventana de 24 horas, con los parámetros mergeados) o TEXTO LIBRE (solo conversaciones activas). Los teléfonos se normalizan solos.', bold_hasta='WhatsApp: ')
    vineta(doc, 'en pruebas, todo envío queda SIMULADO — registrado pero sin salir. Las campañas de prueba sembradas no se pueden enviar.', bold_hasta='Modo Desarrollo: ')
    h2(doc, '2.4 Medir')
    p(doc, 'El botón de recálculo cruza la campaña contra la realidad: en VENTA, los créditos '
           'OTORGADOS del RUT después del envío (ventana de 30 días); en COBRANZA, las cuotas '
           'pagadas después del envío. El informe muestra los estados clicables con export, la '
           'curva de conversión acumulada por día — campaña contra control — y el riesgo crediticio '
           'de los convertidos.')
    captura(doc, 'Campañas Masivas → resultados', 'El informe con la curva campaña vs control y el uplift.')
    flujo(doc, 'AUDIENCIA → ANÁLISIS (OPCIONAL) → DECILES DE CONTROL → ENVÍO POR LOTES → CONVERSIÓN MEDIDA')

    # ── 3. Campañas telefónicas ─────────────────────────────────────────────
    h1(doc, '3. Campañas de Venta telefónicas: el discador')
    p(doc, 'Para trabajar una base con llamadas — por ejemplo, la cartera de otra financiera para '
           'portabilidad — el módulo de Campañas de Venta arma un call center liviano dentro del '
           'sistema: la base se sube, se homologa y el discador reparte las llamadas.')
    ficha(doc,
          'Ejecutivos operan · el administrador de la campaña crea y configura',
          'campanas_ventas (operar) · campanas_ventas_admin (configurar)',
          'Base externa en Excel con RUT y al menos un teléfono',
          'Campañas de Venta (/campanas-ventas/)')
    h2(doc, '3.1 Bases externas, separadas a propósito')
    regla(doc, 'Una base externa NUNCA entra a las tablas normales del sistema: vive en tablas '
               'propias del módulo. Los clientes de otra financiera no son (todavía) nuestros '
               'clientes — mezclarlos contaminaría la fuente única.')
    p(doc, 'Al subir el Excel se homologa a mano: el sistema muestra las columnas del archivo con '
           'ejemplos y se mapea cada una contra el catálogo (cliente, crédito externo, vehículo, '
           'libres). Obligatorio: RUT y al menos un teléfono.')
    h2(doc, '3.2 El discador')
    vineta(doc, 'mi llamada en curso → mis rellamados vencidos → rellamados del pool → pendientes en el orden que fijó el administrador. Cada registro tomado queda bloqueado (20 minutos) para que dos ejecutivos no llamen al mismo cliente.', bold_hasta='La cola: ')
    vineta(doc, 'los teléfonos del cliente se ordenan por contactabilidad según las gestiones previas; un fono malo se marca INHABILITADO y deja de ofrecerse.', bold_hasta='Teléfonos que aprenden: ')
    vineta(doc, 'cada campaña define sus términos de gestión (contacto directo/indirecto/no contacto/inhabilitado), cuáles cierran el caso, cuáles son venta (piden monto) y cuáles agendan rellamado — todo paramétrico.', bold_hasta='Términos por campaña: ')
    captura(doc, 'Campañas de Venta → discador', 'El discador con un registro tomado, sus teléfonos rankeados y los términos de gestión.')
    h2(doc, '3.3 Las métricas del embudo')
    tabla(doc, ('Métrica', 'Fórmula'),
          (('Penetración', 'Gestionados ÷ base total'),
           ('Contactabilidad', 'Contactados ÷ gestionados'),
           ('Conversión', 'Ventas ÷ contactados'),
           ('Conversión real', 'Cruce por RUT contra créditos OTORGADOS después de activada la campaña')),
          (5.0, 11.5))
    p(doc, 'La última fila es la honesta: no basta con que el ejecutivo marque "venta" — el sistema '
           'cruza contra los créditos que efectivamente se cursaron.')

    # ── 4. Los otros brazos ─────────────────────────────────────────────────
    h1(doc, '4. Los otros brazos del contacto')
    p(doc, 'El CRM recibe gestiones de más lugares que las campañas:')
    vineta(doc, 'los correos por tramo y las secuencias de WhatsApp de mora escriben cada envío en la bitácora (ver Manual de Cobranza).', bold_hasta='Cobranza automática: ')
    vineta(doc, 'las conversaciones del bot y las derivaciones a ejecutivos quedan ligadas al cliente; las cotizaciones del bot van a la tabla de cotizaciones (fuente única).', bold_hasta='Facilito: ')
    vineta(doc, 'el motor que le pregunta al dealer por WhatsApp por las cartas que vencen mañana, y tabula el motivo cuando el negocio se perdió — inteligencia comercial que antes no existía.', bold_hasta='Seguimiento de cartas: ')
    vineta(doc, 'las visitas a dealers con su resultado y seguimiento (Ruta AutoFácil) son la versión en terreno de la misma idea: gestión registrada o gestión que no existió.', bold_hasta='Terreno: ')

    # ── 5. Gobierno ─────────────────────────────────────────────────────────
    h1(doc, '5. Gobierno')
    vineta(doc, 'los módulos tienen permisos separados para operar y para administrar; las campañas cerradas quedan como histórico.', bold_hasta='Permisos: ')
    vineta(doc, 'en cobranza, el tope semanal de gestiones es de ley y se controla contra la bitácora única — una campaña de cobranza no puede saltárselo.', bold_hasta='Límites legales: ')
    vineta(doc, 'el Modo Desarrollo simula todos los envíos; las plantillas de WhatsApp pasan por aprobación de Meta; el remitente del correo es el oficial del sistema.', bold_hasta='Canales protegidos: ')
    vineta(doc, 'el costo de los envíos de WhatsApp con plantilla y de los informes DealerNet del análisis previo se gobierna como toda integración (ver documento de Integraciones).', bold_hasta='Costos: ')

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\CRM-Campanas-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
