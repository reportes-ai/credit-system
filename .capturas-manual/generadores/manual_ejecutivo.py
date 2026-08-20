# -*- coding: utf-8 -*-
"""Manual del Ejecutivo Comercial — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DEL EJECUTIVO COMERCIAL', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Del patio del dealer a la comisión en tu liquidación:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'cotizar, evaluar, emitir la carta, cursar y cobrar.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Ejecutivo Comercial'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo sigue la misma estructura: para qué existe el proceso, quién lo hace, qué debe '
           'estar listo antes, el paso a paso pantalla por pantalla, y los recuadros de color: '
           '⚠ OJO (lo que confunde), 🔒 Regla del sistema (lo que no se puede saltar), 🧾 Caso real '
           '(el episodio que originó la regla) y 📸 CAPTURA (dónde va el pantallazo, pendiente en esta '
           'versión). Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. El circuito completo ─────────────────────────────────────────────
    h1(doc, '1. Tu circuito: de la cotización a la comisión')
    p(doc, 'Como ejecutivo comercial tu trabajo recorre el ciclo de vida completo de una operación. Este '
           'manual sigue ese orden:')
    flujo(doc, 'COTIZAR → ¿DÓNDE CURSO? → EVALUAR → CARTA → OTORGAR → FUNDANTES → COMISIÓN')
    vineta(doc, 'precio, pie y plazo se convierten en una cuota concreta para el cliente (cap. 2).', bold_hasta='Cotizar: ')
    vineta(doc, 'antes de armar la carta, saber en qué financiera conviene cursar (cap. 3).', bold_hasta='¿Dónde curso?: ')
    vineta(doc, 'el RUT del cliente pasa por antecedentes, informes y política (cap. 4).', bold_hasta='Evaluar: ')
    vineta(doc, 'el documento que compromete las condiciones al dealer, con firma y QR (cap. 5).', bold_hasta='La carta: ')
    vineta(doc, 'la carta aprobada se convierte en crédito, o se desiste (cap. 6).', bold_hasta='Otorgar: ')
    vineta(doc, 'los documentos que hacen que la financiera libere la plata (cap. 9).', bold_hasta='Fundantes: ')
    vineta(doc, 'tu comisión se calcula sola sobre lo que cursaste (cap. 13).', bold_hasta='La comisión: ')
    p(doc, 'Transversales a todo el circuito: las excepciones comerciales con su código (cap. 7), la '
           'regla del 60% para rescatar rechazos (cap. 8), la ficha del dealer (cap. 10), su base de '
           'vendedores (cap. 11) y el trabajo en terreno con la Ruta AutoFácil (cap. 12).')

    # ── 2. Cotizar ──────────────────────────────────────────────────────────
    h1(doc, '2. Cotizar')
    p(doc, 'La cotización convierte el precio del auto en una cuota con gastos, seguros y CAE. Hay '
           'cuatro puertas de entrada y todas usan el mismo motor de cálculo, así que la cuota es la '
           'misma se cotice donde se cotice.')
    ficha(doc,
          'Ejecutivo comercial · el dealer desde su portal · el cliente vía Facilito (WhatsApp)',
          'Acceso al módulo Cotizaciones',
          'Tasas y parámetros vigentes · UF del día',
          'Cotizaciones · Simulador Rápido · portal del dealer · Facilito')
    paso(doc, 1, 'Ingresar la operación', 'Precio del vehículo, pie, plazo y seguros. El saldo a '
         'financiar se calcula solo.')
    paso(doc, 2, 'Leer la cuota', 'El sistema entrega la cuota con gastos, seguros y CAE.')
    paso(doc, 3, 'Guardar', 'La cotización queda guardada al RUT del cliente: cuando el negocio avance, '
         'la evaluación y la carta parten de ahí.')
    captura(doc, 'Cotizaciones → nueva cotización', 'El formulario con una cotización calculada: cuota, gastos y CAE visibles.')

    # ── 3. Dónde curso ──────────────────────────────────────────────────────
    h1(doc, '3. ¿Dónde curso? — la elección de financiera')
    p(doc, 'Que sepas en el patio, antes de armar la carta, en qué financiera conviene cursar. Antes '
           'esto se decidía de memoria o mirando el cuadro impreso, y el informe de rentabilidad recién '
           'lo desmentía a fin de mes.')
    ficha(doc,
          'Ejecutivo comercial consulta · Gerencia define los umbrales',
          'donde_curso_ver',
          'Tasas vigentes, parámetros del crédito, UF del día y cuadro de Preferencia Financiera. '
          'Entrar una vez con conexión: la app guarda los parámetros en el teléfono',
          'PWA /donde-curso/ (instalable en el celular)')
    paso(doc, 1, 'Instalar', 'Se abre en el celular y se agrega a la pantalla de inicio. Login con la '
         'misma cuenta del sistema.')
    paso(doc, 2, 'Ingresar la operación', 'Valor del vehículo, pie, plazo, tasa (o automática), seguros '
         'incluidos y si el dealer es de patio o de calle.')
    paso(doc, 3, 'Leer el veredicto', 'AUTOFIN, UNIDAD DE CRÉDITO o DECIDES TÚ, con el porqué debajo: '
         'monto financiado, cuota, tasa aplicada y la rentabilidad de cada financiera.')
    captura(doc, 'PWA ¿Dónde Curso?', 'El veredicto con la comparación de rentabilidad entre financieras.')
    h2(doc, '3.1 Cómo decide')
    p(doc, 'Primero elegibilidad: el cuadro de Preferencia Financiera dice qué financiera acepta ese '
           'plazo con ese saldo — la que no puede tomarla no compite, por muy rentable que se vea. '
           'Después rentabilidad: gana la más alta, salvo que la diferencia sea menor al umbral (10% '
           'por defecto), en cuyo caso sale DECIDES TÚ: para la empresa da lo mismo y la decisión es '
           'tuya — sirve para privilegiar la relación con el dealer o la velocidad de curse.')
    advertencia(doc, 'Si aparece el aviso amarillo, la operación deja menos que la rentabilidad mínima '
                     '($100.000 por defecto). Se puede cursar igual, pero conviene revisarla antes.')
    regla(doc, 'El mismo veredicto vive dentro de la carta, bajo el selector de financiera, y define la '
               'excepción: cursar en la financiera conveniente NO es excepción, y con veredicto DECIDES '
               'TÚ tampoco lo es cursar en cualquiera de las dos. Solo se registra excepción de '
               'financiera cuando se elige la MENOS conveniente.')

    # ── 4. Evaluación ───────────────────────────────────────────────────────
    h1(doc, '4. Evaluación crediticia')
    p(doc, 'El RUT del cliente pasa por antecedentes, informes comerciales y la política de crédito. El '
           'resultado es el análisis con que el analista (o el sistema) decide.')
    ficha(doc,
          'Ejecutivo comercial arma la evaluación · Analista de Crédito decide',
          'Acceso a Evaluación Crediticia',
          'Cliente en la base (o se crea) · saldo DealerNet para informes',
          'Evaluación Crediticia')
    paso(doc, 1, 'Buscar por RUT', 'Trae la ficha completa del cliente: datos personales, antecedentes '
         'laborales e información comercial.')
    paso(doc, 2, 'Actualizar antecedentes', 'Si tienen más de 30 días, se cargan nuevos.')
    paso(doc, 3, 'Leer el análisis', 'La IA cruza renta, informes DealerNet y política de crédito, y '
         'entrega el análisis para la decisión.')
    captura(doc, 'Evaluación Crediticia', 'Una evaluación con el análisis de IA y los informes a la vista.')
    advertencia(doc, 'Un informe DealerNet con sello rojo "SIN INFORMACIÓN — Intente más tarde" llegó '
                     'VACÍO: no significa que el cliente no tenga deuda y no sirve para concluir. No '
                     'ocupa el bloqueo de 15 días — se puede volver a pedir de inmediato.')
    advertencia(doc, 'La deuda que informa DealerNet para el producto crediticio viene en MILES de '
                     'pesos. El sistema ya la convierte con su motor; si comparas a mano contra el '
                     'documento original, recuerda el factor.')

    # ── 5. La carta ─────────────────────────────────────────────────────────
    h1(doc, '5. La Carta de Aprobación')
    p(doc, 'La carta es el documento que compromete las condiciones al dealer: monto, tasa, cuotas y '
           'comisión. Sale con correlativo, código de verificación QR y firma electrónica — por eso, '
           'una vez emitida, no se edita: se corrige por reemplazo (ver 5.4).')
    ficha(doc,
          'Ejecutivo comercial emite · el pool de analistas (o Business Suite) aprueba',
          'Emisión de cartas · Rentabilidad AutoFin vs UAC requiere permiso especial',
          'Evaluación hecha · dealer con ficha vigente · vendedor en la base del dealer',
          'Cartas de Aprobación → generador')
    h2(doc, '5.1 Emitir')
    paso(doc, 1, 'Partir del PDF de la financiera', 'El generador autocompleta la carta desde el '
         'documento: monto, tasa, plazo, cuota. Revisar contra el documento siempre — el autofill '
         'sugiere, tú respondes por lo emitido.')
    paso(doc, 2, 'Elegir dealer y vendedor', 'El vendedor se elige de la base del dealer (cap. 11). Es '
         'el dato con que después se mide quién trae cada negocio: elegir bien acá evita el "S/I" que '
         'deja la venta sin dueño.')
    paso(doc, 3, 'Respetar la Preferencia Financiera', 'El sistema bloquea combinaciones fuera de '
         'política y muestra el veredicto de ¿Dónde Curso? bajo el selector de financiera.')
    paso(doc, 4, 'Adjuntar los documentos', 'UNIDAD: Carta Compromiso + Cotización (PDF). AUTOFIN: '
         'Carta de Aprobación (PDF) + pantallazo del sistema AutoFin con la solicitud puntual y su '
         'estado — botón "Capturar pantalla" o pegar con Ctrl+V. Si el pantallazo es un listado o no '
         'se ve el estado, el sistema te avisa al tiro.')
    paso(doc, 5, 'Enviar a revisión', 'La carta queda PENDIENTE en el pool de analistas — o la revisa '
         'Business Suite en el momento si el Revisor Automático está encendido (ver 5.2).')
    captura(doc, 'Cartas de Aprobación → generador', 'El formulario completo con los datos autocompletados y los documentos adjuntos.')
    h2(doc, '5.2 La revisión automática')
    p(doc, 'Si el motor está encendido para tu financiera, Business Suite compara el documento contra '
           'lo digitado (ID, RUT, saldo, plazo, monto, tasa), valida la comisión del dealer contra el '
           'motor y el código de excepción si lo hay. Todo cuadra → carta APROBADA al instante con '
           'checklist firmado, y te llega la campanita. Algo no cuadra → queda con el analista, con un '
           'banner que dice exactamente qué falló — corrígelo antes de preguntar.')
    regla(doc, 'Para AUTOFIN, el pantallazo debe mostrar la solicitud en Revisión Firma o Cursado: '
               'ningún otro estado emite carta. Y su ID debe calzar con la carta.')
    h2(doc, '5.3 La fecha de primera cuota')
    p(doc, 'Al subir la carta AutoFin, el sistema lee el CUADRO DE PAGO del PDF y escribe la primera '
           'fecha de vencimiento en el crédito — solo si estaba vacía. Si la carta viene escaneada o '
           'sin cuadro legible, la operación caerá a Digitación de Datos Faltantes y alguien la '
           'digitará del pagaré. Una fecha inventada corrompe el cálculo de mora: nunca la adivines.')
    h2(doc, '5.4 Corregir una carta emitida')
    p(doc, 'Una carta emitida no se edita: se emite una carta nueva con sufijo -C1, -C2… y la anterior '
           'queda REEMPLAZADA (su QR y firma pasan a no vigentes). Monto, saldo, tasa y cuotas no se '
           'pueden corregir por esta vía — si eso cambió, es otra operación: se anula y se emite carta '
           'nueva. La corrección la hace Administración/Operaciones; pídela con el motivo claro.')

    # ── 6. Cartas vigentes ──────────────────────────────────────────────────
    h1(doc, '6. Cartas Vigentes: otorgar o desistir')
    p(doc, 'La carta aprobada vive 5 días (configurable). En ese plazo el negocio se concreta '
           '(OTORGAR) o se cae (DESISTIR). Si nadie hace nada, el desistimiento automático la vence.')
    ficha(doc,
          'Ejecutivo comercial / Operaciones',
          'Otorgamiento de cartas',
          'Carta APROBADA vigente',
          'Cartas de Aprobación → Cartas Vigentes')
    paso(doc, 1, 'Otorgar', 'Crea el crédito en etapa OTORGADO con todo congelado: calendario, tasa y '
         'la comisión de la carta — la carta manda. En el mismo acto nace el movimiento de cartola del '
         'dealer.')
    paso(doc, 2, 'Desistir', 'Deja la carta DESISTIDA con motivo. La operación queda en la estadística '
         'de conversión.')
    captura(doc, 'Cartas Vigentes', 'La bandeja con una carta por vencer, mostrando los días restantes.')
    advertencia(doc, 'Vencida la carta, no revive: se emite una nueva. Por eso conviene otorgar apenas '
                     'el curse se confirma — una carta otorgada tarde deja la comisión para la cartola '
                     'del mes siguiente.')

    # ── 7. Excepciones ──────────────────────────────────────────────────────
    h1(doc, '7. Excepciones Comerciales: tu código de autoaprobación')
    p(doc, 'Cuando el negocio necesita salirse de la pizarra —bajar la tasa, bajar la comisión del '
           'dealer— no partes pidiendo permiso: simulas la jugada y, si respeta el piso de '
           'rentabilidad, el sistema mismo te entrega un código que aprueba la excepción. Autoservicio '
           'con presupuesto y auditoría.')
    ficha(doc,
          'Ejecutivo Comercial y Supervisor · el código de Gerencia solo el Gerente de Operaciones o '
          'su backup',
          'Acceso al Simulador de Excepciones',
          'Parámetros configurados (piso 75%, rebaja máxima 5% relativo, vigencia 24 h) · tasas y UF '
          'vigentes',
          '/excepciones/ — el simulador oficial vive en ¿Dónde Curso? 2.0')
    h2(doc, '7.1 Simular la jugada')
    paso(doc, 1, 'Armar el negocio', 'Precio + pie entregan saldo, % pie, % a financiar y la grilla '
         '12/24/36/48 con cuota aproximada, dónde conviene cursar y la rentabilidad de cada lado. '
         'Desmarcar un seguro recalcula todo.')
    paso(doc, 2, 'Jugar la excepción', 'Bajar tasa (máximo 5% relativo sobre la pizarra — más allá la '
         'operación se fuerza a AutoFin) o bajar la comisión del dealer marcando calle/parque.')
    paso(doc, 3, 'Respetar el piso', 'La excepción nunca puede rentar menos del 75% de la alternativa '
         'más rentable. Si te pasas, el sistema propone la rebaja de comisión dealer exacta que la '
         'devuelve al piso.')
    captura(doc, 'Simulador de Excepciones', 'Una jugada simulada con la rentabilidad de cada financiera y el piso marcado.')
    h2(doc, '7.2 La escalera de estrellas')
    tabla(doc, ('Nivel', 'Aprueba jugadas que rentan', 'Condición extra'),
          (('1 ⭐', '≥ 75% de la mejor alternativa', '—'),
           ('2 ⭐', '≥ 50% de la mejor alternativa', '—'),
           ('3 ⭐', '≥ 25% de la mejor alternativa', 'Y además ≥ la comisión más alta (dealer o ejecutivo), nunca bajo $200.000'),
           ('👔 Gerencia', 'Fuera de escalera', 'Motivo obligatorio; transferible')),
          (2.6, 7.0, 6.9))
    p(doc, 'Tu presupuesto: 3 estrellas de regalo el primer mes; después, ⌊otorgadas del mes anterior '
           '× 33%⌋. El código vence a las 24 horas — las estrellas vuelven si vence sin usar; al '
           'cierre del mes expiran. El nivel lo decide el SERVIDOR según la rentabilidad real: '
           'declarar un nivel más barato no sirve, se cobran las estrellas del nivel real.')
    h2(doc, '7.3 Usar el código en la carta')
    regla(doc, 'Un código = un solo uso, del ejecutivo dueño (no sirve en la carta de un colega; los '
               'de Gerencia sí son transferibles). El sistema valida saldo precio, primeros dígitos '
               'del RUT y vigencia. La carta queda "aprobada por código del sistema".')
    p(doc, 'Trazabilidad: pestaña "Mis códigos" con los tuyos; el Validador dice de qué operación es '
           'cualquier código, quién lo generó y si ya se usó.')
    flujo(doc, 'SIMULAR → CUMPLE PISO → CÓDIGO → CARTA CON CÓDIGO → USADO')

    # ── 8. Regla 60% ────────────────────────────────────────────────────────
    h1(doc, '8. Rescatar un rechazo: la regla del 60%')
    p(doc, 'Si una financiera rechaza, puedes cursar en la otra en las MISMAS condiciones ofrecidas al '
           'cliente — sin renegociar — siempre que a AutoFácil le quede al menos el 60% del ingreso '
           'total. Vigente desde el 11-08-2026.')
    regla(doc, 'Ingreso Total − Comisión Dealer − Comisión Parque − Comisión Ejecutivo ≥ 60% del '
               'Ingreso Total. Si no llega, solo puede rebajarse la comisión del DEALER hasta '
               'alcanzarlo — parque y ejecutivo no se tocan. Bajo el 60% no se aprueba; no hay '
               'excepción de la excepción.')
    paso(doc, 1, 'Simular', 'En el Simulador de Rentabilidad, con la comisión de dealer que piensas '
         'ofrecer. La pantalla muestra la distribución del ingreso (AutoFácil · dealer · parque · '
         'ejecutivo) y puedes editar la comisión del dealer para ver al instante si cumple.')
    paso(doc, 2, 'Autorizar', 'El Jefe Comercial envía el correo con las condiciones aprobadas.')
    paso(doc, 3, 'Registrar', 'El Analista de Crédito aprueba la excepción en el sistema — ese registro '
         'es el respaldo formal. Fuera del horario de los analistas, autoriza el Jefe Comercial.')
    flujo(doc, 'RECHAZADA → SIMULAR EN LA OTRA → ¿AUTOFÁCIL ≥ 60%? → APROBADA / NO SE APRUEBA')

    # ── 9. Fundantes ────────────────────────────────────────────────────────
    h1(doc, '9. Fundantes: sin documentos no hay plata')
    p(doc, 'Los fundantes respaldan la operación otorgada para que la financiera libere los fondos. '
           'Subirlos a tiempo es la parte tuya; validarlos, la de Operaciones. Cada día de demora es '
           'un día en que ni el dealer recibe su saldo ni la empresa su comisión.')
    ficha(doc,
          'Tú subes · Analista de Operaciones valida',
          'Seguimiento Fundantes — vista Ejecutivo',
          'Crédito OTORGADO · matriz de documentos de la financiera',
          'Seguimiento Fundantes')
    paso(doc, 1, 'Subir', 'Los documentos que exige la matriz de la financiera. El sistema los renombra '
         'solo (tipo + operación). Al completar, Enviar.')
    paso(doc, 2, 'Corregir si rebota', 'Un RECHAZADO interno vuelve con comentario del analista: '
         'corregir y reenviar. Una DEVOLUCIÓN de la financiera (la operación ya estaba cerrada) vuelve '
         'a tus pendientes como POR CORREGIR, con campanita.')
    paso(doc, 3, 'Registrar la gestión', 'En la bitácora de la operación (pinchar el N° OP) queda la '
         'línea de tiempo completa y tus comentarios de gestión. Ahí se lee por qué se demoró cada '
         'operación — escribe ahí, no en correos.')
    captura(doc, 'Seguimiento Fundantes → vista Ejecutivo', 'La matriz de una operación con documentos subidos y uno rechazado con comentario.')
    advertencia(doc, 'Si tienes fundantes PENDIENTES, una vez por semana el sistema te abre un pop-up '
                     'bloqueante con tu operación más antigua: no se cierra sin un comentario de al '
                     'menos 3 palabras, que va a la bitácora. No es castigo — es la rendición semanal '
                     'que antes se pedía por correo.')

    # ── 10. Dealer ──────────────────────────────────────────────────────────
    h1(doc, '10. Incorporar (o modificar) un dealer')
    p(doc, 'Ningún dealer opera sin ficha aprobada: sus comisiones, banco e informes pasan por la '
           'cadena de autorización. Tú armas la ficha; la cadena la revisa; el cliente la firma.')
    ficha(doc,
          'Tú (ficha) · Análisis (nivel 1) · Gerencia (nivel 2, solo sobre pizarra) · el cliente firma',
          'Incorporación de Dealers',
          'Pizarra de comisiones vigente · saldo DealerNet para los informes',
          'Incorporación de Dealers → Nueva')
    paso(doc, 1, 'Armar la ficha', 'Datos, banco, tabla de comisiones (por defecto la pizarra) y '
         'socios. Si el RUT ya existe, entra en modo Modificación con las diferencias en rojo.')
    paso(doc, 2, 'Enviar a autorización', 'El sistema pide solo los informes DealerNet (empresa + '
         'socios) y el análisis de IA. Antecedentes penales disparan alerta grave.')
    paso(doc, 3, 'Esperar la cadena', 'Los niveles autorizan ANTES de la firma. Aprobar sin abrir la '
         'revisión queda marcado.')
    paso(doc, 4, 'Imprimir, firmar y cerrar', 'Se imprime solo la ficha autorizada; firma el cliente; '
         'se sube firmada JUNTO con las cédulas de identidad de los firmantes (hasta 6 archivos — sin '
         'cédulas el sistema no acepta). Al cerrar, el dealer queda creado y su tabla de comisiones '
         'rige desde ese momento.')
    captura(doc, 'Incorporación de Dealers', 'Una ficha en la cadena de autorización con sus niveles.')
    flujo(doc, 'BORRADOR → PEND. AUTORIZACIÓN → AUTORIZADA → FIRMADA → APROBADA / RECHAZADA')
    p(doc, 'Los parques se incorporan por esta misma máquina (Creación de Parque → "Nuevo parque"); '
           'la diferencia es el final: al cerrar se crea el parque con arriendo y comisión en 0, que '
           'se fijan en el mantenedor Arriendos y Comisiones.')
    advertencia(doc, 'Dealer "inactivo" NO significa bloqueado: significa 3 meses sin cursar. No '
                     'habilita ni impide operar — es un estado comercial, no un candado.')

    # ── 11. Vendedores ──────────────────────────────────────────────────────
    h1(doc, '11. Los vendedores del dealer')
    p(doc, 'El vendedor del concesionario es con quien trabajas el día a día. La base de vendedores es '
           'la única fuente de su RUT y su correo, y el selector de la carta sale de ahí: si el '
           'vendedor no está, se agrega ANTES de emitir la carta.')
    vineta(doc, 'Creación/Mantenedor de Dealer → Vendedores por Dealer: nombre, RUT, correo y dealer. Alimenta el selector del generador de cartas.', bold_hasta='Mantener la base: ')
    vineta(doc, 'elegir el vendedor real en cada carta. "S/I" es ausencia de dato — deja la venta sin dueño y ensucia la estadística del dealer.', bold_hasta='Tu parte: ')
    vineta(doc, 'Post Venta → Vendedores con Ventas muestra el aporte de cada vendedor (abre en el último mes cerrado); Cartas → Detalle Mensual lista cada carta con su vendedor, RUT y correo.', bold_hasta='Dónde se mide: ')
    captura(doc, 'Vendedores por Dealer', 'La base de vendedores de un dealer con sus datos de contacto.')

    # ── 12. Ruta AutoFácil ──────────────────────────────────────────────────
    h1(doc, '12. Ruta AutoFácil: las visitas a dealers')
    p(doc, 'La relación con el dealer se cultiva en el patio, no por teléfono. La Ruta AutoFácil '
           'ordena ese trabajo: a quién visitar, cuándo, con qué frecuencia, y qué salió de cada '
           'visita — con la evidencia GPS de que la visita ocurrió.')
    ficha(doc,
          'Ejecutivos comerciales gestionan lo propio · Supervisores y Gerencia ven todo, configuran '
          'y asignan cartera',
          'visitas_dealers (lo propio) · visitas_supervisar (todo + configuración)',
          'Dealers geocodificados (Mapa de Dealers) · configuración de visitas por día y días '
          'habilitados',
          'Dealers → RUTA AUTOFACIL · PWA Terreno (/terreno/)')
    h2(doc, '12.1 El calendario y la bitácora')
    paso(doc, 1, 'Agendar', 'En el calendario mensual se agenda un dealer en una fecha. El sistema '
         'respeta los días habilitados y el tope de visitas por día de cada usuario, y muestra el '
         'cupo usado (X/N).')
    paso(doc, 2, 'Realizar y registrar', 'Tras la visita se registra el resultado — POSITIVO o '
         'NEGATIVO — con comentarios, y opcionalmente un seguimiento con fecha y nota.')
    paso(doc, 3, 'Supervisar', 'Quien supervisa ve el calendario y las gestiones de todo el equipo, '
         'con estadísticas de cumplimiento y positividad por ejecutivo. El resto ve solo lo suyo.')
    captura(doc, 'Dealers → Ruta AutoFácil → Calendario', 'El calendario mensual con visitas agendadas y el cupo del día.')
    h2(doc, '12.2 El planificador de rutas')
    p(doc, 'La pestaña Planificador arma el calendario óptimo: eliges visitas por día, ventana '
           'horaria, duración por visita, filtros (activos / en riesgo / inactivos, comuna, radio '
           'máximo) y punto de partida, y el motor propone la ruta día a día minimizando traslados. '
           'Es una simulación — nada se graba hasta apretar Agendar — y muestra la cobertura: qué '
           'porcentaje de la cartera cubres en 30 días y cuánto tomaría llegar al 100%.')
    vineta(doc, 'un dealer ACTIVO sin créditos cursados en 90 días. Es el filtro más útil del planificador: ahí está la venta que se está enfriando.', bold_hasta='"En riesgo" = ')
    vineta(doc, 'cada dealer del listado muestra cuándo fue la última vez que cursó un crédito — el dato con que se abre la conversación en el patio.', bold_hasta='Última venta: ')
    h2(doc, '12.3 Asignación de cartera (supervisores)')
    p(doc, 'La Asignación Masiva reparte los dealers sin asignar entre los ejecutivos, balanceada por '
           'bloques de comuna (con zonificación de las comunas grandes), filtrando por tipo '
           '(calle/parques/ambos), y agenda automáticamente las visitas del plan entre la fecha de '
           'inicio y el cierre. Regla de oro: un dealer tiene UNA sola asignación activa. Cada '
           'ejecutivo recibe la campanita con su cartera.')
    h2(doc, '12.4 La PWA Terreno: el día en la calle')
    paso(doc, 1, 'Instalar', 'La app "Terreno" se instala desde /terreno/ en el teléfono. Login con '
         'la misma cuenta.')
    paso(doc, 2, 'Mi día', 'La ruta del día en orden de cercanía, con el mapa de marcadores '
         'numerados, el anillo de progreso y tu ubicación.')
    paso(doc, 3, 'Check-in GPS', 'Al llegar donde el dealer, check-in: queda la hora y la distancia '
         'real al dealer. Es la evidencia de la visita.')
    paso(doc, 4, 'Fotos y resultado', 'Hasta 6 fotos por visita (se comprimen solas en el teléfono) '
         'y el registro 👍/👎 con comentarios — el mismo motor del calendario, no un registro aparte.')
    captura(doc, 'PWA Terreno → Mi día', 'La ruta del día con el mapa, marcadores numerados y el anillo de progreso.')
    advertencia(doc, 'La app aguanta quedarse sin señal (guarda el último día y muestra el aviso '
                     '"sin conexión"), pero el check-in y las fotos necesitan conexión para subir. '
                     'Si estás sin señal, registra apenas la recuperes.')
    flujo(doc, 'ASIGNACIÓN → CALENDARIO → VISITA + CHECK-IN GPS → RESULTADO 👍/👎 → SEGUIMIENTO')

    # ── 13. Comisión ────────────────────────────────────────────────────────
    h1(doc, '13. Tu comisión')
    p(doc, 'La comisión se calcula sola, mes a mes: piso + incentivo + ajustes por cumplimiento y '
           'calidad, sobre tus operaciones del mes. Tu supervisor la revisa y aprueba en Comisiones → '
           'Revisión, y de ahí va directo a tu liquidación — sin digitación intermedia.')
    h2(doc, '13.1 Los descuentos por prepago y anulación')
    regla(doc, 'Cláusula novena del anexo, aplicada automática: si una operación ya comisionada se '
               'prepaga dentro de los 3 meses desde el vencimiento de su primera cuota, se devuelve el '
               '100% de la comisión pagada por ella; hasta los 6 meses, el 50%; después, nada. Las '
               'anuladas devuelven el 100% siempre. La reversa se imputa al mes siguiente al del hecho.')
    p(doc, 'Tu liquidación muestra solo la comisión neta a pagar. El detalle de qué operaciones se '
           'revirtieron —con la glosa "Reversa comisión pagada OP… Prepago/Anulación"— vive en el '
           'detalle de cálculo de Revisión de Comisiones y en el informe que te llega. Ningún mes '
           'pagado se reabre, y la comisión nunca queda en negativo: si el descuento la supera, el '
           'saldo queda informado en rojo.')
    captura(doc, 'Comisiones → Revisión (detalle de un ejecutivo)', 'El detalle de cálculo de un mes con piso, incentivo y una reversa visible.')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('La carta no se aprueba sola', 'El Revisor derivó por una diferencia, o el motor está apagado', 'Leer el banner rojo: dice qué no cuadró (cap. 5.2)'),
           ('El código de excepción no entra en la carta', 'Código de otro ejecutivo, vencido (24 h) o ya usado', 'Validador de códigos (cap. 7.3)'),
           ('El sistema no acepta el pantallazo AutoFin', 'Es un listado o no se ve el estado de la solicitud', 'Capturar la solicitud puntual con su estado (cap. 5.1)'),
           ('Mi venta aparece sin vendedor ("S/I")', 'No se eligió vendedor en la carta', 'Agregarlo a la base y elegirlo siempre (cap. 11)'),
           ('El pop-up semanal no me deja trabajar', 'Tienes fundantes pendientes', 'Comentar la gestión real (3+ palabras) y subir los documentos (cap. 9)'),
           ('La operación quedó fuera de la cartola del mes', 'Se otorgó tarde o los fundantes se demoraron', 'Entra sola a la cartola siguiente (cap. 6)'),
           ('DealerNet dice "SIN INFORMACIÓN"', 'El informe llegó vacío', 'Repedirlo de inmediato; NO concluir "sin deuda" (cap. 4)'),
           ('No puedo agendar una visita', 'Día no habilitado o tope diario lleno', 'Ver la configuración de la Ruta (cap. 12)'),
           ('Un dealer no aparece en el planificador', 'Sin coordenadas (no geocodificado)', 'Geocodificarlo en el Mapa de Dealers (cap. 12)')),
          (5.3, 5.9, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('N° OP', 'Correlativo único AutoFácil (AAMM####) de la operación'),
           ('ID Financiera', 'Número de la operación en el sistema de la institución'),
           ('Carta de aprobación', 'Documento verificable (QR + firma electrónica) que compromete las condiciones'),
           ('Pizarra', 'Tabla de comisión por defecto; rige si el dealer no tiene tabla pactada'),
           ('Tabla pactada', 'Comisión negociada con un dealer específico; manda sobre la pizarra'),
           ('Fundantes', 'Documentos que respaldan el curse para que la financiera libere fondos'),
           ('Cartola', 'Estado de cuenta mensual de comisiones del dealer'),
           ('Código de excepción', 'Autorización de una jugada fuera de pizarra, generada por el simulador'),
           ('Estrellas', 'Presupuesto mensual de excepciones del ejecutivo'),
           ('DECIDES TÚ', 'Veredicto de financiera cuando la diferencia de rentabilidad es menor al umbral'),
           ('Parque / calle', 'Ubicación comercial del dealer; cambia el tramo de comisión'),
           ('Facilito', 'El bot de WhatsApp que cotiza y origina puntas')),
          (4.3, 12.2))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-Ejecutivo-Comercial-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
