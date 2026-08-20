# -*- coding: utf-8 -*-
"""Inteligencia Artificial en Business Suite — dónde se usa, beneficios y costos."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'INTELIGENCIA ARTIFICIAL', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'en Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Dónde se usa, cómo se gobierna, qué aporta y cuánto cuesta.', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026 · cifras de uso reales del sistema', size=11,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Filosofía ─────────────────────────────────────────────────────────
    h1(doc, '1. La filosofía: la IA propone, el sistema calcula, la persona decide')
    p(doc, 'Business Suite usa la IA de Anthropic (los modelos Claude) como un empleado más del '
           'equipo: uno que lee rápido, redacta bien y no se aburre de lo repetitivo — pero que '
           'NUNCA tiene la última palabra en lo que importa. Tres reglas fijas gobiernan cada uso:')
    vineta(doc, 'las cifras salen SIEMPRE de los motores del sistema. Facilito cotiza con el cotizador real; el veredicto de una preevaluación lo decide el código con los informes — la IA solo redacta el mensaje. Una cuota, una tasa o un veredicto nunca son "opinión" del modelo.', bold_hasta='La IA no calcula plata: ')
    vineta(doc, 'donde el resultado tiene consecuencias (indicadores previsionales, políticas, aprobaciones), la IA propone y una persona aplica. El Revisor de Cartas es motor determinístico; la IA solo lee documentos que el parser no puede.', bold_hasta='La IA no aplica sola: ')
    vineta(doc, 'cada llamada queda registrada con su funcionalidad, modelo, tokens y costo en dólares. El gasto se mira en vivo, no se estima.', bold_hasta='Todo uso se mide: ')

    # ── 2. Gobernanza ───────────────────────────────────────────────────────
    h1(doc, '2. Cómo se gobierna: el Subsistema de IA')
    p(doc, 'Toda la IA del sistema pasa por un núcleo único (mantenedor Subsistema IA), con estas '
           'propiedades:')
    vineta(doc, 'un interruptor global apaga TODA la IA del sistema de una vez. Además, cada funcionalidad tiene el suyo propio y NACE APAGADA: nada usa IA hasta que el Administrador lo enciende.', bold_hasta='Switch maestro + switch por funcionalidad: ')
    vineta(doc, 'cada funcionalidad usa el modelo que se le asigne en el mantenedor — el barato y rápido (Haiku) para leer documentos y conversar, los mayores (Sonnet/Opus) para análisis financiero y BI. Cambiar el modelo de una funcionalidad es un clic, no un desarrollo.', bold_hasta='Modelo por funcionalidad: ')
    vineta(doc, 'la tabla de precios por modelo (US$ por millón de tokens) es paramétrica: si Anthropic cambia precios, se actualiza la tabla y el costo registrado sigue siendo exacto.', bold_hasta='Precios paramétricos: ')
    vineta(doc, 'las funcionalidades de cara al público tienen límites: el bot tiene tope de informes por conversación y por día; el asistente del portal dealer tiene cuota diaria por dealer.', bold_hasta='Cuotas anti-abuso: ')
    vineta(doc, 'a la IA se le pasa lo mínimo: el asistente del dealer recibe SOLO los datos de ese dealer ya filtrados (no genera consultas); el informe de clima recibe SOLO agregados anónimos, nunca respuestas individuales.', bold_hasta='Datos acotados: ')
    captura(doc, 'Mantenedores → Subsistema IA', 'El mantenedor con las funcionalidades, sus switches y el modelo asignado.')

    # ── 3. Dónde se usa ─────────────────────────────────────────────────────
    h1(doc, '3. Dónde se usa, funcionalidad por funcionalidad')
    p(doc, 'El registro del sistema tiene 24 funcionalidades de IA. Las encendidas hoy:')
    h2(doc, '3.1 Crédito y riesgo')
    tabla(doc, ('Funcionalidad', 'Qué hace', 'Modelo'),
          (('Análisis de informe crediticio', 'Cruza el informe DealerNet del RUT con renta y política, y redacta el análisis que apoya la decisión del analista. También corre solo tras cada preevaluación del bot (Reporte Crediticio Automático)', 'Haiku'),
           ('Evaluación de consistencia / scoring', 'Revisa la coherencia del caso completo (renta declarada vs informes vs política) para el scoring', 'Sonnet'),
           ('Lectura IA de PDFs escaneados (cartas)', 'Cuando el texto del PDF de la financiera no se puede parsear, la IA lee el documento y extrae los datos que el Revisor valida. Igual con los pantallazos de AutoFin', 'Haiku'),
           ('Análisis de liquidaciones de sueldo', 'Extrae líquido e imponible de la liquidación que sube el cliente, con las reglas previsionales paramétricas (AFP, salud, IUSC)', 'Haiku')),
          (4.4, 9.3, 2.8))
    h2(doc, '3.2 Canales y clientes')
    tabla(doc, ('Funcionalidad', 'Qué hace', 'Modelo'),
          (('Bot WhatsApp (Facilito)', 'Conversa con el guion de venta, junta los datos y deriva. Las cifras las calcula el código; la IA redacta. ~$3 pesos por conversación', 'Haiku'),
           ('Revisor de plantillas HSM', 'Revisa las plantillas de WhatsApp antes de mandarlas a aprobación de Meta', 'Sonnet'),
           ('Moderador de Facilbook', 'Modera las publicaciones de la red social interna', 'Haiku')),
          (4.4, 9.3, 2.8))
    h2(doc, '3.3 Finanzas y gerencia')
    tabla(doc, ('Funcionalidad', 'Qué hace', 'Modelo'),
          (('Pregúntale a AutoFácil (BI)', 'BI conversacional: preguntas de negocio en lenguaje natural sobre los datos del sistema', 'Sonnet'),
           ('Pregúntale a Finanzas (BI)', 'La versión financiera-contable, para análisis de estados y cifras', 'Opus'),
           ('Análisis de Cierre (Bitácora)', 'El análisis del cierre contable mensual que acompaña la bitácora', 'default'),
           ('Comentarios al Cierre Mensual', 'Redacta los comentarios del informe de cierre a la matriz', 'default'),
           ('Hechos Relevantes Directorio', 'Propone los hechos relevantes del período para el directorio', 'default'),
           ('Asistente de Asientos Contables', 'Ayuda al contador a armar un asiento manual complejo', 'default'),
           ('Resumen Ejecutivo Diario', 'El resumen del día que llega por correo a gerencia', 'default'),
           ('Indicadores Previred', 'Lee los indicadores previsionales publicados y PROPONE la actualización; un humano revisa y aplica', 'default')),
          (4.4, 9.3, 2.8))
    h2(doc, '3.4 Apagadas hoy (listas para encender)')
    p(doc, 'Registradas y construidas, esperando activación cuando el negocio lo pida: validador '
           'inteligente de carga masiva, análisis de carpeta tributaria (SII) y F22, copiloto de '
           'cobranza, resumen de gestiones CRM, asistente IA del Portal Dealer, informe de encuestas '
           'de clima (Opus), informe del Test de Kuder y revisión de firmas. Encenderlas es un clic '
           'en el mantenedor — sin desarrollo.')

    # ── 4. Beneficios ───────────────────────────────────────────────────────
    h1(doc, '4. Los beneficios, medidos en trabajo')
    tabla(doc, ('Dónde', 'Antes', 'Con IA'),
          (('Revisión de cartas escaneadas', 'El analista transcribía a mano del PDF ilegible', 'La IA extrae y el motor valida; el analista solo ve diferencias'),
           ('Atención WhatsApp', 'Panel externo pagado, respuesta manual en horario', 'Facilito atiende 24/7, cotiza con el motor real y deriva solo lo que importa'),
           ('Análisis crediticio', 'Leer informes DealerNet fila por fila', 'El análisis llega redactado; el analista decide'),
           ('Informe de cierre a la matriz', 'Redactar desde cero cada mes', 'Borrador con las cifras reales; Contabilidad ajusta y firma'),
           ('Indicadores Previred', 'Copiar valores a mano cada mes (riesgo de error en remuneraciones)', 'La IA propone con la fuente a la vista; un humano aplica'),
           ('Preguntas de negocio', 'Pedir un informe a TI y esperar', 'Preguntar en lenguaje natural y obtener la cifra al momento')),
          (4.2, 5.6, 6.7))
    p(doc, 'El patrón común: la IA elimina la parte mecánica (leer, transcribir, redactar el primer '
           'borrador) y le deja a la persona la parte que requiere juicio (decidir, aprobar, firmar).')

    # ── 5. Costos ───────────────────────────────────────────────────────────
    h1(doc, '5. Los costos, en cifras reales')
    p(doc, 'Cada llamada registra sus tokens y su costo en dólares (tabla de uso del subsistema). '
           'Lo acumulado al 17 de agosto de 2026:')
    tabla(doc, ('Mes', 'Análisis', 'Costo USD'),
          (('Junio 2026', '46', 'US$ 2,49'),
           ('Julio 2026', '316', 'US$ 4,64'),
           ('Agosto 2026 (al día 17)', '122', 'US$ 1,12'),
           ('Total histórico', '484', 'US$ 8,25')),
          (6.0, 4.0, 6.5))
    p(doc, 'Y el detalle por funcionalidad (histórico), que muestra dónde se gasta y cuánto cuesta '
           'cada análisis:')
    tabla(doc, ('Funcionalidad', 'Análisis', 'Total USD', 'Costo promedio por análisis'),
          (('Evaluación de consistencia (Sonnet)', '10', 'US$ 2,24', '~US$ 0,22 (≈ $210 pesos)'),
           ('Pregúntale a Finanzas (Opus)', '45', 'US$ 1,90', '~US$ 0,04'),
           ('Pregúntale a AutoFácil (Sonnet)', '93', 'US$ 1,80', '~US$ 0,02'),
           ('Informe crediticio (Haiku)', '144', 'US$ 1,71', '~US$ 0,01'),
           ('Bot WhatsApp (Haiku)', '136', 'US$ 0,40', '~US$ 0,003 (≈ $3 pesos)'),
           ('Lectura de PDFs de cartas (Haiku)', '39', 'US$ 0,15', '~US$ 0,004'),
           ('Resto (Previred, resumen, asientos, moderador…)', '17', 'US$ 0,05', 'centavos')),
          (6.2, 2.4, 2.6, 5.3))
    caja(doc, 'La lectura de negocio',
         'Toda la IA del sistema cuesta hoy entre US$2 y US$5 al mes — menos que un almuerzo. La '
         'jerarquía de modelos es la clave: el 90% del volumen corre en Haiku (centavos por '
         'análisis) y los modelos caros se reservan para las preguntas financieras de fondo. El '
         'costo crece con el uso real, no con licencias: si un mes nadie pregunta, no se paga nada.')
    advertencia(doc, 'El costo por análisis NO es tarifa pactada: depende del tamaño de los '
                     'documentos y del modelo asignado. La tabla de precios por modelo es '
                     'paramétrica en el mantenedor, y el informe de Salud del Sistema trackea el '
                     'gasto en vivo. Si un análisis nuevo sale caro, se ve al tiro — no en la '
                     'factura a fin de mes.')

    # ── 6. Riesgos y resguardos ─────────────────────────────────────────────
    h1(doc, '6. Riesgos y resguardos')
    tabla(doc, ('Riesgo', 'Resguardo en el sistema'),
          (('La IA "inventa" una cifra', 'Las cifras no salen de la IA: los motores calculan, la IA redacta. En Facilito el JSON de la IA se valida antes de actuar'),
           ('Fuga de datos sensibles', 'Datos acotados por diseño: el dealer solo ve lo suyo, el clima solo agregados anónimos, y las claves API viven en el servidor'),
           ('Dependencia del proveedor', 'Si la IA se cae o se apaga, la operación sigue: todo flujo tiene su camino manual (el pool de analistas, la digitación, el panel)'),
           ('Gasto descontrolado', 'Registro por llamada, cuotas anti-abuso, precios paramétricos y monitoreo en vivo'),
           ('Uso malicioso del bot', 'Triggers de riesgo, límites de consulta y despedida automática ante mal uso'),
           ('Decisiones sin responsable', 'La IA nunca aprueba sola nada con consecuencias: propone, y la aprobación queda firmada por una persona (o por un motor determinístico auditado, como el Revisor)')),
          (5.0, 11.5))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Inteligencia-Artificial-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
