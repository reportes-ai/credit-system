# -*- coding: utf-8 -*-
"""Integraciones y APIs Externas — Business Suite: qué, cómo, cuánto cuesta."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'INTEGRACIONES Y APIS EXTERNAS', bold=True, color=AZUL_OSCURO, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Con quién conversa el sistema, para qué, cómo se gobierna cada conexión y cuánto cuesta.',
      size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, '1. El mapa completo')
    p(doc, 'Business Suite se conecta con servicios externos en cinco frentes: infraestructura, datos '
           'financieros y tributarios, evaluación de riesgo, comunicaciones y personas. La regla de '
           'gobierno es una sola: las credenciales viven como variables de entorno en el servidor — '
           'nunca en el código ni en el repositorio — y cada integración degrada con gracia: si se '
           'cae, la operación sigue por su camino manual.')
    tabla(doc, ('Servicio', 'Para qué', 'Costo'),
          (('Render', 'Servidor de producción (hosting + auto-deploy)', 'Plan mensual — ver panel'),
           ('TiDB Cloud', 'Base de datos en la nube', 'Por almacenamiento Y consulta — ver panel'),
           ('GitHub', 'Repositorio + respaldo nocturno de la BD (30 días)', 'Incluido'),
           ('Google Cloud (bucket)', 'Almacén de documentos (gs://autofacil-docs)', '~US$ 0,05/mes (medido)'),
           ('Google Cloud Run', 'Host de contingencia dormido', 'US$ 1/mes dormido · ~US$ 27/mes prorrateado si se promueve (medido)'),
           ('Google Cloud SQL', 'BD de contingencia', 'US$ 2,4/mes detenida · US$ 4,8/día encendida (medido)'),
           ('NIC Chile', 'Dominio autofacilchile.cl', 'Renovación anual — ver panel'),
           ('API CMF', 'UF, UTM, dólar y TMC oficiales', 'Gratis'),
           ('SimpleAPI (SII)', 'Registro de Compras y Ventas (RCV)', 'Plan gratis: 30 consultas/mes; planes pagados si crece'),
           ('Fintoc', 'Conexión bancaria (cartolas para conciliar)', 'Sandbox gratis hoy; plan pagado al producir'),
           ('DealerNet', 'Informes comerciales por RUT (SOAP)', 'Por consulta según plan — mantenedor Costo DealerNet'),
           ('Anthropic (Claude)', 'Todo el subsistema de IA', 'US$ 2–5/mes medido (ver documento de IA)'),
           ('Meta WhatsApp Cloud', 'El número oficial del negocio y Facilito', 'Plantillas HSM pagadas por mensaje; conversación en ventana 24 h gratis'),
           ('Brevo', 'SMTP del correo del sistema (afbs@)', 'Plan por volumen — ver panel'),
           ('Google Geocoding', 'Coordenadas de los dealers para el mapa y las rutas', 'Dentro del crédito gratis mensual de Google (~US$ 200)'),
           ('Workera', 'Reloj control (marcaciones biométricas)', 'Ver contrato/panel')),
          (3.6, 7.1, 5.8))
    advertencia(doc, 'Los costos marcados "ver panel" están deliberadamente sin cifra: la regla del '
                     'proyecto es no inventar montos — se confirman en el panel de cada proveedor. El '
                     'listado consolidado de servicios por escenario (normal / contingencia / '
                     'crecimiento) es un pendiente abierto en docs/PENDIENTES.md, y este documento es '
                     'su primera versión.')

    h1(doc, '2. Infraestructura')
    h2(doc, '2.1 Render: el servidor')
    p(doc, 'La aplicación corre en Render (región Virginia, la misma de la base: latencia de 4 ms). '
           'El deploy es automático: un push a la rama principal queda en producción en 2-3 minutos, '
           'y el badge de versión de las páginas confirma qué build está corriendo. Las claves de '
           'todas las integraciones viven aquí como variables de entorno.')
    h2(doc, '2.2 TiDB Cloud: la base')
    p(doc, 'Base MySQL-compatible en la nube. Particularidad que gobierna el diseño: cobra por '
           'almacenamiento Y POR CONSULTA — por eso los sondeos están cacheados, los documentos no '
           'viven en la base (ver 2.3) y las consultas masivas se diseñan con cuidado. El respaldo '
           'propio corre cada noche por GitHub Actions y guarda 30 días (TiDB solo retiene 1).')
    h2(doc, '2.3 El bucket de documentos')
    p(doc, 'Todo archivo que sube un usuario va al bucket de Google Cloud, no a la base: cuando los '
           'documentos vivían en la BD eran el 95% de su tamaño y el respaldo nocturno arrastraba '
           'todo. La migración dejó la base de 118 MB en 44 MB. El costo del bucket es ridículo '
           '(~US$0,05/mes) y no crece con las consultas.')
    h2(doc, '2.4 La contingencia')
    p(doc, 'Un host de respaldo duerme en Google Cloud Run (US$1/mes) apuntando a la MISMA base de '
           'producción, con los motores automáticos apagados. Si Render cae por horas, se promueve '
           'con un comando (runbook probado). La BD tiene además su propia contingencia en Cloud SQL. '
           'Costo total del seguro: menos de US$4/mes dormido.')

    h1(doc, '3. Datos financieros y tributarios')
    h2(doc, '3.1 API CMF: los indicadores oficiales')
    p(doc, 'UF, UTM, dólar y TMC se sincronizan solos desde la Comisión para el Mercado Financiero. '
           'La TMC alimenta el mantenedor de Tasas (motor único de derivación) y la tasa de mora. '
           'Gratis y oficial. Si la API falla, los indicadores quedan en su último valor y el '
           'monitoreo avisa.')
    h2(doc, '3.2 SimpleAPI: el SII sin digitación')
    p(doc, 'Trae el Registro de Compras y Ventas del SII cada 2 días: las facturas de proveedores '
           'entran solas al libro auxiliar con su asiento (capítulo 9 del Manual de Contabilidad). '
           'Requiere el certificado digital y la clave tributaria cargados en el servidor. El plan '
           'gratis da 30 consultas al mes — suficiente para el ritmo actual con su respaldo '
           'anti-desperdicio: una integración con fallas ya no puede quemarse la cuota reintentando.')
    caja(doc, 'Emisión electrónica (pendiente de decisión)',
         'Para EMITIR facturas desde el sistema (no solo recibirlas) hay que certificarse ante el '
         'SII. El set de certificación cuesta ~3 UF por una vez vía SimpleAPI; la alternativa '
         'SimpleFactura (~$30.000/mes) evita la certificación propia. La evaluación está en el '
         'informe enviado a Finanzas (agosto 2026).')
    h2(doc, '3.3 Fintoc: los bancos')
    p(doc, 'La conexión bancaria trae las cartolas para la conciliación automática. Hoy opera en '
           'sandbox con el matching automático funcionando; el paso a producción activa el plan '
           'pagado del proveedor.')

    h1(doc, '4. Evaluación de riesgo: DealerNet')
    p(doc, 'El servicio de informes comerciales se consulta por RUT vía web service (SOAP). Alimenta '
           'la Evaluación Crediticia, la incorporación de dealers (empresa + socios), la '
           'preevaluación de Facilito y la pre-aprobación del portal dealer.')
    vineta(doc, 'cada informe tiene costo según el plan contratado. El mantenedor "Costo DealerNet" registra el plan, el costo por consulta y recomienda plan según el consumo real.', bold_hasta='Costo por consulta: ')
    vineta(doc, 'los informes se guardan y un RUT consultado hace menos de 15 días no vuelve a pagar: se sirve el informe vigente.', bold_hasta='Caché de 15 días: ')
    vineta(doc, 'el bot tiene tope por conversación y tope global diario, para que nadie queme el saldo desde WhatsApp.', bold_hasta='Límites anti-abuso: ')
    advertencia(doc, 'Dos trampas conocidas del proveedor, ya cableadas en el sistema: la deuda del '
                     'producto crediticio viene en MILES de pesos (el motor la convierte), y una '
                     'respuesta "Ok" vacía NO significa sin deuda — el sistema la marca con sello '
                     'rojo "SIN INFORMACIÓN" y no ocupa el bloqueo de 15 días.')

    h1(doc, '5. Comunicaciones')
    h2(doc, '5.1 Meta WhatsApp Cloud API')
    p(doc, 'El número oficial (+56 9 3246 9071) está conectado DIRECTO a Meta, sin intermediario con '
           'fee. El modelo de cobro de Meta define la operatoria: los mensajes de PLANTILLA (HSM) se '
           'pagan por envío — campañas, cobranza, seguimiento de cartas —, mientras la conversación '
           'dentro de la ventana de 24 horas es gratis. Por eso Facilito conversa gratis y las '
           'plantillas se administran con gestor propio y aprobación de Meta desde el sistema.')
    h2(doc, '5.2 Brevo: el correo del sistema')
    p(doc, 'Todo correo del sistema (avisos, cartolas, liquidaciones, informes) sale por SMTP de '
           'Brevo con la casilla del negocio. Respeta el Modo Desarrollo (en pruebas todo se '
           'redirige a cuentas de prueba) y cada envío queda en Auditoría → Correos Enviados. El '
           'plan se dimensiona por volumen mensual de correos.')

    h1(doc, '6. Personas')
    h2(doc, '6.1 Workera: el reloj control')
    p(doc, 'Las marcaciones (biométrico + app) se consultan por API con credenciales de servicio y '
           'se cruzan con jornadas, vacaciones y feriados para detectar faltas y atrasos ANTES de '
           'pagar remuneraciones. El espejo es bidireccional: una vacación aprobada en el sistema '
           'crea la salida especial en Workera. El RUT es la llave allá (con mapeo de alias cuando '
           'difiere).')
    h2(doc, '6.2 Previred: propone la máquina, aplica el humano')
    p(doc, 'Los indicadores previsionales no tienen API formal: el sistema los lee de la fuente '
           'publicada y la IA PROPONE la actualización con la fuente a la vista; una persona revisa '
           'y aplica. El Archivo Previred mensual (cotizaciones + movimientos de personal) se genera '
           'desde Remuneraciones y se sube al portal.')

    h1(doc, '7. Google y la IA')
    vineta(doc, 'convierte las direcciones de los dealers en coordenadas para el mapa y el planificador de rutas. La clave vive en el servidor (el mapa se pinta con OpenStreetMap, sin clave expuesta). El crédito mensual gratis de Google cubre de sobra el volumen.', bold_hasta='Google Geocoding: ')
    vineta(doc, 'todo el subsistema de IA — 24 funcionalidades gobernadas por mantenedor, US$2-5 al mes medido. El detalle completo está en el documento "Inteligencia Artificial en Business Suite".', bold_hasta='Anthropic (Claude): ')

    h1(doc, '8. Nuestra propia API pública')
    p(doc, 'La integración también corre al revés: Business Suite expone una API pública para '
           'empresas externas (simulación de créditos, entre otros), autenticada por clave de API. '
           'Las claves se emiten y revocan desde el mantenedor de APIs, con su bitácora de uso. Es '
           'la puerta para que un partner cotice contra nuestros motores sin entrar al sistema.')
    captura(doc, 'Mantenedores → APIs', 'El mantenedor de claves de la API pública con su bitácora.')

    h1(doc, '9. Gobierno de las integraciones')
    vineta(doc, 'todas las claves en variables de entorno del servidor; jamás en el código, el repositorio o un documento.', bold_hasta='Credenciales: ')
    vineta(doc, 'el host de contingencia tiene las variables del núcleo pero NO las claves de integración: promovido funciona toda la operación del negocio y quedan apagados indicadores, IA, WhatsApp, DealerNet y SII. Es una degradación deliberada y documentada; cerrarla es cargar esos secretos siguiendo el manual.', bold_hasta='Contingencia degradada: ')
    vineta(doc, 'el sondeo de servicios corre cada 5 minutos y el informe de Salud del Sistema muestra el estado y el gasto (IA en vivo). Una integración caída dispara aviso, no silencio.', bold_hasta='Monitoreo: ')
    vineta(doc, 'toda tarea de fondo que llama servicios externos pasa por el programador central de motores: se puede apagar por ambiente y queda en el log. Un reintento infinito ya no puede quemar una cuota mensual (lección aprendida con el RCV).', bold_hasta='Cuotas protegidas: ')
    p(doc, 'Pendiente abierto (docs/PENDIENTES.md): consolidar el costo de cada servicio en los '
           'cuatro escenarios — operación normal, contingencia, crecimiento de volumen y qué pasa si '
           'se deja de pagar — confirmando cada cifra en su panel. Este documento aporta la '
           'estructura; faltan las cifras de los paneles.', bold=True)

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Integraciones-APIs-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
