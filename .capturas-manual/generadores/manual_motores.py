# -*- coding: utf-8 -*-
"""Motores de Cálculo — Business Suite: fórmulas, concepto y dónde viven sus variables."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def motor(doc, nombre, archivo, concepto, formula, variables, consumidores):
    """Bloque estándar de un motor."""
    h2(doc, nombre)
    p(doc, concepto)
    if formula:
        caja(doc, 'Fórmula', formula, fill='EFF6FF')
    if variables:
        tabla(doc, ('Variable', 'Dónde se configura'), variables, (6.0, 10.5))
    runs(doc, [('Vive en: ', {'bold': True}), (archivo + '.  ', {}),
               ('Consumidores: ', {'bold': True}), (consumidores, {})])

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MOTORES DE CÁLCULO', bold=True, color=AZUL_OSCURO, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Cada magnitud del negocio se calcula en UN solo lugar.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'Fórmulas, concepto y dónde viven sus variables.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026 · verificado contra el código', size=11,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'La máxima que gobierna este documento')
    p(doc, 'Máxima 1 del proyecto: cada magnitud de negocio se calcula en UN SOLO lugar — una función '
           'pura compartida que usan TODOS los consumidores. La cuota que ve el cliente en el portal, '
           'la que cobra la caja y la que imprime el certificado salen de la misma función; si se '
           'arregla, se arregla una vez. Y su regla complementaria: dos magnitudes que se parecen NO '
           'se fusionan — "días de mora" y "costo de la mora" son motores distintos, igual que "saldo '
           'insoluto" y "monto de prepago".')
    p(doc, 'Los motores son isomorfos cuando corresponde (module.exports en el servidor + window.AF_* '
           'en el navegador): el mismo archivo corre en ambos lados. Las variables NUNCA viven en el '
           'código: cada una tiene su mantenedor, y este documento dice cuál.')
    h1(doc, 'Índice')
    toc(doc)

    # ════════ 0. MODELO DE INGRESOS ════════
    h1(doc, '0. El modelo de ingresos: de dónde viene la plata')
    p(doc, 'Antes de las fórmulas, el mapa. AutoFácil tiene tres fuentes de ingreso, y cada una tiene '
           'su motor en este documento:')
    h2(doc, '0.1 Colocación: lo que pagan las financieras por cada crédito cursado')
    vineta(doc, 'nos paga el DIFERENCIAL DE TASAS. AutoFin nos fondea al costo de fondo (tasa del mantenedor menos el spread) y la operación se cursa a la tasa cliente; el ingreso es el valor presente de esa diferencia sobre todo el plazo. Por eso una rebaja de tasa come directamente la rentabilidad, y por eso el motor rechaza una tasa cliente bajo el costo de fondo. → Motor 4.1.', bold_hasta='AutoFin: ')
    vineta(doc, 'nos paga POR VOLUMEN: un porcentaje del saldo precio según el tier de operaciones cursadas en el mes (con corte por plazo en el modelo 2). Más operaciones en el mes mejoran el porcentaje de todas. → Motor 4.1 (uac-tier).', bold_hasta='Unidad de Crédito (UAC): ')
    h2(doc, '0.2 Intermediación de seguros')
    p(doc, 'AutoFin nos traspasa un porcentaje de la prima de cada seguro vendido (RDH, cesantía, '
           'reparaciones). El porcentaje del mes se gana por PENETRACIÓN PAREJA: lo define el tramo '
           'del seguro más débil — vender mucho de uno y poco de otro paga como vender poco de todo. '
           'Cuando el canal informa el porcentaje real del mes, ese valor manda. → Motor 3.7.')
    h2(doc, '0.3 Cartera propia (créditos AUTOFACIL)')
    p(doc, 'En los créditos con recursos propios el ingreso es el del negocio financiero clásico: el '
           'interés de la cuota (tasa cliente sobre el capital, calendario congelado al otorgar), más '
           'el interés por mora y los gastos de cobranza cuando hay atraso (→ motor 2.2), y la '
           'comisión de prepago cuando el cliente salda antes (→ motor 2.3).')
    h2(doc, '0.4 Contra esos ingresos, los repartos')
    p(doc, 'De lo que entra se reparte: comisión al dealer (→ 3.1), comisión y arriendo al parque '
           '(→ 3.8) y remuneración variable al equipo (→ 3.2 a 3.6). La rentabilidad por operación '
           '(→ 4.1) y la regla del 60% miden exactamente eso: cuánto del ingreso total queda en '
           'AutoFácil después de los repartos. El ingreso se devenga contablemente al solicitar la '
           'facturación (neto + IVA débito fiscal) y los montos se validan contra el export del '
           'canal — nunca su columna de comisión dealer, que no es la nuestra.')

    # ════════ I. CRÉDITO Y CUOTA ════════
    h1(doc, '1. El crédito y su cuota')
    motor(doc, '1.1 Cuota francesa y valor presente (AF_RENT_CORE)',
          'api-gateway/public/js/rentabilidad-core.js (isomorfo)',
          'El primitivo de todo el sistema de crédito: dada una plata, una tasa mensual y un plazo, '
          'cuánto se paga al mes (cuota francesa) y cuánto vale hoy una serie de cuotas (valor '
          'presente de la anualidad). Todo lo demás —cotizador, cartas, rentabilidad, aviso de cuota '
          'que no cuadra— se construye encima.',
          'cuota = capital × t × (1+t)ⁿ ÷ ((1+t)ⁿ − 1),  con t = tasa mensual (fracción), n = plazo.\n'
          'VP = cuota × (1 − (1+t)⁻ⁿ) ÷ t  (con t≈0, VP = suma simple de las cuotas).',
          (('Tasa mensual por tramo (≤/> 200 UF)', 'Mantenedores → Tasas (por fecha de vigencia)'),
           ('Umbral del tramo (200 UF)', 'Mantenedores → Parámetros del Crédito (umbral_uf_tramo)')),
          'guardado de operación, recálculo mensual, simulador, cartas, cotizador de créditos.')
    motor(doc, '1.2 Normalización de tasas (tasa-utils)',
          'shared/tasa-utils.js',
          'La tasa se ingresa (o sincroniza desde la CMF) en formato ANUAL; de ahí se derivan la '
          'mensual y el spread implícito del tramo menor. Regla de negocio: TODO a 3 decimales y el '
          'redondeo es SIEMPRE hacia abajo — nunca se redondea una tasa en contra del cliente.',
          'mensual = anual ÷ 12 (3 decimales hacia abajo).\n'
          'costo de fondo = mensual_mayor − spread_mayor.\n'
          'spread_menor = mensual_menor − mensual_mayor + spread_mayor.',
          (('Tasas anuales y spread', 'Mantenedores → Tasas'),
           ('Sincronización TMC', 'Automática desde la API de la CMF (motor tmc-sync)')),
          'mantenedor de Tasas (crear/editar) y la sincronización CMF — mismo criterio en ambos.')
    motor(doc, '1.3 Cotizador "con gastos y todo"',
          'shared/cotizador.js (réplica server-side del simulador /cotizaciones)',
          'Convierte precio + pie + plazo en la cuota real que paga el cliente: al saldo se le suman '
          'los gastos operacionales y los seguros (paquete completo Desgravamen + RDH + Cesantía, '
          'factor actuarial por tramo de plazo), y sobre ese monto financiado se aplica la cuota '
          'francesa con la tasa del tramo. Lo usan el simulador, el bot Facilito y la pre-aprobación '
          'del portal dealer — por eso el bot nunca "inventa" cuotas.',
          'monto financiado = (precio − pie) + gastos operacionales + seguros(tramo plazo).\n'
          'cuota = francesa(monto financiado, tasa tramo 200 UF, plazo).',
          (('Gastos operacionales (prenda, inscripción, GPS, etc.)', 'Mantenedores → Parámetros del Crédito'),
           ('Factores de seguros por tramo (seg_full_drc_*)', 'Mantenedores → Factores de Seguros Clientes'),
           ('Tasa por tramo', 'Mantenedores → Tasas')),
          'simulador /cotizaciones, Simulador Rápido, Facilito (WhatsApp), pre-aprobación del portal dealer.')
    motor(doc, '1.4 UF histórica (getUF)',
          'shared/uf.js',
          'Una sola estrategia de lookup para toda la aplicación: la UF vigente a una fecha es la '
          'ÚLTIMA con fecha ≤ la pedida. Si la fecha es anterior al rango cargado devuelve nulo — no '
          'inventa una UF de otra época. Antes convivían dos estrategias (fecha exacta vs histórica) '
          'y la misma operación daba valores distintos según la pantalla.',
          'UF(fecha) = valor de la última fila de la tabla uf con fecha ≤ pedida.',
          (('Valores de la UF', 'Mantenedores → UF (se sincroniza sola desde la CMF)'),),
          'clasificación mayor/menor 200 UF, gastos de cobranza, certificados, dashboard.')

    # ════════ II. MORA Y PREPAGO ════════
    h1(doc, '2. Mora, cobranza y prepago')
    motor(doc, '2.1 Días de mora y estado de cartera (AF_MORA)',
          'api-gateway/public/js/mora-core.js (isomorfo, con gemelo SQL para los caminos masivos)',
          'Dice CUÁNTO se atrasó un crédito y en qué estado de cartera está. El calendario CONGELADO '
          'de cuotas manda; los días de mora se cuentan desde el vencimiento de la cuota impaga más '
          'antigua. Importante: existe un gemelo en SQL (listados masivos de créditos y cobranza) que '
          'implementa la misma regla por rendimiento — si cambia el motor, cambia el gemelo.',
          'días de mora = hoy − vencimiento de la cuota IMPAGA más antigua.\n'
          'estado: VIGENTE (< umbral mora) / MORA / VENCIDO (≥ umbral vencido) / '
          'PREPAGADO o TERMINADO (todas pagadas).',
          (('Umbral de mora (default 1 día) y de vencido (default 91)', 'Tabla cartera_parametros — mantenedor Estado Cartera'),),
          'listado de créditos, cobranza, portal del cliente, dashboard de cartera.')
    motor(doc, '2.2 El costo de la mora (interés + gastos)',
          'shared/mora-calc.js',
          'Dice CUÁNTO CUESTA el atraso. Son dos magnitudes que no se fusionan: el interés por mora '
          '(diario simple sobre la cuota, a la TMC) y el gasto de cobranza (tramos MARGINALES sobre '
          'la deuda en UF, como la tabla del impuesto a la renta: cada tramo aplica su porcentaje '
          'solo a la porción de deuda que cae en él). Regla clave: la tasa de mora queda FIJA al '
          'otorgamiento (la TMC de ese momento) — no cambia con el mercado.',
          'interés mora = cuota × (TMC del otorgamiento ÷ 30) × días de atraso.\n'
          'gasto cobranza = Σ por tramo: porción de deuda en UF × % del tramo — corre desde el día 21.',
          (('Tasa de mora (modo fijo al otorgar) y tramos de gasto', 'Mantenedores → Parámetros de Cobranza'),
           ('TMC histórica por tramo', 'Mantenedores → Tasas (sync CMF)'),
           ('Día desde el que corren los gastos (default 21)', 'Parámetros de Cobranza (gastos_dias)')),
          'caja, portal del cliente, certificados, motor de prepago, cobranza — mismos pesos en todos.')
    motor(doc, '2.3 Monto de prepago (calcularPrepago)',
          'services/certificados/src/controllers/certificados.controller.js',
          'Cuánto cuesta saldar HOY un crédito completo. No es el "saldo insoluto neto": suma el '
          'capital de las cuotas futuras, las cuotas en mora con su interés y gastos, el interés '
          'corriente desde el último vencimiento y la comisión de prepago (un mes de interés sobre el '
          'capital vigente). Entrega el detalle cuota por cuota y la proyección de los próximos días '
          'hábiles.',
          'prepago = capital vigente + cuotas en mora + interés mora + gastos cobranza\n'
          '        + capital vigente × tasa diaria × días corrientes + capital vigente × tasa mensual.',
          (('Todo hereda de 2.2 (mora) y de Tasas', 'Parámetros de Cobranza · Tasas'),),
          'certificado de prepago, prepago en caja, Aplicación de Fondos (base de la propuesta).')
    motor(doc, '2.4 Provisiones de cartera',
          'motor mensual automático (cierre de provisiones)',
          'Reconocimiento contable del riesgo: cada tramo de días de mora provisiona un porcentaje '
          'del saldo. Si nadie las guarda a mano, el sistema cierra solo el mes anterior y genera el '
          'asiento por la VARIACIÓN contra el mes previo (constitución si subió, liberación si bajó), '
          'con el detalle por crédito auditable.',
          'provisión del mes = Σ créditos: saldo × % del tramo de mora.\n'
          'asiento = provisión(mes) − provisión(mes anterior).',
          (('Tramos y porcentajes de provisión', 'Mantenedores → Parámetros de Cobranza'),
           ('Día del cierre automático', 'Parámetros de Cobranza')),
          'cierre de mes, estados financieros, castigos.')

    # ════════ III. COMISIONES ════════
    h1(doc, '3. Comisiones y remuneración variable')
    motor(doc, '3.1 Comisión del dealer',
          'api-gateway/public/js/comision-dealer.js (isomorfo)',
          'La comisión que se le paga al dealer por cada operación. La jerarquía es estricta: si el '
          'dealer tiene TABLA PACTADA en su ficha, esa manda; si no, rige la PIZARRA general. El '
          'porcentaje depende del tramo de plazo (1-12 / 13-24 / 25-36 / 37+) y de si la operación es '
          'de PARQUE o de CALLE. La llave para encontrar la tabla pactada es el RUT del dealer — por '
          'eso las pantallas exigen elegir el dealer por autocompletado y no escribir el nombre a '
          'mano. El resultado es SIEMPRE BRUTO (IVA incluido, ver 5.1).',
          'comisión = saldo precio × %(tabla del dealer o pizarra)[tramo de plazo][parque/calle].',
          (('Tabla pactada por dealer (com_* y com_parque_*)', 'Mantenedor de Dealers → ficha del dealer'),
           ('Pizarra general por tramo', 'Mantenedores → Parámetros del Crédito'),
           ('Comisión y arriendo del parque', 'Mantenedores → Parques (Arriendos y Comisiones)')),
          'cartas de aprobación, cartolas, otorgadas sin carta, Revisor Automático, rentabilidad.')
    motor(doc, '3.2 Comisión del ejecutivo comercial',
          'shared/comision-ejecutivo.js (función pura)',
          'La remuneración variable del ejecutivo según el anexo 08-2026: piso habilitante de monto, '
          'base por tramo de plazo, ajustes por cruce de seguros (RDH, cesantía, reparaciones — cada '
          'uno solo si supera su umbral), semana corrida al final, y los descuentos de la cláusula '
          'novena por prepagos y anulaciones. El detalle completo, paso a paso y con los valores '
          'vigentes, está en la sección 3 del Diccionario de Datos (v3.0).',
          'incentivo = [Σ monto×tasa por tramo] × (1 + Σ ajustes de seguros) × semana corrida\n'
          '− comisiones ya pagadas de operaciones prepagadas/anuladas (por tramo de meses).',
          (('Todas las variables, versionadas con vigencia por mes', 'Comisiones → Mantenedor Variables (bitácora inmutable)'),),
          'Revisión de Comisiones, libro de Remuneraciones, informe al ejecutivo.')
    h3(doc, 'El bono por cumplimiento de adicionales (detalle)')
    p(doc, 'El "adicional" del ejecutivo son los bonos por venta de seguros sobre el universo NCNU '
           '(operaciones AutoFin sin CORFO). Cada indicador funciona igual: se mide la penetración '
           '(cruce), se compara contra su umbral, y SOLO si lo supera aporta — bajo el umbral el '
           'bono de ese seguro es 0, no proporcional:')
    tabla(doc, ('Indicador', 'Umbral', 'Peso', 'Aporte si supera el umbral'),
          (('Cruce RDH (incluye desgravamen)', '99%', '0,33', 'cruce × 0,33 × factor_max (0,66)'),
           ('Cruce cesantía', '65%', '0,34', 'cruce × 0,34 × factor_max'),
           ('Cruce reparaciones', '55%', '0,33', 'cruce × 0,33 × factor_max'),
           ('Calidad (créditos UNIDAD)', 'meta 3, TODO O NADA', '0,00 (fuera del modelo)', 'Se reactiva dándole peso en el mantenedor')),
          (5.0, 3.4, 3.2, 4.9))
    p(doc, 'Cada bono se paga sobre el incentivo base TOTAL del mes: bono_seguro = incentivo_base × '
           'aporte del indicador. Con los tres cruces al 100% cumplidos, el factor de ajuste llega a '
           'su máximo teórico de 0,66 (factor_max): el adicional puede aumentar la comisión base '
           'hasta en dos tercios. El detalle numérico completo, con ejemplos, está en la sección 3 '
           'del Diccionario de Datos v3.0.')
    motor(doc, '3.3 Semana corrida (art. 45 CT)',
          'shared/semana-corrida.js',
          'No es una constante: la ley paga los domingos y festivos con el promedio de lo devengado '
          'en los días trabajados, así que el factor cambia todos los meses. La jornada pactada de '
          'los ejecutivos es lunes a sábado. Un feriado que cae domingo se cuenta una sola vez.',
          'factor del mes = 1 + (domingos + festivos) ÷ (días lunes-sábado no feriados).',
          (('Feriados de Chile', 'Mantenedores → Feriados'),
           ('Modo calculado vs multiplicador fijo', 'Comisiones → Mantenedor Variables (semana_corrida_calc)')),
          'comisiones de ejecutivos y Bono Jefe Comercial — prohibido usar una constante propia.')
    motor(doc, '3.4 Horas extraordinarias (art. 32 CT)',
          'shared/horas-extras.js',
          'El valor de la hora extra con la fórmula de la Dirección del Trabajo. La base es el SUELDO '
          'BASE pactado — no el total imponible: bonos y comisiones no forman parte de la jornada '
          'ordinaria salvo pacto expreso. La jornada semanal NO está fija en el código porque la ley '
          'la sigue bajando (44 h hoy, 42 en 2027, 40 en 2028). Los trabajadores art. 22 (excluidos '
          'de jornada) no generan horas extra: el motor responde "no aplica" con su motivo, nunca un '
          '0 silencioso.',
          'hora ordinaria = sueldo base × 7 ÷ (30 × jornada semanal).\n'
          'hora extra = hora ordinaria × 1,5 (recargo 50%).',
          (('Jornada de cada colaborador (44/40/especial/art.22)', 'RRHH → Jornada'),
           ('Recargo (mínimo legal 50%)', 'Paramétrico en RRHH')),
          'remuneraciones (cálculo de horas extra del libro).')
    motor(doc, '3.5 Plan Liquidez (AF_LIQUIDEZ)',
          'shared/liquidez-core.js (isomorfo)',
          'El anticipo de comisiones de los dealers Super Partner se "resetea" cada mes al nivel de '
          'la comisión, con tope. La cartola SIEMPRE va por el total de la comisión; el descuento se '
          'aplica solo en la orden de pago, y el abono a la deuda se confirma al pagarse.',
          'A (adelanto del mes) = min(comisión C, tope).\n'
          'descuento = deuda anterior D − A  (negativo = se presta extra).\n'
          'pago neto = C − D + A.  Nueva deuda = A.',
          (('Tope de adelanto por dealer', 'Dealers → Plan Liquidez (contrato del plan)'),),
          'Hoja de Liquidación mensual, ODP del dealer, cuenta corriente del dealer.')
    motor(doc, '3.6 Bono del Jefe Comercial (BSC)',
          'motor del módulo Bono Jefe Comercial',
          'Balanced scorecard sobre el desempeño PROMEDIO del equipo de ejecutivos vigentes del mes: '
          'créditos otorgados (45%), montos otorgados (40%) y nuevos dealers con negocios (15%). Se '
          'promedian las MÉTRICAS del equipo (no los puntajes) y el score entra a una curva '
          'exponencial sobre el sueldo fijo; se suma la semana corrida (motor 3.3). Variables '
          'versionadas con vigencia y bitácora, igual que las de comisiones.',
          'score = Σ pilares(mín/esperado/logrado → puntaje ponderado) → curva exponencial → premio.',
          (('Pilares, umbrales, ponderaciones y curva', 'Bono Jefe Comercial → pestaña Variables (restringida)'),),
          'renta del Jefe Comercial en Remuneraciones.')

    motor(doc, '3.7 Comisión por intermediación de seguros',
          'services/creditos/src/utils/penetracion.js (comisionesSeguro + pctTraspasoMes)',
          'El ingreso de AutoFácil por cada seguro vendido: la prima de la operación multiplicada por '
          'el porcentaje de traspaso del MES. Ese porcentaje no es fijo: sale del tramo alcanzado por '
          'la penetración del seguro MÁS DÉBIL del mes (se toma el mínimo entre RDH, cesantía y '
          'reparaciones — vender parejo importa más que vender mucho de uno). Si AutoFin informa un '
          'porcentaje puntual para un mes, ese override manda. La penetración del mes se calcula real '
          '(motor penetracion.js), no del campo guardado en la operación.',
          'com_seguro = prima del seguro × % traspaso del mes.\n'
          '% traspaso = tramo(min(pen_rdh, pen_cesantia, pen_reparaciones)) — u override del mes.',
          (('% de traspaso por defecto (seg_pct_traspaso_autofin, 30%)', 'Mantenedores → Parámetros del Crédito'),
           ('Tramos penetración → % comisión', 'Tabla de tramos de penetración (paramétrica por tipo de seguro)'),
           ('Overrides mensuales informados por AutoFin', 'Tabla comisiones_seguro_pct_mes')),
          'guardado de operación y recálculo mensual (com_rdh/com_cesantia/com_reparaciones), dashboard (com_seguros), rentabilidad.')
    motor(doc, '3.8 Comisión y arriendo del parque',
          'motor en calcular-operacion.js + comisiones-parques.controller.js',
          'El parque automotriz gana por dos vías: una comisión por cada crédito cursado en el parque '
          '(porcentaje del saldo precio, según el mantenedor de parques — se persiste como com_parque '
          'al calcular la operación) y un arriendo mensual FIJO. El arriendo nunca es un fee por '
          'operación: para efectos de rentabilidad se prorratea entre las otorgadas del mes, y en la '
          'cartola del parque va como línea propia. El parque de la operación viene del Excel de '
          'carga (o de la carta), no del mantenedor: sin ese dato no hay comisión de parque.',
          'com_parque (por operación) = saldo precio × % del parque.\n'
          'cartola del mes = Σ com_parque de las operaciones elegibles + arriendo mensual.\n'
          'arriendo prorrateado (rentabilidad) = arriendo fijo ÷ otorgadas del mes en el parque.',
          (('% de comisión y arriendo por parque', 'Mantenedores → Parques (Arriendos y Comisiones)'),
           ('Universo de la cartola (corte de partida)', 'postventa_config (parques_universo_desde)')),
          'cálculo de la operación, Emisión de Cartolas Parque, rentabilidad, dashboard.')

    # ════════ III-b. RRHH ════════
    h1(doc, '3 bis. RRHH: finiquito y vacaciones')
    motor(doc, '3.9 Finiquito (indemnizaciones)',
          'services/rrhh/src/controllers/contratos.controller.js (con rrhh-core.js)',
          'Cuánto se paga al término de la relación laboral. La BASE es el promedio de las últimas 3 '
          'liquidaciones EMITIDAS (rentas variables); sin liquidaciones, sueldo base + 25% de '
          'gratificación, con aviso para revisar. La base se topa a 90 UF (art. 172). Los años de '
          'servicio: la fracción ≥ 6 meses sube al año siguiente, con tope de 11 años — ambos topes '
          'paramétricos. El feriado proporcional sale de la MISMA cuenta corriente de vacaciones del '
          'formulario (motor único), y los saldos de anticipos y préstamos vigentes se descuentan.',
          'indemnización años = min(años, 11) × min(base, 90 UF)  — solo si la causal la lleva.\n'
          'mes de aviso = base topada — si la causal lo lleva y no hubo aviso previo.\n'
          'feriado proporcional = saldo de la cuenta de vacaciones valorizado (motor 3.10).\n'
          'total = haberes − saldo de anticipos/préstamos pendientes.',
          (('Topes (finiq_tope_anos = 11, finiq_tope_uf = 90)', 'Configuración RRHH (rh_config)'),
           ('Causales y qué lleva cada una (indemniza / mes aviso)', 'Catálogo rh_finiquito_causales (art. 159/161/163…)'),
           ('UF del día del término', 'Motor UF (1.4)')),
          'RRHH → Finiquitos; genera ODP y asientos FINIQUITO_EMITIDO / FINIQUITO_PAGADO.')
    motor(doc, '3.10 Vacaciones: devengo y provisión',
          'api-gateway/public/js/rrhh-core.js (provisionVacaciones) + cuenta corriente vac-cuenta',
          'La cuenta corriente de vacaciones es la única verdad: devenga sola 15 días hábiles en cada '
          'aniversario más el feriado progresivo (art. 68: con 10 años trabajados, 1 día extra por '
          'cada 3 nuevos), y descuenta FIFO del período más antiguo al aprobarse cada solicitud. La '
          'valorización en pesos (para la provisión mensual y el finiquito) usa el mismo motor. Al '
          'cierre de mes se contabiliza la VARIACIÓN del total contra el mes anterior.',
          'devengo anual = 15 días hábiles + progresivo(art. 68).\n'
          'días corridos ≈ hábiles × 1,4.\n'
          'provisión = saldo hábiles valorizado con la base de remuneración (motor rrhh-core).\n'
          'asiento del mes = provisión(mes) − provisión(mes anterior).',
          (('Años previos declarados (feriado progresivo)', 'Ficha del colaborador + certificado AFP en la carpeta'),
           ('Base de remuneración', 'Liquidaciones emitidas (motor base-remuneracion)')),
          'formulario de vacaciones, Saldos del equipo, finiquito (3.9), provisión del cierre de mes.')

    # ════════ IV. RENTABILIDAD ════════
    h1(doc, '4. Rentabilidad y elección de financiera')
    motor(doc, '4.1 Rentabilidad por operación (AF_RENT)',
          'api-gateway/public/js/rentabilidad-core.js + rentabilidad-calc.js (isomorfos)',
          'Cuánto le deja a AutoFácil una operación en cada financiera. En AutoFin, el ingreso por '
          'colocación es el valor presente del diferencial entre la tasa cliente y el costo de fondo; '
          'en UAC, la comisión según el tier. Reglas de negocio grabadas en el motor: la tasa cliente '
          'de la operación MANDA sobre la del mantenedor (salvo que esté bajo el costo de fondo — ahí '
          'es dato inválido y se cae al mantenedor), y las tasas que llegan como fracción (0,02868) '
          'se normalizan a % mensual (2,868).',
          'ingreso AutoFin ≈ VP(cuota a tasa cliente) − VP(cuota a costo de fondo), '
          'con costo de fondo = tasa mantenedor − spread.\n'
          'ingreso UAC = comisión según tier (motor uac-tier.js, el mismo del backend).',
          (('Tasas y spread por tramo y fecha', 'Mantenedores → Tasas'),
           ('Parámetros del negocio', 'Mantenedores → Parámetros del Crédito'),
           ('Tiers UAC', 'Motor uac-tier.js — paramétrico donde definió el negocio')),
          'Simulador de Rentabilidad, cartas (comparación AutoFin vs UAC), ¿Dónde Curso?, informe de rentabilidad.')
    motor(doc, '4.2 Preferencia financiera (¿Dónde Curso?)',
          'api-gateway/public/js/preferencia-financiera.js',
          'Decide dónde conviene cursar. Primero ELEGIBILIDAD: el cuadro de Preferencia Financiera '
          'dice qué financiera acepta ese plazo con ese saldo — la que no puede, no compite. Después '
          'RENTABILIDAD: gana la más alta, salvo que la diferencia sea menor al umbral, en cuyo caso '
          'el veredicto es DECIDES TÚ. Cursar en la financiera conveniente (o cualquiera con DECIDES '
          'TÚ) NO es excepción; solo se registra excepción al elegir la MENOS conveniente.',
          'veredicto = elegibles(cuadro) → max rentabilidad, con empate si |dif| < umbral (10%).\n'
          'aviso amarillo si rentabilidad < mínima ($100.000 default).',
          (('Cuadro de elegibilidad (grilla cuotas × saldo)', 'Mantenedores → Preferencia Financiera'),
           ('Umbral de empate (pref_diff_pct) y rentabilidad mínima (pref_rent_minima)', 'Mantenedores → Parámetros del Crédito')),
          'PWA ¿Dónde Curso?, la carta de aprobación (veredicto bajo el selector), informe de rentabilidad.')
    motor(doc, '4.3 Penetración de seguros',
          'motor penetracion.js',
          'Qué porcentaje de las operaciones ELEGIBLES lleva cada seguro. La clave del motor son los '
          'universos: no toda operación es elegible para todo seguro, así que el denominador se '
          'define por seguro (igual que el NCNU de comisiones excluye CORFO y Unidad). Consistencia: '
          'la penetración del dashboard y la de comisiones usan las mismas definiciones de universo.',
          'penetración(seguro) = operaciones con el seguro ÷ operaciones elegibles para ese seguro.',
          (('Universos elegibles por seguro', 'Definición del motor — cambios se piden a TI con acta de negocio'),),
          'dashboard, comisiones de ejecutivos, informes comerciales.')

    # ════════ V. PLATA Y DOCUMENTOS ════════
    h1(doc, '5. Impuestos y contabilización')
    motor(doc, '5.1 Desglose de impuesto (AF_IMPUESTO)',
          'api-gateway/public/js/desglose-impuesto.js (isomorfo)',
          'La comisión que calcula el sistema es SIEMPRE BRUTA (IVA incluido): el neto y el impuesto '
          'se desagregan HACIA ABAJO, jamás se suma impuesto encima. Equivalencia factura↔boleta: el '
          'costo real de AutoFácil es siempre el NETO — con factura el IVA es crédito fiscal; con '
          'boleta de honorarios, el honorario bruto es el neto de la factura equivalente y se paga el '
          'líquido (honorario − retención).',
          'Factura: neto = bruto ÷ (1 + IVA); impuesto = bruto − neto; se paga el bruto.\n'
          'Boleta: base = comisión ÷ (1 + IVA); retención = base × tasa; se paga base − retención.',
          (('Tasa de IVA y de retención de honorarios', 'Mantenedor de impuestos (indicadores; la retención sube por ley cada año)'),),
          'carga masiva, órdenes de pago de comisión, Tesorería, auxiliar de honorarios y F29.')
    motor(doc, '5.2 Centralización contable (contabilizar)',
          'services/contabilidad/src/motor-asientos.js',
          'Todo movimiento de dinero genera su asiento por este único motor, según la regla '
          'paramétrica del hecho (COMISION_DEV, COMISION_PARQUES, FINIQUITO_EMITIDO…). Devengo y '
          'pago son hechos separados. El motor NUNCA bloquea la operación: si algo falla queda en el '
          'log de eventos (CONTABILIZADO / SIN_REGLA / DESCUADRE / ERROR) y ese log es la lista de '
          'pendientes. Cada asiento es idempotente por referencia: procesar dos veces no duplica.',
          'hecho económico + regla(cuentas debe/haber, impuestos) → comprobante cuadrado.',
          (('Reglas por evento (cuentas, condiciones)', 'Contabilidad → Reglas de Centralización'),
           ('Plan de cuentas', 'Contabilidad → Plan de Cuentas'),),
          'todos los módulos que mueven plata: cartolas, ODP, caja, remuneraciones, finiquitos, RCV.')
    motor(doc, '5.3 Motor de etapa del crédito (setEtapa)',
          'shared/etapa-credito.js',
          'No calcula plata pero decide de quién es: la etapa vive por herencia en TRES columnas y '
          'este motor las escribe juntas en el mismo acto (setEtapa) y las lee por una sola expresión '
          'canónica (ETAPA_SQL). Un UPDATE suelto de una columna vuelve a partir la etapa. El control '
          'desalineados() lista las operaciones partidas.',
          'escribir: setEtapa() → las tres columnas juntas. Leer: ETAPA_SQL (estado manda; si no, respaldo).',
          (('Estados y transiciones del mapa', 'Mantenedores → Estado Créditos'),),
          'otorgamiento, anulación, carga Trinidad, desistimiento automático — y toda pantalla que filtra por etapa.')

    # ── Anexo ──
    h1(doc, 'Anexo A. Mapa motor → mantenedor (resumen)')
    tabla(doc, ('Motor', 'Archivo', 'Mantenedor de sus variables'),
          (('Cuota francesa / VP', 'js/rentabilidad-core.js', 'Tasas · Parámetros del Crédito'),
           ('Normalización de tasas', 'shared/tasa-utils.js', 'Tasas (+ sync CMF)'),
           ('Cotizador completo', 'shared/cotizador.js', 'Parámetros del Crédito · Factores de Seguros · Tasas'),
           ('UF histórica', 'shared/uf.js', 'UF (sync CMF)'),
           ('Días de mora / estado', 'js/mora-core.js', 'Estado Cartera (umbrales)'),
           ('Costo de la mora', 'shared/mora-calc.js', 'Parámetros de Cobranza · Tasas'),
           ('Monto de prepago', 'certificados.controller.js', 'hereda de mora y Tasas'),
           ('Provisiones', 'motor mensual', 'Parámetros de Cobranza'),
           ('Comisión dealer', 'js/comision-dealer.js', 'Ficha del Dealer · Parámetros del Crédito (pizarra) · Parques'),
           ('Comisión ejecutivo', 'shared/comision-ejecutivo.js', 'Comisiones → Mantenedor Variables'),
           ('Semana corrida', 'shared/semana-corrida.js', 'Feriados · variable de modo'),
           ('Horas extras', 'shared/horas-extras.js', 'RRHH → Jornada'),
           ('Comisión por seguros', 'creditos/utils/penetracion.js', 'Parámetros del Crédito (traspaso) · tramos · overrides del canal'),
           ('Comisión y arriendo parque', 'calcular-operacion + cartolas parque', 'Parques (Arriendos y Comisiones) · postventa_config'),
           ('Finiquito', 'rrhh/contratos.controller.js', 'rh_config (topes) · causales · UF'),
           ('Vacaciones / provisión', 'js/rrhh-core.js + vac-cuenta', 'ficha (años previos) · liquidaciones'),
           ('Plan Liquidez', 'shared/liquidez-core.js', 'contrato del plan (tope)'),
           ('Bono Jefe Comercial', 'módulo bono-jefe', 'pestaña Variables (restringida)'),
           ('Rentabilidad', 'js/rentabilidad-core/calc.js', 'Tasas · Parámetros del Crédito'),
           ('Preferencia financiera', 'js/preferencia-financiera.js', 'Preferencia Financiera · Parámetros del Crédito'),
           ('Penetración seguros', 'penetracion.js', 'universos definidos en el motor'),
           ('Desglose impuesto', 'js/desglose-impuesto.js', 'impuestos (IVA / retención)'),
           ('Centralización contable', 'contabilidad/motor-asientos.js', 'Reglas de Centralización · Plan de Cuentas'),
           ('Etapa del crédito', 'shared/etapa-credito.js', 'Estado Créditos')),
          (4.2, 5.4, 6.9))
    p(doc, 'Regla final: si vas a escribir un cálculo nuevo, primero busca en esta lista dónde ya se '
           'computa esa magnitud. Si existe, se reusa o extiende — nunca una segunda copia "rápida". '
           'Si no existe, nace como motor: función pura, un solo hogar, variables en su mantenedor y '
           'su fila en este documento.', bold=True)

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Motores-de-Calculo-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
