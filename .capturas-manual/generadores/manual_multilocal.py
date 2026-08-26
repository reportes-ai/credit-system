# -*- coding: utf-8 -*-
"""Manual Dealers Multi-Local — de la creación del dealer al pago de comisiones (v218.9)."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DEALERS MULTI-LOCAL', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Un dealer, varios locales: de la ficha de incorporación', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'al pago de la comisión, con la tabla correcta en cada parque y en calle.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026 · sistema v218.9', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del manual Dealers Multi-Local (v218.0–v218.9)'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Sigue el orden real del proceso: crear el dealer con sus locales, mantenerlos, emitir la carta '
           'en el local del negocio, y cobrar — cartola, factura, orden de pago y pago. Cada capítulo dice '
           'quién lo hace, en qué pantalla y con qué permisos, más los recuadros de color: ⚠ OJO (lo que '
           'confunde), 🔒 Regla del sistema (lo que no se puede saltar) y 🧾 Caso real. Las rutas se '
           'escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. El concepto ───────────────────────────────────────────────────────
    h1(doc, '1. Qué es un dealer multi-local')
    p(doc, 'Un dealer es UNA sola empresa (un RUT), pero puede vender en varios lugares a la vez: locales '
           'en dos o más parques automotrices y/o venta directa en calle. Cada lugar es un LOCAL del '
           'dealer, y cada local tiene su propia tabla de comisión pactada (4 tramos de plazo: 6–12, '
           '13–24, 25–36 y 37+ cuotas).')
    flujo(doc, 'FICHA CON LOCALES → CARTA EN EL LOCAL → CRÉDITO → CARTOLA POR SECCIONES → FACTURA ÚNICA → ODP → PAGO')
    vineta(doc, 'la comisión de cada operación se calcula con la tabla del local donde CURSÓ '
                '(el parque de la carta), no con la ficha "de hoy" del dealer.', bold_hasta='La regla central: ')
    vineta(doc, 'paga el % de ESE parque y además genera la comisión y el arriendo al dueño del parque (su cartola).', bold_hasta='Operación en un parque: ')
    vineta(doc, 'paga el % de calle, que es MÁS ALTO porque no se le paga a ningún parque.', bold_hasta='Operación de calle: ')
    vineta(doc, 'tabla del local → tabla histórica de la ficha → pizarra de Parámetros de Crédito. '
                'Un tramo vacío usa el siguiente nivel.', bold_hasta='Precedencia del %: ')
    vineta(doc, 'sigue siendo UNA por RUT; si el dealer cursó en varios locales, se corta en secciones con subtotal por local.', bold_hasta='La cartola del dealer: ')
    regla(doc, 'Las comisiones pactadas nunca se cambian "por fuera": se fijan por ficha aprobada (cadena de '
               'autorización, y Gerencia si superan la pizarra) o en el mantenedor Dealers con permiso y auditoría.')

    h2(doc, 'Quién participa')
    tabla(doc, ('Rol', 'Qué hace en este proceso'),
          (('Ejecutivo Comercial', 'Crea la ficha con los locales y sus tablas; digita las cartas eligiendo el local donde cursa cada negocio.'),
           ('Cadena de Autorización', 'Revisa y aprueba la ficha por niveles; Gerencia entra si algún tramo supera la pizarra.'),
           ('Analista de Operaciones', 'Cierra la ficha aprobada (crea el dealer y SELLA sus locales); administra locales en el mantenedor Dealers.'),
           ('Post Venta / Operaciones', 'Emite y envía la cartola, registra la factura, emite la Orden de Pago de comisión.'),
           ('Dealer (externo)', 'Firma la ficha; recibe la cartola y emite UNA factura por el total.'),
           ('Tesorería', 'Paga la ODP (segregación: quien emite NO paga).'),
           ('Contabilidad', 'No digita nada: devengo y pago se contabilizan automáticos (Reglas de Centralización).')),
          (4.2, 11.9))

    # ── 2. Crear el dealer ───────────────────────────────────────────────────
    h1(doc, '2. Crear el dealer con sus locales (ficha de incorporación)')
    ficha(doc, 'Ejecutivo Comercial', 'dealer_inc_crear (crear ficha)', 'Parques registrados en Mantenedores → Arriendos y Comisiones',
          'Creación/Mantenedor de Dealer → Nueva Ficha')
    paso(doc, 1, 'Datos de la empresa', 'RUT, razón social, contactos, socios y cuenta bancaria, como en cualquier ficha.')
    paso(doc, 2, 'Tipo de dealer', 'Parque (solo vende en parques), Ambos (parques + calle) o General (solo calle).')
    paso(doc, 3, 'Parque principal', 'En "Datos del Parque Automotriz" elige el parque principal (de la lista del '
         'mantenedor) con su dirección y comuna.')
    paso(doc, 4, 'Locales en OTROS parques', 'Si el dealer tiene más locales, presiona "+ Agregar otro parque" por cada '
         'uno: parque (de la lista), dirección, comuna y su tabla de comisión (4 tramos %). Los tramos vacíos usan la pizarra.')
    captura(doc, 'Nueva Ficha → Datos del Parque Automotriz', 'Bloque "Locales en OTROS parques" con una fila agregada y su tabla de %.')
    paso(doc, 5, 'Comisión pactada', 'La tabla del principal (y la de calle si es Ambos) va en la sección Comisión '
         'Pactada; la de cada parque adicional va en su propia fila de local.')
    paso(doc, 6, 'Vista previa y envío', 'La ficha impresa incluye la sección "LOCALES ADICIONALES EN OTROS PARQUES" '
         'con cada tabla — eso es lo que firma el dealer. Al enviar, el sistema valida parques registrados y sin '
         'repetidos, y consulta DealerNet de la empresa y los socios.')
    advertencia(doc, 'Si algún tramo va SOBRE la pizarra, la ficha exige comentario de excepción y pasa además por '
                'Gerencia (participación especial). No es un error: es el circuito normal de una negociación especial.')

    h2(doc, '2.1 Aprobación, firma y cierre')
    p(doc, 'Cada nivel de la cadena revisa la ficha —incluidos los locales adicionales con sus tablas— y aprueba o '
           'rechaza con motivo. Aprobada la cadena, el dealer firma la ficha impresa y se sube junto a las cédulas. '
           'El Analista de Operaciones revisa la firma y ejecuta el CIERRE:')
    vineta(doc, 'se crea (o actualiza) el dealer con su número correlativo;')
    vineta(doc, 'el sistema SELLA sus locales: el principal según el tipo, más cada parque adicional con su tabla;')
    vineta(doc, 'desde ese momento las cartas, el cálculo de comisiones y las cartolas ya los usan.')

    # ── 3. Mantener ──────────────────────────────────────────────────────────
    h1(doc, '3. Mantener los locales de un dealer existente')
    p(doc, 'Hay dos caminos, según si el cambio toca comisiones:')
    h2(doc, '3.1 Ficha de Modificación (camino gobernado — cambia comisiones)')
    ficha(doc, 'Ejecutivo Comercial + cadena de autorización', 'dealer_inc_crear', 'Dealer existente',
          'Nueva Ficha → buscador de dealer existente (modo MODIFICACIÓN)')
    paso(doc, 1, 'Cargar el dealer', 'El buscador precarga TODO, incluidos los locales adicionales actuales con sus tablas.')
    paso(doc, 2, 'Ajustar', 'Agrega, corrige o elimina filas de locales; cambia los % que corresponda.')
    paso(doc, 3, 'Tramitar', 'Misma cadena del capítulo 2. Al cerrar, los locales se re-sellan.')
    h2(doc, '3.2 Mantenedor Dealers (ajuste directo, auditado)')
    ficha(doc, 'Analista de Operaciones / Administrador', 'dealers_base_editar o dealer_ficha_revisar', 'Dealer existente',
          'Dealers → pestaña Base → detalle del dealer')
    paso(doc, 1, 'Abrir la sección', 'En el detalle del dealer, sección "Locales y comisión por local".')
    paso(doc, 2, 'Agregar / Editar', 'Ubicación (CALLE o un parque del mantenedor), dirección, comuna, marcar '
         'Principal y los 4 tramos %.')
    paso(doc, 3, 'Desactivar', 'Apaga un local sin borrar nada: las operaciones históricas no se tocan.')
    captura(doc, 'Dealers → detalle → Locales y comisión por local', 'Dealer con dos parques y sus tablas.')
    advertencia(doc, 'En Mantención → Editar Dealer los locales se ven en SOLO LECTURA: es a propósito, para que los '
                'cambios de plata pasen por la ficha de Modificación o por el mantenedor Dealers (ambos auditados).')

    # ── 4. La carta ──────────────────────────────────────────────────────────
    h1(doc, '4. La carta de aprobación elige el LOCAL del negocio')
    ficha(doc, 'Ejecutivo Comercial', 'emitir cartas', 'Dealer con locales sellados',
          'Generador de Carta de Aprobación (/aprobaciones)')
    paso(doc, 1, 'Tipo de carta', 'DEALER PARQUE o DEALER CALLE.')
    paso(doc, 2, 'Parque y dealer', 'Si es PARQUE, al elegir el parque el dealer aparece entre los dealers de ESE '
         'parque si tiene local ahí (aunque su ficha "principal" diga otro). Si es CALLE, aparecen también los '
         'dealers de parque con local calle.')
    paso(doc, 3, 'Participación sugerida', 'Sale de la tabla pactada de ESE local — la etiqueta dice "Tabla pactada '
         'del dealer — PARQUE X". Sin tabla propia, de la pizarra.')
    paso(doc, 4, 'Excepciones', 'Un monto sobre lo pactado del local exige EXCEPCIÓN con comentario, como siempre.')
    paso(doc, 5, 'Guardar', 'El servidor valida que la ubicación sea un local vigente del dealer; si no, el mensaje '
         'dice qué locales tiene y que agregues el nuevo en su ficha.')
    regla(doc, 'Al otorgarse, el crédito hereda el NOMBRE del parque de la carta (o CALLE). Con eso el motor calcula '
               'la comisión del dealer con la tabla correcta y atribuye la comisión y el arriendo al parque correcto. '
               'La participación escrita en la carta manda sobre el cálculo (queda protegida contra recálculos).')
    caso(doc, 'Op 89213 (NIBARO): cursó en Autocenter Maipú pero la ficha ya decía Carmoons — la comisión salía en la '
              'cartola del parque equivocado. Por eso el parque de la OPERACIÓN manda siempre sobre la ficha.')

    # ── 5. Cartola ───────────────────────────────────────────────────────────
    h1(doc, '5. Cartola del dealer: una sola, con secciones por local')
    ficha(doc, 'Post Venta / Operaciones', 'aprob_cartolas', 'Comisiones en estado A PAGAR (fondos recibidos)',
          'Generador de Cartas → pestaña Cartolas')
    paso(doc, 1, 'Armado', 'El sync junta las comisiones A PAGAR del dealer: una cartola por RUT, sin importar la financiera.')
    paso(doc, 2, 'Secciones por local', 'Si el dealer cursó en más de un local, la cartola sale cortada por ubicación: '
         'cada parque con sus operaciones y su subtotal, y Calle al final. Con un solo local se ve igual que siempre.')
    paso(doc, 3, 'Enviar', 'Sale por correo al dealer (comisiones@) con la cartola en PDF adjunta, copia al ejecutivo '
         'y Jefes Comerciales. Estampa el Mes Cartola y marca CARTOLA ENVIADA en Post Venta.')
    captura(doc, 'Cartolas → bloque del dealer', 'Cartola con secciones "PARQUE AUTOMALL / PARQUE CERRILLOS / CALLE" y subtotales.')

    # ── 6. Factura ───────────────────────────────────────────────────────────
    h1(doc, '6. La factura: UNA por el total de la cartola')
    ficha(doc, 'Dealer (emite) · Post Venta (registra)', 'postventa_seguimiento', 'Cartola enviada',
          'Post Venta → Seguimiento → etapa FACTURA RECIBIDA')
    paso(doc, 1, 'El dealer factura', 'UNA factura por el TOTAL BRUTO de la cartola (IVA incluido), aunque tenga '
         'operaciones de varios locales o financieras.')
    paso(doc, 2, 'Registrar UNA vez', 'Post Venta marca FACTURA RECIBIDA en UNA sola operación de la cartola: número, '
         'fecha, monto total y el PDF adjunto.')
    paso(doc, 3, 'Réplica automática', 'El sistema replica la factura a las demás operaciones de la misma cartola '
         '(réplicas de la titular; los montos viven una sola vez). Editar la titular también re-replica.')
    paso(doc, 4, 'Boleta de honorarios', 'Si el dealer boletea, el sistema desagrega la retención: la ODP paga el '
         'líquido y la retención queda para el F29.')
    advertencia(doc, 'No subas el mismo PDF en cada operación: basta una vez. Y si ocurre, la ODP igual adjunta el '
                'archivo UNA sola vez (deduplicación por contenido).')
    caso(doc, 'Factura 40 de Osorio del Valle (ago-2026): registrada dos veces como titular generó dos ODP por la '
              'misma plata. De ahí nacieron el candado anti-duplicado (mismo documento en otra operación se rechaza) '
              'y la regla "un documento = una ODP".')

    # ── 7. ODP y pago ────────────────────────────────────────────────────────
    h1(doc, '7. Orden de Pago y pago de la comisión')
    ficha(doc, 'Post Venta (emite) · Tesorería (paga)', 'pv_com_orden_emitir / pagos según Cajas', 'Factura cuadrada contra la cartola',
          'Post Venta → Comisiones a Pagar / Órdenes de Pago')
    paso(doc, 1, 'Emitir la ODP', 'Una ODP por el total del grupo; el número se asigna recién al EMITIR. Si la factura '
         'no cuadra con la cartola, el sistema bloquea en rojo con la diferencia.')
    paso(doc, 2, 'Correo a Contabilidad', 'La ODP viaja con la factura adjunta y la trazabilidad completa al pie '
         '(carta → otorgamiento → fundantes → factura → orden).')
    paso(doc, 3, 'Pagar', 'Tesorería paga a la cuenta bancaria de la ficha del dealer y marca COMISIÓN PAGADA; el '
         'dealer recibe el aviso automático.')
    paso(doc, 4, 'Contabilidad automática', 'Devengo al registrar el documento (COMISION_DEV) y pago contra banco al '
         'pagar (COMISION_PAGADA). Nada se digita a mano.')
    regla(doc, 'Segregación de funciones: quien emite o envía a pago NO puede pagar. El sistema lo exige según la '
               'matriz paramétrica de Cajas.')

    h2(doc, '7.1 Y el parque también cobra lo suyo')
    p(doc, 'Cada operación cursada en un parque genera además la comisión y el arriendo para el DUEÑO de ese parque. '
           'Como la operación quedó marcada con el parque real, entra sola a la cartola mensual de ESE parque '
           '(Post Venta → Emisión de Cartolas Parque): cartola → factura del parque con cuadratura estricta → ODP '
           'mensual → pago. Un dealer multi-local alimenta las cartolas de todos los parques donde cursó, sin ningún '
           'paso manual extra.')

    # ── 8. Resumen ───────────────────────────────────────────────────────────
    h1(doc, '8. Resumen: quién, dónde, qué')
    tabla(doc, ('#', 'Paso', 'Rol', 'Pantalla'),
          (('1', 'Ficha con locales y tablas por parque', 'Ejecutivo Comercial', 'Nueva Ficha'),
           ('2', 'Aprobación + firma + cierre (sella locales)', 'Cadena + Analista', 'Autorizaciones / Mantención'),
           ('3', 'Mantener locales', 'Analista / Admin', 'Dealers → detalle → Locales'),
           ('4', 'Carta en el local del negocio', 'Ejecutivo Comercial', 'Generador de Cartas'),
           ('5', 'Cartola única con secciones por local', 'Post Venta', 'Cartolas'),
           ('6', 'Factura única por el total', 'Dealer + Post Venta', 'Seguimiento (Factura Recibida)'),
           ('7', 'ODP y pago', 'Post Venta + Tesorería', 'Comisiones a Pagar / ODP'),
           ('8', 'Cartola del parque', 'Post Venta', 'Cartolas Parque')),
          (0.8, 5.6, 4.2, 5.5))

    h2(doc, 'Reglas de oro')
    vineta(doc, 'Un dealer = un RUT = una cartola. Lo que se multiplica son sus LOCALES.')
    vineta(doc, 'La ubicación de la operación la define LA CARTA (dónde cursó), nunca la ficha "de hoy".')
    vineta(doc, 'Precedencia del %: tabla del local → tabla histórica → pizarra. La participación de la carta manda sobre todo.')
    vineta(doc, 'Calle no paga comisión de parque: por eso su % de dealer es más alto.')
    vineta(doc, 'Una factura por cartola, registrada UNA vez; el sistema la replica al resto del grupo.')
    vineta(doc, 'Las comisiones se cambian solo por ficha aprobada o mantenedor con permiso — todo auditado.')
    vineta(doc, 'Todo movimiento se contabiliza solo: si algo aparece SIN_REGLA en el log contable, avisar a TI.')
    p(doc, '')
    p(doc, 'Documentación relacionada: Manual de Procesos cap. 3 (cartola → factura → pago), Config Maestro '
           '(Dealers → Locales y comisión por local) y Documentación Técnica (Dealer multi-local, v218.0–v218.9).',
      italic=True, color=GRIS_SUAVE, size=9.5)
    return doc

if __name__ == '__main__':
    doc = construir()
    SALIDA = r'C:\Users\patri\Documents\credit-system\docs\Manual-Dealers-Multi-Local.docx'
    doc.save(SALIDA)
    print('OK ->', SALIDA, '| capturas marcadas:', len(estilo.CAPTURAS))
