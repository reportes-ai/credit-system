# -*- coding: utf-8 -*-
"""Post Venta y Parques — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'POST VENTA Y PARQUES', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Después del curse: pagarle bien al canal, a tiempo,', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'con respaldo tributario y cuadratura exacta.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Qué es ───────────────────────────────────────────────────────────
    h1(doc, '1. Qué es Post Venta')
    p(doc, 'Cuando una operación se cursa, empieza la mitad del negocio que el cliente no ve: '
           'recibir los fondos de la financiera, entregarle al dealer el saldo del precio del '
           'vehículo, pagarle su comisión (y al parque la suya más el arriendo), con cada paso '
           'fechado, cada documento cuadrado y cada peso contabilizado. Eso es Post Venta: el área '
           'donde el compromiso comercial se convierte en plata bien pagada.')
    p(doc, 'El principio del área: al canal se le paga EXACTO y A TIEMPO — ni un peso de más (por '
           'eso las cuadraturas estrictas) ni un día de atraso injustificado (por eso los SLA y el '
           'seguimiento en vivo). La relación con la red se sostiene ahí.')

    # ── 2. Los tracks ───────────────────────────────────────────────────────
    h1(doc, '2. Los tracks de etapas: la columna vertebral')
    p(doc, 'Cada operación otorgada lleva sus rieles de etapas de seguimiento, por dimensión: el '
           'track del SALDO PRECIO (fundantes recibidos → liberado a pago → orden emitida → enviado '
           'a pago → pagado), el de la COMISIÓN del dealer y el del PARQUE. Cada etapa queda con su '
           'fecha, y esas fechas son las que mandan.')
    regla(doc, 'Las etapas NO se marcan a mano: cada una la estampa el módulo dueño del paso — la '
               'cartola marca las etapas de cartola, el módulo de pago marca ORDEN EMITIDA, ENVIADO '
               'A PAGO y PAGADA. Una etapa marcada a mano rompe la trazabilidad y — como las '
               'cartolas se cortan por estas fechas — puede mover plata de mes.')
    advertencia(doc, 'Consecuencia práctica: si una operación "no aparece" en una cartola o en un '
                     'informe de saldos, lo primero es mirar las FECHAS DE SUS ETAPAS, no los '
                     'campos de la ficha del crédito. El universo de cada documento se corta por '
                     'el track, no por la ficha.')

    # ── 3. Saldos precio ────────────────────────────────────────────────────
    h1(doc, '3. El saldo precio: plata en tránsito')
    p(doc, 'La financiera transfiere el saldo del precio del vehículo y AutoFácil se lo entrega '
           'íntegro al dealer. Somos intermediarios: el monto pasa por nuestras cuentas sin ser '
           'ingreso ni gasto (cuenta de paso contable). El compromiso operacional es el SLA por '
           'categoría del dealer — 24/48/72 horas hábiles según sea Super Partner, Partner o Socio '
           '(recepción en la tarde cuenta desde el día hábil siguiente).')
    vineta(doc, 'definir fondos → armar la nómina → emitir la orden → pagar. Al pagarse, los informes salen solos a Operaciones y al ejecutivo, y el correo al dealer.', bold_hasta='El flujo: ')
    vineta(doc, 'la card de Tesorería reconstruye en vivo la planilla que antes se llevaba a mano: recepción con AM/PM, vencimiento por SLA, número de orden, carga al banco y pago.', bold_hasta='Seguimiento en vivo: ')
    vineta(doc, 'el dealer ve este mismo pipeline en su portal — las cuatro tarjetas del inicio — y deja de llamar a preguntar cuándo le pagan.', bold_hasta='Autoservicio: ')
    captura(doc, 'Tesorería → Saldo Precio en Proceso de Pago', 'La card con operaciones en distintos estados y sus SLA.')

    # ── 4. Circuito dealer ──────────────────────────────────────────────────
    h1(doc, '4. El circuito del dealer: cartola → factura → pago')
    flujo(doc, 'A PAGAR → CARTOLA ENVIADA → FACTURA CUADRADA → ORDEN DE PAGO → PAGADA')
    p(doc, 'La comisión del dealer nace con la carta (la carta manda), entra a su cartola mensual, '
           'el dealer factura el TOTAL exacto de la cartola, y la orden de pago cierra el circuito. '
           'Tres reglas sostienen la limpieza del proceso:')
    vineta(doc, 'una operación cuya etapa no sea OTORGADO — anulada, desistida, rechazada o inexistente — no puede estar en una cartola. Nació de una auditoría que encontró $2,7 MM esperando pago sobre negocios que nunca se cursaron.', bold_hasta='Regla de oro: ')
    vineta(doc, 'las operaciones se agrupan por el RUT del dealer y el nombre sale de su ficha. El mismo RUT escrito de dos formas generaba dos cartolas, dos facturas y dos pagos.', bold_hasta='Una cartola por RUT: ')
    vineta(doc, 'la comisión calculada es BRUTA (IVA incluido) y se desagrega hacia abajo. Con factura se paga el bruto; con boleta de honorarios, el líquido con su retención al F29. Jamás se suma IVA encima.', bold_hasta='El IVA va incluido: ')
    h2(doc, '4.1 Los caminos de excepción')
    vineta(doc, 'las operaciones que entraron por carga masiva no generan carta, así que su comisión no llegaba a ninguna cartola. La pantalla de incorporación las confirma dato a dato contra el motor: las que calzan entran solas, las que difieren pasan por aprobación del analista.', bold_hasta='Otorgadas sin carta: ')
    vineta(doc, 'anular una operación retira su comisión de la cartola en el mismo acto (con doble firma). Si la cartola ya se envió, el movimiento no se toca: se regulariza como descuento, no se borra a espaldas del dealer.', bold_hasta='Anulaciones: ')
    vineta(doc, 'al corregir el dealer de una carta aprobada, la corrección alcanza la carta, sus movimientos no enviados y el crédito — los tres a la vez, para que no queden dos verdades.', bold_hasta='Correcciones: ')
    captura(doc, 'Post Venta → Cartolas', 'Una cartola de dealer con sus movimientos y el estado de la factura.')

    # ── 5. Parques ──────────────────────────────────────────────────────────
    h1(doc, '5. Los parques automotrices')
    p(doc, 'El parque es una contraparte comercial propia: el recinto donde operan varios dealers. '
           'Gana por dos vías — una comisión por cada crédito cursado en el parque y un arriendo '
           'mensual fijo — y su circuito replica el del dealer, agregado por parque.')
    h2(doc, '5.1 La entidad y su incorporación')
    p(doc, 'Un parque se incorpora por la MISMA máquina de la ficha del dealer: informes de la '
           'empresa y los socios, análisis de IA, cadena de niveles, firma del representante y '
           'cierre. La diferencia es el final: al cerrar se crea el parque con arriendo y comisión '
           'en cero, que se fijan en el mantenedor Arriendos y Comisiones.')
    advertencia(doc, 'El parque de cada OPERACIÓN viene del Excel de carga (o de la carta), no del '
                     'mantenedor: si la operación no trae su parque, no genera comisión de parque. '
                     'Y si la carta dice parque y el crédito quedó como calle (o al revés), el '
                     'sistema lo muestra en ámbar y lo decide una persona — esa diferencia cambia '
                     'el tramo de comisión.')
    h2(doc, '5.2 La cartola del parque y su foto')
    p(doc, 'A la cartola del mes T entran las operaciones otorgadas en T o antes (arrastre) cuya '
           'comisión no se ha cobrado y cuyo saldo precio quedó al menos LIBERADO A PAGO al cierre '
           'de T — una operación liberada tarde no se pierde: entra sola al mes siguiente.')
    regla(doc, 'Al aprobar el mes queda una FOTO por operación: qué cartola cobró qué operación. '
               'Esa foto es lo que se imprime, factura y paga — el documento no cambia aunque las '
               'etapas avancen después. Es el mismo principio del snapshot deliberado que rige '
               'cartolas de dealers y finiquitos.')
    paso(doc, 1, 'Aprobar el mes y emitir', 'La cartola lista las operaciones más la línea de '
         'arriendo (que es del parque, no de un crédito).')
    paso(doc, 2, 'Enviar', 'El correo sale del servidor con el PDF adjunto al contacto de finanzas '
         'de la ficha del parque, con copia a operaciones.')
    paso(doc, 3, 'Facturar', 'El parque factura el TOTAL; la cuadratura es estricta — si no calza, '
         'se rechaza con la diferencia. Al cuadrar, FACTURA RECIBIDA en todas las operaciones.')
    paso(doc, 4, 'Pagar', 'La orden de pago mensual nace de la factura (solo con factura en todas '
         'las operaciones) y al pagarse marca COMISIÓN PAGADA.')
    captura(doc, 'Post Venta → Emisión de Cartolas Parque', 'Una cartola de parque con operaciones, arriendo y total.')
    h2(doc, '5.3 Reversas y contabilidad')
    vineta(doc, 'el envío se puede reversar (permite reenviar), la orden no pagada se anula (el correlativo muere, nunca se reutiliza) y el pago se puede revertir. Toda reversa con motivo auditado.', bold_hasta='Reversable con motivo: ')
    vineta(doc, 'al aprobar el mes se devenga el gasto (arriendo y comisión en cuentas separadas contra el pasivo); al pagar, el pasivo se rebaja contra banco. Cada asiento es idempotente: aprobar dos veces no duplica.', bold_hasta='Asientos automáticos: ')
    vineta(doc, 'los avisos (a Contabilidad al emitir, al parque y al equipo al pagar) son paramétricos: texto, destinatarios e interruptor por evento, sin tocar código.', bold_hasta='Avisos configurables: ')
    caso(doc, 'La historia previa al módulo quedó regularizada en agosto 2026: 908 comisiones '
              'pagadas por planilla en 98 cartolas parque-mes (2025 completo y enero-julio 2026), '
              'por $106,5 MM más $13,5 MM de arriendos, marcadas pagadas sin duplicar su '
              'contabilidad — el pago ya había ocurrido y contabilizarlo de nuevo lo habría '
              'duplicado. Las operaciones con saldo precio pendiente quedaron fuera a propósito: '
              'arrastran solas a la cartola que corresponda.')

    # ── 6. La medición ──────────────────────────────────────────────────────
    h1(doc, '6. La medición del canal')
    vineta(doc, 'quién del concesionario trae cada negocio, desde la carta (no desde la carga, que escribe a nuestro ejecutivo). Abre en el último mes cerrado.', bold_hasta='Vendedores con Ventas: ')
    vineta(doc, 'qué dealer rinde bajo su potencial según su categoría — la agenda de la conversación comercial.', bold_hasta='Potencial de Dealers: ')
    vineta(doc, 'la cuadratura mensual contra el export del canal valida los montos de cada operación. Su columna de comisión dealer se ignora SIEMPRE: no es la nuestra.', bold_hasta='Cuadratura Trinidad: ')

    # ── 7. Gobierno ─────────────────────────────────────────────────────────
    h1(doc, '7. Gobierno')
    vineta(doc, 'Analistas y Gerencia de Operaciones operan (cartolas, órdenes); Tesorería y Gerencia de Finanzas pagan. Quien emite no paga.', bold_hasta='Segregación: ')
    vineta(doc, 'las cartolas y comisiones son puntos del checklist de Cierre de Mes: el mes no cierra con el canal impago sin explicación.', bold_hasta='Cierre de mes: ')
    vineta(doc, 'todo el circuito escribe en Auditoría; las cartolas enviadas y las fotos aprobadas son los snapshots que un reclamo de dealer puede revisar meses después.', bold_hasta='Trazabilidad: ')

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\PostVenta-Parques-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
