# -*- coding: utf-8 -*-
"""Te presentamos Business Suite — capítulo de presentación conceptual."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'TE PRESENTAMOS', bold=True, color=AZUL, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'BUSINESS SUITE', bold=True, color=AZUL_OSCURO, size=34, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'El sistema que administra el negocio completo de AutoFácil:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'de la prospección del cliente a la contabilidad, en una sola plataforma.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Qué es ───────────────────────────────────────────────────────────
    h1(doc, '1. Qué es Business Suite')
    p(doc, 'Business Suite es el sistema de gestión integral de AutoFácil, construido a la medida del '
           'negocio de crédito automotriz. En una sola plataforma web viven la venta, la evaluación, '
           'la originación, el pago a la red de dealers, la cobranza, la contabilidad, las personas y '
           'el gobierno del sistema — lo que antes exigía un sistema contable externo, decenas de '
           'planillas Excel y correos que nadie podía auditar.')
    p(doc, 'Opera en producción desde el 1 de agosto de 2026 con toda la operación diaria adentro: '
           'más de 18.000 operaciones administradas, la red completa de dealers, y cada acción '
           'crítica registrada en auditoría. Corre en la nube, no requiere instalación, y funciona '
           'igual en el computador de la oficina, el celular del ejecutivo en terreno y la pantalla '
           'del directorio.')
    h2(doc, '1.1 La idea central')
    caja(doc, 'Un solo lugar, una sola verdad',
         'Cada dato del negocio vive en un solo lugar y cada cálculo se hace en un solo motor. La '
         'cuota que ve el cliente, la que cobra la caja y la que muestra el certificado son EL MISMO '
         'número, porque salen de la misma función. Cuando algo se corrige, se corrige una vez y '
         'queda corregido en todas partes.')
    p(doc, 'Esa idea se sostiene en cuatro máximas de construcción, que son la personalidad del sistema:')
    vineta(doc, 'cada magnitud (cuota, mora, comisión, rentabilidad) tiene una única función que usan todas las pantallas.', bold_hasta='Un solo motor por cálculo: ')
    vineta(doc, 'nombres, direcciones, comisiones y estados tienen un hogar autoritativo; todo lo demás lo referencia, nunca lo copia.', bold_hasta='Una sola fuente de datos: ')
    vineta(doc, 'cada proceso se documenta en el mismo momento en que se construye. La Suite de Documentación es parte del sistema, no un anexo.', bold_hasta='Todo proceso nace documentado: ')
    vineta(doc, 'cada peso que entra o sale genera su asiento contable automático. La contabilidad se construye sola desde la operación.', bold_hasta='Todo movimiento se contabiliza: ')
    h2(doc, '1.2 Paramétrico: el Administrador manda, no el programador')
    p(doc, 'El principio rector es que el usuario Administrador pueda modificar el negocio sin tocar '
           'código: tasas, comisiones, umbrales, estados, plantillas de correo, permisos, plazos y '
           'textos viven en MANTENEDORES — pantallas de configuración con su permiso y su '
           'trazabilidad. El flujo de los procesos está protegido (las validaciones y atribuciones '
           'no se pueden saltar), pero su contenido se ajusta desde la pantalla. Cada mantenedor '
           'termina con un recuadro "Qué afecta este mantenedor" que dice exactamente qué proceso '
           'impacta cada variable.')

    # ── 2. Beneficios ───────────────────────────────────────────────────────
    h1(doc, '2. Los beneficios, en concreto')
    tabla(doc, ('Beneficio', 'Cómo se materializa'),
          (('Consistencia', 'El mismo número en todas las pantallas: motores únicos y fuente única de datos. Se acabaron las planillas que no cuadran'),
           ('Velocidad', 'La carta se aprueba sola cuando cuadra (Revisor Automático); la cartola se arma sola; la contabilidad se centraliza sola; el bot cotiza solo'),
           ('Control', 'Doble firma donde hay plata (anulaciones, castigos, pagos), segregación de funciones en el servidor, y auditoría de cada acción crítica'),
           ('Trazabilidad', 'Toda decisión queda con quién, cuándo y por qué: bitácoras inmutables, firmas electrónicas con huella, documentos verificables con QR'),
           ('Autonomía del negocio', 'Los parámetros se cambian en mantenedores, sin programador y sin ventanas de mantención'),
           ('Menos digitación', 'Los documentos se leen solos (cartas PDF, RCV del SII, pantallazos), lo leído se pre-llena, y lo que falta cae a una cola ordenada'),
           ('Autoservicio', 'El dealer ve sus pagos en su portal, el cliente sus cuotas en el suyo, el colaborador sus solicitudes en su ficha — menos llamadas para preguntar'),
           ('Cumplimiento', 'Ley del Consumidor en cobranza, Código del Trabajo en RRHH, SII en lo tributario: las reglas legales están cableadas en los motores')),
          (3.6, 12.9))

    # ── 3. La operatoria ────────────────────────────────────────────────────
    h1(doc, '3. La operatoria: el viaje de una operación')
    p(doc, 'La mejor forma de entender el sistema es seguir el viaje del negocio. Cada estación del '
           'viaje es un conjunto de módulos, y cada módulo tiene su manual detallado en la Suite de '
           'Documentación.')
    flujo(doc, 'PROSPECTAR → COTIZAR → EVALUAR → APROBAR → CURSAR → PAGAR AL CANAL → COBRAR → CONTABILIZAR')

    h2(doc, '3.1 Prospección y venta')
    vineta(doc, 'gestiones comerciales y campañas masivas por correo o WhatsApp, con deciles de control para medir si la campaña realmente movió la aguja. Cada contacto queda como gestión.', bold_hasta='CRM y Campañas: ')
    vineta(doc, 'el bot de WhatsApp del número oficial: cotiza con el motor real (nunca inventa cifras), pre-evalúa con informes, responde dónde pagar y deriva a ejecutivos cuando corresponde. Atiende 24/7.', bold_hasta='Facilito: ')
    vineta(doc, 'precio, pie y plazo se convierten en cuota con gastos, seguros y CAE. Cuatro puertas (módulo, Simulador Rápido, portal del dealer, Facilito), un solo motor.', bold_hasta='Cotizaciones: ')
    vineta(doc, 'la PWA que le dice al ejecutivo, en el patio, en qué financiera conviene cursar — elegibilidad primero, rentabilidad después. Elimina el "lucro cesante" de cursar donde no convenía.', bold_hasta='¿Dónde Curso?: ')
    h2(doc, '3.2 Evaluación y aprobación')
    vineta(doc, 'el RUT trae la ficha completa; los informes comerciales se piden en línea y la IA cruza renta, deudas y política de crédito para entregar el análisis.', bold_hasta='Evaluación Crediticia: ')
    vineta(doc, 'el documento que compromete las condiciones al dealer, con correlativo, firma electrónica y QR verificable. Se autocompleta desde el PDF de la financiera.', bold_hasta='Cartas de Aprobación: ')
    vineta(doc, 'Business Suite compara el documento contra lo digitado y aprueba sola las cartas que cuadran, con checklist firmado. El analista humano queda para los casos con diferencias.', bold_hasta='Revisor Automático: ')
    vineta(doc, 'cuando el negocio necesita salirse de la pizarra, el ejecutivo simula la jugada y el sistema mismo autoriza con un código si se respeta el piso de rentabilidad. Autoservicio con presupuesto de estrellas y auditoría.', bold_hasta='Excepciones Comerciales: ')
    h2(doc, '3.3 Curse y respaldo')
    vineta(doc, 'la producción del canal entra por Excel con validador de anomalías, sin duplicar y con las diferencias acusadas para que una persona decida.', bold_hasta='Carga Masiva: ')
    vineta(doc, 'lo incompleto cae a una cola con bloqueo por usuario y pre-llenado desde la carta.', bold_hasta='Digitación de Faltantes: ')
    vineta(doc, 'los documentos que la financiera exige para liberar los fondos, con matriz por financiera, validación de Operaciones, devoluciones y rendición semanal automática.', bold_hasta='Seguimiento de Fundantes: ')
    h2(doc, '3.4 Pago al canal (Post Venta)')
    vineta(doc, 'el estado de cuenta mensual de comisiones de cada dealer y cada parque: se emite, se envía, el canal factura el total exacto y la orden de pago cierra el circuito. Cuadratura estricta: si la factura no calza, no avanza.', bold_hasta='Cartolas: ')
    vineta(doc, 'la plata del vehículo que pasa por AutoFácil hacia el dealer, con SLA por categoría (24/48/72 horas hábiles) y seguimiento en vivo en Tesorería.', bold_hasta='Saldos Precio: ')
    vineta(doc, 'toda salida de plata con correlativo único, doble control y segregación de funciones: quien emite no paga.', bold_hasta='Órdenes de Pago: ')
    vineta(doc, 'el anticipo de comisiones de los dealers Super Partner, renovado mes a mes con cadena de aprobación.', bold_hasta='Plan Liquidez: ')
    h2(doc, '3.5 La cartera propia: recaudar')
    vineta(doc, 'el cajero busca el crédito, cobra cuotas individuales o en lote, y el comprobante sale timbrado. La mora y los gastos son los del motor único. Cierre de caja con cuadratura diaria.', bold_hasta='Cajas: ')
    vineta(doc, 'el cliente entra con su RUT y un código al correo, ve sus cuotas, su mora al día y sus comprobantes — los mismos valores de la caja.', bold_hasta='Portal del Cliente: ')
    vineta(doc, 'saldar antes de plazo, al valor del motor o con descuento negociado por cadena de firmas. El orden de condonación está protegido: primero gastos, luego intereses, capital nunca.', bold_hasta='Prepago y Aplicación de Fondos: ')
    h2(doc, '3.6 Cobranza: recuperar cumpliendo la ley')
    vineta(doc, 'correos por tramo y secuencias de WhatsApp actúan solos en los primeros días de atraso, con tope legal de gestiones bloqueado por sistema.', bold_hasta='Mora temprana automática: ')
    vineta(doc, 'gestión humana con compromisos de pago, y la demanda con abogado, tribunal y etapa procesal cuando corresponde. Todo en la bitácora CRM.', bold_hasta='Pre-judicial y judicial: ')
    vineta(doc, 'cada tramo de mora provisiona su porcentaje automáticamente al cierre; el asiento va por la variación mensual.', bold_hasta='Provisiones: ')
    vineta(doc, 'la baja contable del incobrable es deliberadamente manual, con doble firma gerencial. Y si el castigado paga después, el RECUPERO entra por caja como cualquier abono y genera su asiento — la historia del crédito nunca se pierde.', bold_hasta='Castigo y recupero: ')
    h2(doc, '3.7 Contabilidad: el reemplazo del sistema externo')
    vineta(doc, 'cada hecho económico genera su asiento según reglas paramétricas. Devengo y pago son hechos separados; los impuestos van en el asiento; el motor nunca bloquea la operación.', bold_hasta='Centralización automática: ')
    vineta(doc, 'plan de cuentas, comprobantes, libros, balance y los libros históricos importados.', bold_hasta='Contabilidad completa: ')
    vineta(doc, 'las facturas de proveedores entran solas desde el Registro de Compras del SII, con la cuenta propuesta por historial. El F29 sale del sistema.', bold_hasta='RCV del SII: ')
    vineta(doc, 'checklist con responsables, provisiones automáticas, candado contable y acta congelada. Nada retroactivo se toca después del cierre.', bold_hasta='Cierre de Mes: ')
    h2(doc, '3.8 La red de dealers')
    vineta(doc, 'ningún dealer opera sin ficha aprobada: informes, análisis IA, cadena de niveles y firma. Su tabla de comisiones rige desde el cierre.', bold_hasta='Incorporación: ')
    vineta(doc, 'cada dealer ve sus operaciones, el estado de sus pagos, sus cartolas y una pre-aprobación en línea — sin ver jamás datos del cliente.', bold_hasta='Portal del Dealer: ')
    vineta(doc, 'chat en vivo con la mesa, envío de documentos y videollamada. Y el staff puede "ver como dealer" para dar soporte mirando la misma pantalla.', bold_hasta='Atención Remota: ')
    vineta(doc, 'las visitas a terreno con calendario, planificador de rutas óptimas, asignación de cartera y la app móvil con check-in GPS.', bold_hasta='Ruta AutoFácil: ')
    h2(doc, '3.9 Las personas')
    vineta(doc, 'el ciclo laboral completo sin papel: carta oferta, contrato con firma electrónica, remuneraciones con comisiones que llegan solas, vacaciones con cuenta corriente, asistencia contra el reloj control, evaluación de desempeño, y el finiquito con sus topes legales y aviso automático a Previred.', bold_hasta='RRHH: ')
    vineta(doc, 'piso, incentivos y bonos por venta de seguros se calculan solos y van directo a la liquidación, con variables versionadas y bitácora.', bold_hasta='Comisiones: ')
    vineta(doc, 'la Academia enseña cada módulo a su ritmo, los cursos formales quedan en la ficha, el test vocacional apoya la selección, y los concursos hacen del aprendizaje un juego.', bold_hasta='Capacitación: ')
    vineta(doc, 'canal de denuncias confidencial con plazos legales (Ley Karin) y encuestas de clima anónimas con resultados agregados.', bold_hasta='Compliance y clima: ')
    h2(doc, '3.10 Soporte y gobierno del sistema')
    vineta(doc, 'perfiles y funcionalidades definen quién ve y hace qué. La matriz se administra en pantalla; el Administrador ve todo; los permisos sensibles se validan en el servidor.', bold_hasta='Usuarios y permisos: ')
    vineta(doc, 'tickets de TI con SLA y escalamiento, compras internas con workflow de aprobación por niveles, y 15 flujos de escalamiento que suben solos por la jefatura cuando algo no se atiende.', bold_hasta='Soporte interno: ')
    vineta(doc, 'campanitas y correos paramétricos: qué evento avisa, a quién y con qué texto se configura, no se programa.', bold_hasta='Avisos: ')
    vineta(doc, 'cada acción crítica con quién, cuándo y desde dónde; consola SQL de solo lectura para respuestas rápidas; salud del sistema con monitoreo, respaldos diarios y un plan de contingencia ensayado.', bold_hasta='Auditoría y continuidad: ')
    vineta(doc, 'los documentos vivos (usuario, procesos, configuración, técnica) más la biblioteca Word por rol — este documento incluido.', bold_hasta='Documentación: ')

    # ── 4. Cierre ───────────────────────────────────────────────────────────
    h1(doc, '4. Dónde estamos y hacia dónde va')
    p(doc, 'Business Suite pasó de proyecto a columna vertebral de la operación: 30 usuarios diarios, '
           'la red completa de dealers, la contabilidad construyéndose sola y los procesos de '
           'personas dentro del sistema. Lo que sigue es profundidad: cerrar el reemplazo total del '
           'sistema contable externo (emisión electrónica ante el SII), extender el doble factor a '
           'todo el equipo, y seguir convirtiendo cada decisión repetitiva en un motor que decide '
           'solo y deja al equipo lo que de verdad necesita criterio humano.')
    p(doc, 'La regla que resume todo: si algo se hace dos veces igual, el sistema debería hacerlo; '
           'si algo requiere juicio, el sistema debería preparar todo para que la persona decida '
           'con la mejor información. Eso es Business Suite.', bold=True)

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Te-Presentamos-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
