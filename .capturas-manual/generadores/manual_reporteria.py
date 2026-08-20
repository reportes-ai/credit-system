# -*- coding: utf-8 -*-
"""Reportería y Dashboards — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'REPORTERÍA Y DASHBOARDS', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Del dato a la decisión: cómo el negocio se mira a sí mismo.', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── 1. Principios ───────────────────────────────────────────────────────
    h1(doc, '1. Los principios: por qué los números cuadran')
    p(doc, 'La reportería del sistema no es una capa aparte que copia datos: lee las MISMAS tablas y '
           'usa los MISMOS motores que la operación. De ahí salen sus reglas:')
    vineta(doc, 'toda cifra sale de la fuente única y del motor único. La rentabilidad del dashboard es la misma de la carta; la mora del informe es la misma de la caja. Si dos informes difieren, es un bug — no "otra metodología".', bold_hasta='Consistencia = confianza: ')
    vineta(doc, 'los totales y conteos se calculan en el servidor sobre TODO el universo, nunca sumando la página visible. Una tabla paginada muestra 100 filas, pero sus estadísticas son de las 18.000.', bold_hasta='Stats de verdad: ')
    vineta(doc, 'los informes comparativos abren por defecto en el último mes CERRADO: el mes en curso está incompleto y leerlo como caída de ventas es el error clásico. El mes corriente se muestra siempre marcado como parcial.', bold_hasta='El mes en curso engaña: ')
    vineta(doc, 'un mes cerrado tiene candado: sus cifras no cambian aunque se recalcule el presente. Lo que se informó a la matriz sigue siendo lo informado.', bold_hasta='Lo congelado no se mueve: ')
    vineta(doc, 'prácticamente toda vista tiene botón Excel, y los exports respetan los mismos filtros y motores de la pantalla.', bold_hasta='Todo exporta: ')

    # ── 2. Dashboard principal ──────────────────────────────────────────────
    h1(doc, '2. El Dashboard principal')
    p(doc, 'La vista de comando del negocio de colocación: producción aprobada y otorgada, montos, '
           'instituciones, rentabilidad y penetración de seguros, con filtros por fecha y '
           'comparación entre períodos.')
    ficha(doc,
          'Gerencia y jefaturas · las pestañas visibles se configuran por permiso',
          'Configuración de tabs por perfil (dashboard_config)',
          'Producción cargada al día',
          '/dashboard/')
    h2(doc, '2.1 Qué muestra')
    vineta(doc, 'las cartas aprobadas del período, con conversión.', bold_hasta='Tab Aprobados: ')
    vineta(doc, 'la producción cursada: operaciones, montos, split AutoFin/Unidad, mayor/menor 200 UF.', bold_hasta='Tab Otorgados: ')
    vineta(doc, 'ingreso por colocación, seguros y comisiones — con el semáforo paramétrico (verde/amarillo/naranjo).', bold_hasta='Tab Rentabilidades: ')
    h2(doc, '2.2 Las reglas de cálculo que hay que conocer')
    vineta(doc, 'si un número de operación se cargó más de una vez, el dashboard muestra UNA sola (deduplicación en la consulta).', bold_hasta='Deduplicación por operación: ')
    vineta(doc, 'la clasificación se recalcula con la UF de la FECHA DE OTORGAMIENTO de cada crédito, no con la UF de hoy — una operación no cambia de tramo porque la UF subió.', bold_hasta='Mayor/menor 200 UF: ')
    vineta(doc, 'metas de prospección, umbrales de rojo y semáforo de rentabilidad viven en el mantenedor Parámetros del Dashboard — antes estaban en el código.', bold_hasta='Umbrales paramétricos: ')
    captura(doc, 'Dashboard → Otorgados', 'El tab de otorgados con sus KPIs y el filtro de fechas.')

    # ── 3. Autoservicio ─────────────────────────────────────────────────────
    h1(doc, '3. Reportería de autoservicio')
    h2(doc, '3.1 Reportería Tailor Made')
    p(doc, 'Los informes armados a medida del negocio, listos para filtrar y exportar. Es la '
           'estación intermedia: más flexible que un dashboard, más guiada que el constructor.')
    h2(doc, '3.2 Diseño de Consulta: el constructor visual')
    p(doc, 'El poder de armar una consulta propia sin saber SQL, al estilo del diseñador de '
           'consultas de Access: se eligen tablas, campos, filtros y agrupaciones en pantalla, y la '
           'consulta queda guardada para repetirla. Para el usuario avanzado que pregunta cosas que '
           'ningún informe fijo previó.')
    captura(doc, 'Reportería → Diseño de Consulta', 'El constructor con una consulta armada: tablas, campos y filtros.')
    h2(doc, '3.3 Pregúntale a AutoFácil y Pregúntale a Finanzas (BI conversacional)')
    p(doc, 'La pregunta en lenguaje natural: "¿cuántas operaciones cursó el parque X en julio?" y la '
           'respuesta con la cifra real del sistema. Dos sabores: el general (datos de negocio) y el '
           'financiero (estados y cifras contables, con modelo mayor). Gobernados por el subsistema '
           'de IA, con su gasto medido (ver documento de IA). Complemento del autoservicio: la '
           'pregunta puntual ya no espera a que alguien arme el informe.')

    # ── 4. Vistas ejecutivas ────────────────────────────────────────────────
    h1(doc, '4. Vistas ejecutivas y de sala')
    vineta(doc, 'el panel de TV estilo sala de control (/mando): los indicadores del negocio en vivo para la pantalla de la oficina.', bold_hasta='Cuadro de Mando: ')
    vineta(doc, 'la página de inicio personal con widgets por perfil: pendientes, agenda, indicadores del rol. Cada uno ve su día, no el de todos.', bold_hasta='Mi Día: ')
    vineta(doc, 'el correo automático de la mañana con lo esencial del día anterior, redactado por el subsistema de IA sobre cifras reales.', bold_hasta='Resumen Ejecutivo Diario: ')
    vineta(doc, 'el informe mensual de Cierre Contable a la matriz (Ecuador), con sus seis secciones validadas contra la base antes de enviarse.', bold_hasta='Informe a Casa Matriz: ')
    captura(doc, '/mando (Cuadro de Mando TV)', 'El panel de sala con los indicadores en vivo.')

    # ── 5. Informes por área ────────────────────────────────────────────────
    h1(doc, '5. Los informes por área')
    tabla(doc, ('Área', 'Informe', 'Qué responde'),
          (('Comercial', 'Rentabilidad de Créditos', 'Cuánto dejó cada operación y si se cursó donde convenía (audita con los mismos umbrales de ¿Dónde Curso?)'),
           ('Comercial', 'Penetración de Seguros', 'Qué % de las operaciones elegibles lleva cada seguro — el mismo universo del bono del ejecutivo'),
           ('Comercial', 'Vendedores con Ventas', 'Quién del dealer trae cada negocio (abre en el último mes cerrado)'),
           ('Comercial', 'Estadísticas de Ruta', 'Cumplimiento y positividad de las visitas por ejecutivo'),
           ('Comercial', 'Potencial de Dealers', 'Qué dealer rinde bajo su potencial según su categoría'),
           ('Operaciones', 'Historial de Cargas / Diferencias', 'Qué entró, qué se omitió y qué discrepa con el canal'),
           ('Cobranza', 'Cartera y mora por tramo', 'Cómo está la cartera propia, con el motor único de mora'),
           ('Cobranza', 'Score de Mora', 'Mora histórica por segmento — alimenta la política de crédito'),
           ('Tesorería', 'Saldo Precio en Proceso de Pago', 'Qué se debe al canal y su SLA, en vivo desde las etapas'),
           ('Contabilidad', 'Libros, balance y F29', 'La contabilidad completa, construida sola desde la operación'),
           ('RRHH', 'Libro de Remuneraciones / BSC', 'Sueldos, comisiones y el bono del Jefe Comercial con su detalle'),
           ('TI / Gerencia', 'Salud del Sistema', 'Uptime, motores, respaldos y el gasto de IA en vivo')),
          (2.8, 5.0, 8.7))
    p(doc, 'Los informes generados por IA comparten un formato único (el estándar de Reporte IA del '
           'sistema): mismo encabezado, misma estructura, mismo pie — un informe de máquina se '
           'reconoce a la vista y siempre declara sus fuentes.')

    # ── 6. Gobierno ─────────────────────────────────────────────────────────
    h1(doc, '6. El gobierno de la reportería')
    vineta(doc, 'cada informe y dashboard respeta la matriz: quien no tiene el permiso no ve la card, y la visibilidad de datos se acota (un supervisor ve su equipo; Gerencia, todo).', bold_hasta='Permisos: ')
    vineta(doc, 'la reportería tiene su propio techo de peticiones (120/min) porque la base cobra por consulta: un informe pesado repetido sin control es plata.', bold_hasta='Rate limit propio: ')
    vineta(doc, 'los sondeos frecuentes se cachean; los datos que no cambian cada minuto no se consultan cada minuto.', bold_hasta='Caché: ')
    vineta(doc, 'si un informe nuevo necesita una cifra que ya existe en otra pantalla, usa el mismo motor — la máxima 1 aplica a la reportería igual que a la operación. Un informe que "recalcula por su cuenta" es un bug en gestación.', bold_hasta='Sin segundas verdades: ')

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Reporteria-Dashboards-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
