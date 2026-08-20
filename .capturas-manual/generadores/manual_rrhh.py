# -*- coding: utf-8 -*-
"""Manual de Recursos Humanos — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE RECURSOS HUMANOS', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'De la carta oferta al finiquito:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'el ciclo laboral completo, sin papel y con respaldo legal.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Recursos Humanos'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Cada capítulo: para qué existe, quién lo hace, prerequisitos, paso a paso y los recuadros '
           '⚠ OJO / 🔒 Regla del sistema / 🧾 Caso real / 📸 CAPTURA (pendientes en esta versión). '
           'Las rutas se escriben "Módulo → Card".')
    h1(doc, 'Índice')
    toc(doc)

    # ── 1. El ciclo y su respaldo legal ─────────────────────────────────────
    h1(doc, '1. El ciclo laboral paperless y su respaldo legal')
    p(doc, 'Contratos, finiquitos, vacaciones, préstamos y modificaciones viajan sin papel: solicitud, '
           'cadena de aprobación con firma electrónica simple (FES) y ejecución automática. Cada firma '
           'registra identidad por sesión, nombre, cargo, fecha y hora, IP y huella SHA-256 del '
           'contenido — evidencia de que el documento no se alteró después de firmado.')
    h2(doc, '1.1 El marco (resumen para no abogados)')
    tabla(doc, ('Documento', 'Base legal', 'Qué exige'),
          (('Todo documento electrónico', 'Ley 19.799, art. 3 y 5', 'La FES vale como el papel y es admisible como prueba'),
           ('Contrato de trabajo', 'Art. 9 CT', '"Por escrito" — la equivalencia funcional lo cubre; conservar PDF para terceros (bancos, visas)'),
           ('Finiquito', 'Art. 177 CT', 'La FES deja el acuerdo; el poder liberatorio exige ratificación ante ministro de fe o el finiquito electrónico de la DT'),
           ('Descuentos al personal', 'Art. 58 CT', 'Acuerdo escrito (la solicitud FES lo es) y tope 15% de la remuneración'),
           ('Vacaciones', 'Arts. 67-73 CT', 'La solicitud/aprobación FES es el comprobante de feriado'),
           ('Aumentos y ascensos', 'Art. 11 CT', 'Anexo firmado por ambas partes — generar desde Contratos y enviar A firma')),
          (3.6, 3.4, 9.5))
    regla(doc, 'Regla transversal: nadie aprueba ni firma su propia solicitud. Toda decisión queda en '
               'Auditoría además de la firma FES.')

    # ── 2. Contratación ─────────────────────────────────────────────────────
    h1(doc, '2. Contratación y onboarding')
    p(doc, 'Del candidato al contrato firmado sin que lo ofertado y lo contratado se contradigan '
           'jamás: la carta oferta es la fuente única del contrato.')
    ficha(doc,
          'RRHH o jefatura con permiso (carta y contrato) · Gerencia firma · el candidato acepta',
          'rh_colaboradores · contrafirma con rh_aprobar',
          'Cargo creado CON descripción de cargo · catálogo de beneficios al día',
          'RRHH → Contratos y Cartas Oferta')
    paso(doc, 1, 'Carta oferta', 'Candidato, cargo, sueldo y beneficios por checkbox (más otros '
         'libres) → Emitir → imprimir o enviar. Marcar después Aceptada o Rechazada.')
    paso(doc, 2, 'Contratar (solo con carta ACEPTADA)', 'Genera el contrato con la descripción de '
         'cargo y las cláusulas de lo ofertado. Firma en papel o FES: botón "A firma" congela el '
         'documento y avisa al trabajador; él firma en Mi Ficha → Documentos por firmar; contrafirma '
         'la empresa; queda FIRMADO con huella de integridad. Requiere que el trabajador ya tenga '
         'cuenta.')
    paso(doc, 3, 'Alta y onboarding', 'Al contratar se crea solo el checklist de onboarding (usuario, '
         'correo, Workera, Edenred, credencial, ficha Previred, inducción) con responsables y plazos. '
         'Crear el usuario y completar la ficha — los campos con asterisco son los que exige '
         'Previred. El Archivo Previred del mes informa solo la contratación (movimiento 1).')
    captura(doc, 'RRHH → Contratos y Cartas Oferta', 'Una carta oferta emitida y el botón Contratar.')
    flujo(doc, 'BORRADOR → EMITIDA → ACEPTADA → CONTRATADA')
    advertencia(doc, 'Sin descripción de cargo el sistema alega a diario — y el contrato la necesita. '
                     'Crear el cargo completo ANTES de la carta oferta.')

    # ── 3. Remuneraciones ───────────────────────────────────────────────────
    h1(doc, '3. Remuneraciones: el libro de liquidaciones')
    p(doc, 'Todo llega solo al libro: comisiones aprobadas, días trabajados desde ingreso y '
           'licencias, adicionales y descuentos. RRHH revisa la vista previa por persona y emite.')
    ficha(doc,
          'El motor calcula · supervisores aprueban comisiones · RRHH emite el libro',
          'rh_remuneraciones',
          'Comisiones del mes APROBADAS (sin eso el libro avisa en rojo) · indicadores previsionales '
          'al día',
          'RRHH → Remuneraciones')
    paso(doc, 1, 'Revisar', 'Vista previa por persona: sueldo, comisiones, semana corrida, '
         'descuentos.')
    paso(doc, 2, 'EMITIR', 'Congela el mes, envía cada liquidación por correo y alimenta el LRE, el '
         'F29 (código 048) y la contabilidad.')
    paso(doc, 3, 'Pagar', 'Botón Nómina Banco genera el CSV de transferencias en lote para el portal '
         'del banco. Botón Archivo Previred genera las cotizaciones con el movimiento de personal '
         'automático (contrataciones y retiros del mes). El LRE se descarga en Contabilidad → DT.')
    captura(doc, 'RRHH → Remuneraciones', 'El libro de un mes con la vista previa de una liquidación.')
    vineta(doc, 'las horas extras y la semana corrida salen de motores únicos (art. 32 y 45 CT); la base de las horas extras es el sueldo base.', bold_hasta='Motores: ')
    vineta(doc, 'si una operación comisionada se prepaga o anula, la reversa cae al mes siguiente, neta en la liquidación, con el detalle en Revisión de Comisiones (ver Manual del Ejecutivo, cap. 12).', bold_hasta='Reversas de comisión: ')
    vineta(doc, 'las variables del modelo tienen vigencia y bitácora inmutable: cada mes se calcula con la versión que regía ese mes; un mes cerrado no se toca.', bold_hasta='Variables versionadas: ')

    # ── 4. Vacaciones ───────────────────────────────────────────────────────
    h1(doc, '4. Vacaciones y su provisión')
    p(doc, 'La cuenta corriente de vacaciones es la única verdad: el formulario, los saldos de RRHH y '
           'el finiquito leen lo mismo. La deuda queda provisionada mes a mes en la contabilidad.')
    ficha(doc,
          'El colaborador solicita (FES) · su supervisor directo aprueba · RRHH queda notificado y '
          'actúa de respaldo',
          'Aprobación del "reporta a"; RRHH resuelve si no hay jefe definido',
          'Ficha con fecha de ingreso · años previos declarados con certificado AFP (feriado '
          'progresivo)',
          'RRHH → Vacaciones')
    paso(doc, 1, 'Devengo automático', 'Cada aniversario abona 15 días hábiles más el feriado '
         'progresivo si corresponde (art. 68: con +10 años trabajados, 1 día extra por cada 3 '
         'nuevos). Sin intervención de nadie.')
    paso(doc, 2, 'Solicitar', 'El formulario muestra los días disponibles (cuenta corriente + '
         'proporcional en curso) antes de pedir.')
    paso(doc, 3, 'Aprobar', 'El supervisor directo, avisado por campana; su decisión queda firmada '
         'FES. Nadie resuelve su propia solicitud. Al aprobar, los días se cargan solos descontando '
         'del período más antiguo (FIFO), y RRHH recibe la notificación para planilla y reemplazos.')
    paso(doc, 4, 'Provisión al cierre', 'Automática: snapshot del total y asiento por la variación '
         'contra el mes anterior (constitución o liberación).')
    captura(doc, 'RRHH → Vacaciones → cuenta corriente', 'La cuenta de un colaborador con devengos y usos.')
    flujo(doc, 'DEVENGO → SOLICITUD → APROBADA (CARGO FIFO) → PROVISIÓN AL CIERRE')

    # ── 5. Préstamos, anticipos y solicitudes ───────────────────────────────
    h1(doc, '5. Préstamos, anticipos y solicitudes paperless')
    p(doc, 'Dos puertas para lo mismo: RRHH registra directo (Descuentos) o el colaborador solicita y '
           'la cadena aprueba (Solicitudes). En ambos casos las cuotas se descuentan solas de la '
           'liquidación y el saldo se recupera hasta en el finiquito.')
    h2(doc, '5.1 Registro directo (RRHH → Remuneraciones → Descuentos)')
    paso(doc, 1, 'Registrar', 'Anticipo (capital ÷ N meses) o Préstamo (tasa mensual + cuotas, cuota '
         'francesa con vista previa). Las cuotas parten en la próxima remuneración. La tasa se valida '
         'contra la TMC vigente.')
    paso(doc, 2, 'Convenio firmado', 'Al guardar se imprime el convenio de descuento; el colaborador '
         'lo firma y RRHH lo sube a su carpeta. La columna Convenio queda en rojo "SIN FIRMA" hasta '
         'subirlo.')
    paso(doc, 3, 'Descuento automático', 'Cada cuota cae sola en "Otros Dcto." del libro. Anular '
         'detiene las cuotas futuras sin tocar lo descontado. Si la persona se va, el finiquito '
         'precarga el saldo pendiente.')
    h2(doc, '5.2 Solicitudes con cadena (RRHH → Solicitudes)')
    p(doc, 'Anticipos, préstamos, aumentos, ascensos y peticiones libres viajan por pasos: JEFATURA '
           '(el "reporta a") → GERENCIA (si aplica) → RRHH, cada aprobación firmada FES. En el paso '
           'RRHH se pueden ajustar condiciones (tasa tope TMC, cuotas). Las cadenas por tipo son '
           'paramétricas.')
    regla(doc, 'La ejecución es automática al aprobar el último paso: anticipo/préstamo genera el '
               'descuento + la ODP a nombre del colaborador + campana a Tesorería, y al pagarse el '
               'asiento va a cuentas por cobrar al personal (NO gasto). Aumento actualiza el sueldo '
               'base; ascenso, el cargo. Un rechazo llega con su motivo.')
    captura(doc, 'RRHH → Solicitudes → Por Aprobar', 'Una solicitud con su cadena de firmas FES.')
    flujo(doc, 'SOLICITADA → JEFATURA → GERENCIA → RRHH → APROBADA + EJECUTADA / RECHAZADA')

    # ── 6. Asistencia ───────────────────────────────────────────────────────
    h1(doc, '6. Asistencia y jornada (Workera)')
    p(doc, 'Las marcaciones del reloj control (biométrico + app) se cruzan con lo que el sistema sabe '
           '— ausencias, vacaciones, jornada, feriados — para detectar posibles faltas y atrasos '
           'VERIFICADOS antes de pagar remuneraciones.')
    ficha(doc,
          'RRHH revisa y configura · los colaboradores solo marcan · turnos y enrolamiento se '
          'administran en el panel de Workera',
          'rh_asistencia y rh_jornada',
          'Credenciales de la API cargadas · jornada configurada por colaborador · turnos asignados · '
          'feriados al día',
          'RRHH → Jornada · RRHH → Asistencia')
    paso(doc, 1, 'Configurar la jornada', 'Por colaborador: Art. 22 (excluido, no marca), 40 hrs '
         '(Ley 21.561), Especial (part-time) o Externo. Define quién debe marcar.')
    paso(doc, 2, 'Definir horarios y turnos', 'Días y entrada/salida fijas, o turnos rotativos '
         '(ciclos de N semanas). El Calendario de turnos se genera, se graba auditado y se imprime '
         'por persona.')
    paso(doc, 3, 'Revisar el mes', 'La página trae las marcas desde Workera y muestra solo a quienes '
         'deben marcar. Pestañas: Resumen, No enrolados, Atrasos (primera marca vs turno + tolerancia '
         'paramétrica) y Ausencias.')
    paso(doc, 4, 'Verificar antes de descontar', 'Una "posible falta" es un día hábil sin marca y sin '
         'ausencia/vacación que lo cubra. Se confirma con la persona; recién ahí se registra la '
         'ausencia o se descuenta.')
    captura(doc, 'RRHH → Asistencia', 'El resumen del mes con posibles faltas y atrasos.')
    vineta(doc, 'al aprobarse una vacación, licencia o permiso, el sistema crea la salida especial equivalente en Workera — el reloj deja de mostrar inasistencia esos días. Las injustificadas no se espejan.', bold_hasta='Espejo automático: ')
    advertencia(doc, 'El RUT es la llave en Workera y no se edita por API. Si el RUT de la suite '
                     'difiere del de Workera (por ejemplo, un extranjero que obtuvo RUT definitivo), '
                     'se mapea en la configuración — no se corrige allá.')
    flujo(doc, 'MARCA → SYNC API → CRUCE → POSIBLES FALTAS / ATRASOS → VERIFICACIÓN → REMUNERACIONES')

    # ── 7. Desempeño y Kuder ────────────────────────────────────────────────
    h1(doc, '7. Evaluación de desempeño y Test de Kuder')
    p(doc, 'Evaluar formalmente a cada colaborador por período — competencias + objetivos — dejando '
           'registro y feedback, sin planillas sueltas.')
    ficha(doc,
          'RRHH abre el ciclo · cada colaborador se autoevalúa · su jefatura evalúa y da el feedback',
          'Jefaturas asignadas ("reporta a" en la ficha)',
          'Catálogo de competencias revisado · peso competencias/objetivos definido para el ciclo',
          'RRHH → Desempeño')
    paso(doc, 1, 'Abrir el ciclo', 'Nombre, fechas y peso: se genera una evaluación por colaborador '
         'activo con jefatura (los sin jefatura se avisan). Campana y correo de apertura a todos.')
    paso(doc, 2, 'Autoevaluación PRIMERO', 'Nota 1-5 por competencia + comentario. La jefatura no '
         'puede finalizar sin esta autoevaluación (solo RRHH puede forzar la excepción).')
    paso(doc, 3, 'Evaluación de jefatura', 'Nota por competencia viendo la autoevaluación al lado, '
         'objetivos con peso y cumplimiento, fortalezas y oportunidades. Finalizar calcula la nota '
         '(70/30 configurable) y avisa al colaborador.')
    paso(doc, 4, 'Feedback y cierre', 'La jefatura conversa el resultado; el colaborador comenta si '
         'quiere y toma conocimiento — queda CERRADA y en su historial. El 360 opcional agrega la '
         '"Mirada 360" anónima por dimensión, sin alterar la nota.')
    captura(doc, 'RRHH → Desempeño', 'Una evaluación con autoevaluación y evaluación lado a lado.')
    flujo(doc, 'PENDIENTE → AUTOEVAL → EVALUADA → CERRADA')
    h2(doc, '7.1 Test de Kuder')
    p(doc, 'Mide intereses vocacionales en 10 áreas y los contrasta con el perfil esperado del cargo: '
           'en selección (link con token para el candidato, sin cuenta) y en la revisión anual (cada '
           'colaborador rinde desde su ficha). RRHH define el perfil por cargo y lee el match con sus '
           'brechas por área.')

    # ── 8. Capacitación ─────────────────────────────────────────────────────
    h1(doc, '8. Capacitación')
    p(doc, 'Cada curso queda registrado con horas, contenido y resultados, en la ficha de cada '
           'colaborador: evidencia para desempeño, onboarding y eventual franquicia SENCE.')
    paso(doc, 1, 'Crear el curso', 'RRHH → Cursos y Capacitaciones → Nuevo: nombre, fechas y horas '
         '(se calculan solas), relator, contenido, participantes, y si lleva nota (escala 1,0-7,0).')
    paso(doc, 2, 'Tomar asistencia y notas', 'En la bitácora, clic en el curso: check por persona y '
         'nota si corresponde.')
    paso(doc, 3, 'Consultar', 'Cada colaborador ve sus cursos y horas en Mi Ficha → Mis Cursos.')
    vineta(doc, 'cursos estilo "Flash" que enseñan a usar cada módulo de la suite, con progreso por persona. Su contenido se alimenta de la Ayuda de cada pantalla (botón "?"): mantener la ayuda actualizada mantiene la Academia actualizada. Tarea sugerida del onboarding.', bold_hasta='Academia AutoFácil (autocapacitación): ')
    captura(doc, 'RRHH → Cursos y Capacitaciones', 'La bitácora de cursos con asistencia tomada.')

    # ── 9. Clima ────────────────────────────────────────────────────────────
    h1(doc, '9. Encuestas de clima, pulso y eNPS')
    p(doc, 'Medir cómo está el equipo en forma anónima y periódica, para actuar sobre datos y no '
           'sobre rumores.')
    paso(doc, 1, 'Crear y abrir', 'RRHH → Encuestas de Clima: plantillas de Clima, Pulso, eNPS o '
         'Compromiso (14 preguntas estilo Q12 + intención de permanencia), editables, con fecha de '
         'cierre. Al abrir: campana a todos y recordatorio semanal automático a quienes falten.')
    paso(doc, 2, 'Responder', 'Un minuto, totalmente anónimo: solo se marca "ya respondió", sin '
         'vínculo con las respuestas.')
    paso(doc, 3, 'Leer y actuar', 'Promedios y distribución por pregunta, eNPS con semáforo, '
         'comentarios sin autor (ocultos si hay menos de 3 respuestas). El botón de informe IA '
         'redacta el resumen ejecutivo con los agregados anónimos. Comunicar hallazgos y medidas al '
         'equipo — si la gente no ve consecuencias, deja de responder — y cerrar.')
    captura(doc, 'RRHH → Encuestas de Clima → resultados', 'Los resultados agregados con el semáforo eNPS.')
    regla(doc, 'El anonimato es de diseño, no de confianza: las respuestas no guardan el usuario. '
               'Con menos de 3 respuestas los resultados ni se muestran.')

    # ── 10. Compliance ──────────────────────────────────────────────────────
    h1(doc, '10. Canal de Compliance')
    p(doc, 'Recibir e investigar denuncias de acoso, violencia, fraude e irregularidades en forma '
           'confidencial y dentro de los plazos legales (Ley Karin 21.643 y Ley 20.393).')
    ficha(doc,
          'Cualquier colaborador denuncia (anónimo o identificado) · solo los gestores designados '
          'investigan y resuelven',
          'compliance_gestionar (gestores de confianza, idealmente ajenos a la línea denunciada)',
          'Gestores definidos en la matriz de permisos',
          'RRHH → Canal de Compliance')
    paso(doc, 1, 'Denunciar', 'Tipo, contra quién (opcional) y el relato. Checkbox anónima: no se '
         'guarda el usuario. Se entrega un código de seguimiento — es la ÚNICA llave para consultar; '
         'guardarlo.')
    paso(doc, 2, 'Recibir', 'Correo y campana a los gestores con el tipo y el plazo. El relato solo '
         'se lee dentro del canal. Plazo de investigación: 30 días hábiles para todas.')
    paso(doc, 3, 'Investigar y resolver', 'EN INVESTIGACIÓN con comentarios internos o visibles al '
         'denunciante. Las vencidas se marcan y alegan semanalmente. RESUELTA o DESESTIMADA exige '
         'registrar la resolución, visible al consultar el código.')
    regla(doc, 'Si es Ley Karin: además del canal, el protocolo legal — medidas de resguardo '
               'inmediatas e informe a la Dirección del Trabajo cuando corresponda, con asesoría '
               'legal.')
    captura(doc, 'RRHH → Canal de Compliance', 'El formulario de denuncia con la opción anónima.')
    flujo(doc, 'RECIBIDA → EN INVESTIGACIÓN → RESUELTA / DESESTIMADA (30 DÍAS HÁBILES)')

    # ── 11. Desvinculación ──────────────────────────────────────────────────
    h1(doc, '11. Desvinculación: el finiquito')
    p(doc, 'Terminar la relación laboral pagando exactamente lo que corresponde por ley, con '
           'documento listo para ratificar y todo el cierre administrativo automático.')
    ficha(doc,
          'RRHH calcula y ajusta · Gerencia decide la causal · se ratifica ante ministro de fe',
          'rh_remuneraciones',
          'Ficha al día · cuenta de vacaciones cuadrada · liquidaciones emitidas (base del cálculo)',
          'RRHH → Contratos y Cartas Oferta → Finiquitos')
    paso(doc, 1, 'Baja inmediata en Usuarios', 'Al renunciar o ser desvinculada, la persona se '
         'suspende DE INMEDIATO en Usuarios (seguridad), con su fecha de baja. Sigue apareciendo en '
         'el libro del mes (liquidación proporcional en 30avos) y en el selector de Finiquitos.')
    paso(doc, 2, 'Calcular', 'Colaborador + causal + fecha de término: el motor propone indemnización '
         'por años (topes 11 años / 90 UF), mes de aviso y feriado proporcional desde la cuenta de '
         'vacaciones.')
    paso(doc, 3, 'Ajustar', 'Cada monto es editable; se agregan otros haberes y se descuentan '
         'anticipos y préstamos pendientes (el saldo viene precargado).')
    paso(doc, 4, 'Guardar', 'El cierre administrativo es automático: checklist de offboarding, '
         'asiento de devengo del finiquito, anulación de las cuotas futuras de sus créditos internos, '
         'suspensión si seguía activo, y la ODP del pago con campana a Tesorería — pagar TRAS la '
         'ratificación.')
    paso(doc, 5, 'Ratificar', 'Ante notario, inspección del trabajo o el finiquito electrónico de la '
         'DT. La firma FES interna deja el acuerdo, pero el poder liberatorio exige la ratificación.')
    captura(doc, 'RRHH → Finiquitos', 'Un finiquito calculado con su detalle por concepto.')
    regla(doc, 'El aviso a Previred es automático: el archivo del último mes informa el retiro '
               '(movimiento 2). Si no se informa, la AFP asume que la persona sigue trabajando y '
               'presume deuda previsional. Solo si el término se hizo FUERA del módulo hay que '
               'marcarlo a mano en previred.com.')
    flujo(doc, 'CÁLCULO → AJUSTE → GUARDADO → RATIFICADO Y PAGADO')

    # ── Anexos ──────────────────────────────────────────────────────────────
    h1(doc, 'Anexo A. Síntomas frecuentes y su causa')
    tabla(doc, ('Síntoma', 'Causa probable', 'Qué hacer'),
          (('No puedo contratar a un candidato', 'La carta oferta no está ACEPTADA', 'Marcar la respuesta primero (cap. 2)'),
           ('El sistema alega a diario por un cargo', 'Cargo sin descripción', 'Completarla en Contratos → Cargos (cap. 2)'),
           ('No puedo emitir el libro', 'Comisiones del mes sin aprobar', 'Comisiones → Revisión; el aviso rojo lo dice (cap. 3)'),
           ('Una persona aparece con posible falta', 'Día sin marca y sin ausencia que lo cubra', 'Verificar con la persona antes de descontar (cap. 6)'),
           ('Workera muestra inasistencia en vacaciones', 'El tipo no tiene código mapeado para el espejo', 'Configurar el mapeo (cap. 6)'),
           ('La jefatura no puede finalizar una evaluación', 'Falta la autoevaluación', 'Es a propósito; solo RRHH fuerza la excepción (cap. 7)'),
           ('Un descuento aparece "SIN FIRMA"', 'Convenio no subido a la carpeta', 'Imprimir, firmar y subir (cap. 5)'),
           ('La AFP cobra cotizaciones de un desvinculado', 'El retiro no se informó a Previred', 'Solo pasa si el término se hizo fuera del módulo (cap. 11)')),
          (5.5, 5.7, 5.3))

    h1(doc, 'Anexo B. Glosario')
    tabla(doc, ('Término', 'Significado'),
          (('FES', 'Firma electrónica simple, con identidad, fecha, IP y huella SHA-256'),
           ('Carta oferta', 'Fuente única del contrato: lo ofertado es lo contratado'),
           ('Onboarding / offboarding', 'Checklists automáticos de alta y salida'),
           ('Cuenta corriente de vacaciones', 'Registro único de devengos y usos; FIFO al descontar'),
           ('Feriado progresivo', 'Art. 68 CT: +10 años trabajados, 1 día extra por cada 3 nuevos'),
           ('Semana corrida', 'Art. 45 CT; motor único, depende de domingos y festivos del mes'),
           ('LRE', 'Libro de Remuneraciones Electrónico de la DT'),
           ('Movimiento de personal', 'Contrataciones (1) y retiros (2) informados a Previred'),
           ('Posible falta', 'Día hábil sin marca y sin ausencia/vacación aprobada que lo cubra'),
           ('Espejo Workera', 'Salida especial creada automáticamente en el reloj al aprobar permisos'),
           ('Q12 / eNPS', 'Formatos de encuesta de compromiso y recomendación'),
           ('Ley Karin', 'Ley 21.643: acoso y violencia laboral; plazos y protocolo obligatorios'),
           ('Ratificación', 'Firma ante ministro de fe que da poder liberatorio al finiquito')),
          (4.6, 11.9))

    h1(doc, 'Anexo C. Capturas pendientes de esta versión')
    p(doc, 'Recorrer las pantallas en este orden y reemplazar cada recuadro gris:')
    tabla(doc, ('N°', 'Pantalla', 'Qué debe mostrar'),
          tuple((str(n), pant, det) for n, pant, det in estilo.CAPTURAS),
          (1.2, 6.3, 9.0))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Manual-RRHH-Business-Suite.docx'
    d.save(out)
    print('OK ->', out, '| capturas:', len(estilo.CAPTURAS))
