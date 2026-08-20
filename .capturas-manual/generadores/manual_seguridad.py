# -*- coding: utf-8 -*-
"""Seguridad y Contingencia — Business Suite."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'SEGURIDAD Y CONTINGENCIA', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=26)
    p(doc, 'Quién entra, qué puede hacer, qué queda registrado,', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'y qué pasa si algo se cae.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · 17 de agosto de 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz. Distribución restringida.', size=10,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, '1. El modelo: capas que se respaldan entre sí')
    p(doc, 'La seguridad del sistema no depende de una sola barrera sino de capas: quién entra '
           '(autenticación), qué puede hacer (permisos), cómo se protege el proceso (validaciones y '
           'firmas), qué queda registrado (auditoría) y qué pasa si algo falla (contingencia). Un '
           'atacante — o un error honesto — tiene que atravesarlas todas.')
    regla(doc, 'Regla de este documento: acá se describe la ARQUITECTURA de seguridad, nunca los '
               'secretos. Las cuentas especiales, códigos de respaldo y credenciales no se nombran '
               'en ningún documento escrito — quien deba conocerlos, los conoce por otro canal.')

    # ── 2. Seguridad de acceso ──────────────────────────────────────────────
    h1(doc, '2. Seguridad interna: quién entra y qué puede hacer')
    h2(doc, '2.1 Autenticación')
    vineta(doc, 'longitud mínima, complejidad, historial, vencimiento y aviso previo — todo paramétrico en Usuarios → Seguridad y aplicado EN EL SERVIDOR, no como sugerencia de pantalla.', bold_hasta='Política de claves: ')
    vineta(doc, 'tras N intentos fallidos la cuenta se bloquea (paramétrico); el Administrador desbloquea. El login además tiene tope de intentos por minuto por IP.', bold_hasta='Bloqueo por intentos: ')
    vineta(doc, 'el token de sesión vive en la pestaña: cerrarla cierra la sesión. La inactividad también la cierra (minutos paramétricos).', bold_hasta='Sesiones cortas: ')
    vineta(doc, 'existe una cuenta protegida de emergencia para recuperar el acceso si todo lo demás falla. Deliberadamente, ni su nombre ni su resguardo se documentan.', bold_hasta='Acceso de emergencia: ')
    vineta(doc, 'activa en las cuentas que administran la infraestructura; la extensión a todo el equipo (~40 personas) está planificada con aviso y acompañamiento.', bold_hasta='Verificación en dos pasos: ')
    h2(doc, '2.2 Permisos: la matriz manda')
    p(doc, 'El acceso se gobierna por la matriz de Perfiles y Funcionalidades: cada card, botón y '
           'API sensible valida su código de permiso contra la base — con la regla de nunca '
           'hardcodear nombres de perfil en el código. El Administrador tiene bypass; todo el resto '
           've exactamente lo marcado. La visibilidad de datos también se acota: un ejecutivo ve '
           'SOLO sus ejecutivos asignados, un dealer SOLO su cartera, un cliente SOLO su RUT.')
    h2(doc, '2.3 Las fronteras: staff, dealers y clientes')
    p(doc, 'Los tres mundos usan tokens firmados con tipos distintos, y la frontera está en el '
           'middleware: un token de dealer es RECHAZADO por los endpoints internos del staff, y '
           'todo endpoint del portal acota por la identidad del token — el dato de quién pregunta '
           'sale del token, jamás del request. La auditoría de esta frontera (julio 2026) terminó '
           'con cero hallazgos de fuga entre dealers.')
    vineta(doc, 'los enrolamientos responden siempre lo mismo exista o no la cuenta, para no revelar qué RUT está registrado.', bold_hasta='Anti-enumeración: ')
    vineta(doc, 'los links de acceso directo son llaves portadoras regenerables y se guardan hasheados; los tokens de recuperación son de un solo uso y expiran.', bold_hasta='Llaves portadoras: ')
    h2(doc, '2.4 Seguridad perimetral')
    vineta(doc, 'todo el tráfico va cifrado, con HSTS de un año.', bold_hasta='HTTPS forzado: ')
    vineta(doc, 'TODA el API tiene techo de peticiones por usuario (600/min general, más estricto en reportería y contabilidad); el login, por IP. La oficina sale por una sola IP pública, por eso el límite general es por usuario.', bold_hasta='Rate limiting: ')
    vineta(doc, 'las credenciales de integraciones viven como variables de entorno del servidor; el repositorio no contiene ningún secreto y el arranque exige las claves críticas.', bold_hasta='Secretos fuera del código: ')

    # ── 3. Seguridad de los procesos ────────────────────────────────────────
    h1(doc, '3. Seguridad de los procesos: el flujo protegido')
    p(doc, 'La parametrización abre el CONTENIDO del negocio, nunca su ESTRUCTURA: el orden de las '
           'etapas, las validaciones y las atribuciones están protegidas en el servidor.')
    tabla(doc, ('Control', 'Dónde aplica', 'Qué impide'),
          (('Doble firma', 'Anulaciones, castigos, aplicaciones de fondos, incorporaciones con excepción', 'Que una sola persona deje sin efecto plata comprometida'),
           ('Segregación de funciones', 'Todas las órdenes de pago; solicitudes vs aprobaciones', 'Que quien emite también pague; que alguien se apruebe a sí mismo (bloqueo por usuario, en el servidor)'),
           ('Mes cerrado (candado)', 'Toda la operación retroactiva', 'Que un mes ya informado y pagado cambie después'),
           ('Inmutabilidad del crédito otorgado', 'Calendario, tasa y comisión congelados', 'Que la historia financiera se reescriba'),
           ('Snapshot deliberado', 'Cartolas aprobadas, finiquitos, códigos de excepción', 'Que el documento pagado cambie porque el origen avanzó'),
           ('Validación de escritura', 'Operaciones críticas revisan las filas afectadas', 'Que un cambio falle en silencio y se descubra semanas después'),
           ('Motor único', 'Todos los cálculos de plata', 'Que dos pantallas muestren números distintos del mismo hecho')),
          (4.0, 6.2, 6.3))

    # ── 4. Auditoría ────────────────────────────────────────────────────────
    h1(doc, '4. Auditoría: lo incorruptible')
    p(doc, 'La pregunta "¿quién hizo esto y cuándo?" siempre tiene respuesta:')
    vineta(doc, 'las acciones críticas quedan registradas con usuario, fecha, detalle e IP, en el módulo Auditoría — incluyendo la bitácora separada de los dealers (accesos, qué miraron, qué preguntaron).', bold_hasta='Registro de acciones: ')
    vineta(doc, 'la bitácora del Revisor Automático, las versiones de variables de comisiones y del Bono Jefe, y los registros de firmas SOLO AGREGAN FILAS: no existe editar ni borrar desde el sistema. Son los registros que un auditor externo puede tomar como verdad.', bold_hasta='Bitácoras inmutables: ')
    vineta(doc, 'cada firma electrónica registra identidad de sesión, nombre, cargo, fecha, IP y la huella SHA-256 del contenido: se puede demostrar que el documento no se alteró después de firmado.', bold_hasta='Firmas con huella: ')
    vineta(doc, 'cartas, certificados y checklists llevan QR de verificación pública: cualquiera valida que el documento es auténtico y está vigente, sin entrar al sistema.', bold_hasta='Documentos verificables: ')
    vineta(doc, 'todo correo del sistema queda en Auditoría → Correos Enviados.', bold_hasta='Correo trazado: ')
    vineta(doc, 'la consulta directa a la base es de SOLO LECTURA, restringida al perfil más alto, y cada consulta queda auditada.', bold_hasta='Consola SQL: ')
    captura(doc, 'Auditoría', 'El log de acciones con sus filtros y la pestaña de dealers.')

    # ── 5. Workflows ────────────────────────────────────────────────────────
    h1(doc, '5. Workflows y escalamiento: nada se queda dormido')
    p(doc, 'Un control que nadie mira no controla nada. Por eso lo pendiente ESCALA solo:')
    vineta(doc, 'quince flujos de escalamiento vigilan lo que espera decisión (aprobaciones, tickets, solicitudes): pasado el plazo en horas hábiles, el caso sube a la jefatura, con su registro en la bitácora de workflows.', bold_hasta='Escalamiento automático: ')
    vineta(doc, 'las cadenas de aprobación (planes de liquidez, solicitudes RRHH, compras) tienen plazos por nivel y auto-avance al vencer, para que un aprobador ausente no congele el negocio.', bold_hasta='Cadenas con plazo: ')
    vineta(doc, 'campanitas y correos paramétricos avisan a los responsables; las denuncias vencidas alegan semanalmente; el checklist de cierre manda correo diario con los puntos vencidos.', bold_hasta='Avisos que insisten: ')
    vineta(doc, 'los 28 motores automáticos pasan por un programador central con interruptor: se ven, se apagan por ambiente y dejan log. Un reloj suelto e invisible no existe.', bold_hasta='Motores gobernados: ')

    # ── 6. Contingencia ─────────────────────────────────────────────────────
    h1(doc, '6. Redundancia y contingencia: cuando algo se cae')
    h2(doc, '6.1 Los datos')
    tabla(doc, ('Qué', 'Respaldo', 'Retención'),
          (('Base de datos', 'Respaldo automático del proveedor + volcado nocturno propio (02:17) a artefacto privado', '1 día + 30 días'),
           ('Documentos', 'Bucket de Google Cloud, fuera de la base', 'Permanente'),
           ('Código', 'Repositorio Git con historial completo', 'Permanente'),
           ('Configuración', 'Vive en la base (mantenedores) → viaja con el respaldo', '30 días')),
          (3.6, 8.6, 4.3))
    h2(doc, '6.2 El servicio')
    p(doc, 'Si el servidor principal cae por horas, existe un host de contingencia DORMIDO en otra '
           'nube, apuntando a la misma base: se promueve con un comando siguiendo el runbook, y el '
           'arranque en frío toma segundos. La pieza clave es el interruptor central de motores: el '
           'standby atiende peticiones pero NO ejecuta ninguno de los 28 motores automáticos.')
    regla(doc, 'NUNCA encender los motores del standby con el principal vivo: dos procesos contra la '
               'misma base ejecutan cada tarea dos veces y el daño es silencioso — una comisión '
               'aprobada dos veces no se queja. El orden de la vuelta atrás (primero apagar los '
               'motores del standby) está en el runbook.')
    p(doc, 'La base tiene además su propia contingencia (instancia detenida en otra nube), y el '
           'ensayo completo del plan se hizo en agosto 2026 contra el ambiente de pruebas. La '
           'degradación es deliberada: promovido, funciona toda la operación del negocio; las '
           'integraciones externas (IA, WhatsApp, informes, SII) quedan apagadas hasta cargar sus '
           'credenciales según el manual.')
    h2(doc, '6.3 La vigilancia')
    vineta(doc, 'ping de servicios cada 5 minutos con historial por servicio.', bold_hasta='Uptime: ')
    vineta(doc, 'cada error 500 en producción manda correo al administrador (con freno de 10 minutos por ruta) — el sistema avisa antes que el usuario reclame.', bold_hasta='Alertas de error: ')
    vineta(doc, 'el endpoint de salud reporta base, uptime, documentos y motores apagados; el badge de versión confirma qué build corre.', bold_hasta='Salud: ')
    vineta(doc, 'un vigía diario revisa la coherencia de los relojes del sistema (zona horaria, fechas futuras) — la clase de error que descuadra sin ruido.', bold_hasta='Vigía de relojes: ')

    # ── 7. Resumen ──────────────────────────────────────────────────────────
    h1(doc, '7. El resumen para un auditor')
    tabla(doc, ('Pregunta del auditor', 'Respuesta del sistema'),
          (('¿Quién puede hacer qué?', 'Matriz de perfiles y funcionalidades, administrable y auditada; visibilidad acotada por pertenencia'),
           ('¿Puede una persona sola mover plata?', 'No en lo crítico: doble firma y segregación de funciones en el servidor'),
           ('¿Se puede reescribir la historia?', 'No: mes cerrado, créditos otorgados inmutables, snapshots y bitácoras que solo agregan'),
           ('¿Cómo sé que este documento es auténtico?', 'QR de verificación pública + huella SHA-256 de la firma'),
           ('¿Qué pasa si el servidor se cae?', 'Standby ensayado con runbook; respaldos diarios con 30 días de retención'),
           ('¿Y si nadie mira un pendiente?', 'Escala solo por jefatura en horas hábiles'),
           ('¿Dónde están las claves?', 'En el servidor como variables de entorno; jamás en código ni documentos')),
          (6.0, 10.5))

    return doc

if __name__ == '__main__':
    d = construir()
    out = r'C:\Users\patri\Documents\Seguridad-Contingencia-Business-Suite.docx'
    d.save(out)
    print('OK ->', out)
