# -*- coding: utf-8 -*-
"""Parte I: portada, cómo usar el manual, ciclo de vida del crédito y la ETAPA."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from estilo import *

def agregar(doc):
    # ── Portada ──────────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'MANUAL DE OPERACIONES', bold=True, color=AZUL_OSCURO, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'AutoFácil Business Suite', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'De la carga de la producción al pago de la comisión:', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'todos los procesos del área, paso a paso y con sus reglas.', size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.0 · Agosto 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'Documento interno — AutoFácil Crédito Automotriz', size=10,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    # ── Control de versiones ─────────────────────────────────────────────────
    h1(doc, 'Control de versiones')
    tabla(doc, ('Versión', 'Fecha', 'Autor', 'Cambios'),
          (('1.0', 'Agosto 2026', 'Business Suite', 'Emisión inicial del tomo Operaciones'),),
          (2.2, 3.2, 4.0, 7.1))
    h2(doc, 'Cómo usar este manual')
    p(doc, 'Este tomo cubre los procesos del área de Operaciones de principio a fin. Cada capítulo sigue '
           'la misma estructura: para qué existe el proceso, quién lo ejecuta y con qué permiso, qué debe '
           'estar listo antes de partir, el paso a paso con la pantalla de cada acción, las validaciones '
           'que el sistema aplica y los errores que se pueden encontrar en el camino.')
    p(doc, 'Los recuadros de color tienen un significado fijo en todo el documento:')
    vineta(doc, ' recuadro ámbar: una situación que confunde o un dato que engaña. Leerlo evita el error más común del capítulo.', bold_hasta='⚠ OJO —')
    vineta(doc, ' recuadro rojo: algo que el sistema impone y no depende de la configuración ni del criterio del usuario.', bold_hasta='🔒 Regla del sistema —')
    vineta(doc, ' recuadro verde: el episodio que originó la regla, con montos y operaciones reales. Explica el porqué.', bold_hasta='🧾 Caso real —')
    vineta(doc, ' recuadro gris: marca el lugar exacto donde va el pantallazo de esa pantalla. En esta versión están pendientes de captura.', bold_hasta='📸 CAPTURA —')
    p(doc, 'Las rutas de pantalla se escriben con el símbolo →. Por ejemplo, "Post Venta → Cartolas" '
           'significa: entrar al módulo Post Venta desde el inicio y ahí abrir la card Cartolas.')
    advertencia(doc, 'Lo que se ve en pantalla depende del perfil. Si una card o un botón de este manual '
                     'no aparece, no es un error del sistema: el perfil no tiene el permiso. El anexo A '
                     'lista los permisos de cada proceso para pedirlos al Administrador.')

    # ── Índice ───────────────────────────────────────────────────────────────
    h1(doc, 'Índice')
    toc(doc)

    # ── Cap 1: ciclo de vida ─────────────────────────────────────────────────
    h1(doc, '1. El crédito y su ciclo de vida')
    p(doc, 'Todo lo que hace Operaciones gira en torno a un mismo objeto: la operación de crédito. '
           'Entender su ciclo de vida completo —de cotización a comisión pagada— es el mapa que ordena '
           'el resto del manual: cada capítulo es un tramo de este camino.')
    h2(doc, '1.1 El camino feliz')
    flujo(doc, 'COTIZACIÓN → EVALUACIÓN → CARTA APROBADA → OTORGADO → FUNDANTES → CARTOLA → FACTURA → COMISIÓN PAGADA')
    p(doc, 'En palabras: el ejecutivo cotiza y arma la carta de aprobación; el analista (o el Revisor '
           'Automático) la aprueba; al otorgarse nace el crédito con todo congelado; los fundantes '
           'respaldan el curse para que la financiera libere los fondos; la comisión del dealer entra a '
           'su cartola mensual; el dealer factura el total de la cartola; y la orden de pago cierra el '
           'circuito. Cada flecha de ese camino tiene su capítulo.')
    captura(doc, 'Mantenedores → Estado Créditos', 'El mapa de estados brokerage con sus transiciones, tal como está configurado.')
    h2(doc, '1.2 Los desvíos')
    p(doc, 'No toda operación recorre el camino completo. Los desvíos también son procesos formales:')
    vineta(doc, 'la financiera no aprueba. La operación queda en la estadística de conversión y puede intentarse en la otra financiera con la regla del 60% (capítulo 11).', bold_hasta='RECHAZADO: ')
    vineta(doc, 'el cliente se baja. Si la carta estaba aprobada, vence sola a los 5 días (configurable) o se desiste a mano desde Cartas Vigentes.', bold_hasta='DESISTIDO: ')
    vineta(doc, 'la operación ya cursada se deja sin efecto. Es el único desvío con doble firma y retiro de comisión, porque hay plata comprometida (capítulo 15).', bold_hasta='ANULADO: ')
    p(doc, 'Las dos financieras del brokerage son AutoFin y Unidad de Crédito (UAC). La operación es '
           'siempre del negocio AutoFácil: la financiera define quién fondea, no de quién es el cliente.')
    h2(doc, '1.3 Dos dimensiones que no hay que confundir')
    p(doc, 'Una operación tiene ETAPA (dónde va en el camino: DIGITADO, APROBADO, OTORGADO, ANULADO…) y, '
           'si es de cartera propia, además tiene ESTADO DE CARTERA (cómo paga: VIGENTE, EN MORA, '
           'VENCIDO, PREPAGADO, CASTIGADO). Son dimensiones independientes: un crédito EN MORA está, en '
           'etapa, OTORGADO. Las pantallas de Operaciones filtran casi siempre por etapa; las de '
           'Cobranza, por estado de cartera.')

    # ── Cap 2: la ETAPA ──────────────────────────────────────────────────────
    h1(doc, '2. La ETAPA: el dato del que cuelga todo')
    p(doc, 'La etapa decide si una operación cuenta como venta, si entra a cartolas, si genera comisión, '
           'si se le exigen fundantes y si aparece en Post Venta. Que la etapa esté mal no es un error '
           'cosmético: descuadra la plata. Por eso este capítulo va antes que los procesos.')
    ficha(doc,
          'Nadie la escribe "a mano" como acto propio: la fija el proceso dueño de cada paso — otorgar '
          'una carta, anular con doble firma, marcar NO OTORGADO, la carga Trinidad o el desistimiento '
          'automático.',
          'El del proceso que corresponda en cada caso',
          'Que el cambio venga del proceso dueño del paso, respetando mes cerrado',
          'Transversal — no tiene pantalla propia')
    h2(doc, '2.1 El gotcha de las tres columnas')
    p(doc, 'Por herencia histórica la etapa vive en TRES columnas de la tabla de créditos: estado, '
           'estado_credito y estado_eval. Cuando un proceso escribía solo una o dos, la operación '
           'quedaba con la etapa "partida" y cada pantalla veía una verdad distinta.')
    caso(doc, 'La OP 89343 se mostraba otorgada en el listado pero quedaba fuera de Vendedores con '
              'Ventas, de fundantes y de cartolas. La OP 88558 estaba anulada y se seguía mostrando '
              'OTORGADA, porque la columna estado se lee primero. Ambas tenían la etapa partida entre '
              'las tres columnas.')
    p(doc, 'La solución es un motor único: todo proceso escribe la etapa por una sola función que toca '
           'las tres columnas en el mismo acto, y toda pantalla la lee por una sola expresión canónica '
           '(la columna estado manda cuando está poblada; si no, se cae a las otras dos). Así ninguna '
           'pantalla inventa su propio criterio.')
    regla(doc, 'Si en una revisión de datos aparece una operación donde el listado dice una etapa y otra '
               'pantalla dice otra, NO corregirla editando el estado a mano: eso vuelve a partir la '
               'etapa. Reportarla a TI — existe un control permanente (desalineados) que lista las '
               'operaciones partidas, y lo que aparece ahí indica un proceso mal cableado que hay que '
               'arreglar de raíz.')
    h2(doc, '2.2 Dos cosas que NO son un error')
    vineta(doc, 'En cartera propia la columna estado también guarda VIGENTE / EN MORA / VENCIDO / PREPAGADO / CASTIGADO. Describe el pago, no la etapa, y el motor nunca pisa esas palabras.', bold_hasta='ETAPA ≠ ESTADO DE CARTERA. ')
    vineta(doc, 'Las ~16.000 solicitudes de la carga masiva diaria de AutoFin viven con estado vacío a propósito: nunca fueron originación propia. Vacío no es discrepancia.', bold_hasta='El vacío deliberado. ')
    advertencia(doc, 'La base de datos distingue mayúsculas de minúsculas: "Digitado" y "DIGITADO" son '
                     'valores DISTINTOS. Los estados se escriben siempre en MAYÚSCULAS. Una venta se '
                     'perdió por esto: la OP 26080532 quedó con la carta otorgada y el crédito en '
                     '"Digitado", invisible en dashboard y comisiones.')
