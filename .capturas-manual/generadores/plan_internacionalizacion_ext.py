# -*- coding: utf-8 -*-
"""Plan de Internacionalización (versión para el Grupo) — compartible."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import estilo
from estilo import *

def construir():
    estilo.CAPTURAS.clear()
    doc = nuevo_doc()

    # ── Portada ──
    for _ in range(4): doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(LOGO, width=Cm(7))
    p(doc, '', despues=18)
    p(doc, 'PLAN DE INTERNACIONALIZACIÓN', bold=True, color=AZUL_OSCURO, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, 'Business Suite → Ecuador y Bolivia', color=AZUL, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, '', despues=30)
    p(doc, 'Qué viaja tal cual, qué se adapta y qué se apaga — con arquitectura por nodo,', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    p(doc, 'costos de operación, tiempos con holgura y el levantamiento previo por país.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
    for _ in range(6): doc.add_paragraph()
    p(doc, 'Versión 1.1 · Septiembre 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)
    p(doc, 'AutoFácil Crédito Automotriz — documento de trabajo para el Grupo', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS_SUAVE)

    h1(doc, 'Índice')
    toc(doc)

    # ── 1. Resumen ejecutivo ──
    h1(doc, '1. Resumen ejecutivo')
    p(doc, 'Business Suite fue construida bajo un principio rector: TODO lo que es dato de negocio vive en '
           'mantenedores, no en el código. Esa decisión, sostenida durante todo el desarrollo, es lo que hace '
           'viable llevarla a otros países: el corazón operativo (créditos, cartas, seguimiento de pagos, '
           'comisiones, contabilidad, órdenes de pago, permisos, workflows) es país-agnóstico.')
    vineta(doc, 'un fork del sistema por empresa, con su propia base de datos, servidor y dominio. Cero riesgo para la operación chilena.', bold_hasta='Arquitectura: ')
    vineta(doc, 'Ecuador primero (economía dolarizada = menos adaptación) como piloto; Bolivia después reusando el trabajo multipaís.', bold_hasta='Secuencia: ')
    vineta(doc, 'Ecuador operando en ~4 a 5 meses desde el kickoff (incluye levantamiento de 4 semanas y marcha blanca); Bolivia +2,5 a 3,5 meses adicionales.', bold_hasta='Tiempos con holgura: ')
    vineta(doc, 'US$ 40–90 mensuales de infraestructura POR NODO (sin licencias de software: el sistema es propio). La inversión relevante es tiempo de levantamiento y adaptación, no plata.', bold_hasta='Costos: ')
    regla(doc, 'La condición de éxito no es técnica: es el LEVANTAMIENTO. El producto crediticio, la '
               'normativa, los impuestos y la contabilidad de cada país deben entenderse ANTES de tocar '
               'código. Por eso este plan parte con 3–4 semanas de levantamiento en terreno por país.')

    # ── 2. Arquitectura por nodo ──
    h1(doc, '2. Arquitectura: un nodo independiente por empresa')
    flujo(doc, 'REPO FORK (por país) → BD TiDB propia → Render propio → dominio propio → secretos propios')
    p(doc, 'Cada empresa del grupo recibe un NODO completo e independiente. Se descarta el multi-tenant '
           '(varias empresas en una misma base) por tres razones: aislamiento de datos entre sociedades, '
           'normativa distinta por país, y cero riesgo de que un cambio para Ecuador rompa Chile.')
    vineta(doc, 'mismo código base, en un repositorio propio por país. Las mejoras del core chileno se pueden portar al fork cuando convenga (cherry-pick), sin obligación de sincronizar.', bold_hasta='Repo fork: ')
    vineta(doc, 'TiDB Cloud (la misma tecnología probada en Chile) con su respaldo nocturno propio.', bold_hasta='Base de datos: ')
    vineta(doc, 'Render con auto-deploy desde el repo del país; dominio del tipo suite.empresa.com.ec.', bold_hasta='Servidor: ')
    vineta(doc, 'cada nodo tiene sus credenciales (BD, correo, JWT) — nada compartido con Chile.', bold_hasta='Secretos: ')
    advertencia(doc, 'El nodo nace con MOTORES configurables (el interruptor MOTORES=off ya existe): se '
                     'encienden motor por motor a medida que la operación local los valida. Nunca los 28 '
                     'de golpe en una operación que recién parte.')

    # ── 3. Qué viaja, qué se adapta, qué se apaga ──
    h1(doc, '3. Inventario por módulo: viaja / se adapta / se apaga')
    h2(doc, '3.1 Viaja tal cual (país-agnóstico)')
    tabla(doc, ('Módulo', 'Comentario'),
          (('Créditos y originación', 'Digitación, estados paramétricos, carga masiva (se ajusta el formato del archivo local)'),
           ('Cartas de Aprobación', 'Flujo, corrección -C1, QR y firma electrónica de documentos'),
           ('Seguimiento de pagos (Post Venta)', 'Saldos, comisiones, fundantes: etapas y reglas son configurables'),
           ('Contabilidad', 'Plan de cuentas y Reglas de Centralización 100% paramétricos; libro diario/mayor/balance'),
           ('Órdenes de Pago y proveedores', 'Correlativo, workflow, segregación de funciones'),
           ('Compras de Oficina', 'Ciclo mensual completo (catálogo local se siembra)'),
           ('Usuarios, perfiles y permisos', 'Matriz completa + auditoría + solo-lectura'),
           ('Workflows, avisos, correos', 'Todos paramétricos por diseño'),
           ('Dashboard y reportería', 'Tailor Made, tablas dinámicas, diseño de consulta'),
           ('Portal del Dealer y del Cliente', 'Self-service; textos y validaciones de identidad se localizan')),
          (5.5, 11.0))
    h2(doc, '3.2 Se adapta (trabajo real de localización)')
    tabla(doc, ('Ítem', 'Chile hoy', 'Ecuador', 'Bolivia'),
          (('Moneda', 'CLP + UF', 'USD (sin unidad reajustable: SIMPLIFICA)', 'BOB + tipo de cambio BCB'),
           ('Identificación', 'RUT + dígito verificador', 'Cédula / RUC (validador propio)', 'CI / NIT'),
           ('IVA', '19%', '15%', '13%'),
           ('Indicadores', 'API CMF (UF/UTM/TMC/dólar)', 'Se apaga; tasas máximas BCE', 'Tipo de cambio y tasas BCB/ASFI'),
           ('Tasa máxima', 'TMC CMF', 'Tasas máximas por segmento (Junta/BCE)', 'Regulación ASFI'),
           ('Facturación electrónica', 'SII (RCV, DTE)', 'SRI (comprobantes electrónicos)', 'SIN (SIAT)'),
           ('Feriados', 'Tabla Chile', 'Tabla Ecuador', 'Tabla Bolivia'),
           ('Formatos y textos', 'es-CL, $ puntos de miles', 'es-EC, USD', 'es-BO, Bs')),
          (3.4, 4.2, 4.5, 4.4))
    p(doc, 'La adaptación se hace UNA vez como "motor multipaís": un mantenedor País/Moneda/Impuestos del '
           'que leen todos los consumidores (Máxima 2: una sola fuente). Bolivia reusa ese trabajo.')
    h2(doc, '3.3 Se apaga (Chile-specific o de fase posterior)')
    vineta(doc, 'AFP, Isapre, semana corrida, Previred, jornada 40h — es lo MÁS chileno del sistema. El negocio de crédito no lo necesita el día uno. Cada país decide después si levanta su versión local.', bold_hasta='RRHH y Remuneraciones: ')
    vineta(doc, 'topes de gestiones de la Ley 21.320 — se reemplaza por la norma local de cobranza o se deja configurable.', bold_hasta='Cobranza (ley chilena): ')
    vineta(doc, 'DealerNet (buró Chile), SimpleAPI/SII, Workera (reloj de asistencia), integración CMF. Los burós locales (Equifax Ecuador, Infocred Bolivia) son fase 2 con sus propios contratos.', bold_hasta='Integraciones chilenas: ')

    # ── 4. Fase 0: levantamiento por país ──
    h1(doc, '4. Fase 0 — Levantamiento en terreno (la clave del plan)')
    p(doc, 'Antes de escribir una línea de código hay que entender el negocio y el producto de cada país. '
           'Este levantamiento es EL insumo del plan y su duración manda sobre todo lo demás.')
    tabla(doc, ('Semana', 'Foco', 'Entregable'),
          (('1', 'El producto crediticio: tipos de crédito, tasas, plazos, seguros, garantías, cómo se origina y aprueba hoy', 'Ficha del producto por tipo'),
           ('2', 'El flujo operativo: quién hace qué (organigrama), pagos a concesionarios, comisiones, cobranza', 'Mapa de procesos as-is'),
           ('3', 'Normativa e impuestos: regulador, tasa máxima, facturación electrónica, retenciones, libros exigidos', 'Matriz normativa'),
           ('4', 'Contabilidad local: plan de cuentas, políticas, cierres, reportes a casa matriz — y validación del gap con la Suite', 'Documento de brechas + alcance firmado')),
          (1.6, 9.5, 5.0))
    regla(doc, 'El levantamiento termina con un DOCUMENTO DE ALCANCE firmado con la gerencia local: qué '
               'entra en el MVP, qué queda fase 2 y quién es la contraparte operativa. Sin contraparte '
               'local comprometida, la implementación se estanca: la experiencia muestra que es el factor de éxito número uno.')

    # ── 5. Plan de trabajo y tiempos ──
    h1(doc, '5. Plan de trabajo y tiempos (con holgura)')
    h2(doc, '5.1 Ecuador (piloto)')
    tabla(doc, ('Fase', 'Duración', 'Contenido'),
          (('0. Levantamiento', '3–4 semanas', 'En terreno + remoto; documento de alcance'),
           ('1. Motor multipaís + nodo', '2–3 semanas', 'Fork, BD, dominio, mantenedor país/moneda/impuestos, validador de identidad, formatos'),
           ('2. Localización del core', '4–6 semanas', 'Productos, estados, plan de cuentas, reglas contables, correos, seeds; apagado de lo chileno'),
           ('3. UAT con la contraparte', '2–3 semanas', 'Casos reales punta a punta; ajustes'),
           ('4. Marcha blanca', '4 semanas', 'Operación real acompañada; motores encendiéndose uno a uno'),
           ('TOTAL Ecuador', '≈ 4 a 5 meses', 'Desde el kickoff hasta operación estable')),
          (4.3, 3.2, 8.6))
    h2(doc, '5.2 Bolivia (reusa el motor multipaís)')
    tabla(doc, ('Fase', 'Duración', 'Comentario'),
          (('0. Levantamiento', '3–4 semanas', 'Puede traslaparse con la marcha blanca de Ecuador'),
           ('1–3. Nodo + localización + UAT', '4–6 semanas', 'El motor multipaís ya existe; es sembrar y ajustar BOB/SIN/ASFI'),
           ('4. Marcha blanca', '4 semanas', ''),
           ('TOTAL Bolivia', '≈ 2,5 a 3,5 meses', 'Arrancando cuando Ecuador esté en marcha blanca')),
          (4.3, 3.4, 8.4))
    advertencia(doc, 'Las holguras están puestas a propósito: la experiencia dice que el levantamiento '
                     'SIEMPRE descubre reglas de negocio no documentadas, y que la disponibilidad de la '
                     'contraparte local es el cuello de botella real. Mejor prometer 5 meses y llegar en 4.')

    # ── 6. Costos ──
    h1(doc, '6. Costos')
    h2(doc, '6.1 Infraestructura mensual POR NODO (recurrente)')
    tabla(doc, ('Servicio', 'Costo mensual (USD)', 'Nota'),
          (('Render (servidor)', '7 – 25', 'Starter basta para partir; se sube según carga'),
           ('TiDB Cloud (base de datos)', '5 – 25', 'Serverless: paga por uso; la BD chilena completa pesa ~45 MB'),
           ('Bucket documentos (GCS)', '~1', 'Los documentos nunca van a la base'),
           ('Correo transaccional (Brevo)', '0 – 15', 'Plan gratis hasta 300 correos/día'),
           ('Dominio local', '~1–2 (prorrateo anual)', '.com.ec / .com.bo'),
           ('Respaldo GitHub Actions', '0', 'Mismo esquema del dump nocturno chileno'),
           ('TOTAL por nodo', 'US$ 40 – 90 / mes', 'Sin licencias: el software es del grupo')),
          (5.2, 4.2, 6.8))
    h2(doc, '6.2 Costos por una vez y variables')
    vineta(doc, 'no hay. El sistema es propio; no se paga licencia por instalarlo en otra empresa del grupo.', bold_hasta='Licencias: ')
    vineta(doc, 'pasajes y estadía de los viajes de levantamiento y puesta en marcha (Ecuador y Bolivia) — los define el grupo.', bold_hasta='Viajes: ')
    vineta(doc, 'IA (análisis, revisores): ~US$ 4–15/mes por nodo según uso, con el mismo tracking de costos en vivo que tiene Chile.', bold_hasta='Servicios opcionales: ')
    vineta(doc, 'integraciones locales fase 2 (buró de crédito, facturación electrónica SRI/SIN): cada una con su contrato local; se cotizan tras el levantamiento.', bold_hasta='Fase 2: ')
    h2(doc, '6.3 El costo real')
    p(doc, 'La inversión dominante es TIEMPO: el levantamiento, la localización y el acompañamiento de la '
           'marcha blanca. El desarrollo lo realiza el equipo que construyó la Suite, sin necesidad de contratar una fábrica de software externa; sí requiere dedicación protegida durante las fases 1–3 y '
           'una contraparte local con horas comprometidas.')

    # ── 7. Marco contable y regulatorio ──
    h1(doc, '7. Marco contable y regulatorio (IFRS / Basilea)')
    p(doc, 'El módulo de contabilidad de la Suite es AGNÓSTICO del marco normativo: el plan de cuentas, '
           'las reglas de centralización y las provisiones son 100% paramétricos. No hay que "convertir" '
           'el sistema a una norma: se siembra el plan local y sus reglas durante el levantamiento, con '
           'validación de un contador local.')
    h2(doc, '7.1 IFRS (NIIF) — norma de presentación contable')
    tabla(doc, ('País', 'Marco vigente', 'Implicancia para el nodo'),
          (('Chile', 'IFRS / IFRS-PyME', 'El plan de cuentas y las políticas actuales ya están en esta lógica'),
           ('Ecuador', 'IFRS (obligatorio, Superintendencia de Compañías)', 'El plan de cuentas chileno viaja casi directo — la mejor noticia contable del proyecto'),
           ('Bolivia', 'Normas contables locales (Consejo Técnico Nacional); si la empresa es entidad financiera regulada, manda el manual de cuentas de la ASFI', 'El plan de cuentas se siembra según la norma local; validación de contador boliviano en la semana 3–4 del levantamiento')),
          (2.2, 6.8, 7.5))
    h2(doc, '7.2 Basilea — regulación de capital (no de contabilidad)')
    p(doc, 'Basilea aplica a BANCOS y financieras REGULADAS, no a una empresa de crédito automotriz no '
           'bancaria. La pregunta clave del levantamiento es si la empresa de cada país está bajo '
           'supervisión del regulador financiero (ASFI en Bolivia, Superintendencia de Bancos en Ecuador):')
    vineta(doc, 'exigencias de capital mínimo, provisiones por riesgo de crédito NORMADAS por el regulador (tramos dados, no elegidos) y reportes periódicos obligatorios. El alcance del nodo crece y debe cotizarse en el levantamiento.', bold_hasta='Si está regulada: ')
    vineta(doc, 'las provisiones automáticas de la Suite (tramos de mora paramétricos) se calibran a la política interna del grupo, igual que en Chile.', bold_hasta='Si no está regulada: ')
    regla(doc, 'Las provisiones automáticas, el plan de cuentas y las reglas contables de la Suite son '
               'configurables por diseño: sirven bajo IFRS, bajo norma local boliviana o bajo tramos de un '
               'regulador — cambia la SIEMBRA, no el sistema.')

    # ── 8. Riesgos ──
    h1(doc, '8. Riesgos y mitigaciones')
    tabla(doc, ('Riesgo', 'Mitigación'),
          (('Contraparte local con poca disponibilidad', 'Documento de alcance firmado + contraparte nombrada ANTES de partir; hitos quincenales con la gerencia'),
           ('Reglas de negocio no documentadas que aparecen tarde', 'Holguras del plan + UAT con casos reales antes de la marcha blanca'),
           ('Normativa local subestimada (SRI/SIN/ASFI)', 'Semana 3 del levantamiento dedicada + asesor contable local por país'),
           ('Divergencia de código entre nodos', 'Core común documentado; mejoras se portan por cherry-pick deliberado, nunca automático'),
           ('Dependencia de una sola persona', 'Todo documentado bajo la Máxima 3 (los 5 documentos vivos por nodo) + manuales por rol'),
           ('Datos históricos locales', 'Migración tipo INDEXA/AFA ya probada en Chile: se cotiza en el levantamiento según el formato de origen')),
          (7.0, 9.5))

    # ── 9. Siguientes pasos ──
    h1(doc, '9. Siguientes pasos propuestos')
    paso(doc, 1, 'Acordar el alcance del viaje', 'Qué empresa parte (propuesta: Ecuador, por su economía dolarizada), fechas del levantamiento y quién será la contraparte local en cada país.')
    paso(doc, 2, 'Preparar el kit de levantamiento', 'Cuestionarios por área (producto, operación, normativa, contabilidad) para que las 3–4 semanas rindan.')
    paso(doc, 3, 'Demo formal al equipo local', 'Presentación de la Suite en operación, con datos de demostración, a las personas que la operarán en cada país.')
    paso(doc, 4, 'Firmar el documento de alcance del MVP', 'Y recién ahí crear el nodo y partir la localización.')

    return doc
